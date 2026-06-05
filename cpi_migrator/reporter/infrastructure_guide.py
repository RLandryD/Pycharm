"""
reporter/infrastructure_guide.py

Generates filled-in infrastructure setup guides for:
  1. IDoc — WE20 partner profile + WE21 port configuration
  2. RFC  — SM59 destination configuration sheet
  3. JDBC — driver upload + data source config
  4. cTMS — transport route blueprint Dev→QA→Prod

Output: One Excel workbook (one sheet per topic) + Word document
"""
from __future__ import annotations

import logging
from datetime import datetime
from pathlib import Path

logger = logging.getLogger(__name__)


class InfrastructureGuideGenerator:

    def __init__(self, output_dir: str):
        self.output_dir = Path(output_dir) / "infrastructure"
        self.output_dir.mkdir(parents=True, exist_ok=True)

    def generate(
        self,
        assessments: list,
        configs: dict,
        cpi_base_url: str = "",
        project_name: str = "CPI Migration",
        channels: list = None,
    ) -> tuple[Path, Path]:
        # Resolve real channel values (WE20/WE21/SM59 auto-fill). Optional:
        # when no channels are supplied this is an empty map and the sheets
        # fall back to their existing placeholder behaviour.
        autofill = self._build_autofill(assessments, channels)
        xlsx = self._generate_excel(assessments, configs, cpi_base_url, project_name, autofill)
        docx = self._generate_word(assessments, configs, cpi_base_url, project_name, autofill)
        return xlsx, docx

    @staticmethod
    def _build_autofill(assessments, channels) -> dict:
        """Map interface name -> InfraAutofill using channel data, if any."""
        if not channels:
            return {}
        try:
            from reporter.channel_autofill import ChannelAutofill
        except Exception as exc:  # pragma: no cover - defensive
            logger.warning("channel_autofill unavailable: %s", exc)
            return {}
        ca = ChannelAutofill(channels)
        out = {}
        for a in assessments:
            iface = a.interface
            out[iface.name] = ca.for_interface(
                iface.name,
                getattr(iface, "sender_adapter", ""),
                getattr(iface, "receiver_adapter", ""),
            )
        return out

    # ── Excel ─────────────────────────────────────────────────────────

    def _generate_excel(self, assessments, configs, cpi_base_url, project_name, autofill=None) -> Path:
        autofill = autofill or {}
        try:
            import openpyxl
            from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
            from openpyxl.utils import get_column_letter
        except ImportError:
            raise RuntimeError("openpyxl required")

        wb   = openpyxl.Workbook()
        thin = Side(style="thin")
        bdr  = Border(left=thin, right=thin, top=thin, bottom=thin)

        def hdr_cell(ws, row, col, val, colour="1F4E79"):
            c = ws.cell(row=row, column=col, value=val)
            c.fill = PatternFill("solid", fgColor=colour)
            c.font = Font(bold=True, color="FFFFFF")
            c.border = bdr
            c.alignment = Alignment(horizontal="center", wrap_text=True)

        def data_cell(ws, row, col, val, bold=False, colour=None):
            c = ws.cell(row=row, column=col, value=str(val) if val is not None else "")
            c.border = bdr
            c.alignment = Alignment(vertical="top", wrap_text=True)
            if bold:
                c.font = Font(bold=True)
            if colour:
                c.fill = PatternFill("solid", fgColor=colour)

        # ── Sheet 1: IDoc WE20/WE21 ───────────────────────────────────
        ws1 = wb.active
        ws1.title = "IDoc Setup (WE20-WE21)"

        idoc_interfaces = [
            (a, configs.get(a.interface.name))
            for a in assessments
            if a.interface.sender_adapter == "IDoc" or a.interface.receiver_adapter == "IDoc"
        ]

        ws1["A1"] = "IDoc Partner Profile & Port Configuration Guide"
        ws1["A1"].font = Font(bold=True, size=13, color="1F4E79")
        ws1["A2"] = f"Execute in SAP source/target system | Generated: {datetime.now().strftime('%Y-%m-%d')}"
        ws1["A2"].font = Font(italic=True, color="808080")

        # WE20 section
        ws1["A4"] = "STEP 1: Partner Profile (Transaction WE20)"
        ws1["A4"].font = Font(bold=True, size=11)

        we20_headers = ["Interface", "Partner Type", "Partner No.", "Message Type",
                        "IDoc Type", "Receiver Port", "Output Mode", "CPI Endpoint URL"]
        for col, h in enumerate(we20_headers, 1):
            hdr_cell(ws1, 5, col, h, "2E75B6")

        row = 6
        for a, cfg in idoc_interfaces:
            iface = a.interface
            af = autofill.get(iface.name)
            msg_type = ((af.idoc_message_type if af else "")
                        or (cfg.message.idoc_message_type if cfg else "")
                        or iface.message_interface or "[fill]")
            idoc_type = ((af.idoc_type if af else "")
                         or (cfg.message.idoc_type if cfg else "") or "ORDERS05")
            partner   = ((af.idoc_partner_number if af else "")
                         or (cfg.message.idoc_partner_profile if cfg else "")
                         or iface.receiver_system or "[fill]")
            cpi_url   = f"{cpi_base_url}/api/v1/idoc" if cpi_base_url else "[CPI_BASE_URL]/api/v1/idoc"

            row_data = [
                iface.name, "LS (Logical System)", partner,
                msg_type, idoc_type, f"CPI_{iface.name[:10]}_PORT",
                "Transfer IDoc Immed.", cpi_url,
            ]
            for col, val in enumerate(row_data, 1):
                data_cell(ws1, row, col, val)
            row += 1

        if not idoc_interfaces:
            ws1.cell(row=6, column=1, value="No IDoc interfaces detected in this migration.")

        # WE21 section
        ws1.cell(row=row + 1, column=1,
                 value="STEP 2: Port Definition (Transaction WE21)").font = Font(bold=True, size=11)

        we21_headers = ["Port Name", "Port Type", "RFC Destination / URL", "Service Path", "Logical Address"]
        for col, h in enumerate(we21_headers, 1):
            hdr_cell(ws1, row + 2, col, h, "2E75B6")

        r = row + 3
        for a, cfg in idoc_interfaces:
            iface   = a.interface
            cpi_url = cpi_base_url or "[CPI_BASE_URL]"
            row_data = [
                f"CPI_{iface.name[:10]}_PORT",
                "XML File Port (TRFC/SOAP)",
                f"{cpi_url}/api/v1/idoc",
                "/api/v1/idoc",
                cpi_url,
            ]
            for col, val in enumerate(row_data, 1):
                data_cell(ws1, r, col, val)
            r += 1

        for col, width in enumerate([30, 18, 50, 25, 40], 1):
            ws1.column_dimensions[get_column_letter(col)].width = width

        # ── Sheet 2: RFC SM59 ─────────────────────────────────────────
        ws2 = wb.create_sheet("RFC Destinations (SM59)")
        ws2["A1"] = "RFC Destination Configuration (Transaction SM59)"
        ws2["A1"].font = Font(bold=True, size=13, color="1F4E79")
        ws2["A2"] = "Create these RFC destinations in the SAP source system before testing."
        ws2["A2"].font = Font(italic=True, color="808080")

        rfc_interfaces = [
            (a, configs.get(a.interface.name))
            for a in assessments
            if a.interface.sender_adapter == "RFC" or a.interface.receiver_adapter == "RFC"
        ]

        sm59_headers = ["RFC Dest. Name", "Conn. Type", "Target Host", "Path Prefix",
                        "Service No.", "Logon Client", "Auth Method", "SSL Active"]
        for col, h in enumerate(sm59_headers, 1):
            hdr_cell(ws2, 3, col, h, "7030A0")

        row = 4
        for a, cfg in rfc_interfaces:
            iface   = a.interface
            af      = autofill.get(iface.name)
            host    = ((af.rfc_target_host if af else "")
                       or (cpi_base_url.replace("https://", "").split("/")[0] if cpi_base_url else "[CPI_HOST]"))
            dest    = (af.rfc_destination if af else "") or f"CPI_{iface.name[:20]}"
            extra   = getattr(cfg.message if cfg else None, "__dict__", {}).get("extra", {}) if cfg else {}
            row_data = [
                dest,
                "G (HTTP Connection to Ext. Server)",
                host,
                "/api/v1/",
                str(af.port) if (af and af.port) else "443",
                "[not required for HTTP]",
                "Basic / OAuth2",
                "Active (TLS 1.2)",
            ]
            for col, val in enumerate(row_data, 1):
                data_cell(ws2, row, col, val)
            row += 1

        if not rfc_interfaces:
            ws2.cell(row=4, column=1, value="No RFC interfaces detected.")

        for col, width in enumerate([28, 32, 40, 18, 12, 18, 18, 14], 1):
            ws2.column_dimensions[get_column_letter(col)].width = width

        # ── Sheet 3: JDBC ─────────────────────────────────────────────
        ws3 = wb.create_sheet("JDBC Setup")
        ws3["A1"] = "JDBC Driver Upload & Data Source Configuration"
        ws3["A1"].font = Font(bold=True, size=13, color="1F4E79")

        jdbc_interfaces = [
            (a, configs.get(a.interface.name))
            for a in assessments
            if a.interface.sender_adapter == "JDBC" or a.interface.receiver_adapter == "JDBC"
        ]

        # Driver upload guide
        ws3["A3"] = "STEP 1: Upload JDBC Driver"
        ws3["A3"].font = Font(bold=True)
        driver_guide = [
            ["Database", "Driver JAR Name", "Driver Class", "Download URL"],
            ["Microsoft SQL Server", "mssql-jdbc-12.x.x.jre11.jar",
             "com.microsoft.sqlserver.jdbc.SQLServerDriver",
             "https://learn.microsoft.com/en-us/sql/connect/jdbc/download-microsoft-jdbc-driver-for-sql-server"],
            ["Oracle", "ojdbc11.jar", "oracle.jdbc.OracleDriver",
             "https://www.oracle.com/database/technologies/appdev/jdbc-downloads.html"],
            ["MySQL", "mysql-connector-j-8.x.jar", "com.mysql.cj.jdbc.Driver",
             "https://dev.mysql.com/downloads/connector/j/"],
            ["PostgreSQL", "postgresql-42.x.x.jar", "org.postgresql.Driver",
             "https://jdbc.postgresql.org/download/"],
            ["SAP HANA", "ngdbc.jar", "com.sap.db.jdbc.Driver",
             "Available in SAP Software Downloads"],
        ]
        for row_num, row_data in enumerate(driver_guide, 4):
            for col, val in enumerate(row_data, 1):
                if row_num == 4:
                    hdr_cell(ws3, row_num, col, val, "C55A11")
                else:
                    data_cell(ws3, row_num, col, val)

        # Data source config
        ws3["A12"] = "STEP 2: Create Data Source in CPI"
        ws3["A12"].font = Font(bold=True)
        ws3["A13"] = "Monitor → Manage JDBC Material → Data Sources → Add"
        ws3["A13"].font = Font(italic=True)

        ds_headers = ["Interface", "Data Source Name", "Driver Class", "JDBC URL Template",
                      "Schema/Database", "Notes"]
        for col, h in enumerate(ds_headers, 1):
            hdr_cell(ws3, 14, col, h, "C55A11")

        row = 15
        for a, cfg in jdbc_interfaces:
            iface   = a.interface
            af      = autofill.get(iface.name)
            driver  = ((af.jdbc_driver if af else "")
                       or (cfg.message.jdbc_driver if cfg else "") or "[fill driver class]")
            jndi    = ((af.jdbc_url if af else "")
                       or (cfg.message.jdbc_jndi if cfg else ""))
            url_tmpl = (af.jdbc_url if (af and af.jdbc_url) else
                        "jdbc:sqlserver://[HOST]:1433;databaseName=[DB]")
            row_data = [
                iface.name,
                (cfg.message.jdbc_jndi if cfg and cfg.message.jdbc_jndi else f"DS_{iface.name[:20]}"),
                driver or "[fill driver class]",
                url_tmpl,
                "[fill schema/database name]",
                "Test connection after creation. Firewall must allow CPI → DB on DB port.",
            ]
            for col, val in enumerate(row_data, 1):
                data_cell(ws3, row, col, val)
            row += 1

        if not jdbc_interfaces:
            ws3.cell(row=15, column=1, value="No JDBC interfaces detected.")

        for col, width in enumerate([30, 28, 40, 48, 22, 45], 1):
            ws3.column_dimensions[get_column_letter(col)].width = width

        # ── Sheet 4: cTMS Blueprint ───────────────────────────────────
        ws4 = wb.create_sheet("cTMS Transport Blueprint")
        ws4["A1"] = "Cloud Transport Management (cTMS) Setup Blueprint"
        ws4["A1"].font = Font(bold=True, size=13, color="1F4E79")
        ws4["A2"] = "Replace legacy CTS+ with cTMS for DEV → QA → PROD transport."
        ws4["A2"].font = Font(italic=True, color="808080")

        steps = [
            ("STEP 1: Subscribe to cTMS",
             "BTP Cockpit → Global Account → Entitlements → Add cTMS entitlement.\n"
             "Then: Subaccount → Services → Service Marketplace → Cloud Transport Management → Subscribe.",
             "Client Basis"),
            ("STEP 2: Create cTMS service instance",
             "Subaccount → Instances → New Instance → Cloud Transport Management → plan 'standard'.\n"
             "Create a service key to get the API endpoint URL.",
             "Client Basis"),
            ("STEP 3: Create Transport Nodes",
             "cTMS UI → Transport Nodes → Create:\n"
             "  • Node: CPI_DEV  | Content Type: Multi-Target Application\n"
             "  • Node: CPI_QA   | Content Type: Multi-Target Application\n"
             "  • Node: CPI_PROD | Content Type: Multi-Target Application\n"
             "Link each node to its CPI subaccount service instance.",
             "Consultant"),
            ("STEP 4: Create Transport Route",
             "cTMS UI → Transport Routes → Create:\n"
             "  • Route name: DEV_TO_QA_TO_PROD\n"
             "  • Source node: CPI_DEV\n"
             "  • Target nodes: CPI_QA → CPI_PROD (sequential)\n"
             "  • Transport mode: Forward",
             "Consultant"),
            ("STEP 5: Configure CPI to use cTMS",
             "In CPI (each landscape):\n"
             "  Settings → Transport → Configure Transport → Enable cTMS.\n"
             "  Enter cTMS service URL and OAuth credentials from service key.",
             "Consultant"),
            ("STEP 6: Test transport",
             "In CPI DEV:\n"
             "  Design → your package → Transport → Add to Transport Request.\n"
             "  In cTMS: release transport request → it flows to QA automatically.\n"
             "  After QA validation: forward to PROD.",
             "Consultant"),
        ]

        row = 4
        hdr_cell(ws4, row, 1, "Step", "0D6E35")
        hdr_cell(ws4, row, 2, "Instructions", "0D6E35")
        hdr_cell(ws4, row, 3, "Responsible", "0D6E35")
        row += 1

        for step_name, instructions, responsible in steps:
            data_cell(ws4, row, 1, step_name, bold=True)
            data_cell(ws4, row, 2, instructions)
            data_cell(ws4, row, 3, responsible)
            ws4.row_dimensions[row].height = 90
            row += 1

        ws4.column_dimensions["A"].width = 28
        ws4.column_dimensions["B"].width = 70
        ws4.column_dimensions["C"].width = 18

        out = self.output_dir / "infrastructure_guide.xlsx"
        wb.save(out)
        logger.info("Infrastructure guide Excel → %s", out)
        return out

    # ── Word output ───────────────────────────────────────────────────

    def _generate_word(self, assessments, configs, cpi_base_url, project_name, autofill=None) -> Path:
        autofill = autofill or {}
        try:
            from docx import Document
            from docx.shared import Pt
        except ImportError:
            raise RuntimeError("python-docx required")

        doc = Document()
        doc.add_heading("Infrastructure Setup Guide", 0)
        doc.add_heading(project_name, 1)
        doc.add_paragraph(
            f"Generated: {datetime.now().strftime('%Y-%m-%d')}\n"
            "Hand this document to the client's Basis and Security teams. "
            "All items must be completed before iFlow deployment."
        )

        adapters = {a.interface.sender_adapter for a in assessments} | \
                   {a.interface.receiver_adapter for a in assessments}

        if "IDoc" in adapters:
            doc.add_heading("1. IDoc Configuration", 1)
            doc.add_paragraph(
                "Execute in the SAP source system before testing IDoc interfaces."
            )
            doc.add_heading("WE20 — Partner Profile", 2)
            steps = [
                "Open transaction WE20 in the SAP system.",
                "Create partner type LS (Logical System) for each target partner.",
                "Under Outbound Parameters, add the IDoc message type.",
                "Set Receiver Port to the port created in WE21 (next step).",
                "Set Output Mode to 'Transfer IDoc Immediately'.",
                "Save.",
            ]
            for s in steps:
                doc.add_paragraph(s, style="List Number")

            doc.add_heading("WE21 — Port Definition", 2)
            steps2 = [
                "Open transaction WE21.",
                "Create XML File Port.",
                f"Set RFC destination or URL to: {cpi_base_url or '[CPI_BASE_URL]'}/api/v1/idoc",
                "Set Service Path to /api/v1/idoc.",
                "Activate TLS (HTTPS).",
                "Test port with a sample IDoc.",
            ]
            for s in steps2:
                doc.add_paragraph(s, style="List Number")

        if "RFC" in adapters:
            doc.add_heading("2. RFC Destination (SM59)", 1)
            steps3 = [
                "Open transaction SM59 in the SAP source system.",
                "Create new destination, type G (HTTP Connection to External Server).",
                f"Target Host: {cpi_base_url.replace('https://','').split('/')[0] if cpi_base_url else '[CPI_HOST]'}",
                "Port: 443. Path Prefix: /api/v1/",
                "Logon & Security tab: activate SSL, certificate: ANONYM SSL Client.",
                "Test connection — expect HTTP 401 (authentication challenge = reachable).",
            ]
            for s in steps3:
                doc.add_paragraph(s, style="List Number")

        if "JDBC" in adapters:
            doc.add_heading("3. JDBC Setup", 1)
            doc.add_paragraph(
                "Complete in CPI Monitor → Manage JDBC Material."
            )
            steps4 = [
                "Download correct JDBC driver .jar for your database version.",
                "CPI Monitor → Manage JDBC Material → Add JDBC Driver → Upload .jar file.",
                "CPI Monitor → Manage JDBC Material → Data Sources → Add.",
                "Enter: Data Source name, Driver class, JDBC URL, username, password.",
                "Test connection from the Data Sources screen.",
                "Ensure database firewall allows inbound from CPI IP ranges.",
            ]
            for s in steps4:
                doc.add_paragraph(s, style="List Number")

        doc.add_heading("4. Cloud Transport Management (cTMS)", 1)
        doc.add_paragraph(
            "Set up before go-live to enable controlled DEV → QA → PROD promotion."
        )
        ctms_steps = [
            "Add cTMS entitlement to BTP Global Account.",
            "Subscribe to Cloud Transport Management in each subaccount.",
            "Create service instance and service key in each subaccount.",
            "Create Transport Nodes: CPI_DEV, CPI_QA, CPI_PROD.",
            "Create Transport Route: CPI_DEV → CPI_QA → CPI_PROD.",
            "Configure CPI tenant: Settings → Transport → Enable cTMS.",
            "Test with a sample iFlow transport from DEV to QA.",
        ]
        for s in ctms_steps:
            doc.add_paragraph(s, style="List Number")

        out = self.output_dir / "infrastructure_guide.docx"
        doc.save(str(out))
        logger.info("Infrastructure guide Word → %s", out)
        return out
