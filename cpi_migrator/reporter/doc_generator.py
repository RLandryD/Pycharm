"""
reporter/doc_generator.py

Generates a Technical Design Document (TDD) as a Word .docx file
for each migrated CPI interface. Client-ready deliverable.

Sections per document:
  1. Interface Overview
  2. Data Flow (ASCII diagram)
  3. Connectivity & Endpoints
  4. Authentication & Security
  5. Message Processing
  6. Error Handling & Reliability
  7. Clean Core Compliance
  8. Test Cases
  9. Post-Migration Checklist
"""
from __future__ import annotations

import logging
from datetime import datetime
from pathlib import Path
from typing import Optional

logger = logging.getLogger(__name__)


class TDDGenerator:

    def __init__(self, output_dir: str):
        self.output_dir = Path(output_dir) / "docs"
        self.output_dir.mkdir(parents=True, exist_ok=True)

    def generate(
        self,
        assessment,
        cfg,
        clean_core_report=None,
        verification_report=None,
        resolved_destination=None,
    ) -> Path:
        try:
            from docx import Document
            from docx.shared import Pt
            from docx.enum.text import WD_ALIGN_PARAGRAPH
        except ImportError:
            raise RuntimeError("python-docx required: pip install python-docx")

        iface = assessment.interface
        doc   = Document()

        # ── Styles ───────────────────────────────────────────────────
        style = doc.styles["Normal"]
        style.font.name = "Calibri"
        style.font.size = Pt(11)

        # ── Title page ───────────────────────────────────────────────
        title = doc.add_heading(f"Technical Design Document", 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        sub = doc.add_paragraph(iface.name)
        sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
        sub.runs[0].font.size = Pt(16)
        sub.runs[0].bold = True

        meta = doc.add_paragraph()
        meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
        meta.add_run(f"SAP PI/PO → CPI Migration\n")
        meta.add_run(f"Generated: {datetime.now().strftime('%Y-%m-%d')}\n")
        meta.add_run(f"Complexity: {assessment.complexity} | "
                     f"Effort: {assessment.effort_days} day(s)")

        doc.add_page_break()

        # ── 1. Interface Overview ────────────────────────────────────
        doc.add_heading("1. Interface Overview", 1)
        tbl = doc.add_table(rows=1, cols=2)
        tbl.style = "Table Grid"
        hdr = tbl.rows[0].cells
        hdr[0].text = "Property"
        hdr[1].text = "Value"
        self._bold_row(hdr)

        overview_rows = [
            ("Interface Name",     iface.name),
            ("Namespace",          iface.namespace or "—"),
            ("Software Component", iface.software_component or "—"),
            ("Sender System",      iface.sender_system or "—"),
            ("Sender Adapter",     iface.sender_adapter),
            ("Receiver System",    iface.receiver_system or "—"),
            ("Receiver Adapter",   iface.receiver_adapter),
            ("Complexity",         f"{assessment.complexity} (score: {assessment.score})"),
            ("Effort Estimate",    f"{assessment.effort_days} day(s)"),
            ("Pattern",            assessment.recommended_pattern),
            ("Has BPM",            "Yes ⚠" if iface.has_bpm else "No"),
            ("Has Multi-Mapping",  "Yes ⚠" if iface.has_multi_mapping else "No"),
            ("Original Mapping",   iface.mapping_program or "—"),
            ("Description",        iface.description or "—"),
        ]
        for label, value in overview_rows:
            row = tbl.add_row().cells
            row[0].text = label
            row[1].text = str(value)

        # Complexity reasoning
        if assessment.reasoning:
            doc.add_heading("Complexity Reasoning", 2)
            for reason in assessment.reasoning:
                doc.add_paragraph(reason, style="List Bullet")

        # ── 2. Data Flow ─────────────────────────────────────────────
        doc.add_heading("2. Data Flow", 1)
        sender_a   = cfg.sender_adapter if cfg else iface.sender_adapter
        receiver_a = cfg.receiver_adapter if cfg else iface.receiver_adapter
        flow = self._build_flow_diagram(iface, cfg, sender_a, receiver_a)
        doc.add_paragraph(flow).runs[0].font.name = "Courier New"

        # ── 3. Connectivity & Endpoints ──────────────────────────────
        doc.add_heading("3. Connectivity & Endpoints", 1)
        if cfg:
            tbl2 = doc.add_table(rows=1, cols=3)
            tbl2.style = "Table Grid"
            h = tbl2.rows[0].cells
            h[0].text, h[1].text, h[2].text = "Side", "Property", "Value"
            self._bold_row(h)

            conn_rows = [
                ("Sender",   "Adapter",   sender_a),
                ("Sender",   "Address",   cfg.sender_connectivity.address or "⚠ Not set"),
                ("Sender",   "Path",      cfg.sender_connectivity.path or "—"),
                ("Sender",   "Port",      str(cfg.sender_connectivity.port) if cfg.sender_connectivity.port else "Default"),
                ("Receiver", "Adapter",   receiver_a),
                ("Receiver", "Address",   cfg.receiver_connectivity.address or "⚠ Not set"),
                ("Receiver", "Path",      cfg.receiver_connectivity.path or "—"),
                ("Receiver", "Port",      str(cfg.receiver_connectivity.port) if cfg.receiver_connectivity.port else "Default"),
            ]
            for side, prop, val in conn_rows:
                row = tbl2.add_row().cells
                row[0].text, row[1].text, row[2].text = side, prop, val

        # ── 4. Authentication & Security ─────────────────────────────
        doc.add_heading("4. Authentication & Security", 1)
        if cfg:
            for side, auth in [("Sender", cfg.sender_auth), ("Receiver", cfg.receiver_auth)]:
                doc.add_heading(f"4.{1 if side=='Sender' else 2} {side} Authentication", 2)
                tbl3 = doc.add_table(rows=1, cols=2)
                tbl3.style = "Table Grid"
                h = tbl3.rows[0].cells
                h[0].text, h[1].text = "Property", "Value"
                self._bold_row(h)
                auth_rows = [
                    ("Method",            auth.method),
                    ("Credential Alias",  auth.credential_name or "⚠ Not set"),
                    ("Token URL",         auth.token_url or "—"),
                    ("Certificate Alias", auth.certificate_alias or "—"),
                    ("API Key Header",    auth.api_key_header if auth.method == "API Key" else "—"),
                ]
                for prop, val in auth_rows:
                    row = tbl3.add_row().cells
                    row[0].text, row[1].text = prop, str(val)

        # ── 5. Message Processing ─────────────────────────────────────
        doc.add_heading("5. Message Processing", 1)
        if cfg:
            tbl4 = doc.add_table(rows=1, cols=2)
            tbl4.style = "Table Grid"
            h = tbl4.rows[0].cells
            h[0].text, h[1].text = "Property", "Value"
            self._bold_row(h)
            msg_rows = [
                ("Processing Mode",   "Asynchronous" if cfg.message.is_async else "Synchronous"),
                ("Message Format",    cfg.message.format),
                ("Content-Type",      cfg.message.content_type),
                ("Encoding",          cfg.message.encoding),
                ("Namespace",         cfg.message.namespace or "—"),
                ("Mapping Program",   cfg.message.mapping_program or "—"),
                ("XSLT Program",      cfg.message.xslt_program or "—"),
                ("Quality of Service", cfg.runtime.quality_of_service),
            ]
            if cfg.message.idoc_type:
                msg_rows += [
                    ("IDoc Type",        cfg.message.idoc_type),
                    ("IDoc Message Type", cfg.message.idoc_message_type or "—"),
                    ("Partner Profile",  cfg.message.idoc_partner_profile or "⚠ Not set"),
                ]
            for prop, val in msg_rows:
                row = tbl4.add_row().cells
                row[0].text, row[1].text = prop, str(val)

        # ── 6. Error Handling & Reliability ──────────────────────────
        doc.add_heading("6. Error Handling & Reliability", 1)
        if cfg:
            rel = cfg.reliability
            tbl5 = doc.add_table(rows=1, cols=2)
            tbl5.style = "Table Grid"
            h = tbl5.rows[0].cells
            h[0].text, h[1].text = "Setting", "Value"
            self._bold_row(h)
            rel_rows = [
                ("Auto-Retry",            "Enabled" if rel.retry_enabled else "Disabled"),
                ("Max Retry Attempts",    str(rel.retry_max_attempts)),
                ("Retry Delay (sec)",     str(rel.retry_delay_sec)),
                ("Exponential Backoff",   "Yes" if rel.retry_exponential_backoff else "No"),
                ("Dead Letter Queue",     "Enabled" if rel.dead_letter_enabled else "Disabled"),
                ("DLQ Name",             rel.dead_letter_queue or "—"),
                ("Idempotency Check",    "Enabled" if rel.idempotency_enabled else "Disabled"),
                ("Idempotency Header",   rel.idempotency_header if rel.idempotency_enabled else "—"),
                ("Alert on Failure",     "Enabled" if rel.alert_on_failure else "Disabled"),
                ("Alert Address",        rel.alert_address or "—"),
                ("Message Log Level",    rel.log_level),
                ("Store on Failure",     "Yes" if rel.store_message_on_failure else "No"),
                ("Timeout (sec)",        str(cfg.runtime.timeout_sec)),
                ("Max Message Size (MB)", str(cfg.runtime.max_message_mb)),
            ]
            for prop, val in rel_rows:
                row = tbl5.add_row().cells
                row[0].text, row[1].text = prop, str(val)

        # ── 7. Clean Core Compliance ──────────────────────────────────
        doc.add_heading("7. Clean Core Compliance", 1)
        if clean_core_report:
            score_para = doc.add_paragraph()
            score_para.add_run(f"Score: {clean_core_report.score}/100  ").bold = True
            score_para.add_run(f"[{clean_core_report.traffic_light}]  ")
            score_para.add_run(f"RISE Ready: {'Yes ✓' if clean_core_report.rise_ready else 'No ✗'}")

            if clean_core_report.violations:
                doc.add_heading("Violations", 2)
                for v in clean_core_report.violations:
                    p = doc.add_paragraph(style="List Bullet")
                    p.add_run(f"[{v.rule.severity}] {v.rule.name}: ").bold = True
                    p.add_run(v.detail)
                    doc.add_paragraph(f"   Fix: {v.rule.remediation}",
                                      style="List Bullet 2")
            if clean_core_report.passed_rules:
                doc.add_heading("Passed Checks", 2)
                for r in clean_core_report.passed_rules:
                    doc.add_paragraph(r, style="List Bullet")
        else:
            doc.add_paragraph("Clean Core analysis not yet run.")

        # ── 8. Verification Status ────────────────────────────────────
        doc.add_heading("8. Verification Status", 1)
        if verification_report:
            status_p = doc.add_paragraph()
            status_p.add_run(f"Status: {verification_report.status}  ").bold = True
            status_p.add_run(f"Completion: {verification_report.completion_pct:.0f}%")

            if verification_report.gaps:
                doc.add_heading("Gaps / Missing Items", 2)
                for gap in verification_report.gaps:
                    p = doc.add_paragraph(style="List Bullet")
                    p.add_run(f"[{gap.severity}] {gap.category}: ").bold = True
                    p.add_run(gap.description)
                    doc.add_paragraph(f"   Fix: {gap.suggested_fix}",
                                      style="List Bullet 2")
        else:
            doc.add_paragraph("Verification not yet run.")

        # ── 9. Post-Migration Checklist ───────────────────────────────
        doc.add_heading("9. Post-Migration Checklist", 1)
        checklist = [
            "[ ] Import .iflw file in BTP Cockpit → Integration Suite → Design → Import",
            "[ ] Create secure parameters in Monitor → Manage Security Material",
            "[ ] Configure Cloud Connector if on-premise connectivity required",
            "[ ] Set up Communication Arrangement in target S/4HANA system",
            "[ ] Configure partner profiles (WE20) if IDoc adapter used",
            "[ ] Deploy iFlow and run smoke test with mock payload",
            "[ ] Verify message logs in Monitor → Message Processing",
            "[ ] Set up alerting rules in Monitor → Manage Alerts",
            "[ ] Conduct end-to-end test with real payload",
            "[ ] Client sign-off and handover",
        ]
        if cfg and cfg.manual_steps:
            checklist += [f"[ ] {s}" for s in cfg.manual_steps]

        for item in checklist:
            doc.add_paragraph(item, style="List Bullet")

        # ── Save ──────────────────────────────────────────────────────
        safe_name = iface.name.replace(" ", "_").replace("/", "_")
        out_path  = self.output_dir / f"TDD_{safe_name}.docx"
        doc.save(str(out_path))
        logger.info("TDD saved → %s", out_path)
        return out_path

    def generate_all(
        self,
        assessments: list,
        configs: dict,
        clean_core_reports: dict = None,
        verification_reports: dict = None,
    ) -> list[Path]:
        clean_core_reports   = clean_core_reports or {}
        verification_reports = verification_reports or {}
        paths = []
        for a in assessments:
            name = a.interface.name
            cfg  = configs.get(name)
            if not cfg:
                continue
            try:
                p = self.generate(
                    a, cfg,
                    clean_core_report=clean_core_reports.get(name),
                    verification_report=verification_reports.get(name),
                )
                paths.append(p)
            except Exception as exc:
                logger.error("TDD generation failed for %s: %s", name, exc)
        return paths

    # ── Helpers ───────────────────────────────────────────────────────

    @staticmethod
    def _bold_row(cells):
        for cell in cells:
            for para in cell.paragraphs:
                for run in para.runs:
                    run.bold = True
            if not cells[0].paragraphs[0].runs:
                cells[0].paragraphs[0].add_run(cells[0].text).bold = True

    @staticmethod
    def _build_flow_diagram(iface, cfg, sender_adapter, receiver_adapter) -> str:
        sender_sys  = iface.sender_system or "Sender"
        receiver_sys = iface.receiver_system or "Receiver"
        mapping     = f"→ [Mapping: {iface.mapping_program}] " if iface.mapping_program else ""
        bpm         = " [BPM→iFlow] " if iface.has_bpm else ""
        return (
            f"\n"
            f"  ┌─────────────────────┐\n"
            f"  │  {sender_sys:<19}│\n"
            f"  │  [{sender_adapter}]           │\n"
            f"  └──────────┬──────────┘\n"
            f"             │\n"
            f"             ▼\n"
            f"  ┌─────────────────────┐\n"
            f"  │  SAP Cloud          │\n"
            f"  │  Integration (CPI)  │\n"
            f"  │  {mapping}{bpm:<14}│\n"
            f"  └──────────┬──────────┘\n"
            f"             │\n"
            f"             ▼\n"
            f"  ┌─────────────────────┐\n"
            f"  │  {receiver_sys:<19}│\n"
            f"  │  [{receiver_adapter}]           │\n"
            f"  └─────────────────────┘\n"
        )
