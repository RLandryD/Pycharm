"""
reporter/security_inventory.py

Generates a Security Setup Guide covering:
  1. Certificates to import into CPI keystore (TLS 1.2, client certs)
  2. PGP keys to import
  3. Secure parameters (User Credentials) to create
  4. OAuth2 client credential entries to create
  5. SSH keys for SFTP
  6. RBAC role assignments required per adapter type

Output: Excel workbook (one sheet per category) + Word guide
"""
from __future__ import annotations

import logging
from dataclasses import dataclass, field
from datetime import datetime
from pathlib import Path
from typing import Optional

logger = logging.getLogger(__name__)

# Keywords that indicate MLS (message-level security) is required
MLS_KEYWORDS = {
    "payroll", "salary", "wage", "compensation", "bank", "payment",
    "finance", "financial", "credit", "debit", "invoice", "accounting",
    "hr", "human_resource", "employee", "personal", "gdpr", "hipaa",
    "healthcare", "medical", "insurance", "tax", "pension", "benefits",
    "confidential", "sensitive", "restricted", "private", "pii",
}

def detect_security_level(interface_name: str, description: str = "") -> str:
    """Returns 'MLS' if message-level encryption needed, else 'TLS'."""
    text = (interface_name + " " + description).lower().replace("_", " ")
    if any(kw in text for kw in MLS_KEYWORDS):
        return "MLS"
    return "TLS"


@dataclass
class CertificateEntry:
    alias: str
    cert_type: str          # "TLS Client" / "TLS Server" / "PGP Public" / "PGP Private" / "SSH"
    system: str             # which system this cert is for
    interface_names: list[str] = field(default_factory=list)
    import_location: str = ""
    notes: str = ""


@dataclass
class CredentialEntry:
    alias: str
    cred_type: str          # "UserCredentials" / "OAuth2" / "APIKey"
    system: str
    username: str = ""
    token_url: str = ""
    interface_names: list[str] = field(default_factory=list)
    notes: str = ""


@dataclass
class RBACEntry:
    role_collection: str
    description: str
    required_for: str
    assign_to: str          # "CPI_Admin" / "CPI_Developer" / "CPI_Monitor"


# ---------------------------------------------------------------------------
# Role collections required per adapter/feature
# ---------------------------------------------------------------------------

BASE_ROLES = [
    RBACEntry("Integration_Provisioner",
              "Required to activate CPI capability in Integration Suite",
              "All CPI tenants",
              "CPI_Admin"),
    RBACEntry("PI_Administrator",
              "Full admin access to CPI — design, deploy, monitor",
              "All CPI tenants",
              "CPI_Admin"),
    RBACEntry("PI_Integration_Developer",
              "Design and deploy iFlows",
              "All CPI tenants",
              "CPI_Developer"),
    RBACEntry("PI_Business_Expert",
              "Monitor message processing, view logs",
              "All CPI tenants",
              "CPI_Monitor"),
    RBACEntry("PI_Read_Only",
              "Read-only access to design and monitor",
              "All CPI tenants",
              "CPI_Monitor"),
]

ADAPTER_ROLES = {
    "IDoc": RBACEntry(
        "AuthGroup_IntegrationDeveloper",
        "Required for IDoc sender adapter configuration",
        "IDoc adapter usage",
        "CPI_Developer",
    ),
    "AS2": RBACEntry(
        "AuthGroup_BusinessExpert",
        "Required for B2B/Trading Partner management",
        "AS2/AS4/B2B adapters",
        "CPI_Admin",
    ),
    "JDBC": RBACEntry(
        "AuthGroup_Administrator",
        "Required to manage JDBC material (drivers, data sources)",
        "JDBC adapter usage",
        "CPI_Admin",
    ),
}


class SecurityInventoryGenerator:

    def __init__(self, output_dir: str):
        self.output_dir = Path(output_dir) / "security"
        self.output_dir.mkdir(parents=True, exist_ok=True)

    def generate(
        self,
        assessments: list,
        configs: dict,
        channel_configs: list = None,
    ) -> tuple[Path, Path]:
        certs   = self._extract_certificates(assessments, configs, channel_configs or [])
        creds   = self._extract_credentials(assessments, configs)
        roles   = self._build_roles(assessments)
        xlsx    = self._generate_excel(certs, creds, roles)
        docx    = self._generate_word(certs, creds, roles)
        return xlsx, docx

    # ── Extraction ────────────────────────────────────────────────────

    def _extract_certificates(self, assessments, configs, channel_configs) -> list[CertificateEntry]:
        certs = []
        seen  = set()

        # From channel configs (if available)
        for ch in channel_configs:
            if ch.auth_type in ("X509", "Certificate", "x509"):
                alias = ch.credential_name or f"{ch.channel_name}_cert"
                if alias not in seen:
                    seen.add(alias)
                    certs.append(CertificateEntry(
                        alias=alias,
                        cert_type="TLS Client",
                        system=ch.channel_name,
                        import_location="Monitor → Manage Security Material → Keystore",
                        notes="Import as P12/PFX with private key for client certificate auth.",
                    ))

        # From interface configs
        for name, cfg in configs.items():
            for auth, side_label in [(cfg.sender_auth, "Sender"), (cfg.receiver_auth, "Receiver")]:
                if auth.method == "Certificate" and auth.certificate_alias:
                    alias = auth.certificate_alias
                    if alias not in seen:
                        seen.add(alias)
                        certs.append(CertificateEntry(
                            alias=alias,
                            cert_type="TLS Client",
                            system=f"{side_label} of {name}",
                            interface_names=[name],
                            import_location="Monitor → Manage Security Material → Keystore",
                            notes="P12 format required. Obtain from client security team.",
                        ))

        # SFTP SSH keys
        for name, cfg in configs.items():
            if cfg.sender_adapter in ("SFTP",) or cfg.receiver_adapter in ("SFTP",):
                alias = f"SSH_{name}"
                if alias not in seen:
                    seen.add(alias)
                    certs.append(CertificateEntry(
                        alias=alias,
                        cert_type="SSH",
                        system=f"SFTP server for {name}",
                        interface_names=[name],
                        import_location="Monitor → Manage Security Material → SSH Known Hosts / Keys",
                        notes="Add SFTP server to known hosts. "
                              "If key-based auth: import private key as SSH key pair.",
                    ))

        # PGP detection
        for a in assessments:
            note_text = " ".join(a.notes).lower()
            if "pgp" in note_text or "encrypt" in note_text:
                alias = f"PGP_{a.interface.name}"
                if alias not in seen:
                    seen.add(alias)
                    certs.append(CertificateEntry(
                        alias=alias,
                        cert_type="PGP Public",
                        system=a.interface.receiver_system or "Target",
                        interface_names=[a.interface.name],
                        import_location="Monitor → Manage Security Material → PGP Keys",
                        notes="Export recipient's public PGP key as ASCII-armored .asc file. "
                              "For decryption: also import your own private PGP key.",
                    ))

        # Add server TLS cert verification note
        certs.append(CertificateEntry(
            alias="[TLS Server Certs — auto-managed]",
            cert_type="TLS Server",
            system="All HTTPS endpoints",
            import_location="Monitor → Manage Security Material → Keystore (if self-signed)",
            notes="CPI trusts public CAs automatically. "
                  "For self-signed or private CA certs: import the CA root certificate into keystore.",
        ))

        return certs

    def _extract_credentials(self, assessments, configs) -> list[CredentialEntry]:
        creds = []
        seen  = set()

        for name, cfg in configs.items():
            for auth, side in [(cfg.sender_auth, "Sender"), (cfg.receiver_auth, "Receiver")]:
                alias = auth.credential_name
                if not alias:
                    continue
                if alias in seen:
                    # Add interface to existing
                    for c in creds:
                        if c.alias == alias:
                            c.interface_names.append(name)
                    continue
                seen.add(alias)

                if auth.method == "Basic":
                    creds.append(CredentialEntry(
                        alias=alias,
                        cred_type="UserCredentials",
                        system=f"{side} of {name}",
                        username=auth.client_id or "[fill in]",
                        interface_names=[name],
                        notes="Monitor → Manage Security Material → User Credentials → Add.",
                    ))
                elif auth.method == "OAuth2 Client Credentials":
                    creds.append(CredentialEntry(
                        alias=alias,
                        cred_type="OAuth2",
                        system=f"{side} of {name}",
                        token_url=auth.token_url or "[fill in]",
                        interface_names=[name],
                        notes="Monitor → Manage Security Material → OAuth2 Client Credentials → Add. "
                              f"Token URL: {auth.token_url or '[fill in]'}",
                    ))
                elif auth.method == "API Key":
                    creds.append(CredentialEntry(
                        alias=alias,
                        cred_type="APIKey",
                        system=f"{side} of {name}",
                        interface_names=[name],
                        notes=f"Store as User Credentials. Header name: {auth.api_key_header}",
                    ))

        return creds

    def _build_roles(self, assessments) -> list[RBACEntry]:
        roles = list(BASE_ROLES)
        adapters = {a.interface.sender_adapter for a in assessments} | \
                   {a.interface.receiver_adapter for a in assessments}
        for adapter, role in ADAPTER_ROLES.items():
            if adapter in adapters:
                roles.append(role)
        return roles

    # ── Excel output ──────────────────────────────────────────────────

    def _generate_excel(
        self,
        certs: list[CertificateEntry],
        creds: list[CredentialEntry],
        roles: list[RBACEntry],
    ) -> Path:
        try:
            import openpyxl
            from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
            from openpyxl.utils import get_column_letter
        except ImportError:
            raise RuntimeError("openpyxl required")

        wb   = openpyxl.Workbook()
        thin = Side(style="thin")
        border = Border(left=thin, right=thin, top=thin, bottom=thin)

        def write_sheet(ws, headers, rows, title_colour="1F4E79"):
            hfill = PatternFill("solid", fgColor=title_colour)
            hfont = Font(bold=True, color="FFFFFF")
            for col, h in enumerate(headers, 1):
                cell = ws.cell(row=1, column=col, value=h)
                cell.fill = hfill
                cell.font = hfont
                cell.border = border
                cell.alignment = Alignment(horizontal="center")
            for row_num, row in enumerate(rows, 2):
                for col, val in enumerate(row, 1):
                    cell = ws.cell(row=row_num, column=col, value=str(val))
                    cell.border = border
                    cell.alignment = Alignment(vertical="top", wrap_text=True)
            for col in range(1, len(headers) + 1):
                ws.column_dimensions[get_column_letter(col)].width = 28
            ws.freeze_panes = "A2"

        # Certificates sheet
        ws1 = wb.active
        ws1.title = "Certificates & Keys"
        write_sheet(ws1,
            ["Alias", "Type", "System", "Security Level", "Import Location", "Interfaces", "Notes"],
            [(c.alias, c.cert_type, c.system,
              detect_security_level(c.system, c.notes),
              c.import_location,
              ", ".join(c.interface_names), c.notes) for c in certs],
            "1F4E79")

        # Credentials sheet
        ws2 = wb.create_sheet("Credentials")
        write_sheet(ws2,
            ["Alias", "Type", "System", "Username/Token URL", "Interfaces", "Notes"],
            [(c.alias, c.cred_type, c.system,
              c.username or c.token_url, ", ".join(c.interface_names), c.notes)
             for c in creds],
            "0070B8")

        # RBAC sheet
        ws3 = wb.create_sheet("RBAC Roles")
        write_sheet(ws3,
            ["Role Collection", "Description", "Required For", "Assign To"],
            [(r.role_collection, r.description, r.required_for, r.assign_to)
             for r in roles],
            "217346")

        out = self.output_dir / "security_inventory.xlsx"
        wb.save(out)
        logger.info("Security inventory Excel → %s", out)
        return out

    # ── Word output ───────────────────────────────────────────────────

    def _generate_word(self, certs, creds, roles) -> Path:
        try:
            from docx import Document
            from docx.shared import Pt
        except ImportError:
            raise RuntimeError("python-docx required")

        doc = Document()
        doc.add_heading("Security Setup Guide", 0)
        doc.add_paragraph(
            f"Generated: {datetime.now().strftime('%Y-%m-%d')}\n"
            "Complete all items before deploying iFlows to production."
        )

        doc.add_heading("1. Certificates & Keys to Import", 1)
        doc.add_paragraph(
            "Navigate to: CPI Monitor → Manage Security Material"
        )
        for cert in certs:
            p = doc.add_paragraph(style="List Bullet")
            p.add_run(f"[{cert.cert_type}] {cert.alias}").bold = True
            p.add_run(f"\n  System: {cert.system}")
            p.add_run(f"\n  Location: {cert.import_location}")
            if cert.notes:
                p.add_run(f"\n  Note: {cert.notes}")

        doc.add_heading("2. Credentials to Create", 1)
        for cred in creds:
            p = doc.add_paragraph(style="List Bullet")
            p.add_run(f"[{cred.cred_type}] Alias: {cred.alias}").bold = True
            p.add_run(f"\n  System: {cred.system}")
            if cred.notes:
                p.add_run(f"\n  Steps: {cred.notes}")

        doc.add_heading("3. RBAC Role Assignments", 1)
        doc.add_paragraph(
            "Assign role collections in BTP Cockpit → Security → Users."
        )
        for role in roles:
            p = doc.add_paragraph(style="List Bullet")
            p.add_run(f"{role.role_collection}").bold = True
            p.add_run(f" → {role.assign_to}")
            p.add_run(f"\n  {role.description}")

        out = self.output_dir / "security_setup_guide.docx"
        doc.save(str(out))
        logger.info("Security Word guide → %s", out)
        return out
