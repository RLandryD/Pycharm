"""
reporter/interface_request.py

Generates the SAP ISA-M Business Solution Request artifact — the planning
record a consultant would otherwise fill in by hand on the Manage and Provide
Integration Technology Guidance screen.

Schema mirrors the tutorial form exactly:

    BusinessSolutionRequest
      ├── general (solution overview, business requirements)
      ├── interface_requests: list[InterfaceRequest]
      │     ├── sources : list[ApplicationInstance]   (Name | Application | Deployment)
      │     ├── targets : list[ApplicationInstance]
      │     ├── style                                    (Process / Data / User / Thing)
      │     ├── message_flows: list[MessageFlow]   (sources × targets, auto-generated)
      │     │     ├── source_app, target_app, domain (ISA-M taxonomy)
      │     │     ├── questionnaire: 5 steps × N answers
      │     │     └── selected_technology (recommendation + coverage + policy rule)
      │     └── extensions (only in extended mode — strictly separated from canonical)
      └── attachments

Two render modes:

  - CANONICAL: only the fields that exist on the SAP form. Safe to paste into
    the real ISA-M tool. Forward-compatible if SAP changes the schema.
  - EXTENDED:  canonical fields + a clearly-marked `extensions` block
    containing source_interface_id, migration_wave, consultant_owner, and
    client_owner. For internal consultant use.

Sources of truth (no new data captured by this module):
  - Sources / targets / interface name : InterfaceRecord
  - Domain                            : analyzer/domains.derive_domain()
  - Style + technology recommendation : ISAMRecommendation (if available),
                                        otherwise inferred from adapters
  - Source interface ID (extension)   : InterfaceRecord.id

Also emits an "ISA-M Master Data Preload" list — the deduplicated set of
applications referenced across all IRs. The tenant won't accept an IR that
references an application not in master data, so the consultant pre-loads
this list before opening the tool.
"""
from __future__ import annotations

import json
import logging
from dataclasses import asdict, dataclass, field
from datetime import date, datetime
from pathlib import Path
from typing import Optional

logger = logging.getLogger(__name__)


# ---------------------------------------------------------------------------
# Canonical schema — mirrors the SAP form, nothing else
# ---------------------------------------------------------------------------

INTEGRATION_STYLES = ("Process Integration", "Data Integration",
                      "User Integration", "Thing Integration")

BUSINESS_CRITICALITY = ("High", "Medium", "Low")

# The five tutorial steps and their questions. Kept as canonical strings so
# the generated artifact reads identically to the on-screen form.
QUESTIONNAIRE_STEPS = (
    ("Step 1 — Integration Use Case Pattern", (
        "Which integration use case pattern applies?",
    )),
    ("Step 2 — Transformation", (
        "What is the level of transformation?",
        "Apply predefined SAP/partner content?",
        "B2B crowd-based ML mapping?",
    )),
    ("Step 3 — Integration Content", (
        "Use predefined integration content?",
    )),
    ("Step 4 — Connectivity", (
        "What connector type is needed?",
        "Protocol conversion needed?",
        "Manage multiple trading partners?",
    )),
    ("Step 5 — Monitoring and Operations", (
        "Will business users monitor at runtime?",
        "Forward technical errors to receiver?",
    )),
)


@dataclass
class ApplicationInstance:
    """One row in the form's Source/Target application-instance table."""
    name: str                          # e.g. "Maintenance Contracts"
    application: str                   # e.g. "PostgreSQL"
    deployment: str                    # "Cloud" | "On-Premise" | "Edge" | "Unknown"


@dataclass
class QuestionnaireAnswer:
    """One question's answer, plus which message flows it applies to."""
    question: str
    answer: str                        # "" = left empty (No)
    applies_to_flows: list[str] = field(default_factory=list)


@dataclass
class QuestionnaireStep:
    step_title: str
    answers: list[QuestionnaireAnswer] = field(default_factory=list)


@dataclass
class SelectedTechnology:
    technology: str                    # "SAP Integration Suite, Cloud Integration"
    coverage_percent: int              # 0-100
    policy_rule: str                   # explanation text
    instance: str = ""                 # chosen instance name


@dataclass
class MessageFlow:
    """One source→target flow under an Interface Request."""
    flow_id: str                       # e.g. "MSO-API-1"
    source_app: str
    target_app: str
    domain: str                        # canonical ISA-M domain value
    questionnaire: list[QuestionnaireStep] = field(default_factory=list)
    selected_technology: Optional[SelectedTechnology] = None


@dataclass
class InterfaceRequestExtensions:
    """Extension fields — populated only in extended render mode.

    Strictly separated from canonical fields. Each field passed the three
    filter rules (recurring need / already in our data / not in canonical).
    Empty strings are valid placeholders so the schema is stable when future
    UI features populate them.
    """
    source_interface_id: str = ""       # PI/PO InterfaceRecord.id
    migration_wave: str = ""            # populated later by wave-planning UI
    consultant_owner: str = ""          # populated later by ownership UI
    client_owner: str = ""              # populated later by ownership UI


@dataclass
class InterfaceRequest:
    name: str
    style: str = "Process Integration"
    sources: list[ApplicationInstance] = field(default_factory=list)
    targets: list[ApplicationInstance] = field(default_factory=list)
    message_flows: list[MessageFlow] = field(default_factory=list)
    extensions: InterfaceRequestExtensions = field(default_factory=InterfaceRequestExtensions)


@dataclass
class BusinessSolutionGeneral:
    """The General tab of the BSR."""
    data_exchanged: str = ""           # "Which master/transactional data?"
    business_process: str = ""
    planned_go_live: str = ""          # ISO date string
    business_criticality: str = "Medium"


@dataclass
class BusinessSolutionRequest:
    """Top-level artifact — one per migration project."""
    name: str
    general: BusinessSolutionGeneral = field(default_factory=BusinessSolutionGeneral)
    interface_requests: list[InterfaceRequest] = field(default_factory=list)
    attachments: list[str] = field(default_factory=list)
    created: str = field(default_factory=lambda: datetime.now().isoformat(timespec="seconds"))

    # ── Master-data preload helper ─────────────────────────────────────

    def master_data_preload(self) -> list[ApplicationInstance]:
        """Deduplicated list of every application referenced across all IRs.

        The ISA-M tool requires every referenced application to exist in
        tenant master data before an IR can be saved; this is the list the
        consultant pre-loads.
        """
        seen: dict[tuple[str, str, str], ApplicationInstance] = {}
        for ir in self.interface_requests:
            for app in (*ir.sources, *ir.targets):
                key = (app.name, app.application, app.deployment)
                seen.setdefault(key, app)
        return sorted(seen.values(), key=lambda a: (a.application, a.name))

    # ── Serialisation ──────────────────────────────────────────────────

    def to_canonical_dict(self) -> dict:
        """Dict without the extensions block on each IR."""
        d = asdict(self)
        for ir in d.get("interface_requests", []):
            ir.pop("extensions", None)
        return d

    def to_extended_dict(self) -> dict:
        return asdict(self)

    def to_canonical_json(self, indent: int = 2) -> str:
        return json.dumps(self.to_canonical_dict(), indent=indent, default=str)

    def to_extended_json(self, indent: int = 2) -> str:
        return json.dumps(self.to_extended_dict(), indent=indent, default=str)


# ---------------------------------------------------------------------------
# Builders — assemble a BSR from existing data
# ---------------------------------------------------------------------------

def _deployment_from_bucket(bucket: str) -> str:
    """Map analyzer/domains bucket -> form's Deployment column value."""
    return {
        "cloud": "Cloud",
        "on_premise": "On-Premise",
        "edge": "Edge",
    }.get(bucket, "Unknown")


def _infer_style(record) -> str:
    """Pick an Integration Style when no questionnaire was answered.

    Most A2A/B2B/IDoc/RFC flows are Process Integration on the SAP form;
    pure data-replication via JDBC/File bulk loads are Data Integration.
    Anything carrying a User-facing channel is rare in migrations.
    """
    sa = (getattr(record, "sender_adapter", "") or "").upper()
    ra = (getattr(record, "receiver_adapter", "") or "").upper()
    if "JDBC" in (sa, ra) or {"FILE", "FILE"} == {sa, ra} or {"SFTP", "SFTP"} == {sa, ra}:
        return "Data Integration"
    return "Process Integration"


def _app_instance_from_system(system_name: str) -> ApplicationInstance:
    """Best-effort split of a sender/receiver system name into the form's
    Name | Application | Deployment triple, using analyzer/domains bucketing.

    The PI/PO record only has a single 'system' string. We use it for both
    Name and Application by default; the consultant can refine in the tool.
    """
    from analyzer.domains import _bucket  # noqa: WPS437 — internal helper
    bucket = _bucket(system_name)
    return ApplicationInstance(
        name=system_name or "Unknown",
        application=system_name or "Unknown",
        deployment=_deployment_from_bucket(bucket),
    )


def _default_questionnaire(record) -> list[QuestionnaireStep]:
    """Pre-fill the 5-step questionnaire with sensible defaults derived from
    the interface's adapters. Matches the tutorial's MSO example where most
    answers are 'leave empty (No)' except a small handful."""
    sa = (getattr(record, "sender_adapter", "") or "").upper()
    ra = (getattr(record, "receiver_adapter", "") or "").upper()
    is_b2b = any(x in (sa, ra) for x in ("AS2", "AS4"))
    has_mapping = bool(getattr(record, "mapping_program", None))

    steps = []
    # Step 1
    s1 = QuestionnaireStep(step_title=QUESTIONNAIRE_STEPS[0][0])
    s1.answers.append(QuestionnaireAnswer(
        question=QUESTIONNAIRE_STEPS[0][1][0],
        answer="B2B Integration" if is_b2b else "A2A Integration"))
    steps.append(s1)
    # Step 2
    s2 = QuestionnaireStep(step_title=QUESTIONNAIRE_STEPS[1][0])
    s2.answers.append(QuestionnaireAnswer(
        question=QUESTIONNAIRE_STEPS[1][1][0],
        answer="Complex Transformation" if has_mapping else "Simple Transformation"))
    s2.answers.append(QuestionnaireAnswer(question=QUESTIONNAIRE_STEPS[1][1][1], answer=""))
    s2.answers.append(QuestionnaireAnswer(question=QUESTIONNAIRE_STEPS[1][1][2], answer=""))
    steps.append(s2)
    # Step 3 — empty (built from scratch)
    s3 = QuestionnaireStep(step_title=QUESTIONNAIRE_STEPS[2][0])
    s3.answers.append(QuestionnaireAnswer(question=QUESTIONNAIRE_STEPS[2][1][0], answer=""))
    steps.append(s3)
    # Step 4
    s4 = QuestionnaireStep(step_title=QUESTIONNAIRE_STEPS[3][0])
    s4.answers.append(QuestionnaireAnswer(
        question=QUESTIONNAIRE_STEPS[3][1][0], answer="Technology Connectors"))
    s4.answers.append(QuestionnaireAnswer(question=QUESTIONNAIRE_STEPS[3][1][1], answer=""))
    s4.answers.append(QuestionnaireAnswer(
        question=QUESTIONNAIRE_STEPS[3][1][2],
        answer="Yes" if is_b2b else ""))
    steps.append(s4)
    # Step 5 — empty
    s5 = QuestionnaireStep(step_title=QUESTIONNAIRE_STEPS[4][0])
    s5.answers.append(QuestionnaireAnswer(question=QUESTIONNAIRE_STEPS[4][1][0], answer=""))
    s5.answers.append(QuestionnaireAnswer(question=QUESTIONNAIRE_STEPS[4][1][1], answer=""))
    steps.append(s5)
    return steps


def _select_technology(record, isam_rec) -> SelectedTechnology:
    """Derive the SelectedTechnology block for one message flow."""
    if isam_rec is not None:
        tech = getattr(isam_rec, "primary_tool", "") or "SAP Integration Suite, Cloud Integration"
        confidence = getattr(isam_rec, "confidence", 0.8)
        reasoning_list = getattr(isam_rec, "reasoning", [])
        policy = "; ".join(reasoning_list[:2]) if reasoning_list else ""
        return SelectedTechnology(
            technology=tech,
            coverage_percent=int(round(confidence * 100)),
            policy_rule=policy or "Recommendation from ISA-M questionnaire.",
            instance="SAP Integration Suite",
        )
    # No questionnaire run: default to Cloud Integration with the tutorial's
    # 80% coverage rationale (the most common A2A outcome).
    return SelectedTechnology(
        technology="SAP Integration Suite, Cloud Integration",
        coverage_percent=80,
        policy_rule=("Use Cloud Integration when scenarios need mediation "
                     "(mappings, routings, protocol conversion)."),
        instance="SAP Integration Suite",
    )


def build_interface_request(
    record,
    isam_rec=None,
    name_override: str = "",
) -> InterfaceRequest:
    """Build one InterfaceRequest from one PI/PO InterfaceRecord."""
    from analyzer.domains import derive_domain

    source = _app_instance_from_system(getattr(record, "sender_system", "") or "Source")
    target = _app_instance_from_system(getattr(record, "receiver_system", "") or "Target")

    style = getattr(isam_rec, "integration_style", "") if isam_rec else ""
    if style not in INTEGRATION_STYLES:
        style = _infer_style(record)

    ir_name = name_override or getattr(record, "name", "") or "Interface_Request"

    # Message flows = sources × targets. With one source + one target this is
    # one flow; the cross-product happens naturally when callers attach
    # multi-source/target IRs.
    flow_id = f"{ir_name}-1"
    domain_cls = derive_domain(source.name, target.name)
    flow = MessageFlow(
        flow_id=flow_id,
        source_app=source.name,
        target_app=target.name,
        domain=domain_cls.domain,
        questionnaire=_default_questionnaire(record),
        selected_technology=_select_technology(record, isam_rec),
    )
    # Stamp every questionnaire answer with this flow_id (the tutorial models
    # answers as a multi-select of flows — for one flow it's just this id).
    for step in flow.questionnaire:
        for a in step.answers:
            a.applies_to_flows = [flow_id]

    return InterfaceRequest(
        name=ir_name,
        style=style,
        sources=[source],
        targets=[target],
        message_flows=[flow],
        extensions=InterfaceRequestExtensions(
            source_interface_id=str(getattr(record, "id", "") or ""),
        ),
    )


def build_business_solution_request(
    assessments: list,
    isam_recommendation=None,
    project_name: str = "CPI Migration",
    business_process: str = "",
    planned_go_live: str = "",
    business_criticality: str = "Medium",
    data_exchanged: str = "",
) -> BusinessSolutionRequest:
    """Assemble a BSR from a list of MigrationAssessment objects.

    One BSR for the whole migration; one InterfaceRequest per assessment.
    Matches the tutorial structure (BSR > N IRs).
    """
    if business_criticality not in BUSINESS_CRITICALITY:
        business_criticality = "Medium"
    if not planned_go_live:
        planned_go_live = date.today().isoformat()

    bsr = BusinessSolutionRequest(
        name=project_name,
        general=BusinessSolutionGeneral(
            data_exchanged=data_exchanged or "See per-interface message flows.",
            business_process=business_process or "Migration from PI/PO to Cloud Integration",
            planned_go_live=planned_go_live,
            business_criticality=business_criticality,
        ),
    )
    for a in assessments or []:
        record = getattr(a, "interface", None)
        if record is None:
            continue
        bsr.interface_requests.append(
            build_interface_request(record, isam_rec=isam_recommendation))
    return bsr


# ---------------------------------------------------------------------------
# Word renderer
# ---------------------------------------------------------------------------

def render_word(
    bsr: BusinessSolutionRequest,
    output_path: str | Path,
    mode: str = "canonical",
) -> Path:
    """Render the BSR as a Word document.

    mode='canonical' renders only the fields on the SAP form.
    mode='extended'  also renders an Extensions appendix per IR.
    """
    if mode not in ("canonical", "extended"):
        raise ValueError(f"mode must be 'canonical' or 'extended', got {mode!r}")
    try:
        from docx import Document
        from docx.shared import Pt
    except ImportError:  # pragma: no cover - python-docx is a project dep
        raise RuntimeError("python-docx required")

    output_path = Path(output_path)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()

    # Title
    doc.add_heading("Business Solution Request", 0)
    doc.add_paragraph(bsr.name)
    meta = doc.add_paragraph()
    meta.add_run(f"Mode: {mode.upper()} · Generated: {bsr.created}").italic = True

    # General tab
    doc.add_heading("General", 1)
    doc.add_paragraph(f"Data exchanged: {bsr.general.data_exchanged}")
    doc.add_paragraph(f"Business process: {bsr.general.business_process}")
    doc.add_paragraph(f"Planned go-live: {bsr.general.planned_go_live}")
    doc.add_paragraph(f"Business criticality: {bsr.general.business_criticality}")

    # Master-data preload (tutorial flags this as a hard prerequisite)
    doc.add_heading("ISA-M Master Data Preload", 1)
    doc.add_paragraph(
        "The applications below must exist in tenant master data before any "
        "Interface Request can reference them. Pre-load this list first.")
    preload = bsr.master_data_preload()
    if preload:
        table = doc.add_table(rows=1, cols=3)
        hdr = table.rows[0].cells
        hdr[0].text = "Name"
        hdr[1].text = "Application"
        hdr[2].text = "Deployment"
        for app in preload:
            row = table.add_row().cells
            row[0].text = app.name
            row[1].text = app.application
            row[2].text = app.deployment
    else:
        doc.add_paragraph("(none)")

    # Interface Requests
    for i, ir in enumerate(bsr.interface_requests, 1):
        doc.add_heading(f"Interface Request {i}: {ir.name}", 1)
        doc.add_paragraph(f"Integration Style: {ir.style}")

        doc.add_heading("Sources", 2)
        for src in ir.sources:
            doc.add_paragraph(
                f"{src.name} | {src.application} | {src.deployment}",
                style="List Bullet")
        doc.add_heading("Targets", 2)
        for tgt in ir.targets:
            doc.add_paragraph(
                f"{tgt.name} | {tgt.application} | {tgt.deployment}",
                style="List Bullet")

        doc.add_heading("Message Flows", 2)
        for flow in ir.message_flows:
            doc.add_heading(flow.flow_id, 3)
            doc.add_paragraph(
                f"{flow.source_app} → {flow.target_app}  (Domain: {flow.domain})")
            for step in flow.questionnaire:
                doc.add_paragraph(step.step_title).bold = True
                for ans in step.answers:
                    val = ans.answer if ans.answer else "(left empty / No)"
                    doc.add_paragraph(f"{ans.question} — {val}", style="List Bullet")
            if flow.selected_technology:
                st = flow.selected_technology
                doc.add_paragraph("Selected Technology").bold = True
                doc.add_paragraph(
                    f"{st.technology} — coverage {st.coverage_percent}%")
                doc.add_paragraph(f"Policy rule: {st.policy_rule}")
                doc.add_paragraph(f"Instance: {st.instance}")

        if mode == "extended":
            doc.add_heading("Extensions (internal — not part of SAP form)", 2)
            ext = ir.extensions
            doc.add_paragraph(f"Source interface ID: {ext.source_interface_id or '(none)'}")
            doc.add_paragraph(f"Migration wave: {ext.migration_wave or '(unset)'}")
            doc.add_paragraph(f"Consultant owner: {ext.consultant_owner or '(unset)'}")
            doc.add_paragraph(f"Client owner: {ext.client_owner or '(unset)'}")

    doc.save(str(output_path))
    logger.info("Business Solution Request (%s) → %s", mode, output_path)
    return output_path


def render_json(
    bsr: BusinessSolutionRequest,
    output_path: str | Path,
    mode: str = "canonical",
) -> Path:
    """Render the BSR as JSON. Canonical or extended."""
    if mode not in ("canonical", "extended"):
        raise ValueError(f"mode must be 'canonical' or 'extended', got {mode!r}")
    output_path = Path(output_path)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    body = bsr.to_extended_json() if mode == "extended" else bsr.to_canonical_json()
    output_path.write_text(body, encoding="utf-8")
    return output_path
