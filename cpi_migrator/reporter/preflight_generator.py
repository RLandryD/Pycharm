"""
reporter/preflight_generator.py

Generates a client-specific pre-flight infrastructure checklist
based on what the migration analysis found. Output: Excel + Word.

Covers:
  - BTP subaccount and tenant provisioning
  - CPI tenant setup
  - Cloud Connector installation
  - User and role setup
  - Certificate and keystore preparation
  - Network and firewall rules
  - Adapter-specific prerequisites
"""
from __future__ import annotations

import io
import logging
from dataclasses import dataclass, field
from datetime import datetime
from pathlib import Path
from typing import Optional

logger = logging.getLogger(__name__)


@dataclass
class PreflightItem:
    category: str
    task: str
    detail: str
    responsible: str        # "Client Basis" / "Client Security" / "Consultant" / "SAP"
    mandatory: bool = True
    triggered_by: str = "" # what in the migration triggered this item
    doc_link: str = ""


class PreflightGenerator:

    def __init__(self, output_dir: str):
        self.output_dir = Path(output_dir) / "preflight"
        self.output_dir.mkdir(parents=True, exist_ok=True)

    def generate(
        self,
        assessments: list,
        configs: dict,
        target_ids: list[str],
        project_name: str = "CPI Migration",
        channels: list = None,
    ) -> tuple[Path, Path]:
        """Generate Excel and Word preflight checklists. Returns (xlsx_path, docx_path)."""
        items = self._build_items(assessments, configs, target_ids, channels)
        xlsx  = self._generate_excel(items, project_name)
        docx  = self._generate_word(items, project_name)
        return xlsx, docx

    # ── Item builders ─────────────────────────────────────────────────

    def _build_items(self, assessments, configs, target_ids, channels=None) -> list[PreflightItem]:
        items = []
        items += self._btp_base_items(target_ids)
        items += self._cpi_tenant_items()
        items += self._adapter_items(assessments, configs)
        items += self._security_items(assessments, configs)
        items += self._network_items(assessments, configs, target_ids)
        items += self._transport_items()
        items += self._collision_items(channels)
        items += self._transaction_items(assessments, configs)
        return items

    @staticmethod
    def _collision_items(channels) -> list[PreflightItem]:
        if not channels:
            return []
        try:
            from analyzer.endpoint_collision import detect_collisions, collisions_to_preflight_items
        except Exception:  # pragma: no cover - defensive import
            return []
        return collisions_to_preflight_items(detect_collisions(channels))

    @staticmethod
    def _transaction_items(assessments, configs) -> list[PreflightItem]:
        try:
            from analyzer.transaction_advisor import advise_all, advisories_to_preflight_items
        except Exception:  # pragma: no cover
            return []
        records = [a.interface for a in (assessments or []) if getattr(a, "interface", None)]
        return advisories_to_preflight_items(advise_all(records, configs or {}))

    def _btp_base_items(self, target_ids: list[str]) -> list[PreflightItem]:
        items = [
            PreflightItem(
                category="BTP Platform",
                task="Create BTP Global Account",
                detail="Ensure a BTP Global Account exists with sufficient entitlements. "
                       "Contact SAP or your account executive if not yet provisioned.",
                responsible="Client / SAP",
                doc_link="https://help.sap.com/docs/btp/sap-business-technology-platform/getting-started-in-cloud-foundry-environment",
            ),
            PreflightItem(
                category="BTP Platform",
                task="Create BTP Subaccount",
                detail="Create a dedicated subaccount for CPI (separate Dev/QA/Prod subaccounts recommended). "
                       "Region must match your data residency requirements.",
                responsible="Client Basis",
            ),
            PreflightItem(
                category="BTP Platform",
                task="Assign CPI entitlement to subaccount",
                detail="In BTP Cockpit: Global Account → Entitlements → Subaccount Assignments. "
                       "Assign 'Integration Suite' service plan.",
                responsible="Client Basis",
            ),
            PreflightItem(
                category="BTP Platform",
                task="Subscribe to SAP Integration Suite",
                detail="Subaccount → Services → Service Marketplace → Integration Suite → Subscribe.",
                responsible="Client Basis",
            ),
            PreflightItem(
                category="BTP Platform",
                task="Create Global Administrator user",
                detail="Assign role collection 'Integration_Provisioner' to at least one user before "
                       "activating capabilities.",
                responsible="Client Security",
            ),
            PreflightItem(
                category="BTP Platform",
                task="Activate Cloud Integration capability",
                detail="Integration Suite launchpad → Add Capabilities → Cloud Integration. "
                       "This provisions the CPI tenant (takes 10-20 minutes).",
                responsible="Client Basis",
            ),
        ]

        # Cloud Foundry environment
        items.append(PreflightItem(
            category="BTP Platform",
            task="Enable Cloud Foundry environment",
            detail="Subaccount → Cloud Foundry → Enable. Set org name. "
                   "Required for OAuth2 service key creation.",
            responsible="Client Basis",
        ))

        return items

    def _cpi_tenant_items(self) -> list[PreflightItem]:
        return [
            PreflightItem(
                category="CPI Tenant",
                task="Create CPI service instance and service key",
                detail="BTP Cockpit → Instances and Subscriptions → Create instance of "
                       "'Process Integration Runtime' → plan 'api'. Create service key to get "
                       "clientid, clientsecret, tokenurl, baseurl.",
                responsible="Client Basis",
            ),
            PreflightItem(
                category="CPI Tenant",
                task="Note down CPI base URL and OAuth credentials",
                detail="From the service key JSON: copy url (base URL), tokenurl, clientid, clientsecret. "
                       "These go into settings.yaml for this migration tool.",
                responsible="Consultant",
            ),
            PreflightItem(
                category="CPI Tenant",
                task="Verify CPI tenant is accessible",
                detail="Open the CPI tenant URL in a browser and confirm the Integration Suite "
                       "launchpad loads without errors.",
                responsible="Consultant",
            ),
            PreflightItem(
                category="CPI Tenant",
                task="Create integration package for migrated iFlows",
                detail="Design → Create Package. Use naming convention: "
                       "<CompanyCode>_<SourceSystem>_<TargetSystem>_Migration.",
                responsible="Consultant",
            ),
        ]

    def _adapter_items(self, assessments, configs) -> list[PreflightItem]:
        items = []
        adapters = set()
        for a in assessments:
            adapters.add(a.interface.sender_adapter)
            adapters.add(a.interface.receiver_adapter)
        cfg_adapters = set()
        for cfg in configs.values():
            cfg_adapters.add(cfg.sender_adapter)
            cfg_adapters.add(cfg.receiver_adapter)
        all_adapters = adapters | cfg_adapters

        if "IDoc" in all_adapters:
            items += [
                PreflightItem(
                    category="IDoc Setup",
                    task="Configure RFC destination from SAP system to CPI",
                    detail="In SAP source system: SM59 → Create HTTP connection to CPI endpoint. "
                           "Type: HTTP Connection to ABAP System or External Server.",
                    responsible="Client Basis",
                    triggered_by="IDoc adapter detected",
                ),
                PreflightItem(
                    category="IDoc Setup",
                    task="Create partner profile in SAP system (WE20)",
                    detail="WE20 → Create partner profile for CPI as partner. "
                           "Define outbound parameters for each IDoc message type.",
                    responsible="Client Basis",
                    triggered_by="IDoc adapter detected",
                ),
                PreflightItem(
                    category="IDoc Setup",
                    task="Configure port in SAP system (WE21)",
                    detail="WE21 → Create XML port pointing to CPI endpoint URL. "
                           "Use HTTPS with client certificate or basic auth.",
                    responsible="Client Basis",
                    triggered_by="IDoc adapter detected",
                ),
            ]

        if "RFC" in all_adapters:
            items += [
                PreflightItem(
                    category="RFC Setup",
                    task="Install and configure SAP Cloud Connector for RFC",
                    detail="RFC requires Cloud Connector. SCC must map the RFC destination "
                           "as a virtual host. See Cloud Connector Setup Guide output.",
                    responsible="Client Basis",
                    triggered_by="RFC adapter detected",
                ),
                PreflightItem(
                    category="RFC Setup",
                    task="Create RFC connection in Cloud Connector",
                    detail="SCC → Cloud To On-Premise → Add → ABAP System. "
                           "Map virtual host/port to real SAP application server.",
                    responsible="Client Basis",
                    triggered_by="RFC adapter detected",
                ),
            ]

        if "JDBC" in all_adapters:
            items += [
                PreflightItem(
                    category="JDBC Setup",
                    task="Download JDBC driver for target database",
                    detail="Obtain the correct JDBC driver .jar for your database: "
                           "SQL Server (mssql-jdbc), Oracle (ojdbc), MySQL (mysql-connector-java), "
                           "PostgreSQL (postgresql). Must match database version.",
                    responsible="Client / Consultant",
                    triggered_by="JDBC adapter detected",
                ),
                PreflightItem(
                    category="JDBC Setup",
                    task="Upload JDBC driver to CPI",
                    detail="Monitor → Manage JDBC Material → Add JDBC Driver. "
                           "Upload the .jar file obtained above.",
                    responsible="Consultant",
                    triggered_by="JDBC adapter detected",
                ),
                PreflightItem(
                    category="JDBC Setup",
                    task="Create JDBC Data Source in CPI",
                    detail="Monitor → Manage JDBC Material → Data Sources → Add. "
                           "Enter connection URL, driver class, username, password.",
                    responsible="Consultant",
                    triggered_by="JDBC adapter detected",
                ),
                PreflightItem(
                    category="JDBC Setup",
                    task="Open firewall from CPI to database",
                    detail="Database firewall must allow inbound connections from CPI IP ranges. "
                           "Obtain current CPI egress IPs from SAP help portal for your region.",
                    responsible="Client Security",
                    triggered_by="JDBC adapter detected",
                ),
            ]

        if any(a in all_adapters for a in ["File", "FTP", "SFTP"]):
            items.append(PreflightItem(
                category="File/SFTP Setup",
                task="Provision SFTP server or file share accessible from CPI",
                detail="CPI cannot access on-premise file systems directly without Cloud Connector. "
                       "Options: SFTP server in DMZ, Azure Blob, AWS S3, or SFTP via SCC.",
                responsible="Client Basis",
                triggered_by="File/FTP/SFTP adapter detected",
            ))

        if any(a in all_adapters for a in ["AS2", "AS4"]):
            items.append(PreflightItem(
                category="B2B Setup",
                task="Provision B2B/EDI add-on or Integration Advisor",
                detail="AS2/AS4 requires either the B2B add-on license or Integration Advisor. "
                       "Confirm licensing with SAP account executive before migration.",
                responsible="Client / SAP",
                triggered_by="AS2/AS4 adapter detected",
            ))

        if "AMQP" in all_adapters or "JMS" in all_adapters:
            items.append(PreflightItem(
                category="Messaging Setup",
                task="Provision Advanced Event Mesh or activate JMS broker",
                detail="JMS adapter requires either the built-in CPI message broker (limited) "
                       "or SAP Advanced Event Mesh. Check licensing and activate in Integration Suite.",
                responsible="Client Basis",
                triggered_by="JMS/AMQP adapter detected",
            ))

        return items

    def _security_items(self, assessments, configs) -> list[PreflightItem]:
        items = [
            PreflightItem(
                category="Security",
                task="Plan CPI keystore entries",
                detail="Identify all certificates, PGP keys, and SSH keys referenced in "
                       "PI/PO channel configurations. Each must be imported into CPI keystore "
                       "before iFlows can be deployed. See Certificate Inventory output.",
                responsible="Client Security",
            ),
            PreflightItem(
                category="Security",
                task="Create secure parameters (User Credentials) in CPI",
                detail="Monitor → Manage Security Material → User Credentials. "
                       "Create one entry per username/password pair referenced in iFlows. "
                       "Name must match the credential alias in the iFlow configuration.",
                responsible="Consultant",
            ),
            PreflightItem(
                category="Security",
                task="Create OAuth2 client credential entries",
                detail="Monitor → Manage Security Material → OAuth2 Client Credentials. "
                       "Create entries for each OAuth2-authenticated target system.",
                responsible="Consultant",
            ),
        ]

        # Check for PGP
        all_adapter_notes = " ".join(
            str(a.notes) for a in assessments
        ).lower()
        if "pgp" in all_adapter_notes:
            items.append(PreflightItem(
                category="Security",
                task="Import PGP keys into CPI keystore",
                detail="Monitor → Manage Security Material → PGP Keys. "
                       "Import public key for encryption, private key for decryption. "
                       "Keys must be in ASCII-armored format (.asc).",
                responsible="Client Security",
                triggered_by="PGP usage detected",
            ))

        return items

    def _network_items(self, assessments, configs, target_ids) -> list[PreflightItem]:
        items = []
        needs_scc = any(
            a.interface.sender_adapter in ("RFC", "JDBC", "IDoc", "File", "FTP", "SFTP")
            or a.interface.receiver_adapter in ("RFC", "JDBC", "IDoc", "File", "FTP", "SFTP")
            for a in assessments
        )

        if needs_scc:
            items += [
                PreflightItem(
                    category="Cloud Connector",
                    task="Download and install SAP Cloud Connector",
                    detail="Download from tools.hana.ondemand.com. Install on a server in your "
                           "on-premise network with outbound HTTPS access to BTP (port 443). "
                           "Minimum: 2 CPU, 4 GB RAM, Java 11+.",
                    responsible="Client Basis",
                    triggered_by="On-premise adapter detected",
                ),
                PreflightItem(
                    category="Cloud Connector",
                    task="Connect Cloud Connector to BTP subaccount",
                    detail="SCC admin UI (default port 8443) → Add Subaccount → enter "
                           "BTP Region, Subaccount ID, and admin credentials.",
                    responsible="Client Basis",
                    triggered_by="On-premise adapter detected",
                ),
                PreflightItem(
                    category="Cloud Connector",
                    task="Open outbound port 443 from SCC server to BTP",
                    detail="SCC connects outbound to *.hana.ondemand.com on port 443. "
                           "No inbound firewall rules needed — SCC uses reverse tunnel.",
                    responsible="Client Security",
                    triggered_by="On-premise adapter detected",
                ),
                PreflightItem(
                    category="Cloud Connector",
                    task="Add system mappings in Cloud Connector",
                    detail="For each on-premise system: SCC → Cloud To On-Premise → Add. "
                           "Map virtual host/port → real internal host/port.",
                    responsible="Consultant",
                    triggered_by="On-premise adapter detected",
                ),
            ]

        return items

    def _transport_items(self) -> list[PreflightItem]:
        return [
            PreflightItem(
                category="Transport (cTMS)",
                task="Create Dev, QA, and Production subaccounts",
                detail="Separate BTP subaccounts for each landscape are required for cTMS. "
                       "Each needs its own CPI tenant provisioned.",
                responsible="Client Basis",
            ),
            PreflightItem(
                category="Transport (cTMS)",
                task="Subscribe to Cloud Transport Management service",
                detail="BTP Cockpit → Service Marketplace → Cloud Transport Management → Subscribe. "
                       "Available in global account entitlements.",
                responsible="Client Basis",
            ),
            PreflightItem(
                category="Transport (cTMS)",
                task="Create transport nodes for each landscape",
                detail="cTMS → Transport Nodes → Create nodes for Dev, QA, Prod. "
                       "Link each node to the corresponding CPI service instance.",
                responsible="Consultant",
            ),
            PreflightItem(
                category="Transport (cTMS)",
                task="Create transport route Dev → QA → Prod",
                detail="cTMS → Transport Routes → Create route connecting Dev→QA and QA→Prod nodes. "
                       "This replaces the legacy CTS+ landscape configuration.",
                responsible="Consultant",
            ),
        ]

    # ── Excel output ──────────────────────────────────────────────────

    def _generate_excel(self, items: list[PreflightItem], project_name: str) -> Path:
        try:
            import openpyxl
            from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
            from openpyxl.utils import get_column_letter
        except ImportError:
            raise RuntimeError("openpyxl required")

        wb = openpyxl.Workbook()
        ws = wb.active
        ws.title = "Pre-flight Checklist"

        # Title
        ws["A1"] = f"Pre-flight Infrastructure Checklist — {project_name}"
        ws["A1"].font = Font(bold=True, size=14, color="1F4E79")
        ws["A2"] = f"Generated: {datetime.now().strftime('%Y-%m-%d')} | " \
                   f"Complete ALL 🔴 items before importing any iFlow"
        ws["A2"].font = Font(italic=True, color="808080")

        headers = ["✓", "Category", "Task", "Detail", "Responsible", "Triggered By"]
        header_fill = PatternFill("solid", fgColor="1F4E79")
        header_font = Font(bold=True, color="FFFFFF")
        thin = Side(style="thin")
        border = Border(left=thin, right=thin, top=thin, bottom=thin)

        for col, h in enumerate(headers, 1):
            cell = ws.cell(row=3, column=col, value=h)
            cell.fill = header_fill
            cell.font = header_font
            cell.alignment = Alignment(horizontal="center")
            cell.border = border

        category_colours = {
            "BTP Platform":    "DDEEFF",
            "CPI Tenant":      "DDF0FF",
            "IDoc Setup":      "FFF0DD",
            "RFC Setup":       "FFE8DD",
            "JDBC Setup":      "FFDDDD",
            "File/SFTP Setup": "FFFADD",
            "B2B Setup":       "F0DDFF",
            "Messaging Setup": "DDFFE8",
            "Security":        "FFE8F0",
            "Cloud Connector": "E8F0FF",
            "Transport (cTMS)": "F0FFE8",
        }

        for row_num, item in enumerate(items, start=4):
            colour = category_colours.get(item.category, "FFFFFF")
            row_fill = PatternFill("solid", fgColor=colour)

            row_data = [
                "☐",
                item.category,
                item.task,
                item.detail,
                item.responsible,
                item.triggered_by or "—",
            ]
            for col, value in enumerate(row_data, 1):
                cell = ws.cell(row=row_num, column=col, value=value)
                cell.border = border
                cell.alignment = Alignment(vertical="top", wrap_text=True)
                if col > 1:
                    cell.fill = row_fill

        col_widths = [4, 18, 35, 60, 18, 25]
        for col, width in enumerate(col_widths, 1):
            ws.column_dimensions[get_column_letter(col)].width = width

        ws.freeze_panes = "A4"
        ws.auto_filter.ref = f"A3:{get_column_letter(len(headers))}3"

        out = self.output_dir / "preflight_checklist.xlsx"
        wb.save(out)
        logger.info("Pre-flight Excel saved → %s", out)
        return out

    # ── Word output ───────────────────────────────────────────────────

    def _generate_word(self, items: list[PreflightItem], project_name: str) -> Path:
        try:
            from docx import Document
            from docx.shared import Pt, RGBColor
        except ImportError:
            raise RuntimeError("python-docx required")

        doc = Document()
        doc.add_heading(f"Pre-flight Infrastructure Checklist", 0)
        doc.add_heading(project_name, 1)
        p = doc.add_paragraph()
        p.add_run(f"Generated: {datetime.now().strftime('%Y-%m-%d')}\n").italic = True
        p.add_run("Complete ALL items in this checklist before importing any iFlow into CPI. "
                  "Items marked [Client] require action from the client's Basis or Security team.")

        # Group by category
        from itertools import groupby
        items_sorted = sorted(items, key=lambda x: x.category)
        for category, group in groupby(items_sorted, key=lambda x: x.category):
            doc.add_heading(category, 2)
            for item in group:
                p = doc.add_paragraph(style="List Bullet")
                p.add_run(f"[ ] {item.task}\n").bold = True
                p.add_run(f"    {item.detail}\n")
                p.add_run(f"    Responsible: {item.responsible}").italic = True
                if item.triggered_by:
                    p.add_run(f" | Triggered by: {item.triggered_by}").italic = True

        out = self.output_dir / "preflight_checklist.docx"
        doc.save(str(out))
        logger.info("Pre-flight Word saved → %s", out)
        return out
