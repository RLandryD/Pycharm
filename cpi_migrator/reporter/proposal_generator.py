"""
reporter/proposal_generator.py

Generates two commercial documents from the migration assessment:

  1. CLIENT PROPOSAL (public)
     - Interface count and complexity breakdown
     - Recommended price per tier
     - Total project quote (min/max range)
     - Payment milestone schedule
     - Specialist interfaces requiring client decision
     - Estimated timeline in weeks

  2. INTERNAL COST SHEET (private — never share with client)
     - Your effort days per interface
     - Your cost at your day rate
     - Gross margin per tier and total
     - Specialist surcharge breakdown

Both generated as separate Excel files and a combined Word proposal doc.
"""
from __future__ import annotations

import logging
from dataclasses import dataclass, field
from datetime import datetime
from pathlib import Path
from typing import Optional

logger = logging.getLogger(__name__)


# ---------------------------------------------------------------------------
# Pricing configuration
# ---------------------------------------------------------------------------

@dataclass
class PricingConfig:
    """All monetary settings — stored in encrypted profile, never hardcoded."""

    # Client-facing prices per tier (USD)
    price_auto_min:       int = 1000
    price_auto_max:       int = 1500
    price_guided_min:     int = 2000
    price_guided_max:     int = 3000
    price_specialist_min: int = 4000
    price_specialist_max: int = 7000

    # Your internal cost
    your_day_rate_usd:    int = 800     # your cost per day (what you pay yourself)
    currency:             str = "USD"
    currency_symbol:      str = "$"

    # Volume discounts
    discount_pct_50plus:  int = 10      # % discount for 50+ interfaces
    discount_pct_100plus: int = 15      # % discount for 100+ interfaces

    # Risk buffer
    risk_buffer_pct:      int = 15      # % added for unknowns

    # Parallel delivery
    interfaces_per_week:  int = 8       # how many interfaces you deliver per week

    # Your target margin
    target_margin_pct:    int = 60      # % gross margin target

    @classmethod
    def from_profile(cls, profile) -> "PricingConfig":
        """Load pricing from profile extra fields if set."""
        cfg = cls()
        if hasattr(profile, 'extra') and profile.extra:
            pricing = profile.extra.get("pricing", {})
            for field_name in cls.__dataclass_fields__:
                if field_name in pricing:
                    setattr(cfg, field_name, pricing[field_name])
        return cfg


# ---------------------------------------------------------------------------
# Output models
# ---------------------------------------------------------------------------

@dataclass
class InterfaceQuote:
    interface_name: str
    tier: str
    complexity: str
    base_days: float
    client_price_min: int
    client_price_max: int
    your_cost_usd: int
    margin_pct: float
    specialist_surcharge_min: int = 0
    specialist_surcharge_max: int = 0
    notes: str = ""


@dataclass
class ProjectProposal:
    project_name: str
    company_code: str
    generated_date: str
    interface_quotes: list[InterfaceQuote]
    pricing: PricingConfig

    # Totals
    total_interfaces: int = 0
    auto_count: int = 0
    guided_count: int = 0
    specialist_count: int = 0
    total_days: float = 0.0
    total_weeks: float = 0.0
    quote_min_usd: int = 0
    quote_max_usd: int = 0
    your_total_cost: int = 0
    gross_margin_min: int = 0
    gross_margin_max: int = 0
    margin_pct_min: float = 0.0
    margin_pct_max: float = 0.0
    discount_applied_pct: int = 0
    risk_buffer_usd: int = 0


# ---------------------------------------------------------------------------
# Generator
# ---------------------------------------------------------------------------

class ProposalGenerator:

    def __init__(self, output_dir: str):
        self.output_dir = Path(output_dir) / "proposal"
        self.output_dir.mkdir(parents=True, exist_ok=True)

    def generate(
        self,
        assessments: list,
        ceilings: list,
        configs: dict = None,
        pricing: Optional[PricingConfig] = None,
        project_name: str = "CPI Migration",
        company_code: str = "CLIENT",
    ) -> tuple[Path, Path, Path]:
        """
        Generate client proposal Excel, internal cost Excel, and Word doc.
        Returns (client_xlsx, internal_xlsx, proposal_docx)
        """
        pricing   = pricing or PricingConfig()
        configs   = configs or {}
        ceiling_map = {c.interface_name: c for c in ceilings}

        # Build per-interface quotes
        quotes = []
        for a in assessments:
            name    = a.interface.name
            ceiling = ceiling_map.get(name)
            quote   = self._build_quote(a, ceiling, pricing)
            quotes.append(quote)

        # Build proposal totals
        proposal = self._build_proposal(
            quotes, pricing, project_name, company_code
        )

        client_xlsx   = self._generate_client_excel(proposal)
        internal_xlsx = self._generate_internal_excel(proposal)
        docx          = self._generate_proposal_doc(proposal)

        return client_xlsx, internal_xlsx, docx

    # ── Quote builder ─────────────────────────────────────────────────

    def _build_quote(self, assessment, ceiling, pricing: PricingConfig) -> InterfaceQuote:
        from reporter.migration_ceiling import TIER_AUTO, TIER_GUIDED, TIER_SPECIALIST
        tier       = ceiling.tier if ceiling else TIER_GUIDED
        base_days  = assessment.effort_days

        # Client price by tier
        if tier == TIER_AUTO:
            price_min = pricing.price_auto_min
            price_max = pricing.price_auto_max
        elif tier == TIER_GUIDED:
            price_min = pricing.price_guided_min
            price_max = pricing.price_guided_max
        else:
            price_min = pricing.price_specialist_min
            price_max = pricing.price_specialist_max

        # Add specialist surcharge
        surch_min = ceiling.extra_cost_min_usd if ceiling else 0
        surch_max = ceiling.extra_cost_max_usd if ceiling else 0

        # Your internal cost
        your_cost = int(base_days * pricing.your_day_rate_usd)
        if ceiling:
            your_cost += int(ceiling.extra_days_min * pricing.your_day_rate_usd)

        margin = ((price_min - your_cost) / price_min * 100) if price_min > 0 else 0

        return InterfaceQuote(
            interface_name=assessment.interface.name,
            tier=tier,
            complexity=assessment.complexity,
            base_days=base_days,
            client_price_min=price_min,
            client_price_max=price_max,
            your_cost_usd=your_cost,
            margin_pct=round(margin, 1),
            specialist_surcharge_min=surch_min,
            specialist_surcharge_max=surch_max,
            notes="; ".join(
                t.description for t in (ceiling.triggered_by if ceiling else [])
            )[:100],
        )

    def _build_proposal(
        self,
        quotes: list[InterfaceQuote],
        pricing: PricingConfig,
        project_name: str,
        company_code: str,
    ) -> ProjectProposal:
        from reporter.migration_ceiling import TIER_AUTO, TIER_GUIDED, TIER_SPECIALIST

        total       = len(quotes)
        auto_count  = sum(1 for q in quotes if q.tier == TIER_AUTO)
        guided      = sum(1 for q in quotes if q.tier == TIER_GUIDED)
        specialist  = sum(1 for q in quotes if q.tier == TIER_SPECIALIST)
        total_days  = sum(q.base_days for q in quotes)
        total_weeks = round(total_days / (pricing.interfaces_per_week * 1.5), 1)

        # Raw totals
        raw_min = sum(q.client_price_min + q.specialist_surcharge_min for q in quotes)
        raw_max = sum(q.client_price_max + q.specialist_surcharge_max for q in quotes)

        # Volume discount
        discount_pct = 0
        if total >= 100:
            discount_pct = pricing.discount_pct_100plus
        elif total >= 50:
            discount_pct = pricing.discount_pct_50plus

        disc_min = int(raw_min * (1 - discount_pct / 100))
        disc_max = int(raw_max * (1 - discount_pct / 100))

        # Risk buffer
        buffer     = int(disc_min * pricing.risk_buffer_pct / 100)
        quote_min  = disc_min
        quote_max  = disc_max + buffer

        your_total = sum(q.your_cost_usd for q in quotes)
        gm_min     = quote_min - your_total
        gm_max     = quote_max - your_total
        margin_min = round(gm_min / quote_min * 100, 1) if quote_min > 0 else 0
        margin_max = round(gm_max / quote_max * 100, 1) if quote_max > 0 else 0

        return ProjectProposal(
            project_name=project_name,
            company_code=company_code,
            generated_date=datetime.now().strftime("%Y-%m-%d"),
            interface_quotes=quotes,
            pricing=pricing,
            total_interfaces=total,
            auto_count=auto_count,
            guided_count=guided,
            specialist_count=specialist,
            total_days=total_days,
            total_weeks=total_weeks,
            quote_min_usd=quote_min,
            quote_max_usd=quote_max,
            your_total_cost=your_total,
            gross_margin_min=gm_min,
            gross_margin_max=gm_max,
            margin_pct_min=margin_min,
            margin_pct_max=margin_max,
            discount_applied_pct=discount_pct,
            risk_buffer_usd=buffer,
        )

    # ── Client Excel ──────────────────────────────────────────────────

    def _generate_client_excel(self, p: ProjectProposal) -> Path:
        try:
            import openpyxl
            from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
            from openpyxl.utils import get_column_letter
        except ImportError:
            raise RuntimeError("openpyxl required")

        wb   = openpyxl.Workbook()
        sym  = p.pricing.currency_symbol
        thin = Side(style="thin")
        bdr  = openpyxl.styles.Border(left=thin,right=thin,top=thin,bottom=thin)

        # ── Summary sheet ────────────────────────────────────────────
        ws1 = wb.active
        ws1.title = "Project Summary"

        ws1["A1"] = f"SAP PI/PO → CPI Migration — Commercial Proposal"
        ws1["A1"].font = Font(bold=True, size=14, color="1F4E79")
        ws1["A2"] = f"Project: {p.project_name} | Client: {p.company_code}"
        ws1["A3"] = f"Generated: {p.generated_date} | CONFIDENTIAL"
        ws1["A3"].font = Font(italic=True, color="808080")

        summary_rows = [
            ("", ""),
            ("SCOPE", ""),
            ("Total interfaces", p.total_interfaces),
            ("🟢 Auto-migrated",  p.auto_count),
            ("🟡 Guided migration", p.guided_count),
            ("🔴 Specialist required", p.specialist_count),
            ("", ""),
            ("INVESTMENT", ""),
            (f"Total investment (min)", f"{sym}{p.quote_min_usd:,}"),
            (f"Total investment (max)", f"{sym}{p.quote_max_usd:,}"),
            ("Volume discount applied", f"{p.discount_applied_pct}%"
             if p.discount_applied_pct else "—"),
            ("", ""),
            ("TIMELINE", ""),
            ("Estimated effort", f"{p.total_days:.0f} consulting days"),
            ("Estimated duration", f"{p.total_weeks:.0f} weeks"),
            ("", ""),
            ("PAYMENT MILESTONES", ""),
            ("Milestone 1 — Project kick-off (30%)",
             f"{sym}{int(p.quote_min_usd * 0.3):,}"),
            ("Milestone 2 — DEV delivery + testing (40%)",
             f"{sym}{int(p.quote_min_usd * 0.4):,}"),
            ("Milestone 3 — Go-live + handover (30%)",
             f"{sym}{int(p.quote_min_usd * 0.3):,}"),
        ]

        hfill = PatternFill("solid", fgColor="1F4E79")
        for row, (label, value) in enumerate(summary_rows, 4):
            c1 = ws1.cell(row=row, column=1, value=label)
            c2 = ws1.cell(row=row, column=2, value=value)
            if label in ("SCOPE", "INVESTMENT", "TIMELINE", "PAYMENT MILESTONES"):
                c1.font = Font(bold=True, color="1F4E79")
            if label.startswith("Total investment"):
                c1.font = Font(bold=True)
                c2.font = Font(bold=True)

        ws1.column_dimensions["A"].width = 38
        ws1.column_dimensions["B"].width = 22

        # ── Interface detail sheet ────────────────────────────────────
        ws2 = wb.create_sheet("Interface Breakdown")
        headers = ["Interface", "Complexity", "Tier",
                   f"Price (min {sym})", f"Price (max {sym})",
                   "Specialist Surcharge", "Notes"]
        hdr_fill = PatternFill("solid", fgColor="1F4E79")
        for col, h in enumerate(headers, 1):
            c = ws2.cell(row=1, column=col, value=h)
            c.fill = hdr_fill
            c.font = Font(bold=True, color="FFFFFF")
            c.border = bdr

        tier_fills = {
            "AUTO":       PatternFill("solid", fgColor="C6EFCE"),
            "GUIDED":     PatternFill("solid", fgColor="FFEB9C"),
            "SPECIALIST": PatternFill("solid", fgColor="FFC7CE"),
        }

        for row, q in enumerate(p.interface_quotes, 2):
            surch = (f"{sym}{q.specialist_surcharge_min:,}–"
                     f"{sym}{q.specialist_surcharge_max:,}"
                     if q.specialist_surcharge_min else "—")
            row_data = [
                q.interface_name, q.complexity,
                f"{'🟢' if q.tier=='AUTO' else '🟡' if q.tier=='GUIDED' else '🔴'} {q.tier}",
                q.client_price_min, q.client_price_max,
                surch, q.notes,
            ]
            fill = tier_fills.get(q.tier, PatternFill())
            for col, val in enumerate(row_data, 1):
                c = ws2.cell(row=row, column=col, value=val)
                c.border = bdr
                if col == 3:
                    c.fill = fill

        col_widths = [42, 12, 16, 14, 14, 22, 45]
        for col, width in enumerate(col_widths, 1):
            ws2.column_dimensions[get_column_letter(col)].width = width
        ws2.freeze_panes = "A2"

        # ── Specialist decisions sheet ────────────────────────────────
        specialist_quotes = [q for q in p.interface_quotes if q.tier == "SPECIALIST"]
        if specialist_quotes:
            ws3 = wb.create_sheet("⚠ Specialist Decisions")
            ws3["A1"] = "Interfaces requiring client decision before migration"
            ws3["A1"].font = Font(bold=True, size=12, color="CC0000")
            ws3["A2"] = ("Review each interface below. Select an option or "
                         "mark as 'Defer' / 'Retire'.")

            dec_headers = ["Interface", "Issue", "Options", "Client Decision",
                           f"Extra Cost ({sym})"]
            for col, h in enumerate(dec_headers, 1):
                c = ws3.cell(row=4, column=col, value=h)
                c.fill = PatternFill("solid", fgColor="CC0000")
                c.font = Font(bold=True, color="FFFFFF")
                c.border = bdr

            row = 5
            for q in specialist_quotes:
                ws3.cell(row=row, column=1, value=q.interface_name)
                ws3.cell(row=row, column=2, value=q.notes)
                ws3.cell(row=row, column=3,
                         value="See options in migration report")
                ws3.cell(row=row, column=4,
                         value="[ ] Escalate  [ ] Defer  [ ] Retire")
                ws3.cell(row=row, column=5,
                         value=f"{sym}{q.specialist_surcharge_min:,}–"
                               f"{sym}{q.specialist_surcharge_max:,}"
                               if q.specialist_surcharge_min else "—")
                for col in range(1, 6):
                    ws3.cell(row=row, column=col).border = bdr
                row += 1

            for col, width in enumerate([42, 35, 28, 30, 18], 1):
                ws3.column_dimensions[get_column_letter(col)].width = width

        out = self.output_dir / f"proposal_client_{p.company_code}.xlsx"
        wb.save(out)
        logger.info("Client proposal Excel → %s", out)
        return out

    # ── Internal Excel ────────────────────────────────────────────────

    def _generate_internal_excel(self, p: ProjectProposal) -> Path:
        try:
            import openpyxl
            from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
            from openpyxl.utils import get_column_letter
        except ImportError:
            raise RuntimeError("openpyxl required")

        sym  = p.pricing.currency_symbol
        wb   = openpyxl.Workbook()
        thin = Side(style="thin")
        bdr  = openpyxl.styles.Border(left=thin,right=thin,top=thin,bottom=thin)

        ws = wb.active
        ws.title = "Internal Cost Sheet"

        ws["A1"] = f"INTERNAL — DO NOT SHARE WITH CLIENT"
        ws["A1"].font = Font(bold=True, size=14, color="CC0000")
        ws["A2"] = f"Project: {p.project_name} | Day rate: {sym}{p.pricing.your_day_rate_usd}"

        # Margin summary
        summary = [
            ("Total quote to client (min)", f"{sym}{p.quote_min_usd:,}"),
            ("Total quote to client (max)", f"{sym}{p.quote_max_usd:,}"),
            ("Your total cost",             f"{sym}{p.your_total_cost:,}"),
            ("Gross margin (min)",          f"{sym}{p.gross_margin_min:,} ({p.margin_pct_min}%)"),
            ("Gross margin (max)",          f"{sym}{p.gross_margin_max:,} ({p.margin_pct_max}%)"),
            ("Risk buffer included",        f"{sym}{p.risk_buffer_usd:,}"),
            ("Volume discount applied",     f"{p.discount_applied_pct}%"),
            ("Your day rate",               f"{sym}{p.pricing.your_day_rate_usd}"),
            ("Target margin",               f"{p.pricing.target_margin_pct}%"),
        ]
        for row, (label, value) in enumerate(summary, 4):
            ws.cell(row=row, column=1, value=label)
            ws.cell(row=row, column=2, value=value)
            if "margin" in label.lower():
                ws.cell(row=row, column=2).font = Font(bold=True, color="0D6E35")

        ws.column_dimensions["A"].width = 35
        ws.column_dimensions["B"].width = 25

        # Per-interface breakdown
        ws2 = wb.create_sheet("Per Interface")
        headers = ["Interface", "Tier", "Days", f"Your Cost ({sym})",
                   f"Client Min ({sym})", f"Client Max ({sym})",
                   "Margin %", "Specialist Surcharge", "Notes"]
        hfill = PatternFill("solid", fgColor="7030A0")
        for col, h in enumerate(headers, 1):
            c = ws2.cell(row=1, column=col, value=h)
            c.fill = hfill
            c.font = Font(bold=True, color="FFFFFF")
            c.border = bdr

        for row, q in enumerate(p.interface_quotes, 2):
            margin_colour = ("0D6E35" if q.margin_pct >= p.pricing.target_margin_pct
                             else "CC0000")
            surch = (f"{sym}{q.specialist_surcharge_min:,}–"
                     f"{sym}{q.specialist_surcharge_max:,}"
                     if q.specialist_surcharge_min else "—")
            row_data = [
                q.interface_name, q.tier, q.base_days,
                q.your_cost_usd, q.client_price_min, q.client_price_max,
                f"{q.margin_pct}%", surch, q.notes,
            ]
            for col, val in enumerate(row_data, 1):
                c = ws2.cell(row=row, column=col, value=val)
                c.border = bdr
                if col == 7:
                    c.font = Font(color=margin_colour)

        col_widths = [40, 14, 8, 14, 16, 16, 10, 24, 40]
        for col, width in enumerate(col_widths, 1):
            ws2.column_dimensions[get_column_letter(col)].width = width
        ws2.freeze_panes = "A2"

        out = self.output_dir / f"proposal_internal_{p.company_code}.xlsx"
        wb.save(out)
        logger.info("Internal cost sheet → %s", out)
        return out

    # ── Word proposal doc ─────────────────────────────────────────────

    def _generate_proposal_doc(self, p: ProjectProposal) -> Path:
        try:
            from docx import Document
            from docx.shared import Pt, RGBColor
        except ImportError:
            raise RuntimeError("python-docx required")

        sym = p.pricing.currency_symbol
        doc = Document()

        doc.add_heading("SAP Integration Suite Migration", 0)
        doc.add_heading("Commercial Proposal", 1)

        meta = doc.add_paragraph()
        meta.add_run(f"Prepared for: {p.company_code}\n")
        meta.add_run(f"Project: {p.project_name}\n")
        meta.add_run(f"Date: {p.generated_date}\n")
        meta.add_run("CONFIDENTIAL").bold = True

        doc.add_heading("Executive Summary", 1)
        doc.add_paragraph(
            f"This proposal covers the migration of {p.total_interfaces} "
            f"integration interfaces from SAP PI/PO to SAP Cloud Integration "
            f"(CPI) on SAP BTP. The migration has been assessed using automated "
            f"tooling to classify each interface by complexity and migration "
            f"feasibility."
        )

        doc.add_heading("Scope Overview", 1)
        scope_data = [
            ("Total interfaces assessed", str(p.total_interfaces)),
            ("🟢 Auto-migration (tool-handled)", str(p.auto_count)),
            ("🟡 Guided migration (consultant)", str(p.guided_count)),
            ("🔴 Specialist required", str(p.specialist_count)),
            ("Estimated effort", f"{p.total_days:.0f} days"),
            ("Estimated duration", f"{p.total_weeks:.0f} weeks"),
        ]
        tbl = doc.add_table(rows=1, cols=2)
        tbl.style = "Table Grid"
        tbl.rows[0].cells[0].text = "Item"
        tbl.rows[0].cells[1].text = "Value"
        for label, value in scope_data:
            row = tbl.add_row()
            row.cells[0].text = label
            row.cells[1].text = value

        doc.add_heading("Investment", 1)
        doc.add_paragraph(
            f"Total project investment: "
            f"**{sym}{p.quote_min_usd:,} – {sym}{p.quote_max_usd:,}**"
        )
        doc.paragraphs[-1].runs[0].bold = True

        if p.discount_applied_pct:
            doc.add_paragraph(
                f"A volume discount of {p.discount_applied_pct}% has been "
                f"applied for {p.total_interfaces}+ interfaces."
            )

        doc.add_heading("Payment Milestones", 2)
        milestones = [
            ("Milestone 1 — Project kick-off (30%)",
             f"{sym}{int(p.quote_min_usd*0.3):,}"),
            ("Milestone 2 — DEV delivery + testing complete (40%)",
             f"{sym}{int(p.quote_min_usd*0.4):,}"),
            ("Milestone 3 — Go-live + handover (30%)",
             f"{sym}{int(p.quote_min_usd*0.3):,}"),
        ]
        for ms, amount in milestones:
            doc.add_paragraph(f"{ms}: {amount}", style="List Bullet")

        if p.specialist_count > 0:
            doc.add_heading("Specialist Interfaces — Client Decision Required", 1)
            doc.add_paragraph(
                f"{p.specialist_count} interface(s) require a client decision "
                f"before migration can proceed. These interfaces exceed the "
                f"standard migration scope due to technical complexity. "
                f"Options for each are detailed in the attached decision matrix."
            )
            specialist_quotes = [q for q in p.interface_quotes
                                  if q.tier == "SPECIALIST"]
            for q in specialist_quotes:
                p_para = doc.add_paragraph(style="List Bullet")
                p_para.add_run(q.interface_name).bold = True
                if q.notes:
                    p_para.add_run(f" — {q.notes}")

        doc.add_heading("What's Included", 1)
        included = [
            "Migration assessment and complexity scoring for all interfaces",
            "Generated iFlow XML stubs ready to import into CPI",
            "Technical Design Document (TDD) per interface",
            "Pre-flight infrastructure checklist",
            "Security setup guide (certificates, credentials, RBAC)",
            "Infrastructure guide (IDoc/RFC/JDBC/cTMS configuration)",
            "Test harness with mock payloads per interface",
            "Clean Core compliance report",
            "Groovy mapping scripts (stubs with logic to complete)",
            "Deployment to CPI DEV tenant",
            "30 days post-go-live support for migrated interfaces",
        ]
        for item in included:
            doc.add_paragraph(item, style="List Bullet")

        doc.add_heading("Not Included", 1)
        not_included = [
            "SAP BTP subaccount provisioning and licensing",
            "Business logic completion for custom Java mappings",
            "BPM/ccBPM redesign (quoted separately per interface)",
            "External partner coordination for B2B/EDI interfaces",
            "SAP system configuration (WE20, SM59, Communication Arrangements)",
            "User acceptance testing by client business teams",
        ]
        for item in not_included:
            doc.add_paragraph(item, style="List Bullet")

        out = self.output_dir / f"proposal_{p.company_code}.docx"
        doc.save(str(out))
        logger.info("Proposal Word doc → %s", out)
        return out

    # ── Task split documents: client tasks vs consultant tasks ──────────────

    def generate_task_documents(
        self, assessments: list, ceilings: list = None,
        interventions=None, project_name: str = "CPI Migration",
        company_code: str = "CLIENT",
    ) -> tuple:
        """Generate two task-list Word docs:
          - what the CLIENT needs to do (prerequisites, access, decisions)
          - what I (the consultant) need to do (build/configure/test)

        Derived from adapter advisories (client-side setup like Cloud
        Connector, credentials) and the intervention estimate (consultant
        build tasks). Returns (client_tasks_path, consultant_tasks_path).
        """
        from docx import Document
        try:
            from analyzer.adapter_advisor import advise_all
        except Exception:
            advise_all = None

        ceiling_map = {c.interface_name: c for c in (ceilings or [])}

        # ---- Collect client-side tasks from adapter advisories ----
        client_tasks = set()
        consultant_tasks = set()
        ifaces = [a.interface for a in assessments]

        if advise_all:
            adv = advise_all(ifaces)
            for a in adv.get("advisories", []):
                for gap in a.gaps:
                    client_tasks.add(f"[{a.pi_adapter}] {gap}")
                if a.severity == "warning":
                    # warnings usually require a client setup action
                    client_tasks.add(
                        f"[{a.pi_adapter}] Confirm/setup: {a.cpi_adapter}")
                for note in a.notes:
                    if any(k in note.lower() for k in
                           ("cloud connector", "credential", "certificate",
                            "security material", "partner", "entitlement")):
                        client_tasks.add(f"[{a.pi_adapter}] {note}")

        # Generic client prerequisites
        client_tasks.update({
            "Provide CPI tenant access (or a service key) for deployment",
            "Confirm endpoint URLs and hostnames for all sender/receiver systems",
            "Provide / create credentials in CPI Security Material (no plain passwords)",
            "Confirm which interfaces are in scope and their go-live priority",
            "Perform User Acceptance Testing (UAT) on migrated interfaces",
        })

        # ---- Consultant tasks (per interface + from interventions) ----
        for a in assessments:
            name = a.interface.name
            consultant_tasks.add(f"Build & configure iFlow: {name}")
            consultant_tasks.add(f"Complete mapping logic: {name}")
        consultant_tasks.update({
            "Externalize all endpoint/credential parameters",
            "Add and test error handling per interface",
            "Unit-test each iFlow before client handover",
            "Package and deploy to the client tenant",
        })
        if interventions:
            try:
                for name, iv in (interventions.items()
                                 if hasattr(interventions, "items") else []):
                    for task in getattr(iv, "tasks", [])[:3]:
                        consultant_tasks.add(
                            f"{name}: {getattr(task, 'description', str(task))}")
            except Exception:
                pass

        def _write_doc(title, intro, tasks, fname):
            doc = Document()
            doc.add_heading(title, level=0)
            doc.add_paragraph(f"Project: {project_name}")
            doc.add_paragraph(intro)
            doc.add_heading("Action items", level=1)
            for t in sorted(tasks):
                doc.add_paragraph(t, style="List Bullet")
            out = self.output_dir / fname
            doc.save(str(out))
            return out

        client_path = _write_doc(
            "What You (the Client) Need to Do",
            "These items are the client's responsibility and are prerequisites "
            "or decisions required for the migration to proceed.",
            client_tasks, f"client_tasks_{company_code}.docx")
        consultant_path = _write_doc(
            "What I (the Consultant) Will Do",
            "These are the migration tasks the consultant performs.",
            consultant_tasks, f"consultant_tasks_{company_code}.docx")
        logger.info("Task docs → %s , %s", client_path, consultant_path)
        return client_path, consultant_path


# ── Effort-model integration (additive; testable without the workbench UI) ──
def quote_from_effort(
    interface_name: str,
    effort,                       # reporter.effort_model.EffortBreakdown
    pricing: "PricingConfig" = None,
    tier: str = "",
    complexity: str = "",
) -> "InterfaceQuote":
    """Build an InterfaceQuote from an EffortBreakdown (the Option-X effort
    model: base + gap*multiplier + optional hypercare).

    This is the 'repoint' of the proposal onto the effort model: instead of a
    single assessment.effort_days, the quote reflects the full breakdown
    (adjusted low/high days, hypercare). Pure function, fully unit-testable.
    """
    pricing = pricing or PricingConfig()
    days_low = effort.total_days_low
    days_high = effort.total_days_high
    your_cost = int(days_high * pricing.your_day_rate_usd)
    # client price: midpoint days * a tier-agnostic client rate band
    client_min = int(days_low * pricing.your_day_rate_usd * 1.5)
    client_max = int(days_high * pricing.your_day_rate_usd * 2.0)
    margin = ((client_min - your_cost) / client_min * 100) if client_min > 0 else 0
    return InterfaceQuote(
        interface_name=interface_name,
        tier=tier or "GUIDED",
        complexity=complexity or effort.mode or "",
        base_days=round(days_high, 1),
        client_price_min=client_min,
        client_price_max=client_max,
        your_cost_usd=your_cost,
        margin_pct=round(margin, 1),
        specialist_surcharge_min=0,
        specialist_surcharge_max=0,
        notes=f"effort: {days_low:.1f}-{days_high:.1f}d "
              f"(x{effort.multiplier}"
              + (", +hypercare" if effort.hypercare_enabled else "") + ")",
    )
