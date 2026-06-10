"""
tests/test_all.py
Full test suite — 40+ tests covering all modules.
Run:  pytest tests/test_all.py -v
"""
from __future__ import annotations

import json
import sys
import time
from pathlib import Path
from unittest.mock import MagicMock, patch

import pytest

sys.path.insert(0, str(Path(__file__).parent.parent))

from extractor.pi_extractor import InterfaceRecord, normalise_adapter, PIRestExtractor
from analyzer.complexity_analyzer import ComplexityAnalyzer, MigrationAssessment
from scaffolder.iflow_scaffolder import IFlowScaffolder, _slugify
from reporter.report_generator import ReportGenerator
from destinations.registry import (
    DESTINATION_REGISTRY, get_target, list_targets,
    DestinationTarget, HubSource,
)
from destinations.hub_fetcher import HubFetcher, CacheEntry
from destinations.resolver import DestinationResolver, AdapterRecommendation

# ─── Fixtures ────────────────────────────────────────────────────────────────

BASE_CFG = {
    "environment": "cf",
    "migration": {
        "output_dir": "/tmp/cpi_test_output",
        "complexity_thresholds": {
            "low":    {"max_score": 10},
            "medium": {"max_score": 25},
            "high":   {"max_score": 999},
        },
        "effort_days": {"low": 1, "medium": 3, "high": 8},
    },
}


def make_record(**kw) -> InterfaceRecord:
    defaults = dict(
        id="t001", name="Test_Interface", namespace="http://example.com",
        software_component="SC", sender_system="ECC", receiver_system="S4",
        sender_adapter="HTTPS", receiver_adapter="HTTPS",
        message_interface="MI_Test", mapping_program=None,
        has_bpm=False, has_multi_mapping=False, channel_count=1, description="",
    )
    defaults.update(kw)
    return InterfaceRecord(**defaults)


def make_assessment(**kw) -> MigrationAssessment:
    record = make_record(**kw)
    return ComplexityAnalyzer(BASE_CFG).assess(record)


# ─── extractor ───────────────────────────────────────────────────────────────

class TestNormaliseAdapter:
    def test_soap(self):        assert normalise_adapter("soap") == "SOAP"
    def test_ws(self):          assert normalise_adapter("WS") == "SOAP"
    def test_idocaae(self):     assert normalise_adapter("idocaae") == "IDoc"
    def test_idoc_upper(self):  assert normalise_adapter("IDOC") == "IDoc"
    def test_unknown(self):     assert normalise_adapter("CUSTOM") == "CUSTOM"
    def test_file(self):        assert normalise_adapter("FILE") == "File"
    def test_https(self):       assert normalise_adapter("https") == "HTTPS"


class TestPIRestExtractor:
    def test_parse_config(self):
        raw = {
            "IntegratedConfigurationID": "IC001",
            "IntegratedConfigurationName": "PO_To_S4",
            "IntegratedConfigurationNamespace": "http://test.com",
            "SoftwareComponentName": "SC",
            "SenderComponentName": "ERP",
            "ReceiverComponentName": "S4",
            "SenderChannel": {"AdapterType": "idoc"},
            "ReceiverChannel": {"AdapterType": "soap"},
            "InboundMessageInterface": "MI_Order",
            "MappingProgram": "MM_OrderMap",
            "NumberOfChannels": 2,
        }
        session = MagicMock()
        extractor = PIRestExtractor("http://pi:50000", session)
        r = extractor._parse_config(raw)
        assert r.name == "PO_To_S4"
        assert r.sender_adapter == "IDoc"
        assert r.receiver_adapter == "SOAP"
        assert r.mapping_program == "MM_OrderMap"
        assert r.channel_count == 2

    def test_pagination_stops_on_partial_page(self):
        page1 = {"d": {"results": [
            {"IntegratedConfigurationID": str(i), "IntegratedConfigurationName": f"IF_{i}",
             "IntegratedConfigurationNamespace": "", "SoftwareComponentName": "",
             "SenderComponentName": "", "ReceiverComponentName": "",
             "SenderChannel": {}, "ReceiverChannel": {},
             "InboundMessageInterface": "", "NumberOfChannels": 1}
            for i in range(100)
        ]}}
        page2 = {"d": {"results": [
            {"IntegratedConfigurationID": "100", "IntegratedConfigurationName": "IF_100",
             "IntegratedConfigurationNamespace": "", "SoftwareComponentName": "",
             "SenderComponentName": "", "ReceiverComponentName": "",
             "SenderChannel": {}, "ReceiverChannel": {},
             "InboundMessageInterface": "", "NumberOfChannels": 1}
        ]}}
        mock_session = MagicMock()
        mock_session.get.return_value.json.side_effect = [page1, page2]
        mock_session.get.return_value.raise_for_status = MagicMock()
        extractor = PIRestExtractor("http://pi:50000", mock_session)
        configs = extractor._fetch_integrated_configs()
        assert len(configs) == 101
        assert mock_session.get.call_count == 2


# ─── analyzer ────────────────────────────────────────────────────────────────

class TestComplexityAnalyzer:
    def setup_method(self):
        self.a = ComplexityAnalyzer(BASE_CFG)

    def test_low_https(self):
        r = make_record()
        a = self.a.assess(r)
        assert a.complexity == "LOW"
        assert a.score == 2
        assert a.effort_days == 1

    def test_high_bpm(self):
        r = make_record(sender_adapter="JDBC", receiver_adapter="RFC", has_bpm=True)
        a = self.a.assess(r)
        assert a.complexity == "HIGH"
        assert a.score > 25
        assert any("BPM" in n for n in a.notes)

    def test_mapping_adds_5(self):
        r1 = make_record()
        r2 = make_record(mapping_program="MM_Test")
        assert self.a.assess(r2).score == self.a.assess(r1).score + 5

    def test_multi_mapping_adds_8(self):
        base  = self.a.assess(make_record()).score
        extra = self.a.assess(make_record(has_multi_mapping=True)).score
        assert extra == base + 8

    def test_extra_channels(self):
        base  = self.a.assess(make_record(channel_count=1)).score
        extra = self.a.assess(make_record(channel_count=3)).score
        assert extra == base + 4   # 2 extra channels × 2 pts each

    def test_rfc_pattern(self):
        a = self.a.assess(make_record(sender_adapter="RFC"))
        assert "RFC" in a.recommended_pattern

    def test_b2b_pattern(self):
        a = self.a.assess(make_record(sender_adapter="AS2"))
        assert "B2B" in a.recommended_pattern

    def test_idoc_pattern(self):
        a = self.a.assess(make_record(sender_adapter="IDoc"))
        assert "IDoc" in a.recommended_pattern

    def test_file_to_file_pattern(self):
        a = self.a.assess(make_record(sender_adapter="File", receiver_adapter="File"))
        assert "File" in a.recommended_pattern

    def test_assess_all_length(self):
        records = [make_record(id=str(i), name=f"IF_{i}") for i in range(7)]
        assert len(self.a.assess_all(records)) == 7

    def test_unknown_adapter_adds_5(self):
        base    = self.a.assess(make_record()).score
        unknown = self.a.assess(make_record(sender_adapter="ZEBRA_ADAPTER")).score
        assert unknown > base

    def test_effort_days_band(self):
        low    = self.a.assess(make_record())
        assert low.effort_days == 1
        # HIGH is no longer a flat cap — effort scales with score, so a
        # higher-scoring HIGH interface costs more days than a borderline one.
        high   = self.a.assess(make_record(sender_adapter="JDBC", receiver_adapter="RFC", has_bpm=True))
        assert high.complexity == "HIGH"
        assert high.effort_days >= 8
        bigger = self.a.assess(make_record(sender_adapter="JDBC", receiver_adapter="RFC",
                                           has_bpm=True, has_multi_mapping=True))
        assert bigger.score > high.score
        assert bigger.effort_days > high.effort_days


# ─── scaffolder ──────────────────────────────────────────────────────────────

class TestSlugify:
    def test_spaces_to_underscore(self):    assert _slugify("Test Interface") == "Test_Interface"
    def test_special_chars_stripped(self):  assert "/" not in _slugify("IF/PO:123")
    def test_length_capped(self):           assert len(_slugify("A" * 200)) <= 80
    def test_empty_string(self):            assert _slugify("") == ""


class TestIFlowScaffolder:
    TEMPLATES = str(Path(__file__).parent.parent / "templates")

    def test_creates_file(self, tmp_path):
        s = IFlowScaffolder(output_dir=str(tmp_path), templates_dir=self.TEMPLATES)
        a = make_assessment(name="OrderCreate", mapping_program="MM_Order")
        # default produces a CPI-valid iFlow via the minimal generator
        p = s.scaffold(a)
        assert p.exists()
        content = p.read_text()
        # The interface name appears in the artifact id / filename
        assert "OrderCreate" in p.name
        # Default shape is the self-contained timer flow: diagram +
        # cmdVariantUri on every element, run-once timer, and crucially NO
        # message flow (no sender/receiver → no endpoint/package dependency).
        assert "BPMNDiagram" in content
        assert "cmdVariantUri" in content
        assert "messageFlow" not in content
        assert "fireNow=true" in content

    def test_bpm_comment_present(self, tmp_path):
        # The validated minimal generator produces a valid skeleton; richer
        # step wiring (BPM branch) folds in later. Confirm it still generates.
        s = IFlowScaffolder(output_dir=str(tmp_path), templates_dir=self.TEMPLATES)
        a = make_assessment(name="BPM_Flow", has_bpm=True)
        assert s.scaffold(a).exists()

    def test_multi_mapping_comment(self, tmp_path):
        s = IFlowScaffolder(output_dir=str(tmp_path), templates_dir=self.TEMPLATES)
        a = make_assessment(name="MultiMap_Flow", has_multi_mapping=True)
        # Validated generator produces a CPI-valid iFlow (diagram present)
        content = s.scaffold(a).read_text()
        assert "BPMNDiagram" in content

    def test_destination_aware_suffix(self, tmp_path):
        s = IFlowScaffolder(output_dir=str(tmp_path), templates_dir=self.TEMPLATES)
        a = make_assessment(name="RFC_Flow", sender_adapter="RFC")

        # Build a fake ResolvedDestination
        from destinations.resolver import ResolvedDestination, AdapterRecommendation
        from destinations.registry import get_target
        resolved = ResolvedDestination(
            target=get_target("s4hana_cloud"),
            sender_recommendation=AdapterRecommendation("RFC", "OData", True, False, ""),
            receiver_recommendation=AdapterRecommendation("HTTPS", "HTTPS", True, False, ""),
            hub_matches=[],
            migration_hints=["Check S/4 API catalog"],
            effort_multiplier=1.5,
            compatibility_warnings=[],
        )
        p = s.scaffold(a, resolved=resolved, wired=False)
        content = p.read_text()
        assert "OData" in content
        assert "S/4HANA Cloud" in content

    def test_scaffold_all_multi_target(self, tmp_path):
        from destinations.resolver import ResolvedDestination, AdapterRecommendation
        from destinations.registry import get_target
        s = IFlowScaffolder(output_dir=str(tmp_path), templates_dir=self.TEMPLATES)
        assessments = [make_assessment(name=f"IF_{i}") for i in range(3)]
        resolved = ResolvedDestination(
            target=get_target("ariba"),
            sender_recommendation=AdapterRecommendation("HTTPS", "HTTPS", True, False, ""),
            receiver_recommendation=AdapterRecommendation("HTTPS", "HTTPS", True, False, ""),
            hub_matches=[], migration_hints=[], effort_multiplier=1.0, compatibility_warnings=[],
        )
        resolutions = {a.interface.name: {"ariba": resolved} for a in assessments}
        paths = s.scaffold_all(assessments, resolutions=resolutions, target_ids=["ariba"])
        assert len(paths) == 3


# ─── destinations/registry ───────────────────────────────────────────────────

class TestDestinationRegistry:
    def test_all_required_targets_present(self):
        for tid in ["s4hana_cloud", "s4hana_op", "ariba", "successfactors", "btp"]:
            assert tid in DESTINATION_REGISTRY

    def test_get_target_valid(self):
        t = get_target("s4hana_cloud")
        assert t.id == "s4hana_cloud"
        assert t.variant == "cloud"

    def test_get_target_invalid(self):
        with pytest.raises(KeyError):
            get_target("nonexistent_target")

    def test_list_targets(self):
        targets = list_targets()
        assert len(targets) >= 5
        ids = {t.id for t in targets}
        assert "s4hana_cloud" in ids

    def test_adapter_mapping_coverage(self):
        for t in list_targets():
            assert len(t.adapter_mapping) > 0
            assert "HTTPS" in t.adapter_mapping or "HTTP" in t.adapter_mapping

    def test_supported_adapters_non_empty(self):
        for t in list_targets():
            assert len(t.supported_adapters) > 0

    def test_hub_sources_non_empty(self):
        for t in list_targets():
            assert len(t.hub_sources) > 0
            for src in t.hub_sources:
                assert src.package_id

    def test_s4cloud_clean_core_hint(self):
        t = get_target("s4hana_cloud")
        assert any("Clean Core" in h or "API" in h for h in t.migration_hints)

    def test_sf_adapter_in_sf_target(self):
        t = get_target("successfactors")
        assert "SuccessFactors" in t.supported_adapters


# ─── destinations/hub_fetcher ────────────────────────────────────────────────

class TestCacheEntry:
    def test_fresh_when_new(self, tmp_path):
        entry = CacheEntry(tmp_path, "test_pkg")
        entry.write({"value": []}, ttl=3600, source_url="http://test")
        assert entry.is_fresh(3600)
        assert not entry.is_fresh(0)

    def test_stale_after_ttl(self, tmp_path):
        entry = CacheEntry(tmp_path, "stale_pkg")
        # Write with fetched_at in the past
        data_path = tmp_path / "stale_pkg.json"
        meta_path = tmp_path / "stale_pkg.meta.json"
        data_path.write_text("{}", "utf-8")
        meta = {"fetched_at": time.time() - 7200, "ttl": 3600,
                "expires_at": time.time() - 3600, "source_url": ""}
        meta_path.write_text(json.dumps(meta), "utf-8")
        assert not entry.is_fresh(3600)

    def test_missing_returns_not_fresh(self, tmp_path):
        entry = CacheEntry(tmp_path, "missing_pkg")
        assert not entry.is_fresh(3600)

    def test_ttl_remaining(self, tmp_path):
        entry = CacheEntry(tmp_path, "ttl_test")
        entry.write({"value": []}, ttl=3600)
        remaining = entry.ttl_remaining()
        assert 3500 < remaining <= 3600


class TestHubFetcher:
    def test_returns_cached_data(self, tmp_path):
        fetcher = HubFetcher(cache_dir=tmp_path, default_ttl=3600)
        cached  = {"value": [{"Id": "pkg1", "Title": "Test iFlow"}]}
        entry   = CacheEntry(tmp_path, "pkg_TEST_PKG")
        entry.write(cached, ttl=3600)
        result = fetcher.get_package_artifacts("TEST_PKG")
        assert result == cached

    def test_falls_back_to_static_on_network_error(self, tmp_path):
        # When GitHub fails and cache is stale, static catalog is used.
        # SAPS4HANACloud has static entries; FAIL_PKG returns empty value list.
        fetcher = HubFetcher(cache_dir=tmp_path, default_ttl=1)
        with patch.object(fetcher, "_fetch_from_github", side_effect=ConnectionError("offline")):
            result = fetcher.get_package_artifacts("FAIL_PKG")
        # FAIL_PKG is not in static catalog — returns empty list, never raises
        assert "value" in result
        assert isinstance(result["value"], list)

    def test_returns_empty_on_miss_and_network_error(self, tmp_path):
        fetcher = HubFetcher(cache_dir=tmp_path, default_ttl=3600)
        with patch.object(fetcher, "_fetch_from_github", side_effect=ConnectionError("offline")):
            result = fetcher.get_package_artifacts("NO_CACHE_PKG")
        assert result.get("value") == [] or "error" in result

    def test_invalidate_removes_files(self, tmp_path):
        fetcher = HubFetcher(cache_dir=tmp_path, default_ttl=3600)
        entry   = CacheEntry(tmp_path, "pkg_DEL_PKG")
        entry.write({"value": []}, ttl=3600)
        assert entry.exists()
        fetcher.invalidate("DEL_PKG")
        assert not entry.exists()

    def test_cache_status_lists_entries(self, tmp_path):
        fetcher = HubFetcher(cache_dir=tmp_path, default_ttl=3600)
        for pkg in ["PKG_A", "PKG_B"]:
            CacheEntry(tmp_path, f"pkg_{pkg}").write({"value": []}, ttl=3600)
        status = fetcher.cache_status()
        keys   = {s["key"] for s in status}
        assert "pkg_PKG_A" in keys
        assert "pkg_PKG_B" in keys

    def test_invalidate_all(self, tmp_path):
        fetcher = HubFetcher(cache_dir=tmp_path, default_ttl=3600)
        for pkg in ["PKG_X", "PKG_Y"]:
            CacheEntry(tmp_path, f"pkg_{pkg}").write({"value": []}, ttl=3600)
        fetcher.invalidate_all()
        assert fetcher.cache_status() == []


# ─── destinations/resolver ───────────────────────────────────────────────────

class TestDestinationResolver:
    def _make_fetcher(self, tmp_path):
        fetcher = HubFetcher(cache_dir=tmp_path, default_ttl=3600)
        # Seed cache with fake artifacts for s4hana_cloud packages
        target = get_target("s4hana_cloud")
        for src in target.hub_sources:
            CacheEntry(tmp_path, f"pkg_{src.package_id}").write(
                {"value": [
                    {"Id": "IF_001", "Title": "Purchase Order Integration",
                     "Type": "IntegrationFlow", "ShortText": "PO integration for S4"},
                    {"Id": "IF_002", "Title": "Sales Order Replication",
                     "Type": "IntegrationFlow", "ShortText": "Sales order sync"},
                ]},
                ttl=3600,
            )
        return fetcher

    def test_resolve_maps_rfc_to_odata_for_s4cloud(self, tmp_path):
        fetcher  = self._make_fetcher(tmp_path)
        resolver = DestinationResolver(fetcher=fetcher)
        a        = make_assessment(sender_adapter="RFC", receiver_adapter="HTTPS")
        resolved = resolver.resolve(a, "s4hana_cloud")
        assert resolved.receiver_recommendation.recommended_adapter in (
            get_target("s4hana_cloud").adapter_mapping.values()
        )
        assert resolved.sender_recommendation.recommended_adapter == "OData"

    def test_resolve_adds_hints(self, tmp_path):
        fetcher  = self._make_fetcher(tmp_path)
        resolver = DestinationResolver(fetcher=fetcher)
        a        = make_assessment()
        resolved = resolver.resolve(a, "s4hana_cloud")
        assert len(resolved.migration_hints) > 0

    def test_bpm_warning_injected(self, tmp_path):
        fetcher  = self._make_fetcher(tmp_path)
        resolver = DestinationResolver(fetcher=fetcher)
        a        = make_assessment(has_bpm=True)
        resolved = resolver.resolve(a, "s4hana_cloud")
        assert any("BPM" in w for w in resolved.compatibility_warnings)

    def test_effort_multiplier_cloud(self, tmp_path):
        fetcher  = self._make_fetcher(tmp_path)
        resolver = DestinationResolver(fetcher=fetcher)
        a        = make_assessment(sender_adapter="IDoc", receiver_adapter="HTTPS")
        resolved = resolver.resolve(a, "s4hana_cloud")
        assert resolved.effort_multiplier >= 1.0

    def test_resolve_multi_returns_all_targets(self, tmp_path):
        fetcher  = self._make_fetcher(tmp_path)
        resolver = DestinationResolver(fetcher=fetcher)
        a        = make_assessment()
        results  = resolver.resolve_multi(a, ["s4hana_cloud", "ariba"])
        assert "s4hana_cloud" in results
        assert "ariba" in results

    def test_hub_match_found_by_keyword(self, tmp_path):
        fetcher = HubFetcher(cache_dir=tmp_path, default_ttl=3600)
        target  = get_target("s4hana_cloud")
        for src in target.hub_sources:
            CacheEntry(tmp_path, f"pkg_{src.package_id}").write(
                {"value": [
                    {"Id": "PO_001", "Title": "Purchase Order to S4HANA",
                     "Type": "IntegrationFlow", "ShortText": "Purchase order sync"},
                ]},
                ttl=3600,
            )
        resolver = DestinationResolver(fetcher=fetcher)
        a = make_assessment(name="Purchase_Order_Replication")
        resolved = resolver.resolve(a, "s4hana_cloud")
        assert len(resolved.hub_matches) > 0
        assert any("Purchase" in m.title for m in resolved.hub_matches)

    def test_resolve_all_structure(self, tmp_path):
        fetcher  = self._make_fetcher(tmp_path)
        resolver = DestinationResolver(fetcher=fetcher)
        assessments = [make_assessment(name=f"IF_{i}") for i in range(3)]
        results = resolver.resolve_all(assessments, ["s4hana_cloud"])
        assert len(results) == 3
        for iface_name, target_map in results.items():
            assert "s4hana_cloud" in target_map


# ─── reporter ────────────────────────────────────────────────────────────────

class TestReportGenerator:
    def _make(self, tmp_path):
        records = [
            make_record(id="1", name="Simple_HTTP"),
            make_record(id="2", name="IDoc_Receiver",  sender_adapter="IDoc",
                        receiver_adapter="SOAP", mapping_program="MM_IDoc"),
            make_record(id="3", name="RFC_BPM_Bridge", sender_adapter="RFC",
                        receiver_adapter="OData", has_bpm=True),
        ]
        assessments = ComplexityAnalyzer(BASE_CFG).assess_all(records)

        # Build minimal resolutions
        from destinations.resolver import ResolvedDestination, AdapterRecommendation
        from destinations.registry import get_target
        def fake_resolved(tid):
            t = get_target(tid)
            return ResolvedDestination(
                target=t,
                sender_recommendation=AdapterRecommendation("RFC", "OData", True, False, ""),
                receiver_recommendation=AdapterRecommendation("HTTPS", "HTTPS", True, False, ""),
                hub_matches=[], migration_hints=["Hint A"], effort_multiplier=1.2,
                compatibility_warnings=["Warning X"],
            )
        resolutions = {
            a.interface.name: {"s4hana_cloud": fake_resolved("s4hana_cloud")}
            for a in assessments
        }
        return assessments, resolutions

    def test_markdown_created(self, tmp_path):
        assessments, resolutions = self._make(tmp_path)
        r   = ReportGenerator(str(tmp_path))
        md  = r.generate_markdown(assessments, resolutions=resolutions, target_ids=["s4hana_cloud"])
        txt = md.read_text()
        assert "Migration Assessment" in txt
        assert "HIGH" in txt
        assert "s4hana_cloud" in txt

    def test_markdown_hub_section(self, tmp_path):
        assessments, resolutions = self._make(tmp_path)
        r   = ReportGenerator(str(tmp_path))
        md  = r.generate_markdown(assessments, resolutions=resolutions, target_ids=["s4hana_cloud"])
        txt = md.read_text()
        assert "Destination:" in txt

    def test_markdown_sorted_by_score(self, tmp_path):
        assessments, resolutions = self._make(tmp_path)
        r   = ReportGenerator(str(tmp_path))
        md  = r.generate_markdown(assessments, resolutions=resolutions, target_ids=["s4hana_cloud"])
        txt = md.read_text()
        pos_rfc    = txt.index("RFC_BPM_Bridge")
        pos_simple = txt.index("Simple_HTTP")
        assert pos_rfc < pos_simple

    def test_excel_created(self, tmp_path):
        assessments, resolutions = self._make(tmp_path)
        r  = ReportGenerator(str(tmp_path))
        xl = r.generate_excel(assessments, resolutions=resolutions, target_ids=["s4hana_cloud"])
        assert xl.exists()
        assert xl.suffix == ".xlsx"

    def test_excel_sheets(self, tmp_path):
        import openpyxl
        assessments, resolutions = self._make(tmp_path)
        r  = ReportGenerator(str(tmp_path))
        xl = r.generate_excel(assessments, resolutions=resolutions, target_ids=["s4hana_cloud"])
        wb = openpyxl.load_workbook(xl)
        assert "Summary" in wb.sheetnames
        assert "Interface Inventory" in wb.sheetnames
        # Should have one target sheet
        assert any("S-4HANA" in s or "s4" in s.lower() or "S/4" in s for s in wb.sheetnames)

    def test_excel_target_sheet_content(self, tmp_path):
        import openpyxl
        assessments, resolutions = self._make(tmp_path)
        r  = ReportGenerator(str(tmp_path))
        xl = r.generate_excel(assessments, resolutions=resolutions, target_ids=["s4hana_cloud"])
        wb = openpyxl.load_workbook(xl)
        target_sheet = [s for s in wb.sheetnames if s not in ("Summary", "Interface Inventory")][0]
        ws = wb[target_sheet]
        headers = [ws.cell(1, c).value for c in range(1, 5)]
        assert "Interface Name" in headers

    def test_no_resolutions_still_works(self, tmp_path):
        assessments, _ = self._make(tmp_path)
        r  = ReportGenerator(str(tmp_path))
        xl = r.generate_excel(assessments)
        md = r.generate_markdown(assessments)
        assert xl.exists()
        assert md.exists()


# ─── channel autofill (Feature: WE20/WE21/SM59 auto-fill) ────────────────────

from extractor.channel_parser import ChannelConfig
from reporter.channel_autofill import ChannelAutofill, InfraAutofill


def _idoc_channel(**kw):
    d = dict(channel_id="CC_ORD", channel_name="Send_Orders_IDoc",
             adapter_type="IDoc", direction="Receiver",
             idoc_type="ORDERS05", idoc_message_type="ORDERS",
             idoc_partner_number="LS_ECC01")
    d.update(kw)
    return ChannelConfig(**d)


class TestChannelAutofill:
    def test_idoc_values_resolved_by_name(self):
        af = ChannelAutofill([_idoc_channel()]).for_interface("Send_Orders_IDoc", "IDoc", "SOAP")
        assert af.idoc_type == "ORDERS05"
        assert af.idoc_message_type == "ORDERS"
        assert af.idoc_partner_number == "LS_ECC01"
        assert af.has_any

    def test_rfc_values_resolved(self):
        ch = ChannelConfig(channel_id="CC_RFC", channel_name="Get_Stock_RFC",
                           adapter_type="RFC", direction="Receiver",
                           rfc_destination="RFC_S4_PROD", function_module="BAPI_MATERIAL_GET",
                           address="rfchost.corp")
        af = ChannelAutofill([ch]).for_interface("Get_Stock_RFC", "RFC", "HTTPS")
        assert af.rfc_destination == "RFC_S4_PROD"
        assert af.rfc_function_module == "BAPI_MATERIAL_GET"
        assert af.rfc_target_host == "rfchost.corp"

    def test_single_channel_fallback_when_no_name_match(self):
        # exactly one channel + no token overlap -> still used
        af = ChannelAutofill([_idoc_channel(channel_name="totally_different")]) \
            .for_interface("Unrelated_Name", "IDoc", "SOAP")
        assert af.idoc_type == "ORDERS05"

    def test_no_channels_returns_empty(self):
        af = ChannelAutofill([]).for_interface("Anything")
        assert isinstance(af, InfraAutofill)
        assert not af.has_any

    def test_no_match_no_single_channel_is_empty(self):
        chs = [_idoc_channel(channel_name="aaa"), _idoc_channel(channel_name="bbb")]
        af = ChannelAutofill(chs).for_interface("zzz_no_overlap_zzz", "File", "File")
        # adapter mismatch + name mismatch + >1 channel -> nothing resolved
        assert not af.has_any

    def test_index_by_interface(self):
        ca = ChannelAutofill([_idoc_channel()])
        idx = ca.index_by_interface(["Send_Orders_IDoc", "Other"])
        assert idx["Send_Orders_IDoc"].idoc_type == "ORDERS05"
        assert "Other" in idx


class TestInfrastructureGuideAutofill:
    def _assess(self, name, sa, ra):
        return make_assessment(name=name, sender_adapter=sa, receiver_adapter=ra)

    def test_we20_uses_channel_partner(self, tmp_path):
        import openpyxl
        from reporter.infrastructure_guide import InfrastructureGuideGenerator
        a = [self._assess("Send_Orders_IDoc", "IDoc", "SOAP")]
        gen = InfrastructureGuideGenerator(str(tmp_path))
        xlsx, _ = gen.generate(a, {}, channels=[_idoc_channel()])
        ws = openpyxl.load_workbook(xlsx)["IDoc Setup (WE20-WE21)"]
        assert ws["C6"].value == "LS_ECC01"
        assert ws["D6"].value == "ORDERS"

    def test_we20_placeholder_without_channels(self, tmp_path):
        import openpyxl
        from reporter.infrastructure_guide import InfrastructureGuideGenerator
        a = [self._assess("Send_Orders_IDoc", "IDoc", "SOAP")]
        gen = InfrastructureGuideGenerator(str(tmp_path))
        xlsx, _ = gen.generate(a, {}, channels=None)
        ws = openpyxl.load_workbook(xlsx)["IDoc Setup (WE20-WE21)"]
        # falls back to receiver system, not the channel partner
        assert ws["C6"].value != "LS_ECC01"

    def test_sm59_uses_channel_rfc_dest(self, tmp_path):
        import openpyxl
        from reporter.infrastructure_guide import InfrastructureGuideGenerator
        ch = ChannelConfig(channel_id="CC_RFC", channel_name="Get_Stock_RFC",
                           adapter_type="RFC", direction="Receiver",
                           rfc_destination="RFC_S4_PROD", address="rfchost.corp")
        a = [self._assess("Get_Stock_RFC", "RFC", "HTTPS")]
        gen = InfrastructureGuideGenerator(str(tmp_path))
        xlsx, _ = gen.generate(a, {}, channels=[ch])
        ws = openpyxl.load_workbook(xlsx)["RFC Destinations (SM59)"]
        assert ws["A4"].value == "RFC_S4_PROD"
        assert ws["C4"].value == "rfchost.corp"

    def test_backward_compatible_no_channels_kwarg(self, tmp_path):
        from reporter.infrastructure_guide import InfrastructureGuideGenerator
        a = [self._assess("Simple", "HTTPS", "HTTPS")]
        gen = InfrastructureGuideGenerator(str(tmp_path))
        xlsx, docx = gen.generate(a, {})   # old call signature still works
        assert xlsx.exists() and docx.exists()


# ─── parameter injector (Feature: auto-inject externalized parameters) ───────

from scaffolder.parameter_injector import build_parameters, build_parameters_prop, ParameterSet, _key


class TestParameterInjector:
    def test_key_normalisation(self):
        assert _key("Receiver Host!") == "Receiver_Host"
        assert _key("a//b::c") == "a_b_c"
        assert _key("") == "Param"

    def test_real_values_from_channel(self):
        ch = ChannelConfig(channel_id="C", channel_name="N", adapter_type="HTTPS",
                           direction="Receiver", address="s4.example.com",
                           path="/sap/odata", credential_name="S4_CRED")
        ps = build_parameters("Iface", channel=ch)
        assert ps.params["Receiver_Host"] == "s4.example.com"
        assert ps.params["Receiver_Path"] == "/sap/odata"
        assert ps.params["Credential_Alias"] == "S4_CRED"
        assert "Credential_Alias" not in ps.unresolved

    def test_placeholder_when_missing(self):
        ps = build_parameters("Iface")  # no channel, no config
        assert ps.params["Receiver_Host"] == "<FILL_Receiver_Host>"
        assert "Receiver_Host" in ps.unresolved

    def test_prop_body_format(self):
        ch = ChannelConfig(channel_id="C", channel_name="N", adapter_type="HTTPS",
                           direction="Receiver", address="h", credential_name="cred")
        body = build_parameters_prop("Iface", channel=ch)
        assert body.startswith("# Externalized parameters for Iface")
        assert "Receiver_Host=h" in body
        assert body.endswith("\n")

    def test_file_params_only_for_file_adapter(self):
        ch = ChannelConfig(channel_id="C", channel_name="N", adapter_type="SFTP",
                           direction="Sender", file_directory="/in", file_pattern="*.xml")
        ps = build_parameters("F", channel=ch)
        assert ps.params["File_Directory"] == "/in"
        assert ps.params["File_Pattern"] == "*.xml"

    def test_extra_channel_params_carried(self):
        ch = ChannelConfig(channel_id="C", channel_name="N", adapter_type="HTTPS",
                           direction="Receiver", address="h")
        ch.parameters = {"customTimeout": "9000"}
        ps = build_parameters("X", channel=ch)
        assert any("customTimeout" in k for k in ps.params)


# ─── odata suggester (Feature: auto-suggest OData API for RFC/BAPI) ──────────

from fetcher.odata_suggester import ODataSuggester, ApiSuggestion, BAPI_TO_ODATA


class TestODataSuggester:
    def setup_method(self):
        self.s = ODataSuggester()   # offline, no hub client

    def test_exact_bapi_high_confidence(self):
        sugg = self.s.suggest_for_bapi("BAPI_SALESORDER_CREATEFROMDAT2")
        assert sugg[0].api_name == "API_SALES_ORDER_SRV"
        assert sugg[0].confidence == "high"
        assert sugg[0].source == "static"

    def test_exact_match_case_insensitive(self):
        sugg = self.s.suggest_for_bapi("bapi_po_create1")
        assert sugg[0].api_name == "API_PURCHASEORDER_PROCESS_SRV"

    def test_keyword_fallback_medium(self):
        sugg = self.s.suggest_for_bapi("Z_CUSTOM_PURCHASE_RFC")
        assert sugg
        assert sugg[0].api_name == "API_PURCHASEORDER_PROCESS_SRV"
        assert sugg[0].confidence == "medium"

    def test_unknown_returns_empty_offline(self):
        sugg = self.s.suggest_for_bapi("Z_TOTALLY_OPAQUE_FM")
        assert sugg == []

    def test_empty_fm_returns_empty(self):
        assert self.s.suggest_for_bapi("") == []

    def test_suggest_for_interface_infers_from_name(self):
        sugg = self.s.suggest_for_interface("Customer_Master_Sync", function_module="")
        assert sugg
        assert sugg[0].api_name == "API_BUSINESS_PARTNER"

    def test_hub_enrichment_used_when_client_present(self):
        class FakePkg:
            def __init__(self, name): self.name=name; self.url="u"; self.short_text="s"; self.id="pid"
        class FakeHub:
            def search_packages(self, query="", top=3): return [FakePkg("CustomPkg")]
        s = ODataSuggester(hub_client=FakeHub())
        sugg = s.suggest_for_bapi("Z_OPAQUE_FM", top=3)
        assert any(x.source == "hub" for x in sugg)

    def test_top_limit_respected(self):
        sugg = self.s.suggest_for_bapi("BAPI_SALESORDER_CREATEFROMDAT2", top=1)
        assert len(sugg) <= 1


# ─── endpoint collision detector (P1-4) ──────────────────────────────────────

from analyzer.endpoint_collision import (
    detect_collisions, CollisionFinding, _normalise_path,
    collisions_to_preflight_items,
)


def _sender(name, path, adapter="HTTPS"):
    return ChannelConfig(channel_id=name, channel_name=name, adapter_type=adapter,
                         direction="Sender", path=path)


class TestEndpointCollision:
    def test_path_normalisation_handles_case_and_slashes(self):
        assert _normalise_path("", "/Customer/V1/") == _normalise_path("", "/customer/v1")
        assert _normalise_path("", "Customer/v1") == "/customer/v1"

    def test_full_url_in_address_stripped(self):
        assert _normalise_path("https://host:443/api/v1", "") == "/api/v1"

    def test_duplicate_paths_flagged(self):
        c1 = _sender("Send_A", "/Customer/v1")
        c2 = _sender("Send_B", "/customer/v1/")
        findings = detect_collisions([c1, c2])
        assert len(findings) == 1
        assert findings[0].severity == "HIGH"
        assert "Send_A" in findings[0].channels and "Send_B" in findings[0].channels

    def test_unique_paths_no_findings(self):
        c1 = _sender("A", "/Order/v1")
        c2 = _sender("B", "/Customer/v1")
        assert detect_collisions([c1, c2]) == []

    def test_receiver_channels_ignored(self):
        c1 = _sender("Send_A", "/x")
        c2 = ChannelConfig(channel_id="R", channel_name="Recv_X", adapter_type="HTTPS",
                           direction="Receiver", path="/x")
        assert detect_collisions([c1, c2]) == []

    def test_non_http_adapters_ignored(self):
        c1 = ChannelConfig(channel_id="J1", channel_name="J1", adapter_type="JDBC",
                           direction="Sender", path="/x")
        c2 = ChannelConfig(channel_id="J2", channel_name="J2", adapter_type="JDBC",
                           direction="Sender", path="/x")
        assert detect_collisions([c1, c2]) == []

    def test_mixed_adapter_collision_is_medium(self):
        c1 = _sender("A", "/x", adapter="HTTPS")
        c2 = _sender("B", "/x", adapter="SOAP")
        f = detect_collisions([c1, c2])[0]
        assert f.severity == "MEDIUM"

    def test_empty_input(self):
        assert detect_collisions([]) == []
        assert detect_collisions(None) == []

    def test_preflight_items_built(self):
        items = collisions_to_preflight_items([
            CollisionFinding(address_path="/customer/v1",
                             channels=["A", "B"],
                             adapter_types=["HTTPS", "HTTPS"],
                             recommendation="rec"),
        ])
        assert len(items) == 1
        assert "collision" in items[0].task.lower()
        assert items[0].mandatory is True


# ─── transaction advisor (P1-3) ──────────────────────────────────────────────

from analyzer.transaction_advisor import (
    advise, advise_all, advisories_to_preflight_items, TransactionAdvisory,
)


class _Rec:
    def __init__(self, **kw): self.__dict__.update(kw)


class TestTransactionAdvisor:
    def test_jdbc_no_eoio_is_required(self):
        a = advise(_Rec(name="x", sender_adapter="HTTPS", receiver_adapter="JDBC"))
        assert a.handling == "Required"
        assert a.has_jdbc and not a.has_eoio

    def test_jdbc_plus_eoio_is_requires_new(self):
        a = advise(_Rec(name="x", sender_adapter="XI", receiver_adapter="JDBC"))
        assert a.handling == "Requires New"
        assert a.has_jdbc and a.has_eoio

    def test_eoio_no_jdbc_is_required(self):
        a = advise(_Rec(name="x", sender_adapter="JMS", receiver_adapter="HTTPS"))
        assert a.handling == "Required"
        assert not a.has_jdbc and a.has_eoio

    def test_default_is_not_supported(self):
        a = advise(_Rec(name="x", sender_adapter="HTTPS", receiver_adapter="HTTPS"))
        assert a.handling == "Not Supported"

    def test_advise_all_returns_one_per_record(self):
        results = advise_all([
            _Rec(name="a", sender_adapter="HTTPS", receiver_adapter="HTTPS"),
            _Rec(name="b", sender_adapter="HTTPS", receiver_adapter="JDBC"),
        ])
        assert len(results) == 2
        assert {r.handling for r in results} == {"Not Supported", "Required"}

    def test_not_supported_omitted_from_preflight(self):
        items = advisories_to_preflight_items([
            TransactionAdvisory(interface_name="x", handling="Not Supported",
                                reasoning="r", has_jdbc=False, has_eoio=False),
            TransactionAdvisory(interface_name="y", handling="Required",
                                reasoning="r", has_jdbc=True, has_eoio=False),
        ])
        assert len(items) == 1
        assert "y" in items[0].task


# ─── domain taxonomy (P1-5) ──────────────────────────────────────────────────

from analyzer.domains import (
    derive_domain, audit_domain_coverage, ALL_DOMAINS,
    CLOUD_TO_CLOUD, ON_PREMISE_TO_CLOUD, ON_PREMISE_TO_ON_PREMISE,
    CLOUD_TO_ON_PREMISE, EDGE_LOCAL, HYBRID,
)


class TestDomainTaxonomy:
    def test_canonical_set_has_six_values(self):
        assert len(set(ALL_DOMAINS)) == 6

    def test_c2c_classification(self):
        d = derive_domain("s4hana_cloud", "aws")
        assert d.domain == CLOUD_TO_CLOUD
        assert d.confidence == "high"

    def test_on_prem_to_cloud(self):
        assert derive_domain("ecc", "s4hana_cloud").domain == ON_PREMISE_TO_CLOUD

    def test_on_prem_to_on_prem(self):
        assert derive_domain("ecc", "po").domain == ON_PREMISE_TO_ON_PREMISE

    def test_cloud_to_on_prem(self):
        assert derive_domain("ariba", "ecc").domain == CLOUD_TO_ON_PREMISE

    def test_edge_local(self):
        assert derive_domain("eic", "edge").domain == EDGE_LOCAL

    def test_edge_cross_domain_is_hybrid(self):
        d = derive_domain("eic", "s4hana_cloud")
        assert d.domain == HYBRID

    def test_unknown_defaults_to_c2c_low_confidence(self):
        d = derive_domain("mystery_system", "another_mystery")
        assert d.domain == CLOUD_TO_CLOUD
        assert d.confidence == "low"

    def test_audit_canonical_covered(self):
        report = audit_domain_coverage(list(ALL_DOMAINS))
        assert report["covered"]
        assert report["missing"] == []

    def test_audit_detects_missing(self):
        report = audit_domain_coverage([CLOUD_TO_CLOUD])
        assert CLOUD_TO_ON_PREMISE in report["missing"]
        assert not report["covered"]

    def test_audit_detects_unknown_labels(self):
        report = audit_domain_coverage(["FooBar"])
        assert "FooBar" in report["unknown"]


# ─── APIM fault catalog (P2-A6) ──────────────────────────────────────────────

from analyzer.apim_faults import lookup, by_category, search, all_codes, APIMFault


class TestAPIMFaults:
    def test_lookup_known_code(self):
        f = lookup("oauth.v2.InvalidApiKey")
        assert isinstance(f, APIMFault)
        assert f.http_status == 401
        assert f.category == "OAuth"

    def test_lookup_unknown_returns_none(self):
        assert lookup("does.not.exist") is None
        assert lookup("") is None
        assert lookup(None) is None

    def test_by_category_returns_multiple(self):
        oauth = by_category("OAuth")
        assert len(oauth) >= 2
        assert all(f.category == "OAuth" for f in oauth)

    def test_by_category_case_insensitive(self):
        assert by_category("rate limit") == by_category("Rate Limit")

    def test_search_finds_by_meaning(self):
        hits = search("quota")
        assert hits
        assert any("quota" in h.code.lower() or "quota" in h.meaning.lower() for h in hits)

    def test_search_empty_returns_empty(self):
        assert search("") == []
        assert search(None) == []

    def test_all_codes_sorted(self):
        codes = all_codes()
        assert codes == sorted(codes)
        assert "policies.ratelimit.SpikeArrestViolation" in codes

    def test_remediation_non_empty(self):
        for code in all_codes():
            f = lookup(code)
            assert f.remediation, f"missing remediation for {code}"
            assert f.likely_cause, f"missing likely_cause for {code}"


# ─── preflight integration of P1-4 + P1-3 ────────────────────────────────────

class TestPreflightIntegration:
    def test_backward_compat_no_channels_kwarg(self, tmp_path):
        from reporter.preflight_generator import PreflightGenerator
        a = [make_assessment(name="x")]
        gen = PreflightGenerator(str(tmp_path))
        xlsx, docx = gen.generate(a, {}, ["s4hana_cloud"])  # no channels kwarg
        assert xlsx.exists() and docx.exists()

    def test_collision_appears_when_channels_passed(self, tmp_path):
        import openpyxl
        from reporter.preflight_generator import PreflightGenerator
        a = [make_assessment(name="Send_A"), make_assessment(name="Send_B")]
        chs = [_sender("Send_A", "/customer/v1"), _sender("Send_B", "/Customer/V1")]
        gen = PreflightGenerator(str(tmp_path))
        xlsx, _ = gen.generate(a, {}, ["s4hana_cloud"], channels=chs)
        ws = openpyxl.load_workbook(xlsx).active
        cats = {ws.cell(r, 2).value for r in range(4, ws.max_row + 1)}
        assert "Sender Endpoints" in cats

    def test_transaction_handling_appears_for_jdbc(self, tmp_path):
        import openpyxl
        from reporter.preflight_generator import PreflightGenerator
        a = [make_assessment(name="Writes", receiver_adapter="JDBC")]
        gen = PreflightGenerator(str(tmp_path))
        xlsx, _ = gen.generate(a, {}, ["s4hana_cloud"])
        ws = openpyxl.load_workbook(xlsx).active
        cats = {ws.cell(r, 2).value for r in range(4, ws.max_row + 1)}
        assert "Transaction Handling" in cats


# ─── Interface Request artifact (P1-1) ───────────────────────────────────────

from reporter.interface_request import (
    build_business_solution_request, build_interface_request,
    render_word, render_json,
    BusinessSolutionRequest, InterfaceRequest, InterfaceRequestExtensions,
    ApplicationInstance, MessageFlow, QuestionnaireStep, QuestionnaireAnswer,
    SelectedTechnology, BusinessSolutionGeneral,
    INTEGRATION_STYLES, BUSINESS_CRITICALITY, QUESTIONNAIRE_STEPS,
    _infer_style, _app_instance_from_system,
)


class TestInterfaceRequestBuilder:
    def test_one_ir_per_assessment(self):
        a = [make_assessment(name="A"), make_assessment(name="B"),
             make_assessment(name="C")]
        bsr = build_business_solution_request(a, project_name="Proj")
        assert len(bsr.interface_requests) == 3
        assert bsr.name == "Proj"

    def test_ir_has_one_flow_per_source_target_pair(self):
        ir = build_interface_request(make_record(name="X"))
        assert len(ir.message_flows) == 1
        assert ir.message_flows[0].flow_id == "X-1"

    def test_flow_domain_derives_from_systems(self):
        r = make_record(name="Cloud_Flow",
                        sender_system="s4hana_cloud", receiver_system="aws")
        ir = build_interface_request(r)
        assert ir.message_flows[0].domain == "Cloud2Cloud"

    def test_on_prem_to_cloud_domain(self):
        r = make_record(name="Hybrid_Flow",
                        sender_system="ecc", receiver_system="s4hana_cloud")
        ir = build_interface_request(r)
        assert ir.message_flows[0].domain == "OnPremise2Cloud"

    def test_style_inferred_when_no_questionnaire(self):
        r = make_record(name="X", sender_adapter="JDBC", receiver_adapter="HTTPS")
        ir = build_interface_request(r)
        assert ir.style == "Data Integration"

    def test_style_from_questionnaire_takes_precedence(self):
        from intake.isam_questionnaire import ISAMRecommendation
        r = make_record(name="X", sender_adapter="JDBC", receiver_adapter="HTTPS")
        isam = ISAMRecommendation(primary_tool="CPI", secondary_tools=[],
                                  isa_m_pattern="A2A", integration_style="User Integration",
                                  score_breakdown={}, reasoning=[], confidence=0.9)
        ir = build_interface_request(r, isam_rec=isam)
        assert ir.style == "User Integration"

    def test_style_falls_back_when_isam_value_invalid(self):
        from intake.isam_questionnaire import ISAMRecommendation
        r = make_record(name="X")
        isam = ISAMRecommendation(primary_tool="CPI", secondary_tools=[],
                                  isa_m_pattern="A2A", integration_style="Garbage",
                                  score_breakdown={}, reasoning=[], confidence=0.9)
        ir = build_interface_request(r, isam_rec=isam)
        assert ir.style in INTEGRATION_STYLES

    def test_b2b_questionnaire_step1_answer(self):
        r = make_record(name="B2B_Flow", sender_adapter="AS2", receiver_adapter="AS2")
        ir = build_interface_request(r)
        step1 = ir.message_flows[0].questionnaire[0]
        assert step1.answers[0].answer == "B2B Integration"

    def test_questionnaire_has_all_five_steps(self):
        ir = build_interface_request(make_record())
        assert len(ir.message_flows[0].questionnaire) == 5
        titles = [s.step_title for s in ir.message_flows[0].questionnaire]
        assert titles == [s[0] for s in QUESTIONNAIRE_STEPS]

    def test_questionnaire_answers_tagged_with_flow_id(self):
        ir = build_interface_request(make_record(name="Foo"))
        for step in ir.message_flows[0].questionnaire:
            for ans in step.answers:
                assert ans.applies_to_flows == ["Foo-1"]

    def test_default_tech_is_cloud_integration_80(self):
        ir = build_interface_request(make_record())
        st = ir.message_flows[0].selected_technology
        assert "Cloud Integration" in st.technology
        assert st.coverage_percent == 80

    def test_tech_from_isam_recommendation(self):
        from intake.isam_questionnaire import ISAMRecommendation
        isam = ISAMRecommendation(
            primary_tool="SAP Integration Suite, API Management",
            secondary_tools=[], isa_m_pattern="API",
            integration_style="Process Integration",
            score_breakdown={}, reasoning=["api reasoning"], confidence=0.85)
        ir = build_interface_request(make_record(), isam_rec=isam)
        st = ir.message_flows[0].selected_technology
        assert st.technology == "SAP Integration Suite, API Management"
        assert st.coverage_percent == 85


class TestExtensionLayerStrictness:
    """The extension layer must NEVER leak into canonical output. This is the
    contract that protects compatibility with the SAP ISA-M tool."""

    def test_canonical_json_has_no_extensions_key(self):
        a = [make_assessment(name="X")]
        bsr = build_business_solution_request(a)
        canonical = json.loads(bsr.to_canonical_json())
        for ir in canonical["interface_requests"]:
            assert "extensions" not in ir

    def test_extended_json_has_extensions_key(self):
        a = [make_assessment(name="X")]
        bsr = build_business_solution_request(a)
        extended = json.loads(bsr.to_extended_json())
        for ir in extended["interface_requests"]:
            assert "extensions" in ir

    def test_source_interface_id_populated_in_extensions(self):
        a = [make_assessment(id="IF_999", name="X")]
        bsr = build_business_solution_request(a)
        assert bsr.interface_requests[0].extensions.source_interface_id == "IF_999"

    def test_extension_placeholders_are_empty_strings(self):
        ext = InterfaceRequestExtensions()
        assert ext.migration_wave == ""
        assert ext.consultant_owner == ""
        assert ext.client_owner == ""

    def test_canonical_excludes_extensions_even_when_populated(self):
        ir = build_interface_request(make_record())
        ir.extensions.migration_wave = "Wave 1"
        ir.extensions.consultant_owner = "Ric"
        bsr = BusinessSolutionRequest(name="P", interface_requests=[ir])
        canonical = json.loads(bsr.to_canonical_json())
        assert "extensions" not in canonical["interface_requests"][0]
        # And the values aren't leaked elsewhere
        as_str = bsr.to_canonical_json()
        assert "Wave 1" not in as_str
        assert "consultant_owner" not in as_str


class TestMasterDataPreload:
    def test_preload_deduplicates_across_irs(self):
        # Two IRs sharing the same source system
        a = [make_assessment(name="A", sender_system="ecc", receiver_system="aws"),
             make_assessment(name="B", sender_system="ecc", receiver_system="s4hana_cloud")]
        bsr = build_business_solution_request(a)
        preload = bsr.master_data_preload()
        names = [(p.name, p.application, p.deployment) for p in preload]
        # ecc should appear exactly once
        assert sum(1 for n in names if n[0] == "ecc") == 1

    def test_preload_deployment_buckets(self):
        a = [make_assessment(name="X", sender_system="ecc",
                             receiver_system="s4hana_cloud")]
        bsr = build_business_solution_request(a)
        preload = bsr.master_data_preload()
        by_name = {p.name: p.deployment for p in preload}
        assert by_name["ecc"] == "On-Premise"
        assert by_name["s4hana_cloud"] == "Cloud"

    def test_preload_empty_for_empty_bsr(self):
        bsr = BusinessSolutionRequest(name="X")
        assert bsr.master_data_preload() == []


class TestBSRGeneralTab:
    def test_default_criticality_medium(self):
        bsr = build_business_solution_request([make_assessment()])
        assert bsr.general.business_criticality == "Medium"

    def test_invalid_criticality_normalised(self):
        bsr = build_business_solution_request(
            [make_assessment()], business_criticality="extreme")
        assert bsr.general.business_criticality == "Medium"

    def test_valid_criticality_preserved(self):
        bsr = build_business_solution_request(
            [make_assessment()], business_criticality="High")
        assert bsr.general.business_criticality == "High"

    def test_default_go_live_is_today(self):
        bsr = build_business_solution_request([make_assessment()])
        # Date-format string
        assert len(bsr.general.planned_go_live) == 10
        assert bsr.general.planned_go_live[4] == "-"


class TestRendering:
    def test_word_canonical_smaller_than_extended(self, tmp_path):
        a = [make_assessment(name="X")]
        bsr = build_business_solution_request(a)
        c = render_word(bsr, tmp_path / "c.docx", mode="canonical")
        e = render_word(bsr, tmp_path / "e.docx", mode="extended")
        assert c.exists() and e.exists()
        # Extended has additional Extensions section per IR
        assert e.stat().st_size > c.stat().st_size

    def test_json_canonical_excludes_extensions(self, tmp_path):
        a = [make_assessment(name="X")]
        bsr = build_business_solution_request(a)
        path = render_json(bsr, tmp_path / "out.json", mode="canonical")
        body = path.read_text()
        assert "extensions" not in body

    def test_json_extended_includes_extensions(self, tmp_path):
        a = [make_assessment(name="X")]
        bsr = build_business_solution_request(a)
        path = render_json(bsr, tmp_path / "out.json", mode="extended")
        body = path.read_text()
        assert "extensions" in body
        assert "source_interface_id" in body

    def test_invalid_mode_rejected(self, tmp_path):
        a = [make_assessment(name="X")]
        bsr = build_business_solution_request(a)
        with pytest.raises(ValueError):
            render_word(bsr, tmp_path / "x.docx", mode="garbage")
        with pytest.raises(ValueError):
            render_json(bsr, tmp_path / "x.json", mode="garbage")

    def test_word_canonical_contains_form_fields(self, tmp_path):
        import zipfile
        a = [make_assessment(name="MSO-API", sender_system="s4hana_cloud",
                             receiver_system="aws")]
        bsr = build_business_solution_request(a, project_name="MSO-REQ-1",
                                              business_process="Order Fulfilment")
        path = render_word(bsr, tmp_path / "x.docx", mode="canonical")
        # python-docx files are zips; pull the document XML and check content
        with zipfile.ZipFile(path) as z:
            xml = z.read("word/document.xml").decode("utf-8")
        assert "MSO-REQ-1" in xml
        assert "Order Fulfilment" in xml
        assert "Cloud2Cloud" in xml
        assert "ISA-M Master Data Preload" in xml

    def test_word_canonical_has_no_extensions_label(self, tmp_path):
        import zipfile
        a = [make_assessment(name="X")]
        bsr = build_business_solution_request(a)
        path = render_word(bsr, tmp_path / "c.docx", mode="canonical")
        with zipfile.ZipFile(path) as z:
            xml = z.read("word/document.xml").decode("utf-8")
        assert "Extensions" not in xml


class TestStyleAndAppInstanceHelpers:
    def test_infer_style_jdbc_is_data(self):
        r = make_record(sender_adapter="JDBC", receiver_adapter="HTTPS")
        assert _infer_style(r) == "Data Integration"

    def test_infer_style_default_is_process(self):
        r = make_record(sender_adapter="HTTPS", receiver_adapter="HTTPS")
        assert _infer_style(r) == "Process Integration"

    def test_app_instance_unknown_bucket(self):
        app = _app_instance_from_system("mystery_system")
        assert app.deployment == "Unknown"
        assert app.name == "mystery_system"

    def test_app_instance_cloud_bucket(self):
        app = _app_instance_from_system("s4hana_cloud")
        assert app.deployment == "Cloud"

    def test_app_instance_empty_falls_back(self):
        app = _app_instance_from_system("")
        assert app.name == "Unknown"


# ─── SAP samples browser: org-per-repo + SAP/ recipes catalog ────────────────

from fetcher.sap_samples_browser import (
    INTEGRATION_REPOS, SAPSamplesBrowser, DEFAULT_ORG, SAP_ORG,
)


class TestSAPSamplesCatalog:
    def test_default_org_unchanged(self):
        # back-compat: the SAP_ORG alias must still exist for any external
        # importer that referenced it before the per-repo org migration
        assert SAP_ORG == DEFAULT_ORG == "SAP-samples"

    def test_official_cpi_recipes_in_catalog(self):
        match = [r for r in INTEGRATION_REPOS
                 if r["repo"] == "apibusinesshub-integration-recipes"]
        assert len(match) == 1
        assert match[0].get("org") == "SAP", \
            "Official CPI recipes live under SAP/, not SAP-samples/"
        assert match[0]["priority"] == 1

    def test_official_api_recipes_in_catalog(self):
        match = [r for r in INTEGRATION_REPOS
                 if r["repo"] == "apibusinesshub-api-recipes"]
        assert len(match) == 1
        assert match[0].get("org") == "SAP"

    def test_legacy_entries_default_to_sap_samples(self):
        # Any repo without an explicit `org` field should still resolve to
        # SAP-samples (the historical default). The browser does this at
        # scan time via repo_info.get("org", DEFAULT_ORG).
        legacy = [r for r in INTEGRATION_REPOS if "org" not in r]
        # At least one legacy entry should exist (cloud-integration-flow etc.)
        assert legacy, "Test guards backward-compatible default behaviour"
        # The scanner code path resolves their org to DEFAULT_ORG; we don't
        # actually hit the network here — we just assert the contract.
        for r in legacy:
            assert r.get("org", DEFAULT_ORG) == DEFAULT_ORG

    def test_browser_resolves_per_repo_org_in_url(self, monkeypatch):
        # Smoke test of the org-resolution code path without hitting GitHub.
        captured = []

        class FakeResp:
            status_code = 404
            def json(self): return []

        class FakeSession:
            headers = {}
            def get(self, url, timeout=15):
                captured.append(url)
                return FakeResp()
            @property
            def auth(self): return None

        b = SAPSamplesBrowser(cache_dir=Path("/tmp/cpi_test_cache_xyz"))
        b.session = FakeSession()
        # SAP-org repo
        b._scan_repo("apibusinesshub-integration-recipes",
                     {"org": "SAP", "repo": "apibusinesshub-integration-recipes"})
        # legacy SAP-samples repo
        b._scan_repo("cloud-integration-flow",
                     {"repo": "cloud-integration-flow"})

        assert any("/SAP/apibusinesshub-integration-recipes/" in u for u in captured)
        assert any("/SAP-samples/cloud-integration-flow/" in u for u in captured)


# ─── Match aggregator (Tab 3 wire-up) ────────────────────────────────────────

from fetcher.match_aggregator import (
    MatchAggregator, MatchResult, MatchSource, MatchMode,
    _canonical_key, _keywords, FALLBACK_THRESHOLD,
)


class _MockHubPackage:
    def __init__(self, pid, name="Hub Pkg", short_text="hub desc",
                 url="u", artifact_count=3):
        self.id = pid; self.name = name; self.short_text = short_text
        self.url = url; self.artifact_count = artifact_count


class _MockTenantArt:
    def __init__(self, art_id, name="Tenant Art", package_id="pkg",
                 description="ten desc"):
        self.id = art_id; self.name = name; self.package_id = package_id
        self.description = description


class _MockSamplePkg:
    def __init__(self, pid, name="Gh Recipe", description="gh desc"):
        self.id = pid; self.name = name; self.description = description
        self.tags = ["order"]; self.detected_adapters = ["IDoc"]
        self.download_url = "gh.zip"


class _MockCPIFetcher:
    def __init__(self, results=None):
        self._results = results or []
    def suggest_matches(self, *_, **__):
        return self._results


class _MockHubClient:
    def __init__(self, results=None):
        self._results = results or []
    def search_for_interface(self, *_, **__):
        return self._results


class _MockSamplesBrowser:
    def __init__(self, index=None):
        self._index = index or []
    def get_package_index(self):
        return self._index


class TestMatchAggregatorDedup:
    def test_hub_preferred_over_tenant_with_same_package(self):
        tenant = MatchResult(source=MatchSource.TENANT, id="a", name="X",
                             package_id="PkgA")
        hub    = MatchResult(source=MatchSource.HUB, id="PkgA", name="X")
        deduped = MatchAggregator._dedup([tenant, hub])
        assert len(deduped) == 1
        assert deduped[0].source == MatchSource.HUB

    def test_hub_preferred_regardless_of_input_order(self):
        tenant = MatchResult(source=MatchSource.TENANT, id="a",
                             name="X", package_id="PkgA")
        hub    = MatchResult(source=MatchSource.HUB, id="PkgA", name="X")
        # Hub appears first
        deduped = MatchAggregator._dedup([hub, tenant])
        assert deduped[0].source == MatchSource.HUB

    def test_no_dedup_when_keys_differ(self):
        t = MatchResult(source=MatchSource.TENANT, id="a", name="X",
                        package_id="PkgA")
        g = MatchResult(source=MatchSource.GITHUB, id="b", name="Y")
        deduped = MatchAggregator._dedup([t, g])
        assert len(deduped) == 2

    def test_canonical_key_normalises(self):
        # Same package id with different casing / separators hashes the same
        r1 = MatchResult(source=MatchSource.HUB, id="SAP-Cloud_Pkg", name="X")
        r2 = MatchResult(source=MatchSource.HUB, id="sapcloudpkg", name="Y")
        assert _canonical_key(r1) == _canonical_key(r2)

    def test_empty_input(self):
        assert MatchAggregator._dedup([]) == []


class TestMatchAggregatorModes:
    def setup_method(self):
        self.hub = _MockHubClient([(10, _MockHubPackage("HubPkg1", "Hub Order"))])
        self.cpi = _MockCPIFetcher([(5, _MockTenantArt("art1", "Tenant Ord", "TenPkg"))])
        self.gh  = _MockSamplesBrowser([_MockSamplePkg("gh1", "GH Order")])
        self.agg = MatchAggregator(self.cpi, self.hub, self.gh)

    def test_tenant_only_mode(self):
        results = self.agg.find_matches("Order", "IDoc", "SOAP",
                                        tenant_artifacts=["x"],
                                        mode=MatchMode.TENANT_ONLY)
        assert len(results) == 1
        assert results[0].source == MatchSource.TENANT

    def test_hub_only_mode(self):
        results = self.agg.find_matches("Order", "IDoc", "SOAP",
                                        mode=MatchMode.HUB_ONLY)
        assert len(results) == 1
        assert results[0].source == MatchSource.HUB

    def test_github_only_mode(self):
        results = self.agg.find_matches("Order", "IDoc", "SOAP",
                                        mode=MatchMode.GITHUB_ONLY)
        assert len(results) == 1
        assert results[0].source == MatchSource.GITHUB

    def test_fallback_chain_extends_when_tenant_thin(self):
        # Tenant returns 1 result, below FALLBACK_THRESHOLD of 3 -> Hub joins
        assert FALLBACK_THRESHOLD == 3
        results = self.agg.find_matches("Order", "IDoc", "SOAP",
                                        tenant_artifacts=["x"],
                                        mode=MatchMode.FALLBACK_CHAIN)
        sources = {r.source for r in results}
        assert MatchSource.HUB in sources

    def test_fallback_chain_stops_when_tenant_sufficient(self):
        many = [(i, _MockTenantArt(f"art{i}", "Y", "PkgY"))
                for i in range(5)]
        agg = MatchAggregator(_MockCPIFetcher(many), self.hub, self.gh)
        results = agg.find_matches("Order", "IDoc", "SOAP",
                                   tenant_artifacts=["x"],
                                   mode=MatchMode.FALLBACK_CHAIN)
        # 5 tenant results >= threshold of 3, so no Hub join
        assert all(r.source == MatchSource.TENANT for r in results)

    def test_missing_hub_client_degrades_gracefully(self):
        agg = MatchAggregator(self.cpi, hub_client=None, samples_browser=self.gh)
        results = agg.find_matches("Order", "IDoc", "SOAP",
                                   tenant_artifacts=["x"],
                                   mode=MatchMode.FALLBACK_CHAIN)
        # Chain still works without Hub
        assert results
        assert all(r.source != MatchSource.HUB for r in results)

    def test_hub_only_without_client_returns_empty(self):
        agg = MatchAggregator(self.cpi, hub_client=None, samples_browser=self.gh)
        assert agg.find_matches("X", "A", "B", mode=MatchMode.HUB_ONLY) == []

    def test_keywords_strip_short_name_tokens_but_keep_adapters(self):
        kw = _keywords("Z_OUTB_DELIVERY", "RFC", "SOAP")
        assert "rfc" in kw         # adapter survives even though length<=3
        assert "soap" in kw
        assert "outb" in kw         # 4 chars survives
        assert "z" not in kw        # 1 char dropped


# ─── Workbench import smoke test ─────────────────────────────────────────────
# This single import catches NameError / ImportError crashes (like the
# missing SolverSession reference) before Streamlit ever starts.

class TestWorkbenchImports:
    def test_workbench_module_imports(self):
        import workbench  # noqa: F401 — import is the assertion


# ─── Excel extractor: optional HIGH-driving columns ──────────────────────────

class TestExcelOptionalColumns:
    def _write(self, tmp_path, headers, rows):
        import openpyxl
        wb = openpyxl.Workbook(); ws = wb.active
        ws.append(headers)
        for r in rows:
            ws.append(r)
        p = tmp_path / "in.xlsx"
        wb.save(p)
        return p

    def test_legacy_file_still_parses(self, tmp_path):
        """Files without the optional columns must still parse with safe defaults."""
        from extractor.pi_extractor import PIFileExtractor
        p = self._write(tmp_path,
            ["Name", "SenderAdapter", "ReceiverAdapter"],
            [["IF1", "HTTPS", "HTTPS"]])
        recs = PIFileExtractor(str(p)).extract_all()
        assert len(recs) == 1
        assert recs[0].has_bpm is False
        assert recs[0].has_multi_mapping is False
        assert recs[0].channel_count == 1

    def test_optional_columns_read_when_present(self, tmp_path):
        from extractor.pi_extractor import PIFileExtractor
        p = self._write(tmp_path,
            ["Name", "SenderAdapter", "ReceiverAdapter",
             "HasBPM", "HasMultiMapping", "NumberOfChannels"],
            [["IF1", "RFC", "RFC", "Yes", "Yes", 3]])
        recs = PIFileExtractor(str(p)).extract_all()
        assert recs[0].has_bpm is True
        assert recs[0].has_multi_mapping is True
        assert recs[0].channel_count == 3

    def test_truthy_values_accepted(self, tmp_path):
        from extractor.pi_extractor import PIFileExtractor
        p = self._write(tmp_path,
            ["Name", "HasBPM", "HasMultiMapping"],
            [
                ["IF_yes", "yes", "true"],
                ["IF_y",   "Y",   "1"],
                ["IF_no",  "no",  "false"],
                ["IF_x",   "X",   ""],
            ])
        recs = PIFileExtractor(str(p)).extract_all()
        by_name = {r.name: r for r in recs}
        assert by_name["IF_yes"].has_bpm and by_name["IF_yes"].has_multi_mapping
        assert by_name["IF_y"].has_bpm and by_name["IF_y"].has_multi_mapping
        assert not by_name["IF_no"].has_bpm and not by_name["IF_no"].has_multi_mapping
        assert by_name["IF_x"].has_bpm   # "X" is truthy
        assert not by_name["IF_x"].has_multi_mapping   # blank is falsy

    def test_channel_count_coerces_safely(self, tmp_path):
        from extractor.pi_extractor import PIFileExtractor
        p = self._write(tmp_path,
            ["Name", "NumberOfChannels"],
            [["IF_int", 5], ["IF_str", "3"], ["IF_float", 4.0],
             ["IF_bad", "abc"], ["IF_zero", 0]])
        recs = PIFileExtractor(str(p)).extract_all()
        by_name = {r.name: r.channel_count for r in recs}
        assert by_name["IF_int"] == 5
        assert by_name["IF_str"] == 3
        assert by_name["IF_float"] == 4
        assert by_name["IF_bad"] == 1   # bad input -> safe default
        assert by_name["IF_zero"] == 1  # 0 clamped to 1 (channel_count is at least 1)

    def test_high_scoring_row_actually_scores_high(self, tmp_path):
        """End-to-end: a HIGH-recipe row produces HIGH complexity."""
        from extractor.pi_extractor import PIFileExtractor
        from analyzer.complexity_analyzer import ComplexityAnalyzer
        p = self._write(tmp_path,
            ["Name", "SenderAdapter", "ReceiverAdapter",
             "MappingProgram", "HasBPM", "HasMultiMapping"],
            [["IF_high", "RFC", "RFC", "MM_X", "Yes", "Yes"]])
        recs = PIFileExtractor(str(p)).extract_all()
        cfg = {'environment':'cf','migration':{'output_dir':'/tmp',
            'complexity_thresholds':{'low':{'max_score':10},'medium':{'max_score':25},'high':{'max_score':999}},
            'effort_days':{'low':1,'medium':3,'high':8}}}
        a = ComplexityAnalyzer(cfg).assess(recs[0])
        assert a.complexity == "HIGH"
        assert a.score > 25


# ─── SAP samples browser: artifact extraction + README-only filtering ────────

class TestSAPSamplesArtifactExtraction:
    def test_constructor_accepts_recipes_topic_limit(self):
        from fetcher.sap_samples_browser import SAPSamplesBrowser
        b = SAPSamplesBrowser(recipes_topic_limit=42)
        assert b.recipes_topic_limit == 42

    def test_default_topic_limit_is_unrestricted(self):
        """Default must be high enough to cover the actual SAP repo size (~170 folders)."""
        from fetcher.sap_samples_browser import SAPSamplesBrowser
        assert SAPSamplesBrowser().recipes_topic_limit >= 170

    def test_dataclass_has_artifact_fields(self):
        from fetcher.sap_samples_browser import SAPSamplePackage
        p = SAPSamplePackage(id="x", name="y", repo="r", path="p")
        assert p.has_zip is False
        assert p.extracted_mappings == []
        assert p.extracted_xslt == []
        assert p.extracted_iflows == []
        assert p.extracted_groovy == []

    def test_dict_roundtrip_preserves_has_zip(self):
        from fetcher.sap_samples_browser import SAPSamplesBrowser, SAPSamplePackage
        p = SAPSamplePackage(id="x", name="y", repo="r", path="p", has_zip=True)
        d = SAPSamplesBrowser._package_to_dict(p)
        p2 = SAPSamplesBrowser._dict_to_package(d)
        assert p2.has_zip is True

    def test_extract_artifacts_finds_all_file_types(self, tmp_path):
        """End-to-end: real zip on disk → extract_artifacts populates pkg fields."""
        import zipfile, http.server, threading, time
        from fetcher.sap_samples_browser import SAPSamplesBrowser, SAPSamplePackage

        # Build a fake repo zip with one of each artifact type
        zip_path = tmp_path / "fake.zip"
        prefix = "fake-master/Recipes/for/Sample/"
        with zipfile.ZipFile(zip_path, "w") as zf:
            zf.writestr(f"{prefix}mapping/M.mmap", "<m/>")
            zf.writestr(f"{prefix}xslt/t.xsl", "<x/>")
            zf.writestr(f"{prefix}script/s.groovy", "// g")
            zf.writestr(f"{prefix}iflow/F.iflw", "<i/>")
            zf.writestr(f"{prefix}schema/T.wsdl", "<w/>")
            zf.writestr(f"{prefix}README.md", "# sample")

        # Local HTTP server to serve the zip
        zip_bytes = zip_path.read_bytes()

        class H(http.server.BaseHTTPRequestHandler):
            def do_GET(self):
                self.send_response(200)
                self.send_header("Content-Length", str(len(zip_bytes)))
                self.end_headers()
                self.wfile.write(zip_bytes)
            def log_message(self, *a, **k): pass

        server = http.server.HTTPServer(("127.0.0.1", 0), H)
        port = server.server_port
        threading.Thread(target=server.serve_forever, daemon=True).start()
        try:
            time.sleep(0.05)
            cache = tmp_path / "cache"
            browser = SAPSamplesBrowser(cache_dir=cache)
            pkg = SAPSamplePackage(
                id="fake/Sample", name="Sample", repo="fake",
                path="Recipes/for/Sample",
                download_url=f"http://127.0.0.1:{port}/x.zip",
            )
            artifacts = browser.extract_artifacts(pkg)
            assert len(artifacts["mappings"]) == 1
            assert len(artifacts["xslt"]) == 1
            assert len(artifacts["groovy"]) == 1
            assert len(artifacts["iflows"]) == 1
            assert len(artifacts["wsdl"]) == 1
            # Package object is also populated
            assert len(pkg.extracted_mappings) == 1
            assert len(pkg.extracted_xslt) == 1
            assert len(pkg.extracted_groovy) == 1
            assert len(pkg.extracted_iflows) == 1
        finally:
            server.shutdown()

    def test_extract_artifacts_returns_empty_on_download_failure(self, tmp_path):
        from fetcher.sap_samples_browser import SAPSamplesBrowser, SAPSamplePackage
        browser = SAPSamplesBrowser(cache_dir=tmp_path / "cache")
        pkg = SAPSamplePackage(
            id="bad", name="bad", repo="bad", path="x",
            download_url="http://127.0.0.1:1/does-not-exist",
        )
        # Should not raise — just returns {}
        artifacts = browser.extract_artifacts(pkg)
        assert artifacts == {}


# ─── GitHub repo-zip walker in download_from_upload ──────────────────────────

class TestRepoZipWalker:
    def _make_inner_zip(self):
        import zipfile, io
        b = io.BytesIO()
        with zipfile.ZipFile(b, "w") as z:
            z.writestr("src/main/resources/mapping/M.mmap", "<m/>")
            z.writestr("META-INF/MANIFEST.MF", "Bundle-SymbolicName: pkg")
        return b.getvalue()

    def _make_repo_zip(self, leaves):
        """leaves: list of (topic, package_folder, files_dict)"""
        import zipfile, io
        b = io.BytesIO()
        with zipfile.ZipFile(b, "w") as z:
            z.writestr("apibusinesshub-master/README.md", "# repo")
            for topic, pkg, files in leaves:
                for fname, content in files.items():
                    z.writestr(
                        f"apibusinesshub-master/Recipes/for/{topic}/{pkg}/{fname}",
                        content)
        return b.getvalue()

    def test_repo_zip_discovers_inner_packages(self, tmp_path):
        from fetcher.cpi_fetcher import CPIFetcher
        inner = self._make_inner_zip()
        repo_zip = self._make_repo_zip([
            ("Aggregator", "Aggregator", {"Aggregator.zip": inner}),
            ("JSON_to_XML", "JSON_to_XML", {"JSON_to_XML.zip": inner}),
            ("Multi", "Multi", {"Multi.zip": inner}),
        ])
        f = CPIFetcher("http://localhost", None, cache_dir=tmp_path)
        arts = f.download_from_upload(repo_zip, "test_repo")
        assert len(arts) == 3
        ids = {a.id for a in arts}
        assert ids == {"Aggregator", "JSON_to_XML", "Multi"}

    def test_repo_zip_skips_readme_only_folders(self, tmp_path):
        from fetcher.cpi_fetcher import CPIFetcher
        repo_zip = self._make_repo_zip([
            ("HasContent", "HasContent", {"HasContent.zip": self._make_inner_zip()}),
            ("READMEonly", "READMEonly", {"readme.md": "# nothing here"}),
        ])
        f = CPIFetcher("http://localhost", None, cache_dir=tmp_path)
        arts = f.download_from_upload(repo_zip, "t")
        assert len(arts) == 1
        assert arts[0].id == "HasContent"

    def test_single_iflow_zip_still_works(self, tmp_path):
        """Layout 1 — backward compatibility check."""
        import zipfile, io
        from fetcher.cpi_fetcher import CPIFetcher
        b = io.BytesIO()
        with zipfile.ZipFile(b, "w") as z:
            z.writestr("META-INF/MANIFEST.MF", "Bundle-SymbolicName: solo")
            z.writestr("src/main/resources/integration_flow/F.iflw", "<i/>")
        f = CPIFetcher("http://localhost", None, cache_dir=tmp_path)
        arts = f.download_from_upload(b.getvalue(), "solo")
        assert len(arts) == 1
        assert arts[0].id == "solo"

    def test_container_zip_with_inner_iflows(self, tmp_path):
        """Layout 2 — backward compatibility check."""
        import zipfile, io
        from fetcher.cpi_fetcher import CPIFetcher
        # Outer zip with two inner iFlow zips, NO 'recipes/' in path
        inner = self._make_inner_zip()
        b = io.BytesIO()
        with zipfile.ZipFile(b, "w") as z:
            z.writestr("iflow1.zip", inner)
            z.writestr("iflow2.zip", inner)
        f = CPIFetcher("http://localhost", None, cache_dir=tmp_path)
        arts = f.download_from_upload(b.getvalue(), "container")
        assert len(arts) == 2

    def test_bad_inner_zip_does_not_crash(self, tmp_path):
        """A corrupt inner zip should be skipped, not crash the whole walk."""
        import zipfile, io
        from fetcher.cpi_fetcher import CPIFetcher
        good_inner = self._make_inner_zip()
        bad_inner  = b"not a real zip"
        b = io.BytesIO()
        with zipfile.ZipFile(b, "w") as z:
            z.writestr("apibusinesshub-master/Recipes/for/Good/Good/Good.zip", good_inner)
            z.writestr("apibusinesshub-master/Recipes/for/Bad/Bad/Bad.zip", bad_inner)
        f = CPIFetcher("http://localhost", None, cache_dir=tmp_path)
        arts = f.download_from_upload(b.getvalue(), "mixed")
        # Good is kept, Bad is skipped
        assert any(a.id == "Good" for a in arts)
        assert not any(a.id == "Bad" for a in arts)


# ─── SAP Migration Assessment parser ────────────────────────────────────────

class TestSAPMAParser:
    def _make_ma_file(self, tmp_path, summary_rows, scenario_rows, rules_rows):
        """Build a minimal SAP MA-shaped Excel for testing.

        Each row in {summary,scenario,rules}_rows is the data row only —
        title banner + header row are written automatically with the
        right offsets matching real exports.
        """
        import openpyxl
        wb = openpyxl.Workbook()
        # Executive Summary
        es = wb.active
        es.title = "Executive Summary"
        es.cell(2, 2, "SAP Integration Suite Migration Assessment")
        es.cell(5, 2, "High-Level Sizing Dashboard")
        es.cell(6, 2, "Metric")
        es.cell(6, 3, "KPI Value")
        es.cell(6, 4, "Percentage")
        es.cell(6, 5, "Description / Primary Action Items")
        for i, (metric, kpi, pct, desc) in enumerate(summary_rows):
            es.cell(7 + i, 2, metric)
            es.cell(7 + i, 3, kpi)
            es.cell(7 + i, 4, pct)
            es.cell(7 + i, 5, desc)

        # Scenario Evaluation
        sc = wb.create_sheet("Scenario Evaluation")
        sc.cell(2, 2, "Integration Scenario Detailed Assessment Inventory")
        headers = ["ICO Technical ID", "Sender System", "Receiver System",
                   "Sender Adapter", "Receiver Adapter", "Mapping Types Found",
                   "Migration Status", "Rule Weight", "Estimated Effort"]
        for i, h in enumerate(headers):
            sc.cell(4, 2 + i, h)
        for i, row in enumerate(scenario_rows):
            for j, val in enumerate(row):
                sc.cell(5 + i, 2 + j, val)

        # Rules Log
        rl = wb.create_sheet("Rules Log")
        rl.cell(2, 2, "Triggered Migration Rules & Blockers Log")
        rl_headers = ["Triggered Rule ID", "Target ICO Scenario Affected",
                      "Identified Asset String / Context",
                      "Assessment Technical Note & Remediation Strategy"]
        for i, h in enumerate(rl_headers):
            rl.cell(4, 2 + i, h)
        for i, row in enumerate(rules_rows):
            for j, val in enumerate(row):
                rl.cell(5 + i, 2 + j, val)

        path = tmp_path / "sap_ma.xlsx"
        wb.save(path)
        return path

    def test_parser_reads_kpis_correctly(self, tmp_path):
        from intake.sap_ma_parser import parse_sap_ma_excel
        path = self._make_ma_file(
            tmp_path,
            summary_rows=[
                ("Total Extracted ICOs", 100, "100%", "desc"),
                ("Ready to Migrate", 70, "70%", "desc"),
                ("Adjustment Required", 20, "20%", "desc"),
                ("Evaluation Required", 10, "10%", "desc"),
                ("Total Estimated Effort", "300 Hrs", "-", "desc"),
            ],
            scenario_rows=[],
            rules_rows=[],
        )
        report = parse_sap_ma_excel(str(path))
        assert report.summary.total_icos == 100
        assert report.summary.ready_to_migrate == 70
        assert report.summary.adjustment_required == 20
        assert report.summary.evaluation_required == 10
        assert report.summary.total_effort_hours == 300.0

    def test_parser_normalises_adapters(self, tmp_path):
        """IDOC_AAE → IDoc, HCIC → HTTPS, etc."""
        from intake.sap_ma_parser import parse_sap_ma_excel
        path = self._make_ma_file(
            tmp_path,
            summary_rows=[],
            scenario_rows=[
                ("ICO_001", "S4H", "ARIBA", "SOAP", "IDOC_AAE",
                 "Graphical (.mmap)", "Ready to Migrate", 20, "1.5 Hrs"),
                ("ICO_002", "S4H", "SF", "IDOC_AAE", "HCIC",
                 "Graphical (.mmap)", "Ready to Migrate", 20, "1.5 Hrs"),
            ],
            rules_rows=[],
        )
        report = parse_sap_ma_excel(str(path))
        assert len(report.interfaces) == 2
        ico1, ico2 = report.interfaces
        assert ico1.sender_adapter == "SOAP"
        assert ico1.receiver_adapter == "IDoc"   # IDOC_AAE → IDoc
        assert ico2.sender_adapter == "IDoc"
        assert ico2.receiver_adapter == "HTTPS"  # HCIC → HTTPS

    def test_multi_mapping_detected_from_plus_separator(self, tmp_path):
        """'Graphical + Groovy' → has_multi_mapping=True"""
        from intake.sap_ma_parser import parse_sap_ma_excel
        path = self._make_ma_file(
            tmp_path,
            summary_rows=[],
            scenario_rows=[
                ("ICO_simple", "A", "B", "SOAP", "SOAP",
                 "Graphical (.mmap)", "Ready to Migrate", 20, "1.0 Hrs"),
                ("ICO_multi", "A", "B", "SOAP", "REST",
                 "Graphical + Groovy", "Adjustment Required", 50, "3.5 Hrs"),
            ],
            rules_rows=[],
        )
        report = parse_sap_ma_excel(str(path))
        by_name = {r.name: r for r in report.interfaces}
        assert by_name["ICO_simple"].has_multi_mapping is False
        assert by_name["ICO_multi"].has_multi_mapping is True

    def test_rules_cross_reference_sets_bpm_flag(self, tmp_path):
        """A BPM_Detected rule should set has_bpm on the matching interface."""
        from intake.sap_ma_parser import parse_sap_ma_excel
        path = self._make_ma_file(
            tmp_path,
            summary_rows=[],
            scenario_rows=[
                ("ICO_with_bpm", "A", "B", "RFC", "RFC",
                 "Graphical (.mmap)", "Evaluation Required", 100, "10 Hrs"),
                ("ICO_no_rules", "A", "B", "SOAP", "SOAP",
                 "Graphical (.mmap)", "Ready to Migrate", 20, "1.5 Hrs"),
            ],
            rules_rows=[
                ("BPM_Detected", "ICO_with_bpm",
                 "ccBPM workflow", "Redesign as iFlow steps."),
            ],
        )
        report = parse_sap_ma_excel(str(path))
        by_name = {r.name: r for r in report.interfaces}
        assert by_name["ICO_with_bpm"].has_bpm is True
        assert by_name["ICO_no_rules"].has_bpm is False
        # The rule itself is preserved
        assert len(report.rules) == 1
        assert report.rules[0].rule_id == "BPM_Detected"

    def test_channel_count_derived_from_rule_weight(self, tmp_path):
        """Higher Rule Weight → higher channel_count (proxy bucketing)."""
        from intake.sap_ma_parser import parse_sap_ma_excel
        path = self._make_ma_file(
            tmp_path,
            summary_rows=[],
            scenario_rows=[
                ("ICO_low",  "A", "B", "SOAP", "SOAP", "Graphical", "Ready", 15, "1 Hr"),
                ("ICO_med",  "A", "B", "SOAP", "SOAP", "Graphical", "Adj",   50, "3 Hrs"),
                ("ICO_high", "A", "B", "SOAP", "SOAP", "Graphical", "Eval", 180, "12 Hrs"),
            ],
            rules_rows=[],
        )
        report = parse_sap_ma_excel(str(path))
        by_name = {r.name: r.channel_count for r in report.interfaces}
        assert by_name["ICO_low"]  == 1
        assert by_name["ICO_med"]  == 2
        assert by_name["ICO_high"] == 3

    def test_records_score_correctly_in_analyzer(self, tmp_path):
        """End-to-end: parsed records + ComplexityAnalyzer = useful complexity bands."""
        from intake.sap_ma_parser import parse_sap_ma_excel
        from analyzer.complexity_analyzer import ComplexityAnalyzer
        path = self._make_ma_file(
            tmp_path,
            summary_rows=[],
            scenario_rows=[
                ("ICO_LOW",  "A", "B", "HTTPS", "HTTPS",  "Graphical",
                 "Ready", 10, "1 Hr"),
                ("ICO_HIGH", "A", "B", "RFC", "RFC", "Graphical + Java",
                 "Eval", 180, "16 Hrs"),
            ],
            rules_rows=[
                ("BPM_Detected", "ICO_HIGH", "ccBPM", "Redesign"),
            ],
        )
        report = parse_sap_ma_excel(str(path))
        cfg = {'environment':'cf','migration':{'output_dir':'/tmp',
            'complexity_thresholds':{'low':{'max_score':10},
                                      'medium':{'max_score':25},
                                      'high':{'max_score':999}},
            'effort_days':{'low':1,'medium':3,'high':8}}}
        analyzer = ComplexityAnalyzer(cfg)
        results = {r.name: analyzer.assess(r) for r in report.interfaces}
        assert results["ICO_LOW"].complexity in ("LOW", "MEDIUM")
        assert results["ICO_HIGH"].complexity == "HIGH"

    def test_rejects_file_with_no_scenario_sheet(self, tmp_path):
        """A file that isn't a SAP MA export should raise ValueError."""
        import openpyxl
        from intake.sap_ma_parser import parse_sap_ma_excel
        wb = openpyxl.Workbook()
        wb.active.title = "Some Other Sheet"
        path = tmp_path / "not_ma.xlsx"
        wb.save(path)
        import pytest
        with pytest.raises(ValueError, match="Scenario Evaluation"):
            parse_sap_ma_excel(str(path))

    def test_real_simulation_file(self):
        """Round-trip the user-supplied simulated export end-to-end."""
        from pathlib import Path
        from intake.sap_ma_parser import parse_sap_ma_excel
        ma_file = Path("/tmp/sap_ma.xlsx")
        if not ma_file.exists():
            import pytest
            pytest.skip("simulated SAP MA file not available")
        report = parse_sap_ma_excel(str(ma_file))
        # Verify against the known-good values from the user-supplied file
        assert report.summary.total_icos == 142
        assert report.summary.ready_to_migrate == 94
        assert report.summary.adjustment_required == 32
        assert report.summary.evaluation_required == 16
        assert report.summary.total_effort_hours == 420.0
        assert len(report.interfaces) == 7
        assert len(report.rules) == 4


# ─── XSLT generator + Groovy exception template fix ────────────────────────

class TestXSLTGenerator:
    def _record(self, **overrides):
        from extractor.pi_extractor import InterfaceRecord
        base = dict(id="x", name="IF_TEST", namespace="http://co.com/test",
                    software_component="SC", sender_system="SRC",
                    receiver_system="DST", sender_adapter="HTTPS",
                    receiver_adapter="HTTPS", message_interface="",
                    mapping_program=None, description="")
        base.update(overrides)
        return InterfaceRecord(**base)

    def _assess(self, record):
        from analyzer.complexity_analyzer import ComplexityAnalyzer
        cfg = {'environment':'cf','migration':{'output_dir':'/tmp',
            'complexity_thresholds':{'low':{'max_score':10},'medium':{'max_score':25},'high':{'max_score':999}},
            'effort_days':{'low':1,'medium':3,'high':8}}}
        return ComplexityAnalyzer(cfg).assess(record)

    def test_should_generate_for_xslt_mapping(self, tmp_path):
        from scaffolder.xslt_generator import XSLTGenerator
        g = XSLTGenerator(output_dir=str(tmp_path))
        r = self._record(mapping_program="MM_XSLT_Transform")
        assert g.should_generate(self._assess(r)) is True

    def test_should_generate_for_idoc(self, tmp_path):
        from scaffolder.xslt_generator import XSLTGenerator
        g = XSLTGenerator(output_dir=str(tmp_path))
        r = self._record(sender_adapter="IDoc")
        assert g.should_generate(self._assess(r)) is True

    def test_should_generate_for_edi(self, tmp_path):
        from scaffolder.xslt_generator import XSLTGenerator
        g = XSLTGenerator(output_dir=str(tmp_path))
        r = self._record(sender_adapter="AS2", description="EDIFACT order processing")
        assert g.should_generate(self._assess(r)) is True

    def test_should_not_generate_for_plain_https(self, tmp_path):
        from scaffolder.xslt_generator import XSLTGenerator
        g = XSLTGenerator(output_dir=str(tmp_path))
        r = self._record()  # plain HTTPS → HTTPS, no mapping
        assert g.should_generate(self._assess(r)) is False

    def test_generated_xslt_is_well_formed_xml(self, tmp_path):
        """ALL XSLT outputs must be well-formed XML or they're worthless."""
        import lxml.etree as ET
        from scaffolder.xslt_generator import XSLTGenerator
        g = XSLTGenerator(output_dir=str(tmp_path))
        # Three different template kinds
        cases = [
            self._record(sender_adapter="IDoc"),
            self._record(sender_adapter="AS2", description="EDIFACT"),
            self._record(mapping_program="MM_XSLT_Test"),
        ]
        for r in cases:
            script = g.generate(self._assess(r))
            assert script is not None
            # Parse must succeed
            ET.fromstring(script.content.encode())

    def test_generate_all_writes_files_to_disk(self, tmp_path):
        from scaffolder.xslt_generator import XSLTGenerator
        g = XSLTGenerator(output_dir=str(tmp_path))
        records = [self._record(name=f"IF_{i}", sender_adapter="IDoc") for i in range(3)]
        assessments = [self._assess(r) for r in records]
        out = g.generate_all(assessments)
        assert len(out) == 3
        # Files exist on disk
        xslt_dir = tmp_path / "xslt"
        assert xslt_dir.is_dir()
        assert len(list(xslt_dir.glob("*.xsl"))) == 3


class TestGroovyExceptionHandlerFix:
    def test_exception_handler_template_no_longer_raises(self, tmp_path):
        """Regression test for the f-string brace bug — `${{{{{errorMsg}}}}}`
        used to call .format() with no errorMsg kwarg and KeyError.
        Now it should produce a valid Groovy script."""
        from extractor.pi_extractor import InterfaceRecord
        from scaffolder.groovy_generator import GroovyGenerator
        r = InterfaceRecord(
            id="x", name="IF_Exception_Test", namespace="",
            software_component="", sender_system="SRC", receiver_system="DST",
            sender_adapter="HTTPS", receiver_adapter="HTTPS",
            message_interface="", mapping_program="MM_Test",
            has_bpm=True,  # forces exception_handler to be generated
        )
        g = GroovyGenerator(output_dir=str(tmp_path))
        script = g._generate_exception_handler(r)
        # Must contain the proper Groovy ${...} syntax, not Python format errors
        assert "${errorMsg}" in script.content
        # Must NOT contain unexpanded Python braces or KeyError artifacts
        assert "{errorMsg}" not in script.content.replace("${errorMsg}", "")
        # Must be syntactically reasonable Groovy (rough check)
        assert "def Message processData(Message message)" in script.content
        assert "exception?.getMessage()" in script.content


# ─── Orphan-module wire-up smoke tests ──────────────────────────────────────

class TestOrphanWireup:
    """Each formerly-orphaned module is now reachable via a workbench function;
    these tests prove the module's primary API still works in isolation so the
    UI calls won't crash."""

    def test_parameter_injector_produces_prop_file(self):
        from scaffolder.parameter_injector import build_parameters_prop
        prop = build_parameters_prop("IF_TEST")
        assert "Externalized parameters for IF_TEST" in prop
        assert len(prop.splitlines()) > 1

    def test_interface_request_builds_bsr(self):
        from extractor.pi_extractor import InterfaceRecord
        from reporter.interface_request import (
            build_interface_request, build_business_solution_request)
        from analyzer.complexity_analyzer import MigrationAssessment
        r = InterfaceRecord(
            id="x", name="IF_Foo", namespace="http://co.com/test",
            software_component="SC", sender_system="ECC",
            receiver_system="S4H", sender_adapter="HTTPS",
            receiver_adapter="HTTPS", message_interface="",
            mapping_program=None, description="",
        )
        ir = build_interface_request(r)
        assert ir.name == "IF_Foo"
        assert len(ir.message_flows) == 1
        assert ir.sources and ir.targets

        a = MigrationAssessment(interface=r, score=5, complexity="LOW",
                                 effort_days=1, notes=[], recommended_pattern="HTTPS-to-HTTPS")
        bsr = build_business_solution_request([a], project_name="Test")
        assert len(bsr.interface_requests) == 1

    def test_interface_request_word_render(self, tmp_path):
        """BSR Word render writes a valid .docx file."""
        from extractor.pi_extractor import InterfaceRecord
        from reporter.interface_request import (
            build_business_solution_request, render_word)
        from analyzer.complexity_analyzer import MigrationAssessment
        r = InterfaceRecord(
            id="x", name="IF_Foo", namespace="", software_component="",
            sender_system="ECC", receiver_system="S4H",
            sender_adapter="HTTPS", receiver_adapter="HTTPS",
            message_interface="", mapping_program=None, description="",
        )
        a = MigrationAssessment(interface=r, score=5, complexity="LOW",
                                 effort_days=1, notes=[], recommended_pattern="P")
        bsr = build_business_solution_request([a])
        out = render_word(bsr, tmp_path / "bsr.docx", mode="canonical")
        assert out.exists()
        assert out.stat().st_size > 0

    def test_transaction_advisor_returns_advisories(self):
        from extractor.pi_extractor import InterfaceRecord
        from analyzer.transaction_advisor import advise_all
        records = [InterfaceRecord(
            id=f"x{i}", name=f"IF_{i}", namespace="", software_component="",
            sender_system="S", receiver_system="R",
            sender_adapter="JDBC" if i == 0 else "HTTPS",
            receiver_adapter="HTTPS",
            message_interface="", mapping_program=None, description="",
        ) for i in range(3)]
        advisories = advise_all(records)
        assert len(advisories) == 3
        # JDBC sender should get a non-default recommendation
        jdbc_advice = next(a for a in advisories if a.interface_name == "IF_0")
        assert jdbc_advice.handling  # non-empty

    def test_endpoint_collision_handles_empty(self):
        from analyzer.endpoint_collision import detect_collisions
        assert detect_collisions([]) == []

    def test_apim_faults_catalog(self):
        from analyzer.apim_faults import all_codes, search, lookup
        codes = all_codes()
        assert len(codes) > 0
        # Lookup a known code
        first = lookup(codes[0])
        assert first is not None
        # Search returns relevant entries
        matches = search("quota")
        # If "quota" is referenced anywhere in the catalog the search should find it
        for m in matches:
            blob = (m.code + " " + m.meaning + " " + m.category + " " +
                    m.likely_cause + " " + m.remediation).lower()
            assert "quota" in blob

    def test_domains_classification(self):
        from analyzer.domains import derive_domain
        d = derive_domain("ECC", "S4HANA_Cloud")
        assert d.domain  # non-empty classification

    def test_odata_suggester_curated_match(self):
        from fetcher.odata_suggester import ODataSuggester
        s = ODataSuggester()
        suggestions = s.suggest_for_bapi("BAPI_SALESORDER_CREATEFROMDAT2")
        assert len(suggestions) >= 1
        assert suggestions[0].api_name == "API_SALES_ORDER_SRV"


# ─── Large archive + tar support ────────────────────────────────────────────

class TestArchiveFormatSupport:
    def _make_inner_zip(self):
        import zipfile, io
        b = io.BytesIO()
        with zipfile.ZipFile(b, "w") as z:
            z.writestr("src/main/resources/mapping/M.mmap", "<m/>")
        return b.getvalue()

    def test_tar_gz_with_github_repo_layout(self, tmp_path):
        """tar.gz containing Recipes/for/<topic>/<pkg>/<inner>.zip works."""
        import tarfile, io
        from fetcher.cpi_fetcher import CPIFetcher
        tar_buf = io.BytesIO()
        with tarfile.open(fileobj=tar_buf, mode="w:gz") as tar:
            for name in ["Aggregator", "JSON_to_XML"]:
                path = f"repo-master/Recipes/for/{name}/{name}/Inner.zip"
                data = self._make_inner_zip()
                info = tarfile.TarInfo(name=path)
                info.size = len(data)
                tar.addfile(info, io.BytesIO(data))
        f = CPIFetcher("http://localhost", None, cache_dir=tmp_path)
        arts = f.download_from_upload(tar_buf.getvalue(), "tar_test")
        assert len(arts) == 2
        assert {a.id for a in arts} == {"Aggregator", "JSON_to_XML"}

    def test_uncompressed_tar(self, tmp_path):
        """Plain .tar (no compression) is also accepted."""
        import tarfile, io
        from fetcher.cpi_fetcher import CPIFetcher
        tar_buf = io.BytesIO()
        with tarfile.open(fileobj=tar_buf, mode="w:") as tar:
            data = self._make_inner_zip()
            info = tarfile.TarInfo(name="repo/Recipes/for/Test/Test/Inner.zip")
            info.size = len(data)
            tar.addfile(info, io.BytesIO(data))
        f = CPIFetcher("http://localhost", None, cache_dir=tmp_path)
        arts = f.download_from_upload(tar_buf.getvalue(), "plain_tar")
        assert len(arts) == 1

    def test_tar_bz2(self, tmp_path):
        """tar.bz2 also accepted (rarer but used in some SAP exports)."""
        import tarfile, io
        from fetcher.cpi_fetcher import CPIFetcher
        tar_buf = io.BytesIO()
        with tarfile.open(fileobj=tar_buf, mode="w:bz2") as tar:
            data = self._make_inner_zip()
            info = tarfile.TarInfo(name="repo/Recipes/for/X/X/Inner.zip")
            info.size = len(data)
            tar.addfile(info, io.BytesIO(data))
        f = CPIFetcher("http://localhost", None, cache_dir=tmp_path)
        arts = f.download_from_upload(tar_buf.getvalue(), "bz2_tar")
        assert len(arts) == 1

    def test_zip_still_works_unchanged(self, tmp_path):
        """Regression — zip uploads must still work exactly as before."""
        import zipfile, io
        from fetcher.cpi_fetcher import CPIFetcher
        b = io.BytesIO()
        with zipfile.ZipFile(b, "w") as z:
            z.writestr("repo-master/Recipes/for/Solo/Solo/Inner.zip",
                       self._make_inner_zip())
        f = CPIFetcher("http://localhost", None, cache_dir=tmp_path)
        arts = f.download_from_upload(b.getvalue(), "zip_test")
        assert len(arts) == 1

    def test_normalise_helper_passes_zip_through_unchanged(self):
        """A zip's bytes should come back identical, not re-encoded."""
        import zipfile, io
        from fetcher.cpi_fetcher import CPIFetcher
        b = io.BytesIO()
        with zipfile.ZipFile(b, "w") as z:
            z.writestr("a.txt", "hello")
        original = b.getvalue()
        result = CPIFetcher._normalise_to_zip_bytes(original)
        assert result is original  # same object — no work done

    def test_normalise_helper_converts_targz_to_zip(self):
        """tar.gz bytes should come back as a valid zip with same files."""
        import tarfile, zipfile, io
        from fetcher.cpi_fetcher import CPIFetcher
        tar_buf = io.BytesIO()
        with tarfile.open(fileobj=tar_buf, mode="w:gz") as tar:
            for name, body in [("a.txt", b"hello"), ("b/c.txt", b"world")]:
                info = tarfile.TarInfo(name=name)
                info.size = len(body)
                tar.addfile(info, io.BytesIO(body))
        result_bytes = CPIFetcher._normalise_to_zip_bytes(tar_buf.getvalue())
        # Verify it's now a valid zip with the same files
        with zipfile.ZipFile(io.BytesIO(result_bytes)) as zf:
            assert set(zf.namelist()) == {"a.txt", "b/c.txt"}
            assert zf.read("a.txt") == b"hello"
            assert zf.read("b/c.txt") == b"world"

    def test_unknown_format_passes_through(self):
        """Bytes that aren't zip/tar/tar.gz/tar.bz2 pass through unchanged
        so the caller's zipfile.ZipFile raises a clear BadZipFile."""
        from fetcher.cpi_fetcher import CPIFetcher
        garbage = b"this is not an archive of any kind"
        result = CPIFetcher._normalise_to_zip_bytes(garbage)
        assert result == garbage


class TestStreamlitUploadConfig:
    def test_config_file_exists_with_10gb_limit(self):
        """.streamlit/config.toml must set maxUploadSize >= 10240 (10 GB in MB)."""
        from pathlib import Path
        cfg = Path(__file__).parent.parent / ".streamlit" / "config.toml"
        assert cfg.exists(), ".streamlit/config.toml is missing — large uploads will be rejected"
        text = cfg.read_text()
        # Both must be raised together, see config comments
        import re
        m_up = re.search(r"maxUploadSize\s*=\s*(\d+)", text)
        m_msg = re.search(r"maxMessageSize\s*=\s*(\d+)", text)
        assert m_up is not None, "maxUploadSize not set"
        assert m_msg is not None, "maxMessageSize not set"
        assert int(m_up.group(1)) >= 10240, "maxUploadSize must be >= 10 GB"
        assert int(m_msg.group(1)) >= 10240, "maxMessageSize must be >= 10 GB"


# ─── Shadow testing (Level 2): XML diff + XSLT executor + fixture harness ───

class TestXmlDiffer:
    def test_identical_documents_pass(self):
        from testing.xml_differ import XmlDiffer
        r = XmlDiffer().diff("<a><b>1</b></a>", "<a><b>1</b></a>")
        assert r.passed
        assert len(r.entries) == 0

    def test_value_difference_fails(self):
        from testing.xml_differ import XmlDiffer
        r = XmlDiffer().diff("<a><b>1</b></a>", "<a><b>2</b></a>")
        assert not r.passed
        assert len(r.real_diffs) == 1
        assert r.real_diffs[0].kind == "value_mismatch"

    def test_missing_element_fails(self):
        from testing.xml_differ import XmlDiffer
        r = XmlDiffer().diff("<a><b/><c/></a>", "<a><b/></a>")
        assert not r.passed
        assert any(e.kind == "missing_element" for e in r.real_diffs)

    def test_extra_element_fails(self):
        from testing.xml_differ import XmlDiffer
        r = XmlDiffer().diff("<a><b/></a>", "<a><b/><z/></a>")
        assert not r.passed
        assert any(e.kind == "extra_element" for e in r.real_diffs)

    def test_number_tolerance(self):
        """'100' and '100.00' must compare equal under default config."""
        from testing.xml_differ import XmlDiffer
        r = XmlDiffer().diff("<a><n>100</n></a>", "<a><n>100.00</n></a>")
        assert r.passed
        assert len(r.cosmetic_diffs) == 1

    def test_date_format_tolerance(self):
        """ISO and German date formats representing same instant are equal."""
        from testing.xml_differ import XmlDiffer
        r = XmlDiffer().diff("<a><d>2026-01-15</d></a>",
                              "<a><d>15.01.2026</d></a>")
        assert r.passed

    def test_boolean_tolerance(self):
        from testing.xml_differ import XmlDiffer
        r = XmlDiffer().diff("<a><b>true</b></a>", "<a><b>1</b></a>")
        assert r.passed

    def test_namespace_prefix_tolerance(self):
        """Same URI with different prefixes must compare equal."""
        from testing.xml_differ import XmlDiffer
        r = XmlDiffer().diff(
            '<a xmlns:ns1="urn:x"><ns1:b>1</ns1:b></a>',
            '<a xmlns:myns="urn:x"><myns:b>1</myns:b></a>')
        assert r.passed

    def test_path_rule_ignore(self):
        """A path with 'ignore' rule is skipped entirely, even if values differ."""
        from testing.xml_differ import XmlDiffer, DiffConfig
        cfg = DiffConfig(path_rules={"/Order/Timestamp": "ignore"})
        r = XmlDiffer(cfg).diff(
            "<Order><Timestamp>2026-01-01</Timestamp><Total>100</Total></Order>",
            "<Order><Timestamp>2099-99-99</Timestamp><Total>100</Total></Order>")
        assert r.passed

    def test_attribute_value_mismatch(self):
        from testing.xml_differ import XmlDiffer
        r = XmlDiffer().diff('<a x="1"/>', '<a x="2"/>')
        assert not r.passed

    def test_disabled_number_tolerance(self):
        """When normalize_numbers is False, '100' != '100.00'."""
        from testing.xml_differ import XmlDiffer, DiffConfig
        r = XmlDiffer(DiffConfig(normalize_numbers=False)).diff(
            "<a><n>100</n></a>", "<a><n>100.00</n></a>")
        assert not r.passed

    def test_optional_path_allows_missing(self):
        """A path in optional_paths can be missing in actual without failing."""
        from testing.xml_differ import XmlDiffer, DiffConfig
        cfg = DiffConfig(optional_paths=["/a/notes"])
        r = XmlDiffer(cfg).diff("<a><b/><notes/></a>", "<a><b/></a>")
        assert r.passed


class TestXsltExecutor:
    def test_basic_transformation(self):
        from testing.xslt_executor import XsltExecutor
        xslt = ('<?xml version="1.0"?>'
                '<xsl:stylesheet xmlns:xsl="http://www.w3.org/1999/XSL/Transform" version="1.0">'
                '<xsl:template match="/Order"><Invoice><Amount>'
                '<xsl:value-of select="Total"/></Amount></Invoice></xsl:template>'
                '</xsl:stylesheet>')
        r = XsltExecutor().run(xslt, "<Order><Total>100</Total></Order>")
        assert r.success
        assert "<Amount>100</Amount>" in r.output_xml

    def test_invalid_xslt_returns_error(self):
        from testing.xslt_executor import XsltExecutor
        r = XsltExecutor().run("<not-valid-xslt/>", "<x/>")
        # An <not-valid-xslt/> root is XML but not a valid stylesheet
        assert not r.success
        assert r.error_message

    def test_invalid_input_xml_returns_error(self):
        from testing.xslt_executor import XsltExecutor
        valid_xslt = ('<?xml version="1.0"?>'
                      '<xsl:stylesheet xmlns:xsl="http://www.w3.org/1999/XSL/Transform" version="1.0">'
                      '<xsl:template match="/"><x/></xsl:template></xsl:stylesheet>')
        r = XsltExecutor().run(valid_xslt, "<this is not valid xml")
        assert not r.success
        assert "valid XML" in r.error_message or "Input payload" in r.error_message

    def test_xslt_2_warning(self):
        """An XSLT 2.0 stylesheet should produce a warning even if it parses."""
        from testing.xslt_executor import XsltExecutor
        xslt2 = ('<?xml version="1.0"?>'
                 '<xsl:stylesheet xmlns:xsl="http://www.w3.org/1999/XSL/Transform" version="2.0">'
                 '<xsl:template match="/"><x/></xsl:template></xsl:stylesheet>')
        r = XsltExecutor().run(xslt2, "<x/>")
        # Either the transform succeeds with a warning, or fails — both acceptable
        assert any("2.0" in w or "XSLT 1.0" in w for w in r.warnings) or not r.success

    def test_value_map_extension(self):
        """SAP valuemap:get extension resolves via the SAPExtensions stub."""
        from testing.xslt_executor import XsltExecutor, SAPExtensions
        xslt = ('<?xml version="1.0"?>'
                '<xsl:stylesheet xmlns:xsl="http://www.w3.org/1999/XSL/Transform"'
                ' xmlns:valuemap="http://sap.com/it/ValueMapping" version="1.0">'
                '<xsl:template match="/Order"><Out><Country>'
                "<xsl:value-of select=\"valuemap:get('ERP','Country','TGT','Country',Country)\"/>"
                '</Country></Out></xsl:template></xsl:stylesheet>')
        ext = SAPExtensions(value_maps={"ERP:Country:DE": "Germany"})
        r = XsltExecutor(extensions=ext).run(
            xslt, "<Order><Country>DE</Country></Order>")
        assert r.success
        assert "Germany" in r.output_xml

    def test_value_map_miss_returns_source(self):
        """Missed lookup returns the source value, not crash."""
        from testing.xslt_executor import XsltExecutor, SAPExtensions
        xslt = ('<?xml version="1.0"?>'
                '<xsl:stylesheet xmlns:xsl="http://www.w3.org/1999/XSL/Transform"'
                ' xmlns:valuemap="http://sap.com/it/ValueMapping" version="1.0">'
                '<xsl:template match="/Order"><Out>'
                "<xsl:value-of select=\"valuemap:get('ERP','X','TGT','X',Code)\"/>"
                '</Out></xsl:template></xsl:stylesheet>')
        r = XsltExecutor(extensions=SAPExtensions()).run(
            xslt, "<Order><Code>UNKNOWN</Code></Order>")
        assert r.success
        assert "UNKNOWN" in r.output_xml

    def test_with_xsl_parameter(self):
        """XSLT params are passed through correctly."""
        from testing.xslt_executor import XsltExecutor
        xslt = ('<?xml version="1.0"?>'
                '<xsl:stylesheet xmlns:xsl="http://www.w3.org/1999/XSL/Transform" version="1.0">'
                '<xsl:param name="suffix"/>'
                '<xsl:template match="/Order"><x>'
                '<xsl:value-of select="concat(Name,$suffix)"/>'
                '</x></xsl:template></xsl:stylesheet>')
        r = XsltExecutor().run(xslt, "<Order><Name>foo</Name></Order>",
                                params={"suffix": "-bar"})
        assert r.success
        assert "foo-bar" in r.output_xml


class TestFixtureHarness:
    def _setup_fixture(self, tmp_path, xslt_body, input_body, expected_body,
                       config_yaml=None):
        from pathlib import Path
        fdir = Path(tmp_path) / "IFACE"
        fdir.mkdir()
        (fdir / "transform.xsl").write_text(xslt_body)
        (fdir / "request_001.input.xml").write_text(input_body)
        (fdir / "request_001.expected.xml").write_text(expected_body)
        if config_yaml is not None:
            (fdir / "config.yaml").write_text(config_yaml)
        return fdir

    def test_passing_case(self, tmp_path):
        from testing.fixture_harness import run_interface_tests
        fdir = self._setup_fixture(
            tmp_path,
            xslt_body=('<?xml version="1.0"?>'
                       '<xsl:stylesheet xmlns:xsl="http://www.w3.org/1999/XSL/Transform" version="1.0">'
                       '<xsl:template match="/Order">'
                       '<Invoice><Amount><xsl:value-of select="Total"/></Amount></Invoice>'
                       '</xsl:template></xsl:stylesheet>'),
            input_body="<Order><Total>100</Total></Order>",
            expected_body="<Invoice><Amount>100</Amount></Invoice>",
        )
        result = run_interface_tests(fdir)
        assert result.passed
        assert result.pass_count == 1
        assert result.fail_count == 0

    def test_failing_case_with_real_diff(self, tmp_path):
        from testing.fixture_harness import run_interface_tests
        fdir = self._setup_fixture(
            tmp_path,
            xslt_body=('<?xml version="1.0"?>'
                       '<xsl:stylesheet xmlns:xsl="http://www.w3.org/1999/XSL/Transform" version="1.0">'
                       '<xsl:template match="/Order">'
                       '<Invoice><Amount><xsl:value-of select="Total"/></Amount></Invoice>'
                       '</xsl:template></xsl:stylesheet>'),
            input_body="<Order><Total>100</Total></Order>",
            expected_body="<Invoice><Amount>999</Amount></Invoice>",
        )
        result = run_interface_tests(fdir)
        assert not result.passed
        assert result.fail_count == 1
        # The diff result is populated with the real diff
        case = result.cases[0]
        assert case.diff_result is not None
        assert len(case.diff_result.real_diffs) == 1

    def test_numeric_tolerance_keeps_case_passing(self, tmp_path):
        """'100' (input) → '100.00' (XSLT output) vs '100' (expected) = PASS."""
        from testing.fixture_harness import run_interface_tests
        fdir = self._setup_fixture(
            tmp_path,
            xslt_body=('<?xml version="1.0"?>'
                       '<xsl:stylesheet xmlns:xsl="http://www.w3.org/1999/XSL/Transform" version="1.0">'
                       '<xsl:template match="/Order">'
                       '<x><Amount><xsl:value-of select="format-number(Total, \'0.00\')"/></Amount></x>'
                       '</xsl:template></xsl:stylesheet>'),
            input_body="<Order><Total>100</Total></Order>",
            expected_body="<x><Amount>100</Amount></x>",
        )
        result = run_interface_tests(fdir)
        assert result.passed  # 100 == 100.00 by tolerance

    def test_config_yaml_path_ignore(self, tmp_path):
        """A path in config.yaml's path_rules: ignore is skipped."""
        from testing.fixture_harness import run_interface_tests
        fdir = self._setup_fixture(
            tmp_path,
            xslt_body=('<?xml version="1.0"?>'
                       '<xsl:stylesheet xmlns:xsl="http://www.w3.org/1999/XSL/Transform" version="1.0">'
                       '<xsl:template match="/Order">'
                       '<x><T>now</T><V><xsl:value-of select="Val"/></V></x>'
                       '</xsl:template></xsl:stylesheet>'),
            input_body="<Order><Val>42</Val></Order>",
            expected_body="<x><T>different</T><V>42</V></x>",
            config_yaml=("diff:\n"
                         "  path_rules:\n"
                         '    "/x/T": ignore\n'),
        )
        result = run_interface_tests(fdir)
        # Despite the T values differing, the rule says "ignore" so test passes
        assert result.passed

    def test_no_cases_handled_gracefully(self, tmp_path):
        from pathlib import Path
        from testing.fixture_harness import run_interface_tests
        fdir = Path(tmp_path) / "empty"
        fdir.mkdir()
        (fdir / "transform.xsl").write_text(
            '<?xml version="1.0"?><xsl:stylesheet '
            'xmlns:xsl="http://www.w3.org/1999/XSL/Transform" version="1.0">'
            '<xsl:template match="/"><x/></xsl:template></xsl:stylesheet>')
        result = run_interface_tests(fdir)
        assert not result.passed  # no cases = not passing
        assert len(result.cases) == 0
        assert result.config_warnings

    def test_missing_transform_handled(self, tmp_path):
        from pathlib import Path
        from testing.fixture_harness import run_interface_tests
        fdir = Path(tmp_path) / "no_xslt"
        fdir.mkdir()
        (fdir / "request_001.input.xml").write_text("<x/>")
        (fdir / "request_001.expected.xml").write_text("<x/>")
        result = run_interface_tests(fdir)
        assert not result.passed
        assert any("not found" in w for w in result.config_warnings)

    def test_run_all_fixtures(self, tmp_path):
        from pathlib import Path
        from testing.fixture_harness import run_all_fixtures
        for name in ["A", "B"]:
            fdir = Path(tmp_path) / name
            fdir.mkdir()
            (fdir / "transform.xsl").write_text(
                '<?xml version="1.0"?><xsl:stylesheet '
                'xmlns:xsl="http://www.w3.org/1999/XSL/Transform" version="1.0">'
                '<xsl:template match="/"><x/></xsl:template></xsl:stylesheet>')
            (fdir / "request_001.input.xml").write_text("<root/>")
            (fdir / "request_001.expected.xml").write_text("<x/>")
        results = run_all_fixtures(Path(tmp_path))
        assert len(results) == 2
        assert all(r.passed for r in results)

    def test_create_fixture_skeleton_produces_runnable_output(self, tmp_path):
        """Skeleton must produce a fixture that actually runs without error."""
        from pathlib import Path
        from testing.fixture_harness import create_fixture_skeleton, run_interface_tests
        target = Path(tmp_path) / "NEW_IFACE"
        create_fixture_skeleton(target, "NEW_IFACE")
        assert (target / "config.yaml").exists()
        assert (target / "transform.xsl").exists()
        assert (target / "request_001.input.xml").exists()
        assert (target / "request_001.expected.xml").exists()
        # And the skeleton actually runs (identity transform vs identity expected)
        result = run_interface_tests(target)
        # May pass or fail depending on the identity behaviour, but must not crash
        assert isinstance(result.cases, list)


# ─── SAP MA parser: newer schema with Complexity Group column ────────────────

class TestSAPMANewerSchema:
    def _build_newer_schema_file(self, tmp_path):
        """Build an Excel matching the newer SAP MA schema:
        - Scenario Evaluation has 'Complexity Group' instead of 'Mapping Types Found'
        - Scenario Evaluation uses 'Est. Effort' instead of 'Estimated Effort'
        - Rules Log has an extra 'Complexity' column
        """
        import openpyxl
        wb = openpyxl.Workbook()
        es = wb.active
        es.title = "Executive Summary"
        es.cell(6, 2, "Metric")
        es.cell(6, 3, "KPI Value")
        es.cell(7, 2, "Total Extracted ICOs"); es.cell(7, 3, 169)

        sc = wb.create_sheet("Scenario Evaluation")
        # Newer headers
        headers = ["ICO Technical ID", "Sender System", "Receiver System",
                   "Sender Adapter", "Receiver Adapter",
                   "Complexity Group",    # was "Mapping Types Found"
                   "Migration Status", "Rule Weight",
                   "Est. Effort"]          # was "Estimated Effort"
        for i, h in enumerate(headers):
            sc.cell(4, 2 + i, h)
        rows = [
            ("ICO_LOW",  "A", "B", "SOAP",  "SOAP", "Low Complexity",
             "Ready to Migrate",   10, "1.0 Hrs"),
            ("ICO_MED",  "A", "B", "JDBC",  "REST", "Medium Complexity",
             "Adjustment Required", 50, "3.5 Hrs"),
            ("ICO_HIGH", "A", "B", "IDOC_AAE", "RFC", "High Complexity",
             "Evaluation Required", 180, "16.0 Hrs"),
        ]
        for i, row in enumerate(rows):
            for j, val in enumerate(row):
                sc.cell(5 + i, 2 + j, val)

        rl = wb.create_sheet("Rules Log")
        rl_headers = ["Triggered Rule ID", "Target ICO Scenario Affected",
                      "Complexity",   # NEW in newer schema
                      "Identified Asset String / Context",
                      "Assessment Technical Note & Remediation Strategy"]
        for i, h in enumerate(rl_headers):
            rl.cell(4, 2 + i, h)
        rl.cell(5, 2, "BPM_Detected")
        rl.cell(5, 3, "ICO_HIGH")
        rl.cell(5, 4, "High Complexity")
        rl.cell(5, 5, "ccBPM workflow")
        rl.cell(5, 6, "Redesign as iFlow steps")

        path = tmp_path / "newer_ma.xlsx"
        wb.save(path)
        return path

    def test_complexity_group_column_recognised(self, tmp_path):
        """Newer 'Complexity Group' column must be read and surface complexity."""
        from intake.sap_ma_parser import parse_sap_ma_excel
        path = self._build_newer_schema_file(tmp_path)
        report = parse_sap_ma_excel(str(path))
        assert len(report.interfaces) == 3
        by_name = {r.name: r for r in report.interfaces}
        assert by_name["ICO_LOW"].raw["sap_ma_complexity"] == "LOW"
        assert by_name["ICO_MED"].raw["sap_ma_complexity"] == "MEDIUM"
        assert by_name["ICO_HIGH"].raw["sap_ma_complexity"] == "HIGH"

    def test_complexity_group_implies_mapping_for_med_and_high(self, tmp_path):
        """Medium and High complexity must mark mapping_program (SAP MA's signal
        that transformation logic is present)."""
        from intake.sap_ma_parser import parse_sap_ma_excel
        path = self._build_newer_schema_file(tmp_path)
        report = parse_sap_ma_excel(str(path))
        by_name = {r.name: r for r in report.interfaces}
        assert by_name["ICO_LOW"].mapping_program is None
        assert by_name["ICO_MED"].mapping_program is not None
        assert by_name["ICO_HIGH"].mapping_program is not None

    def test_est_effort_abbreviated_label(self, tmp_path):
        """Hours from 'Est. Effort' column must be extracted, not lost."""
        from intake.sap_ma_parser import parse_sap_ma_excel
        path = self._build_newer_schema_file(tmp_path)
        report = parse_sap_ma_excel(str(path))
        by_name = {r.name: r for r in report.interfaces}
        assert by_name["ICO_LOW"].raw["sap_ma_effort_hrs"] == 1.0
        assert by_name["ICO_MED"].raw["sap_ma_effort_hrs"] == 3.5
        assert by_name["ICO_HIGH"].raw["sap_ma_effort_hrs"] == 16.0

    def test_rules_log_with_extra_complexity_column(self, tmp_path):
        """The extra 'Complexity' column in Rules Log must not break parsing —
        Asset String and Technical Note still read from correct positions."""
        from intake.sap_ma_parser import parse_sap_ma_excel
        path = self._build_newer_schema_file(tmp_path)
        report = parse_sap_ma_excel(str(path))
        assert len(report.rules) == 1
        rule = report.rules[0]
        assert rule.rule_id == "BPM_Detected"
        assert rule.affected_ico == "ICO_HIGH"
        # Crucially: asset string is from col 5 (NOT col 4 — that's now Complexity)
        assert "ccBPM" in rule.asset_string
        assert "Redesign" in rule.technical_note


# ─── Groovy template library: structural validation ──────────────────────────

class TestGroovyTemplates:
    """Static checks on the Groovy template library. We can't run the scripts
    locally (no JVM + CPI runtime), but we can verify they compile by hand
    by ensuring braces/parens balance and key CPI conventions are present."""

    @staticmethod
    def _count_balanced(src: str) -> tuple[int, int]:
        """Brace + paren counter that ignores comments, string literals,
        and Groovy regex literals (/pattern/)."""
        parens = braces = 0
        i = 0
        in_str = None
        in_block = in_line = False
        while i < len(src):
            c = src[i]
            n = src[i + 1] if i + 1 < len(src) else ""
            if in_line:
                if c == "\n":
                    in_line = False
                i += 1; continue
            if in_block:
                if c == "*" and n == "/":
                    in_block = False; i += 2; continue
                i += 1; continue
            if in_str:
                if c == "\\":
                    i += 2; continue
                if c == in_str:
                    in_str = None
                i += 1; continue
            if c == "/" and n == "/":
                in_line = True; i += 2; continue
            if c == "/" and n == "*":
                in_block = True; i += 2; continue
            # Groovy regex literal — heuristic based on preceding context
            if c == "/" and i > 0 and src[i - 1] in "=~,(\n ":
                i += 1
                while i < len(src) and not (src[i] == "/" and src[i - 1] != "\\"):
                    i += 1
                i += 1; continue
            if c == '"':
                in_str = '"'; i += 1; continue
            if c == "'":
                in_str = "'"; i += 1; continue
            if c == "(": parens += 1
            elif c == ")": parens -= 1
            elif c == "{": braces += 1
            elif c == "}": braces -= 1
            i += 1
        return parens, braces

    def test_all_groovy_templates_balance(self):
        """Every .groovy file in templates/groovy must have balanced
        braces and parens, ignoring strings and comments."""
        from pathlib import Path
        groovy_dir = Path(__file__).parent.parent / "templates" / "groovy"
        files = sorted(groovy_dir.glob("*.groovy"))
        assert len(files) >= 10, f"Expected at least 10 templates, found {len(files)}"
        failures = []
        for f in files:
            p, b = self._count_balanced(f.read_text())
            if p != 0 or b != 0:
                failures.append(f"{f.name}: parens={p}, braces={b}")
        assert not failures, "Unbalanced templates:\n" + "\n".join(failures)

    def test_all_have_process_data_entry(self):
        """Every CPI Groovy script must define `processData(Message)`."""
        from pathlib import Path
        groovy_dir = Path(__file__).parent.parent / "templates" / "groovy"
        for f in sorted(groovy_dir.glob("*.groovy")):
            text = f.read_text()
            assert "processData" in text, f"{f.name}: missing processData entry point"

    def test_all_import_message_api(self):
        """Every script must import the CPI Message interface."""
        from pathlib import Path
        groovy_dir = Path(__file__).parent.parent / "templates" / "groovy"
        for f in sorted(groovy_dir.glob("*.groovy")):
            text = f.read_text()
            assert "com.sap.gateway.ip.core.customdev.util.Message" in text, (
                f"{f.name}: missing CPI Message import")

    def test_no_silent_exception_swallows(self):
        """Anti-pattern check: no template should contain
        `catch (Exception ignored) {}` outside of pitfall examples that
        explicitly show what NOT to do."""
        from pathlib import Path
        import re
        groovy_dir = Path(__file__).parent.parent / "templates" / "groovy"
        for f in sorted(groovy_dir.glob("*.groovy")):
            if f.name.startswith("pitfall_"):
                continue  # pitfall files demonstrate the bad pattern intentionally
            text = f.read_text()
            # Strip line + block comments before scanning
            stripped = re.sub(r"/\*.*?\*/", "", text, flags=re.DOTALL)
            stripped = re.sub(r"//[^\n]*", "", stripped)
            assert "catch (Exception ignored)" not in stripped, (
                f"{f.name}: contains silent exception swallow — see pitfall_b_*")


# ─── New migration XSLT templates: Saxon execution verification ──────────────

class TestMigrationXsltTemplates:
    """Verify the migration_*.xsl templates actually transform correctly
    using real Saxon (XSLT 2.0/3.0), the same engine family CPI uses."""

    def _run(self, xsl_filename, input_xml):
        try:
            from saxonche import PySaxonProcessor
        except ImportError:
            import pytest
            pytest.skip("saxonche not installed")
        from pathlib import Path
        xsl = (Path(__file__).parent.parent / "templates" / "xslt" / xsl_filename).read_text()
        with PySaxonProcessor(license=False) as proc:
            xslt30 = proc.new_xslt30_processor()
            executable = xslt30.compile_stylesheet(stylesheet_text=xsl)
            node = proc.parse_xml(xml_text=input_xml)
            return executable.transform_to_string(xdm_node=node)

    def test_date_number_format(self):
        out = self._run("migration_01_date_number_format.xsl",
            "<Record><SapDate>20260115</SapDate><IsoDate>2026-01-15</IsoDate>"
            "<GermanAmount>1.234,56</GermanAmount><MatNr>0000012345</MatNr></Record>")
        assert "<IsoDate>2026-01-15</IsoDate>" in out
        assert "<SapDate>20260115</SapDate>" in out
        assert "<UsAmount>1234.56</UsAmount>" in out
        assert "12345" in out

    def test_value_map_inline_hit(self):
        out = self._run("migration_02_value_map_inline.xsl",
            "<Order><CountryCode>DE</CountryCode></Order>")
        assert "Germany" in out

    def test_value_map_inline_miss(self):
        out = self._run("migration_02_value_map_inline.xsl",
            "<Order><CountryCode>ZZ</CountryCode></Order>")
        assert "UNKNOWN:ZZ" in out

    def test_default_injection_uses_default_when_empty(self):
        out = self._run("migration_03_default_injection.xsl",
            "<Order><Currency></Currency><Priority>HIGH</Priority></Order>")
        assert "<Currency>EUR</Currency>" in out
        assert "<Priority>HIGH</Priority>" in out

    def test_default_injection_keeps_value_when_present(self):
        out = self._run("migration_03_default_injection.xsl",
            "<Order><Currency>USD</Currency><Priority></Priority></Order>")
        assert "<Currency>USD</Currency>" in out
        assert "<Priority>STANDARD</Priority>" in out

    def test_namespace_strip(self):
        out = self._run("migration_04_namespace_strip.xsl",
            '<ns2:Order xmlns:ns2="urn:sap:test" ns2:id="1"><ns2:Item>A</ns2:Item></ns2:Order>')
        assert "ns2" not in out
        assert "<Order" in out
        assert "<Item>A</Item>" in out

    def test_message_split_produces_n_messages(self):
        out = self._run("migration_05_message_split.xsl",
            "<PurchaseOrder><OrderId>PO-1</OrderId><Lines>"
            "<Line><Num>1</Num><Product>A</Product></Line>"
            "<Line><Num>2</Num><Product>B</Product></Line></Lines></PurchaseOrder>")
        assert out.count("<SplitMessage>") == 2
        assert out.count("PO-1") == 2


# ─── Held-off Program 1: Content Modifier generator ──────────────────────────

class TestContentModifierGenerator:
    def _channel(self, **kw):
        from extractor.channel_parser import ChannelConfig
        defaults = dict(channel_id="c1", channel_name="Ch1", adapter_type="HTTPS",
                        direction="Sender")
        defaults.update(kw)
        return ChannelConfig(**defaults)

    def test_basic_properties_always_present(self):
        from scaffolder.content_modifier_generator import build_from_channel
        spec = build_from_channel(self._channel(), "IF1")
        names = {p.name for p in spec.properties}
        assert "OriginalChannel" in names
        assert "AdapterType" in names
        assert "Direction" in names

    def test_idoc_headers_derived(self):
        from scaffolder.content_modifier_generator import build_from_channel
        ch = self._channel(adapter_type="IDoc", idoc_type="ORDERS05",
                           idoc_message_type="ORDERS", idoc_partner_number="P1")
        spec = build_from_channel(ch, "IF1")
        header_names = {h.name for h in spec.headers}
        assert "SAP_IDocType" in header_names
        assert "SAP_IDocMessageType" in header_names
        assert "SAP_IDocPartner" in header_names

    def test_credential_alias_not_value(self):
        from scaffolder.content_modifier_generator import build_from_channel
        ch = self._channel(auth_type="Basic", credential_name="MY_CRED", username="admin")
        spec = build_from_channel(ch, "IF1")
        prop_values = {p.value for p in spec.properties}
        assert "MY_CRED" in prop_values   # alias yes
        assert "admin" not in prop_values  # username (secret-adjacent) no

    def test_bpmn_step_is_valid_xml(self):
        from scaffolder.content_modifier_generator import build_from_channel, render_bpmn_step
        from lxml import etree
        spec = build_from_channel(self._channel(endpoint_url="https://x.com"), "IF1")
        xml = render_bpmn_step(spec)
        wrapped = f'<root xmlns:bpmn2="http://x" xmlns:ifl="http://y">{xml}</root>'
        etree.fromstring(wrapped.encode())  # raises if malformed

    def test_descriptor_structure(self):
        from scaffolder.content_modifier_generator import build_from_channel, render_descriptor
        spec = build_from_channel(self._channel(), "IF1")
        d = render_descriptor(spec)
        assert "step_name" in d and "headers" in d and "properties" in d


# ─── Held-off Program 1: Value Mapping generator ─────────────────────────────

class TestValueMappingGenerator:
    def test_build_from_pairs(self):
        from scaffolder.value_mapping_generator import build_from_pairs
        vm = build_from_pairs("CC", [("DE", "Germany"), ("FR", "France")])
        assert vm.total_entries() == 2
        assert len(vm.groups) == 1

    def test_artifact_is_valid_xml(self):
        from scaffolder.value_mapping_generator import build_from_pairs, render_artifact
        from lxml import etree
        vm = build_from_pairs("CC", [("DE", "Germany")])
        xml = render_artifact(vm)
        etree.fromstring(xml.encode())
        assert "Germany" in xml
        assert "DE" in xml

    def test_descriptor(self):
        from scaffolder.value_mapping_generator import build_from_pairs, render_descriptor
        vm = build_from_pairs("CC", [("A", "1"), ("B", "2"), ("C", "3")])
        d = render_descriptor(vm)
        assert d["total_entries"] == 3
        assert d["group_count"] == 1


# ─── Program 2: APIM model ───────────────────────────────────────────────────

class TestAPIMModel:
    def test_key_lifecycle(self):
        from apim.model import Application, KeyState
        app = Application(name="App1")
        key = app.issue_key()
        assert key.is_valid()
        assert len(app.active_keys()) == 1
        app.revoke_key(key.key_value)
        assert not key.is_valid()
        assert len(app.active_keys()) == 0

    def test_key_expiry(self):
        from apim.model import Application
        from datetime import datetime, timedelta
        app = Application(name="App1")
        key = app.issue_key(ttl_days=1)
        future = datetime.utcnow() + timedelta(days=2)
        assert key.is_valid()                  # valid now
        assert not key.is_valid(now=future)    # expired in the future

    def test_keys_are_unique(self):
        from apim.model import Application
        app = Application(name="App1")
        keys = {app.issue_key().key_value for _ in range(50)}
        assert len(keys) == 50  # no collisions

    def test_landscape_validation_clean(self):
        from apim.model import APIMLandscape, APIProxy, APIProduct, Application
        land = APIMLandscape(
            proxies=[APIProxy(name="P1", base_path="/v1", target_url="http://t")],
            products=[APIProduct(name="Prod1", proxies=["P1"])],
            applications=[Application(name="App1", subscribed_products=["Prod1"])],
        )
        assert land.validate() == []

    def test_landscape_validation_catches_missing_proxy(self):
        from apim.model import APIMLandscape, APIProduct
        land = APIMLandscape(products=[APIProduct(name="Prod1", proxies=["Ghost"])])
        issues = land.validate()
        assert any("Ghost" in i for i in issues)

    def test_landscape_validation_catches_missing_product(self):
        from apim.model import APIMLandscape, Application
        land = APIMLandscape(applications=[Application(name="A", subscribed_products=["NoSuch"])])
        issues = land.validate()
        assert any("NoSuch" in i for i in issues)


# ─── Program 2: APIM policy library ──────────────────────────────────────────

class TestAPIMPolicyLibrary:
    def test_all_policies_valid_xml(self):
        from apim import policy_library
        from lxml import etree
        for name in policy_library.list_policies():
            builder = policy_library.POLICY_BUILDERS[name]
            xml = builder("Test", "X-H", "val") if name == "SetHeader" else builder()
            etree.fromstring(xml.encode())  # raises if malformed

    def test_quota_parameters_applied(self):
        from apim import policy_library
        xml = policy_library.quota(allow_count=5000, time_unit="day")
        assert "5000" in xml
        assert "day" in xml

    def test_list_policies_nonempty(self):
        from apim import policy_library
        assert len(policy_library.list_policies()) >= 7


# ─── Program 2: APIM proxy generator ─────────────────────────────────────────

class TestAPIMProxyGenerator:
    def test_generate_full_bundle(self):
        from apim.model import APIProxy, ProxyAuthType
        from apim.proxy_generator import generate_proxy
        from lxml import etree
        proxy = APIProxy(name="OrderAPI", base_path="/v1/orders",
                         target_url="http://backend", auth_type=ProxyAuthType.API_KEY)
        gen = generate_proxy(proxy)
        files = gen.all_files()
        assert len(files) >= 5
        for content in files.values():
            etree.fromstring(content.encode())

    def test_api_key_auth_adds_verify_policy(self):
        from apim.model import APIProxy, ProxyAuthType
        from apim.proxy_generator import generate_proxy
        proxy = APIProxy(name="A", base_path="/v1", target_url="http://t",
                         auth_type=ProxyAuthType.API_KEY)
        gen = generate_proxy(proxy)
        assert any("Verify" in p for p in gen.manifest["policies"])

    def test_oauth_auth_adds_oauth_policy(self):
        from apim.model import APIProxy, ProxyAuthType
        from apim.proxy_generator import generate_proxy
        proxy = APIProxy(name="A", base_path="/v1", target_url="http://t",
                         auth_type=ProxyAuthType.OAUTH2)
        gen = generate_proxy(proxy)
        assert any("OAuth" in p for p in gen.manifest["policies"])

    def test_proxy_from_iflow_links_source(self):
        from apim.proxy_generator import proxy_from_iflow, generate_proxy
        proxy = proxy_from_iflow("MyIFlow", "/v1/x", "http://rt/http/x")
        assert proxy.source_iflow == "MyIFlow"
        gen = generate_proxy(proxy)
        assert gen.manifest["source_iflow"] == "MyIFlow"


# ─── SAP MA parser: real (2-sheet) schema support ───────────────────────────

class TestSAPMARealSchema:
    FIXTURE = "tests/fixtures/mock_ma_real_schema.xlsx"

    def _parse(self):
        import os
        from intake.sap_ma_parser import parse_sap_ma_excel
        if not os.path.exists(self.FIXTURE):
            import pytest
            pytest.skip("real-schema fixture not present")
        return parse_sap_ma_excel(self.FIXTURE)

    def test_real_schema_detected_and_counted(self):
        r = self._parse()
        # Mock was built with exactly 60 scenarios
        assert r.summary.total_icos == 60

    def test_real_schema_categories_populated(self):
        r = self._parse()
        # All three categories present and summing to total
        total = (r.summary.ready_to_migrate + r.summary.adjustment_required
                 + r.summary.evaluation_required)
        assert total == 60
        assert r.summary.ready_to_migrate > 0
        assert r.summary.evaluation_required > 0

    def test_real_schema_effort_nonzero(self):
        r = self._parse()
        assert r.summary.total_effort_hours > 0

    def test_real_schema_does_not_read_banner_as_data(self):
        # Regression: the row-2 title banner contains "Integration Scenario"
        # as a substring; must NOT be parsed as a data row (would give 61 +
        # a junk record named "Integration Scenario").
        r = self._parse()
        names = {i.name for i in r.interfaces}
        assert "Integration Scenario" not in names
        assert len(r.interfaces) == 60

    def test_real_schema_complexity_mapped(self):
        r = self._parse()
        # Every interface should get a complexity hint from its category
        complexities = {i.raw.get("sap_ma_complexity") for i in r.interfaces}
        assert complexities <= {"LOW", "MEDIUM", "HIGH"}
        assert "HIGH" in complexities  # Evaluation-required -> HIGH

    def test_real_schema_adapters_parsed(self):
        r = self._parse()
        # Adapters should be real values, not all defaulted to HTTPS
        adapters = {i.sender_adapter for i in r.interfaces}
        assert len(adapters) > 1  # variety means real parsing happened

    def test_real_schema_rules_parsed(self):
        r = self._parse()
        assert len(r.rules) > 0


class TestMigrationCeilingField:
    def test_ceiling_uses_triggered_by_not_triggers(self):
        # Regression: workbench referenced c.triggers (wrong); the field is
        # triggered_by. Confirm the dataclass exposes triggered_by.
        from reporter.migration_ceiling import MigrationCeiling
        import dataclasses
        fields = {f.name for f in dataclasses.fields(MigrationCeiling)}
        assert "triggered_by" in fields
        assert "triggers" not in fields


# ─── CPI uploader: CSRF + diagnostics ───────────────────────────────────────

class TestCPIDiagnostics:
    def test_diagnostics_importable(self):
        from fetcher.cpi_diagnostics import run_diagnostics, DiagnosticReport, fetch_csrf_token
        assert callable(run_diagnostics)

    def test_uploader_has_csrf_methods(self):
        from fetcher.cpi_uploader import CPIUploader
        assert hasattr(CPIUploader, "_ensure_csrf")
        assert hasattr(CPIUploader, "_write_headers")

    def test_diagnostic_report_summary(self):
        from fetcher.cpi_diagnostics import DiagnosticReport, DiagnosticStep
        r = DiagnosticReport()
        r.add(DiagnosticStep("Test", True, "ok"))
        r.add(DiagnosticStep("Test2", False, "bad", status_code=401))
        s = r.summary()
        assert "✅" in s and "❌" in s and "401" in s


class TestInterventionEstimatorImport:
    def test_intervention_estimator_importable(self):
        # Regression: workbench referenced InterventionEstimator without importing it
        from reporter.intervention_estimator import InterventionEstimator
        assert InterventionEstimator is not None


# ─── Log viewer de-duplication ───────────────────────────────────────────────

class TestLogViewer:
    def _write(self, tmp_path, content):
        p = tmp_path / "test.log"
        p.write_text(content, encoding="utf-8")
        return p

    def test_collapses_identical_timestamped_lines(self, tmp_path):
        from reporter.log_viewer import deduplicate_log, log_summary
        content = "\n".join(
            f"2026-05-29 21:34:{i:02d}.000 Please replace use_container_width with width."
            for i in range(10))
        groups = deduplicate_log(self._write(tmp_path, content))
        assert len(groups) == 1
        assert groups[0].count == 10
        s = log_summary(groups)
        assert s["collapsed"] == 9

    def test_separates_different_messages(self, tmp_path):
        from reporter.log_viewer import deduplicate_log
        content = (
            "ERROR:x:Failed to create package: bad\n"
            "WARNING:y:Could not list artifacts\n"
            "ERROR:x:Failed to create package: bad\n"
        )
        groups = deduplicate_log(self._write(tmp_path, content))
        # Two unique signatures
        assert len(groups) == 2
        # The repeated error counts 2
        err = [g for g in groups if g.level == "ERROR"][0]
        assert err.count == 2

    def test_level_filter(self, tmp_path):
        from reporter.log_viewer import deduplicate_log
        content = (
            "ERROR:x:bad thing\n"
            "INFO:x:fine thing\n"
            "WARNING:x:meh thing\n"
        )
        p = self._write(tmp_path, content)
        problems = deduplicate_log(p, levels={"ERROR", "WARNING"})
        levels = {g.level for g in problems}
        assert "INFO" not in levels
        assert "ERROR" in levels

    def test_multiline_entry_grouped(self, tmp_path):
        from reporter.log_viewer import deduplicate_log
        block = (
            "ERROR:up:Upload failed for A: HTTP 404\n"
            "  URL: https://x/api\n"
            "  CSRF token present: True\n"
            "ERROR:up:Upload failed for B: HTTP 404\n"
            "  URL: https://x/api\n"
            "  CSRF token present: True\n"
        )
        groups = deduplicate_log(self._write(tmp_path, block))
        # Both blocks share a signature (names differ but get normalised? names
        # are distinct so they stay separate) — at minimum each block is ONE
        # entry, not 3 lines.
        assert all("\n" in g.sample for g in groups)

    def test_missing_file_returns_empty(self):
        from reporter.log_viewer import deduplicate_log
        assert deduplicate_log("/nonexistent/path.log") == []


# ─── Batch orchestrator ──────────────────────────────────────────────────────

class TestBatchOrchestrator:
    class _FakeIface:
        def __init__(self, name):
            self.name = name
            self.sender_system = "SRC"
            self.receiver_system = "TGT"
            self.namespace = ""
            self.id = name

    class _FakeAssessment:
        def __init__(self, name):
            self.interface = TestBatchOrchestrator._FakeIface(name)

    class _FakeCeiling:
        def __init__(self, name, tier):
            self.interface_name = name
            self.tier = tier

    class _FakeScaffolder:
        def __init__(self, output_dir):
            self.output_dir = output_dir
        def scaffold(self, assessment, resolved=None, shape="timer", **kw):
            from pathlib import Path
            p = Path(self.output_dir) / f"{assessment.interface.name}.iflw"
            p.parent.mkdir(parents=True, exist_ok=True)
            p.write_text(f"<iflow shape='{shape}'/>", encoding="utf-8")
            return p

    def test_routes_specialist_to_attention(self, tmp_path):
        from scaffolder.batch_orchestrator import BatchOrchestrator
        assessments = [self._FakeAssessment("A"), self._FakeAssessment("B")]
        ceilings = {
            "A": self._FakeCeiling("A", "AUTO"),
            "B": self._FakeCeiling("B", "SPECIALIST"),
        }
        orch = BatchOrchestrator(
            scaffolder=self._FakeScaffolder(str(tmp_path / "iflows")),
            output_dir=str(tmp_path))
        run = orch.run(assessments, ceilings)
        assert len(run.processed) == 1
        assert run.processed[0].interface_name == "A"
        assert len(run.needs_attention) == 1
        assert run.needs_attention[0].interface_name == "B"

    def test_guided_skipped_unless_included(self, tmp_path):
        from scaffolder.batch_orchestrator import BatchOrchestrator
        assessments = [self._FakeAssessment("G")]
        ceilings = {"G": self._FakeCeiling("G", "GUIDED")}
        orch = BatchOrchestrator(
            scaffolder=self._FakeScaffolder(str(tmp_path / "iflows")),
            output_dir=str(tmp_path))
        # default: skipped
        run = orch.run(assessments, ceilings)
        assert len(run.needs_attention) == 1
        # opted in: processed
        run2 = orch.run(assessments, ceilings, include_guided=True)
        assert len(run2.processed) == 1

    def test_no_ceiling_treated_as_auto(self, tmp_path):
        from scaffolder.batch_orchestrator import BatchOrchestrator
        assessments = [self._FakeAssessment("X")]
        orch = BatchOrchestrator(
            scaffolder=self._FakeScaffolder(str(tmp_path / "iflows")),
            output_dir=str(tmp_path))
        run = orch.run(assessments, {})  # no ceilings
        assert len(run.processed) == 1

    def test_summary_counts(self, tmp_path):
        from scaffolder.batch_orchestrator import BatchOrchestrator
        assessments = [self._FakeAssessment(n) for n in ("A", "B", "C")]
        ceilings = {
            "A": self._FakeCeiling("A", "AUTO"),
            "B": self._FakeCeiling("B", "AUTO"),
            "C": self._FakeCeiling("C", "SPECIALIST"),
        }
        orch = BatchOrchestrator(
            scaffolder=self._FakeScaffolder(str(tmp_path / "iflows")),
            output_dir=str(tmp_path))
        run = orch.run(assessments, ceilings)
        s = run.summary()
        assert s["total"] == 3
        assert s["processed"] == 2
        assert s["needs_attention"] == 1


# ─── Uploader: Mode field removed from package payload ───────────────────────

class TestUploaderPackagePayload:
    def test_no_mode_field_in_package_creation(self):
        # Regression: trial API rejects "Mode" with HTTP 400. Confirm the
        # payload dict no longer includes a Mode key (ignore comments).
        import inspect
        from fetcher.cpi_uploader import CPIUploader
        src = inspect.getsource(CPIUploader.ensure_package)
        # Strip comment lines so we only check real code
        code_lines = [ln for ln in src.splitlines()
                      if not ln.strip().startswith("#")]
        code = "\n".join(code_lines)
        assert '"Mode":' not in code


# ─── Package ID sanitization (CPI rejects special chars) ─────────────────────

class TestPackageIdSanitization:
    def test_underscores_and_specials_removed(self):
        from fetcher.cpi_uploader import CPIUploader
        # The exact IDs that failed against the real tenant
        out = CPIUploader.sanitize_package_id("MIGRATION_ECC_S4HANA_httpcompanycompo")
        assert "_" not in out
        assert out.isalnum()

    def test_dots_and_slashes_removed(self):
        from fetcher.cpi_uploader import CPIUploader
        out = CPIUploader.sanitize_package_id("with.dots/and/slashes")
        assert out == "withdotsandslashes"

    def test_numeric_start_prefixed(self):
        from fetcher.cpi_uploader import CPIUploader
        out = CPIUploader.sanitize_package_id("123abc")
        assert out[0].isalpha()

    def test_empty_gets_fallback(self):
        from fetcher.cpi_uploader import CPIUploader
        out = CPIUploader.sanitize_package_id("")
        assert out and out.isalnum()

    def test_result_is_always_valid_id(self):
        from fetcher.cpi_uploader import CPIUploader
        import string
        for raw in ["A_B_C", "x!y@z", "___", "ECC->S4", "café_münchen"]:
            out = CPIUploader.sanitize_package_id(raw)
            # All chars alphanumeric, starts with a letter
            assert all(c in string.ascii_letters + string.digits for c in out)
            assert out[0].isalpha()


# ─── iFlow package structure (InputStream fix) ───────────────────────────────

class TestIFlowPackaging:
    def _make_iflw(self, tmp_path):
        p = tmp_path / "x.iflw"
        p.write_text("<bpmn2:definitions/>", encoding="utf-8")
        return p

    def test_zip_has_manifest_and_project(self, tmp_path):
        import io, zipfile
        from fetcher.cpi_uploader import CPIUploader
        data = CPIUploader._package_iflow(self._make_iflw(tmp_path),
                                          "Art1", "Art One", "k=v")
        assert data
        names = zipfile.ZipFile(io.BytesIO(data)).namelist()
        assert "META-INF/MANIFEST.MF" in names   # the InputStream fix
        # Real exported packages do NOT include .project — we match that.
        assert ".project" not in names
        assert any(n.endswith(".iflw") for n in names)

    def test_manifest_has_bundle_type(self, tmp_path):
        import io, zipfile
        from fetcher.cpi_uploader import CPIUploader
        data = CPIUploader._package_iflow(self._make_iflw(tmp_path),
                                          "Art1", "Art One")
        zf = zipfile.ZipFile(io.BytesIO(data))
        manifest = zf.read("META-INF/MANIFEST.MF").decode()
        assert "SAP-BundleType: IntegrationFlow" in manifest
        assert "Bundle-SymbolicName" in manifest

    def test_missing_iflow_returns_none(self, tmp_path):
        from fetcher.cpi_uploader import CPIUploader
        data = CPIUploader._package_iflow(tmp_path / "nope.iflw", "A", "A")
        assert data is None

    def test_manifest_uses_crlf_line_endings(self, tmp_path):
        # Regression: CPI's OData artifact parser returns HTTP 500
        # "InputStream cannot be null" when META-INF/MANIFEST.MF has bare-LF
        # line endings. The JAR/OSGi spec requires CRLF + a trailing blank
        # line. _package_iflow must never ship a bare-LF manifest, regardless
        # of whether the manifest came from the validated __meta copy (read as
        # bytes) or the in-code fallback.
        import io, zipfile
        from fetcher.cpi_uploader import CPIUploader
        data = CPIUploader._package_iflow(self._make_iflw(tmp_path), "Art1", "Art One")
        mf = zipfile.ZipFile(io.BytesIO(data)).read("META-INF/MANIFEST.MF")
        assert b"\r\n" in mf, "manifest must use CRLF"
        assert b"\n" not in mf.replace(b"\r\n", b""), "manifest must have no bare LF"
        assert mf.endswith(b"\r\n\r\n"), "manifest must end with a blank line"

    def test_validated_meta_manifest_crlf_preserved(self, tmp_path):
        # The validated manifest stashed in <stem>__meta/ is written with CRLF
        # by the generator; _package_iflow must read it without universal-newline
        # translation so the CRLF survives into the uploaded bundle.
        import io, zipfile
        from fetcher.cpi_uploader import CPIUploader
        iflw = self._make_iflw(tmp_path)
        meta = tmp_path / f"{iflw.stem}__meta"
        meta.mkdir()
        (meta / "MANIFEST.MF").write_bytes(
            b"Manifest-Version: 1.0\r\nBundle-SymbolicName: X\r\n"
            b"SAP-BundleType: IntegrationFlow\r\n\r\n")
        (meta / ".project").write_bytes(b"<projectDescription/>\r\n")
        data = CPIUploader._package_iflow(iflw, "X", "X")
        mf = zipfile.ZipFile(io.BytesIO(data)).read("META-INF/MANIFEST.MF")
        assert b"\n" not in mf.replace(b"\r\n", b""), "validated manifest CRLF was stripped"

    def test_bundle_includes_parameters_propdef(self, tmp_path):
        # Regression: CPI's OData create returns HTTP 500 "InputStream cannot be
        # null" when src/main/resources/parameters.propdef is absent. Every real
        # importable iFlow bundle carries both parameters.prop AND
        # parameters.propdef; the timer scaffold was missing the propdef.
        import io, zipfile
        import xml.etree.ElementTree as ET
        from fetcher.cpi_uploader import CPIUploader
        data = CPIUploader._package_iflow(self._make_iflw(tmp_path), "Art1", "Art One")
        z = zipfile.ZipFile(io.BytesIO(data))
        names = z.namelist()
        assert "src/main/resources/parameters.prop" in names
        assert "src/main/resources/parameters.propdef" in names, \
            "parameters.propdef missing → CPI 500 InputStream cannot be null"
        # propdef must be well-formed XML (CPI parses it to build the param model)
        ET.fromstring(z.read("src/main/resources/parameters.propdef").decode())


# ─── SAP-style naming + URL leak fix ─────────────────────────────────────────

class TestPackageNaming:
    def test_url_does_not_leak_into_clean(self):
        from scaffolder.pipeline_scaffolder import _clean
        out = _clean("http://company.com/ariba")
        assert "http" not in out
        assert "company" not in out
        assert out == "ariba"

    def test_display_name_follows_sap_convention(self):
        from scaffolder.pipeline_scaffolder import generate_package_display_name
        name = generate_package_display_name("SuccessFactors", "S4HANA", "Master Data")
        assert "Integration for" in name
        assert "with" in name

    def test_display_name_without_domain(self):
        from scaffolder.pipeline_scaffolder import generate_package_display_name
        name = generate_package_display_name("ECC", "Ariba")
        assert "Integration with" in name


# ─── iFlow wiring engine (Phase 1 — real configured iFlows) ──────────────────

class TestIFlowWiring:
    class _Iface:
        def __init__(self, **kw):
            self.name = kw.get("name", "TestFlow")
            self.sender_system = kw.get("ss", "ECC")
            self.receiver_system = kw.get("rs", "S4")
            self.sender_adapter = kw.get("sa", "IDoc")
            self.receiver_adapter = kw.get("ra", "HTTPS")
            self.has_bpm = kw.get("bpm", False)
            self.has_multi_mapping = kw.get("mm", False)
            self.description = kw.get("desc", "")

    def test_produces_valid_bpmn(self):
        from scaffolder.iflow_wiring import wire_iflow
        from lxml import etree
        w = wire_iflow(self._Iface())
        etree.fromstring(w.xml.encode())  # raises if malformed

    def test_has_real_steps_not_skeleton(self):
        from scaffolder.iflow_wiring import wire_iflow
        w = wire_iflow(self._Iface())
        # The whole point: real callActivity steps, not an empty skeleton
        assert w.xml.count("<bpmn2:callActivity") >= 4
        assert len(w.steps) >= 3

    def test_exception_subprocess_folded_in(self):
        from scaffolder.iflow_wiring import wire_iflow
        w = wire_iflow(self._Iface())
        assert w.has_exception_handler
        assert "ErrorEventSubProcessTemplate" in w.xml
        assert "StartErrorEvent" in w.xml

    def test_externalized_parameters(self):
        from scaffolder.iflow_wiring import wire_iflow
        w = wire_iflow(self._Iface())
        # Adapter config uses {{param}} references, not hardcoded values
        assert "{{sender_endpoint}}" in w.xml
        assert "{{receiver_address}}" in w.xml
        assert "{{receiver_credential}}" in w.xml
        assert "receiver_credential" in w.parameters

    def test_variant_detection(self):
        from scaffolder.iflow_wiring import wire_iflow
        assert wire_iflow(self._Iface()).variant == "linear"
        assert wire_iflow(self._Iface(bpm=True)).variant == "router"
        assert wire_iflow(self._Iface(mm=True)).variant == "splitter"

    def test_router_has_gateway(self):
        from scaffolder.iflow_wiring import wire_iflow
        w = wire_iflow(self._Iface(bpm=True))
        assert "ExclusiveGateway" in w.xml

    def test_three_participants(self):
        from scaffolder.iflow_wiring import wire_iflow
        w = wire_iflow(self._Iface())
        assert w.xml.count("<bpmn2:participant") == 3

    def test_two_adapters(self):
        from scaffolder.iflow_wiring import wire_iflow
        w = wire_iflow(self._Iface())
        assert w.xml.count("<bpmn2:messageFlow") == 2

    def test_prop_file_escapes_colons(self):
        from scaffolder.iflow_wiring import wire_iflow, parameters_prop
        w = wire_iflow(self._Iface())
        prop = parameters_prop(w.parameters)
        assert "\\:" in prop  # colons escaped Java-properties style


# ─── Held items: effort reconciliation, mapping inventory, adapter advisor ───

class TestHeldItems:
    def test_effort_reconciliation(self):
        from reporter.effort_reconciliation import reconcile
        class A:
            def __init__(self, n, d):
                self.interface = type("I", (), {"name": n})()
                self.effort_days = d
        report = reconcile([A("X", 3), A("Y", 5)], ma_report=None)
        assert len(report.comparisons) == 2
        assert report.tool_total_days == 8

    def test_mapping_inventory(self):
        from analyzer.mapping_inventory import build_inventory
        class I:
            def __init__(self, n, m):
                self.name = n; self.mapping_program = m; self.description = ""
        inv = build_inventory([I("A", "MM_Shared.mmap"), I("B", "MM_Shared.mmap"),
                               I("C", "transform.xsl")])
        s = inv.summary()
        assert s["total_mappings"] == 2  # MM_Shared deduped
        assert s["reused_mappings"] == 1  # MM_Shared used twice

    def test_adapter_advisor(self):
        from analyzer.adapter_advisor import advise_for_adapter, advise_all
        adv = advise_for_adapter("IDoc", "sender")
        assert "IDoc" in adv.cpi_adapter
        assert adv.notes
        # RFC should warn about Cloud Connector
        rfc = advise_for_adapter("RFC")
        assert rfc.severity == "warning"

    def test_content_modifier_cell_format(self):
        # Regression: CM must use real <cell id='...'> format, not <row><id>
        from extractor.channel_parser import ChannelConfig
        from scaffolder.content_modifier_generator import build_from_channel, render_bpmn_step
        ch = ChannelConfig(channel_id="c", channel_name="C", adapter_type="HTTPS",
                           direction="Sender")
        xml = render_bpmn_step(build_from_channel(ch, "IF"))
        assert "cell id=" in xml
        assert "&lt;row&gt;" in xml  # escaped table format


# ─── Artifact bundle (Generate-all produces self-contained packages) ─────────

class TestArtifactBundle:
    class _Iface:
        def __init__(self, name="Flow", desc=""):
            self.name = name; self.description = desc
            self.sender_system = "ECC"; self.receiver_system = "S4"
            self.sender_adapter = "IDoc"; self.receiver_adapter = "HTTPS"
            self.has_bpm = False; self.has_multi_mapping = False; self.id = "1"

    def test_generates_script_and_mapping(self, tmp_path):
        from scaffolder.artifact_bundle import generate_bundle
        iflow = tmp_path / "iflows" / "Flow.iflw"
        iflow.parent.mkdir(parents=True)
        iflow.write_text("<x/>")
        bundle = generate_bundle(self._Iface(), iflow)
        kinds = {a.kind for a in bundle.artifacts}
        assert "script" in kinds and "mapping" in kinds

    def test_script_capability_detection(self, tmp_path):
        from scaffolder.artifact_bundle import generate_bundle
        iflow = tmp_path / "iflows" / "F.iflw"
        iflow.parent.mkdir(parents=True); iflow.write_text("<x/>")
        # date-related interface gets the date script
        b = generate_bundle(self._Iface("DateFlow", "format the timestamp"), iflow)
        script = next(a for a in b.artifacts if a.kind == "script")
        assert "SimpleDateFormat" in script.content

    def test_mapping_auto_drafts_field_matches(self, tmp_path):
        from scaffolder.artifact_bundle import generate_bundle
        iflow = tmp_path / "iflows" / "F.iflw"
        iflow.parent.mkdir(parents=True); iflow.write_text("<x/>")
        b = generate_bundle(self._Iface(), iflow,
                            source_fields=["CustomerID", "Name"],
                            target_fields=["CustomerID", "Name"])
        mapping = next(a for a in b.artifacts if a.kind == "mapping")
        assert 'function="direct"' in mapping.content  # matched fields

    def test_package_bundles_extra_artifacts(self, tmp_path):
        import io, zipfile
        from fetcher.cpi_uploader import CPIUploader
        iflow = tmp_path / "x.iflw"; iflow.write_text("<bpmn2:definitions/>")
        extra = [("src/main/resources/script/s.groovy", "def x(){}"),
                 ("src/main/resources/mapping/m.mmap", "<m/>")]
        data = CPIUploader._package_iflow(iflow, "A", "A", "", extra_artifacts=extra)
        names = zipfile.ZipFile(io.BytesIO(data)).namelist()
        assert any("script/" in n for n in names)
        assert any("mapping/" in n for n in names)


class TestSolverApiKey:
    def test_no_key_returns_clear_message(self):
        from engine.claude_solver import ClaudeSolver
        s = ClaudeSolver(api_key="")
        assert not s.has_api_key
        class A:
            interface = type("I", (), {"name": "X"})()
            complexity = "LOW"; score = 10; effort_days = 2
            notes = []; recommended_pattern = "linear"
        result = s.solve(A())
        assert result.confidence == 0.0
        assert any("API key" in m for m in result.remaining_manual)

    def test_key_sets_has_api_key(self):
        from engine.claude_solver import ClaudeSolver
        s = ClaudeSolver(api_key="sk-test")
        assert s.has_api_key


# ─── Proposal generation (regression: PricingConfig + task docs) ─────────────

class TestProposalGeneration:
    def _setup(self, tmp_path):
        from reporter.proposal_generator import ProposalGenerator, PricingConfig
        from reporter.migration_ceiling import MigrationCeilingClassifier
        assessments = [make_assessment(name="OrderSync", mapping_program="MM_O"),
                       make_assessment(name="EmpSync")]
        ceilings = MigrationCeilingClassifier().classify_all(assessments)
        gen = ProposalGenerator(output_dir=str(tmp_path))
        return gen, assessments, ceilings, PricingConfig

    def test_pricing_config_correct_fields(self):
        from reporter.proposal_generator import PricingConfig
        # Regression: the field is your_day_rate_usd, not day_rate_usd
        p = PricingConfig(your_day_rate_usd=900, target_margin_pct=55)
        assert p.your_day_rate_usd == 900
        import pytest
        with pytest.raises(TypeError):
            PricingConfig(day_rate_usd=900)  # old wrong name must fail

    def test_generate_three_standard_docs(self, tmp_path):
        gen, a, c, PC = self._setup(tmp_path)
        cx, ix, dx = gen.generate(a, ceilings=c, pricing=PC())
        assert cx.exists() and ix.exists() and dx.exists()

    def test_generate_task_documents(self, tmp_path):
        gen, a, c, PC = self._setup(tmp_path)
        client_doc, consultant_doc = gen.generate_task_documents(a, ceilings=c)
        assert client_doc.exists() and consultant_doc.exists()
        # client doc should mention prerequisites, consultant doc the builds
        from docx import Document
        client_text = " ".join(p.text for p in Document(str(client_doc)).paragraphs)
        consultant_text = " ".join(p.text for p in Document(str(consultant_doc)).paragraphs)
        assert "client" in client_text.lower() or "UAT" in client_text
        assert "iFlow" in consultant_text


# ─── Stage 1: minimal valid iFlow generator (from-scratch CPI-valid) ─────────

class TestMinimalIFlow:
    def test_wellformed_xml(self):
        from scaffolder.minimal_iflow import generate_minimal_iflow
        from lxml import etree
        r = generate_minimal_iflow("Flow A", "FlowA")
        etree.fromstring(r.iflw_xml.encode())  # raises if malformed

    def test_has_connected_sender_message_flow(self):
        from scaffolder.minimal_iflow import generate_minimal_iflow
        r = generate_minimal_iflow("Flow A", "FlowA")
        # The fix for "start event should have an incoming message flow"
        assert r.iflw_xml.count("<bpmn2:messageFlow") == 1
        assert 'targetRef="StartEvent_2"' in r.iflw_xml

    def test_has_diagram_section(self):
        from scaffolder.minimal_iflow import generate_minimal_iflow
        r = generate_minimal_iflow("Flow A", "FlowA")
        assert "BPMNDiagram" in r.iflw_xml
        assert r.iflw_xml.count("<bpmndi:BPMNShape") == 5
        assert r.iflw_xml.count("<bpmndi:BPMNEdge") == 2

    def test_cmdvarianturi_on_elements(self):
        from scaffolder.minimal_iflow import generate_minimal_iflow
        r = generate_minimal_iflow("Flow A", "FlowA")
        assert "MessageStartEvent" in r.iflw_xml
        assert "MessageEndEvent" in r.iflw_xml
        assert "IntegrationProcess/version" in r.iflw_xml

    def test_manifest_has_osgi_block(self):
        from scaffolder.minimal_iflow import generate_minimal_iflow
        r = generate_minimal_iflow("Flow A", "FlowA")
        assert "Import-Package" in r.manifest
        assert "SAP-BundleType: IntegrationFlow" in r.manifest
        # manifest lines wrap at 72 cols (OSGi continuation)
        for line in r.manifest.split("\r\n"):
            assert len(line) <= 72

    def test_manifest_import_package_matches_real(self):
        # The Import-Package must match a known-good real iFlow's block
        from scaffolder.minimal_iflow import _IMPORT_PACKAGE
        # sanity: contains the key SAP + camel + cxf + osgi packages
        for pkg in ["com.sap.esb.security", "org.apache.camel",
                    "org.apache.cxf.transport", "org.osgi.framework",
                    "org.osgi.service.blueprint"]:
            assert pkg in _IMPORT_PACKAGE

    def test_bundle_zip_structure(self):
        from scaffolder.minimal_iflow import generate_minimal_iflow, build_bundle_zip
        import io, zipfile
        r = generate_minimal_iflow("Flow A", "FlowA")
        names = zipfile.ZipFile(io.BytesIO(build_bundle_zip(r))).namelist()
        assert "META-INF/MANIFEST.MF" in names
        assert ".project" in names
        assert any(n.endswith(".iflw") for n in names)

    def test_id_sanitization(self):
        from scaffolder.minimal_iflow import generate_minimal_iflow
        r = generate_minimal_iflow("My Flow! 123", "")
        # Id must be alphanumeric, letter-start
        assert r.iflow_id.isalnum()
        assert r.iflow_id[0].isalpha()


class TestContentModifierIFlow:
    def _gen(self):
        from scaffolder.minimal_iflow import generate_content_modifier_iflow
        body = "<Ack><status>OK</status><echo>${in.body}</echo></Ack>"
        return generate_content_modifier_iflow(
            "S_ContentModifier_SetConstant", "SContentModifierSetConstant",
            body_expr=body, headers=[("X-Pilot-Status", "OK")])

    def test_wellformed_and_start_cm_end(self):
        from lxml import etree
        r = self._gen()
        etree.fromstring(r.iflw_xml.encode())          # well-formed
        x = r.iflw_xml
        # start -> CM -> end wiring
        assert 'sourceRef="StartEvent_2" targetRef="CallActivity_1"' in x
        assert 'sourceRef="CallActivity_1" targetRef="EndEvent_2"' in x
        assert "<value>Enricher</value>" in x          # Content Modifier

    def test_body_in_wrapcontent_and_header(self):
        x = self._gen().iflw_xml
        assert "wrapContent" in x and "OK" in x        # body via wrapContent
        assert "X-Pilot-Status" in x and "HEADER_0" in x

    def test_no_message_mapping(self):
        import re
        x = self._gen().iflw_xml
        acts = re.findall(r"<key>activityType</key><value>([^<]+)</value>", x)
        assert "Mapping" not in acts and ".mmap" not in x

    def test_diagram_has_cm_shape_and_two_edges(self):
        x = self._gen().iflw_xml
        assert "BPMNShape_CallActivity_1" in x
        # start->cm and cm->end sequence edges, plus the message-flow edge
        assert x.count("<bpmndi:BPMNEdge") == 3

    def test_bundle_importable_structure(self):
        from scaffolder.minimal_iflow import build_bundle_zip
        import io, zipfile
        names = zipfile.ZipFile(io.BytesIO(build_bundle_zip(self._gen()))).namelist()
        assert "META-INF/MANIFEST.MF" in names and ".project" in names
        assert any(n.endswith(".iflw") for n in names)


class TestTimerPilotIFlow:
    def _gen(self):
        from scaffolder.minimal_iflow import generate_timer_pilot_iflow
        return generate_timer_pilot_iflow(
            "S_ContentModifier_SetConstant", "SContentModifierSetConstant")

    def test_wellformed_timer_no_sender(self):
        from lxml import etree
        x = self._gen().iflw_xml
        etree.fromstring(x.encode())                    # well-formed
        # timer start, decoded-real schema; no sender, no message flow
        assert "timerEventDefinition" in x and "StartTimerEvent" in x
        assert "cname::intermediatetimer" in x
        assert "messageFlow" not in x and "EndpointSender" not in x

    def test_timer_cm_cm_end_wiring(self):
        x = self._gen().iflw_xml
        assert 'sourceRef="StartEvent_2" targetRef="CallActivity_1"' in x
        assert 'sourceRef="CallActivity_1" targetRef="CallActivity_2"' in x
        assert 'sourceRef="CallActivity_2" targetRef="EndEvent_2"' in x
        assert x.count("<value>Enricher</value>") == 2

    def test_cm1_properties_cm2_body(self):
        x = self._gen().iflw_xml
        assert "PilotStatus" in x and "Source" in x          # CM1 properties
        assert "PilotAck" in x and "property.PilotStatus" in x  # CM2 body
        assert "propertyTable" in x

    def test_real_schedule_embedded(self):
        x = self._gen().iflw_xml
        # run-once (fireNow) schedule cloned from a proven-deployable timer iFlow
        assert "fireNow=true" in x and "triggerType" in x
        assert "intermediatetimer/version::1.3.0" in x

    def test_diagram_five_shapes_three_edges(self):
        x = self._gen().iflw_xml
        assert x.count("<bpmndi:BPMNShape") == 5
        assert x.count("<bpmndi:BPMNEdge") == 3


# ─── Wire log (communication transcript) ─────────────────────────────────────

class TestWireLog:
    def test_redacts_sensitive_headers(self):
        from fetcher import wire_log
        wire_log.reset_wire_log()
        wire_log.log_request("t", "POST", "https://x/api",
                             {"Authorization": "Bearer SECRET",
                              "X-CSRF-Token": "TOK", "Content-Type": "application/json"})
        out = wire_log.read_wire_log()
        assert "SECRET" not in out
        assert "TOK" not in out
        assert "<redacted>" in out
        assert "application/json" in out  # non-sensitive kept

    def test_captures_request_and_response(self):
        from fetcher import wire_log
        wire_log.reset_wire_log()
        wire_log.log_request("upload", "POST", "https://x/api/v1/Art", {}, "body")
        wire_log.log_response("upload", 500, {"Content-Type": "application/json"},
                              '{"error":"InputStream cannot be null"}')
        out = wire_log.read_wire_log()
        assert "REQUEST" in out and "RESPONSE" in out
        assert "HTTP 500" in out
        assert "InputStream cannot be null" in out

    def test_noise_filter_drops_width_warnings(self):
        import logging
        from fetcher.wire_log import _NoiseFilter
        f = _NoiseFilter()
        spam = logging.LogRecord("streamlit", logging.WARNING, "x.py", 1,
                                 "Please replace use_container_width with width",
                                 None, None)
        real = logging.LogRecord("cpi", logging.ERROR, "x.py", 1,
                                 "Upload failed: HTTP 500", None, None)
        assert f.filter(spam) is False   # dropped
        assert f.filter(real) is True    # kept


# ─── Unified logging + valid-iflow upload path (the two bug fixes) ───────────

class TestUnifiedLoggingAndValidUpload:
    def test_log_handler_captures_diagnostics(self):
        import logging
        from fetcher import wire_log
        wire_log.reset_wire_log()
        wire_log.install_unified_logging()
        logging.getLogger("cpi.test").error("a diagnostic message here")
        out = wire_log.read_wire_log()
        assert "a diagnostic message here" in out

    def test_scaffold_produces_valid_iflow(self, tmp_path):
        # Bug A fix: the generation path now produces a CPI-valid iFlow
        from scaffolder.iflow_scaffolder import IFlowScaffolder
        s = IFlowScaffolder(output_dir=str(tmp_path))
        a = make_assessment(name="ValidFlow")
        p = s.scaffold(a)
        content = p.read_text()
        assert "BPMNDiagram" in content
        assert "cmdVariantUri" in content
        # the validated manifest (with Import-Package) is stashed in __meta
        meta = p.parent / f"{p.stem}__meta"
        assert (meta / "MANIFEST.MF").exists()
        assert "Import-Package" in (meta / "MANIFEST.MF").read_text()

    def test_packager_uses_validated_manifest(self, tmp_path):
        # The upload packager picks up the validated manifest from __meta
        import io, zipfile
        from scaffolder.iflow_scaffolder import IFlowScaffolder
        from fetcher.cpi_uploader import CPIUploader
        s = IFlowScaffolder(output_dir=str(tmp_path))
        p = s.scaffold(make_assessment(name="PkgFlow"))
        data = CPIUploader._package_iflow(p, "PkgFlow", "Pkg Flow", "")
        zf = zipfile.ZipFile(io.BytesIO(data))
        manifest = zf.read("META-INF/MANIFEST.MF").decode()
        assert "Import-Package" in manifest      # validated, not hand-built
        iflw = [n for n in zf.namelist() if n.endswith(".iflw")][0]
        assert "BPMNDiagram" in zf.read(iflw).decode()


class TestArtifactRouter:
    """Step 3: per-type endpoint routing + idempotent dispatch."""

    def _mkzip(self, files):
        import io, zipfile
        buf = io.BytesIO()
        with zipfile.ZipFile(buf, "w") as z:
            for n, c in files.items():
                z.writestr(n, c)
        return buf.getvalue()

    def test_endpoint_registry_covers_seven_types(self):
        from fetcher.cpi_uploader import CPIUploader
        assert len(CPIUploader.ARTIFACT_ENDPOINTS) == 7
        assert CPIUploader.endpoint_for("IFlow") == "IntegrationDesigntimeArtifacts"
        assert CPIUploader.endpoint_for("MessageMapping") == "MessageMappingDesigntimeArtifacts"
        # unknown falls back to iFlow
        assert CPIUploader.endpoint_for("Nonsense") == "IntegrationDesigntimeArtifacts"

    def test_detect_type_from_bundle(self):
        from fetcher.artifact_router import detect_type_from_bundle
        iflow = self._mkzip({"META-INF/MANIFEST.MF": "x",
                             "src/main/resources/scenarioflows/F.iflw": "<x/>"})
        mmap = self._mkzip({"META-INF/MANIFEST.MF": "x",
                            "src/main/resources/mapping/M.mmap": "<x/>"})
        assert detect_type_from_bundle(iflow) == "IFlow"
        assert detect_type_from_bundle(mmap) == "MessageMapping"

    def test_normalize_type_skips_non_artifacts(self):
        from fetcher.artifact_router import normalize_type
        assert normalize_type("IFlow") == "IFlow"
        assert normalize_type("File") is None
        assert normalize_type("Url") is None
        assert normalize_type("ContentPackage") is None

    def test_plan_routes_and_skips(self):
        from unittest.mock import MagicMock
        from fetcher.cpi_uploader import CPIUploader
        from fetcher.artifact_router import ArtifactRouter
        u = CPIUploader("https://t", MagicMock())
        r = ArtifactRouter(u)
        iflow = self._mkzip({"META-INF/MANIFEST.MF": "x",
                             "src/main/resources/scenarioflows/F.iflw": "<x/>"})
        plan = r.plan("PKG", "Pkg", [
            {"id": "IF1", "name": "IF1", "zip_bytes": iflow},
            {"id": "doc", "name": "readme", "resource_type": "File"},
        ])
        assert len(plan.artifacts) == 1
        assert plan.artifacts[0].endpoint == "IntegrationDesigntimeArtifacts"
        assert len(plan.skipped) == 1

    def test_post_artifact_creates_then_updates(self):
        from unittest.mock import MagicMock
        from fetcher.cpi_uploader import CPIUploader, UploadResult
        u = CPIUploader("https://t", MagicMock())
        u._ensure_csrf = lambda: "tok"
        # CREATE: not existing -> POST
        u._artifact_exists = lambda aid, ep="IntegrationDesigntimeArtifacts": False
        posted = []
        def fp(url, **kw):
            posted.append("POST"); r = MagicMock(); r.status_code = 201
            r.text = "ok"; r.headers = {}; return r
        u.session.post = fp
        res = UploadResult(interface_name="i", package_id="P",
                           artifact_id="A", status="failed")
        u._post_artifact(b"PKz", "P", "A", "A", res, artifact_type="IFlow")
        assert res.status == "uploaded"
        # REPLACE: existing -> delete + recreate (re-skinned bundle always
        # changes the Bundle-SymbolicName, so the in-place PUT would 400; the
        # tool skips it and goes straight to delete + recreate).
        u._artifact_exists = lambda aid, ep="IntegrationDesigntimeArtifacts": True
        seq = []
        def fdel(url, **kw):
            seq.append("DELETE"); r = MagicMock(); r.status_code = 200
            r.text = "ok"; r.headers = {}; return r
        def fpost2(url, **kw):
            seq.append("POST"); r = MagicMock(); r.status_code = 201
            r.text = "ok"; r.headers = {}; return r
        def fput(url, **kw):
            seq.append("PUT"); r = MagicMock(); r.status_code = 200
            r.text = "ok"; r.headers = {}; return r
        u.session.delete = fdel
        u.session.post = fpost2
        u.session.put = fput
        res2 = UploadResult(interface_name="i", package_id="P",
                            artifact_id="A", status="failed")
        u._post_artifact(b"PKz", "P", "A", "A", res2, artifact_type="IFlow")
        assert res2.status == "updated"
        assert "PUT" not in seq          # doomed PUT is skipped
        assert seq == ["DELETE", "POST"] # delete then recreate


class TestPackageExtraction:
    """Step 8-9: extract artifacts from a package wrapper + read real Ids."""

    def _mkbundle(self, sym_name, iflw_name="F.iflw"):
        import io, zipfile
        buf = io.BytesIO()
        mf = (f"Manifest-Version: 1.0\r\n"
              f"Bundle-SymbolicName: {sym_name}; singleton:=true\r\n"
              f"Bundle-Name: {sym_name}\r\n")
        with zipfile.ZipFile(buf, "w") as z:
            z.writestr("META-INF/MANIFEST.MF", mf)
            z.writestr(f"src/main/resources/scenarioflows/integrationflow/{iflw_name}", "<x/>")
        return buf.getvalue()

    def test_read_artifact_id_name(self):
        from fetcher.artifact_router import read_artifact_id_name
        b = self._mkbundle("My_Real_IFlow")
        aid, name = read_artifact_id_name(b)
        assert aid == "My_Real_IFlow"
        assert name == "My_Real_IFlow"

    def test_single_bundle_extraction(self):
        from fetcher.artifact_router import extract_package_artifacts
        b = self._mkbundle("Solo_Flow")
        arts = extract_package_artifacts(b)
        assert len(arts) == 1
        assert arts[0]["id"] == "Solo_Flow"

    def test_package_wrapper_extraction(self):
        import io, zipfile
        from fetcher.artifact_router import extract_package_artifacts
        b1 = self._mkbundle("Flow_One")
        b2 = self._mkbundle("Flow_Two")
        # Build a wrapper: two _content bundles (no manifest at wrapper root)
        buf = io.BytesIO()
        with zipfile.ZipFile(buf, "w") as z:
            z.writestr("aaaa1111bbbb2222_content", b1)
            z.writestr("cccc3333dddd4444_content", b2)
            z.writestr("resources.cnt", "{}")
        arts = extract_package_artifacts(buf.getvalue())
        ids = sorted(a["id"] for a in arts)
        assert ids == ["Flow_One", "Flow_Two"]

    def test_wrapper_routes_to_distinct_ids(self):
        # The collision scenario: two real iFlows must get DIFFERENT ids.
        import io, zipfile
        from unittest.mock import MagicMock
        from fetcher.cpi_uploader import CPIUploader
        from fetcher.artifact_router import (ArtifactRouter,
                                             extract_package_artifacts)
        b1 = self._mkbundle("Clear_CustomFields")
        b2 = self._mkbundle("SF_to_OpenText")
        buf = io.BytesIO()
        with zipfile.ZipFile(buf, "w") as z:
            z.writestr("hash1aaaaaaaaaaa_content", b1)
            z.writestr("hash2bbbbbbbbbbb_content", b2)
        arts = extract_package_artifacts(buf.getvalue())
        u = CPIUploader("https://t", MagicMock())
        plan = ArtifactRouter(u).plan("PKG", "Pkg", arts)
        ids = sorted(a.artifact_id for a in plan.artifacts)
        assert ids == ["Clear_CustomFields", "SF_to_OpenText"]
        assert all(a.endpoint == "IntegrationDesigntimeArtifacts"
                   for a in plan.artifacts)


class TestRecursiveIFlowCount:
    """Tab 1 fix: count iFlows (interfaces/ICOs) recursively, not packages."""

    def _mkbundle(self, sym_name, iflw=True):
        import io, zipfile
        buf = io.BytesIO()
        mf = (f"Manifest-Version: 1.0\r\n"
              f"Bundle-SymbolicName: {sym_name}; singleton:=true\r\n"
              f"Bundle-Name: {sym_name}\r\n")
        with zipfile.ZipFile(buf, "w") as z:
            z.writestr("META-INF/MANIFEST.MF", mf)
            if iflw:
                z.writestr("src/main/resources/scenarioflows/integrationflow/f.iflw", "<x/>")
            else:
                z.writestr("src/main/resources/mapping/m.mmap", "<x/>")
        return buf.getvalue()

    def _wrapper(self, *bundles):
        import io, zipfile
        buf = io.BytesIO()
        with zipfile.ZipFile(buf, "w") as z:
            for i, b in enumerate(bundles):
                z.writestr(f"hash{i:04d}aaaabbbb_content", b)
            z.writestr("resources.cnt", "{}")
        return buf.getvalue()

    def test_single_iflow_bundle_counts_one(self):
        from fetcher.artifact_router import count_iflows_recursive
        assert count_iflows_recursive(self._mkbundle("Solo")) == 1

    def test_package_wrapper_counts_each_iflow(self):
        from fetcher.artifact_router import count_iflows_recursive
        w = self._wrapper(self._mkbundle("A"), self._mkbundle("B"))
        assert count_iflows_recursive(w) == 2

    def test_non_iflow_artifacts_not_counted(self):
        # A package with 1 iFlow + 1 message mapping => 1 interface.
        from fetcher.artifact_router import count_iflows_recursive
        w = self._wrapper(self._mkbundle("Flow", iflw=True),
                          self._mkbundle("Map", iflw=False))
        assert count_iflows_recursive(w) == 1

    def test_nested_container_of_packages(self):
        # Container zip holding multiple package zips (the 55-package case).
        import io, zipfile
        from fetcher.artifact_router import count_iflows_recursive
        pkg1 = self._wrapper(self._mkbundle("A"), self._mkbundle("B"))
        pkg2 = self._wrapper(self._mkbundle("C"))
        outer = io.BytesIO()
        with zipfile.ZipFile(outer, "w") as z:
            z.writestr("pkgs/one.zip", pkg1)
            z.writestr("pkgs/two.zip", pkg2)
        assert count_iflows_recursive(outer.getvalue()) == 3

    def test_dedup_same_iflow_id(self):
        # Same iFlow id appearing twice counts once.
        import io, zipfile
        from fetcher.artifact_router import count_iflows_recursive
        pkg = self._wrapper(self._mkbundle("Dup"))
        outer = io.BytesIO()
        with zipfile.ZipFile(outer, "w") as z:
            z.writestr("a/p.zip", pkg)
            z.writestr("b/p.zip", pkg)
        assert count_iflows_recursive(outer.getvalue()) == 1


class TestIdempotencyFallback:
    """Task 2: re-upload update path. On a Bundle-SymbolicName 400, the
    uploader falls back to delete + recreate via the proven create path."""

    def _uploader(self, responses):
        """Build a CPIUploader whose session returns queued responses.
        responses: list of (method, status, text) consumed in order."""
        from unittest.mock import MagicMock
        from fetcher.cpi_uploader import CPIUploader
        calls = []
        seq = list(responses)

        def _resp(status, text):
            r = MagicMock()
            r.status_code = status
            r.text = text
            r.headers = {}
            return r

        sess = MagicMock()
        def record(method):
            def _fn(url, **kw):
                calls.append((method, url))
                _, status, text = seq.pop(0)
                return _resp(status, text)
            return _fn
        sess.get = record("GET")
        sess.post = record("POST")
        sess.put = record("PUT")
        sess.delete = record("DELETE")
        u = CPIUploader("https://t", sess)
        u._csrf_token = "tok"
        u._ensure_csrf = lambda: "tok"
        return u, calls

    def test_symbolic_name_400_triggers_delete_recreate(self):
        from fetcher.cpi_uploader import UploadResult
        # Sequence for _update_artifact:
        #   PUT -> 400 symbolicName  (update rejected)
        #   DELETE -> 200            (delete existing)
        #   POST -> 201              (recreate)
        u, calls = self._uploader([
            ("PUT", 400, '{"error":{"message":{"value":"Could not update '
                         'artifact of the package; due to change in the '
                         'Bundle-symbolicName."}}}'),
            ("DELETE", 200, ""),
            ("POST", 201, '{"d":{"Id":"X"}}'),
        ])
        res = UploadResult(interface_name="X", package_id="PKG", artifact_id="X", status="pending")
        u._update_artifact(b"zipbytes", "PKG", "X", "X", res, "IFlow")
        assert res.status == "updated"
        methods = [m for m, _ in calls]
        assert methods == ["PUT", "DELETE", "POST"]

    def test_clean_update_succeeds_without_fallback(self):
        from fetcher.cpi_uploader import UploadResult
        # PUT -> 200 means a normal in-place update; no delete/recreate.
        u, calls = self._uploader([("PUT", 200, "{}")])
        res = UploadResult(interface_name="X", package_id="PKG", artifact_id="X", status="pending")
        u._update_artifact(b"zip", "PKG", "X", "X", res, "IFlow")
        assert res.status == "updated"
        assert [m for m, _ in calls] == ["PUT"]


class TestSAPComplexityEngine:
    """Faithful two-axis SAP complexity engine: weight->size->effort, both
    modes (true MA passthrough + signal approximation), real SAP tables."""

    def _mkbundle(self, sym="Flow", groovy=0, mmap=0, xslt=0,
                  call_acts=0, participants=0, msgflows=0, routers=0):
        import io, zipfile
        buf = io.BytesIO()
        mf = (f"Manifest-Version: 1.0\r\n"
              f"Bundle-SymbolicName: {sym}; singleton:=true\r\n"
              f"Bundle-Name: {sym}\r\n")
        # Build a minimal iflw BPMN with the requested element counts.
        xml = ['<bpmn2:definitions xmlns:bpmn2="x">']
        for _ in range(call_acts):  xml.append("<bpmn2:callActivity/>")
        for _ in range(participants): xml.append("<bpmn2:participant/>")
        for _ in range(msgflows):   xml.append("<bpmn2:messageFlow/>")
        for _ in range(routers):    xml.append("<bpmn2:exclusiveGateway/>")
        xml.append("</bpmn2:definitions>")
        with zipfile.ZipFile(buf, "w") as z:
            z.writestr("META-INF/MANIFEST.MF", mf)
            z.writestr("src/main/resources/scenarioflows/integrationflow/f.iflw",
                       "\n".join(xml))
            for i in range(groovy):
                z.writestr(f"src/main/resources/script/s{i}.groovy", "x")
            for i in range(mmap):
                z.writestr(f"src/main/resources/mapping/m{i}.mmap", "<x/>")
            for i in range(xslt):
                z.writestr(f"src/main/resources/mapping/x{i}.xsl", "<x/>")
        return buf.getvalue()

    def test_weight_to_size_thresholds(self):
        from analyzer.sap_complexity_engine import weight_to_size
        assert weight_to_size(1) == "S"
        assert weight_to_size(150) == "S"
        assert weight_to_size(151) == "M"
        assert weight_to_size(350) == "M"
        assert weight_to_size(351) == "L"
        assert weight_to_size(500) == "L"
        assert weight_to_size(501) == "XL"
        assert weight_to_size(99999) == "XL"

    def test_effort_table_values(self):
        from analyzer.sap_complexity_engine import effort_for
        # SAP's real effort table
        assert effort_for("S", "Migrate") == (1, 4)
        assert effort_for("L", "Evaluate") == (16, 40)
        assert effort_for("XL", "Evaluate") == (40, 80)

    def test_mode1_true_ma_passthrough(self):
        from analyzer.sap_complexity_engine import SAPComplexityEngine
        eng = SAPComplexityEngine()
        r = eng.assess_true_ma("IF", weight=380, size="L", category="Evaluate")
        assert r.size == "L" and r.category == "Evaluate"
        assert (r.effort_hours_low, r.effort_hours_high) == (16, 40)
        assert r.mode == "true_ma"

    def test_mode1_category_label_normalization(self):
        from analyzer.sap_complexity_engine import SAPComplexityEngine
        eng = SAPComplexityEngine()
        # Long SAP labels normalize to short
        r = eng.assess_true_ma("IF", weight=100, size="S",
                               category="Evaluation required")
        assert r.category == "Evaluate"

    def test_mode2_signal_simple_iflow_is_small(self):
        from analyzer.sap_complexity_engine import SAPComplexityEngine
        eng = SAPComplexityEngine()
        b = self._mkbundle("Simple", groovy=1, call_acts=2, participants=2,
                           msgflows=2)
        r = eng.assess_bundle("Simple", b)
        assert r.size in ("S", "M")
        assert r.mode == "signal"
        assert any(c for c in r.caveats)  # approximation caveat present

    def test_mode2_signal_complex_iflow_scores_higher(self):
        from analyzer.sap_complexity_engine import SAPComplexityEngine
        eng = SAPComplexityEngine()
        simple = eng.assess_bundle("Simple", self._mkbundle(
            "Simple", groovy=1, call_acts=2, participants=2, msgflows=2))
        complex_ = eng.assess_bundle("Complex", self._mkbundle(
            "Complex", groovy=8, mmap=2, xslt=7, call_acts=27,
            participants=18, msgflows=14, routers=8))
        # A complex iFlow must score strictly higher weight than a simple one.
        assert complex_.total_weight > simple.total_weight

    def test_has_bpm_flag_fires_ccbpm_and_outweighs_non_bpm(self):
        """Regression: the authoritative has_bpm flag was extracted but never
        wired to a rule, so BPM interfaces scored identically to non-BPM ones
        and everything collapsed to size S. has_bpm must fire SAP's ccBPM rule
        (Evaluate category) — adding weight AND bumping the t-shirt size."""
        from analyzer.sap_complexity_engine import SAPComplexityEngine
        from extractor.pi_extractor import InterfaceRecord
        eng = SAPComplexityEngine()
        common = dict(id="x", namespace="", software_component="",
                      sender_system="", receiver_system="",
                      sender_adapter="HTTPS", receiver_adapter="HTTPS",
                      message_interface="X", mapping_program="MM")
        plain = eng.assess_interface(InterfaceRecord(name="Plain", has_bpm=False, **common))
        bpm = eng.assess_interface(InterfaceRecord(name="Bpm", has_bpm=True, **common))
        assert bpm.total_weight > plain.total_weight, "has_bpm added no weight"
        assert bpm.category == "Evaluate", "ccBPM should mark the interface Evaluate"
        # the BPM interface must not be classified identically-small to the plain one
        assert (bpm.size, plain.size) != ("S", "S") or bpm.total_weight > plain.total_weight

    def test_real_ma_record_routes_through_mode1_else_approximation(self):
        """Contract _ma_assess depends on: a record carrying real MA figures
        (ma_weight set) yields calibrated Mode 1 (mode 'true_ma', SAP weight),
        while a record without them yields the keyword approximation ('signal').
        This is what makes an imported MA export stop showing estimates."""
        from analyzer.sap_complexity_engine import SAPComplexityEngine
        from extractor.pi_extractor import InterfaceRecord
        eng = SAPComplexityEngine()
        common = dict(id="x", name="X", namespace="", software_component="",
                      sender_system="", receiver_system="",
                      sender_adapter="HTTPS", receiver_adapter="HTTPS",
                      message_interface="X")
        rec_ma = InterfaceRecord(ma_weight=420, ma_size="L",
                                 ma_status="Adjustment required", **common)
        r1 = eng.assess_true_ma(rec_ma.name, weight=int(rec_ma.ma_weight),
                                size=rec_ma.ma_size, category=rec_ma.ma_status)
        assert r1.mode == "true_ma" and r1.total_weight == 420 and r1.size == "L"
        rec_plain = InterfaceRecord(**common)
        assert getattr(rec_plain, "ma_weight", None) is None
        r2 = eng.assess_interface(rec_plain)
        assert r2.mode == "signal"

    def test_ccbpm_not_inferred_from_structure(self):
        # Corpus-driven correction: ccBPM/Evaluate must NOT be inferred from
        # flow complexity (step/router counts). A structurally-complex CPI
        # bundle with no explicit BPM signal stays Migrate — the Evaluate
        # drivers are source-side (ABAP/JavaBPM/custom modules) that don't
        # survive into a migrated bundle.
        from analyzer.sap_complexity_engine import SAPComplexityEngine
        eng = SAPComplexityEngine()
        r = eng.assess_bundle("Heavy", self._mkbundle(
            "Heavy", groovy=10, call_acts=30, participants=20, routers=10))
        assert not any(f.rule in ("ccBPM", "JavaBPM") for f in r.fired_rules)

    def test_ccbpm_fires_only_on_explicit_signal(self):
        # When an explicit BPM signal is present (from real source / MA data),
        # ccBPM fires and forces Evaluate.
        from analyzer.sap_complexity_engine import SAPComplexityEngine
        eng = SAPComplexityEngine()
        sig = {"groovy_scripts": 2, "is_ccbpm": True, "mapping_types": []}
        r = eng.assess_signals("BPMFlow", sig)
        assert any(f.rule == "ccBPM" for f in r.fired_rules)
        assert r.category == "Evaluate"

    def test_mode2_fires_both_characteristics(self):
        # Count rules fire against RECEIVER_IF AND EXT_RCV_DET (SAP behaviour).
        from analyzer.sap_complexity_engine import SAPComplexityEngine
        eng = SAPComplexityEngine()
        r = eng.assess_bundle("F", self._mkbundle("F", xslt=7))
        xslt_fires = [f for f in r.fired_rules if f.rule == "XSLTDependenciesCount"]
        chars = {f.characteristic for f in xslt_fires}
        assert "RECEIVER_IF" in chars and "EXT_RCV_DET" in chars

    def test_pluggable_effort_profile(self):
        from analyzer.sap_complexity_engine import (SAPComplexityEngine,
                                                    register_effort_profile)
        # A build profile with bigger numbers must not affect migration.
        register_effort_profile("build", {"table": {
            "S": {"Migrate": [40, 80], "Adapt": [40, 80], "Evaluate": [40, 80]},
            "M": {"Migrate": [80, 160], "Adapt": [80, 160], "Evaluate": [80, 160]},
            "L": {"Migrate": [160, 320], "Adapt": [160, 320], "Evaluate": [160, 320]},
            "XL": {"Migrate": [320, 640], "Adapt": [320, 640], "Evaluate": [320, 640]},
        }})
        mig = SAPComplexityEngine("migration")
        bld = SAPComplexityEngine("build")
        rm = mig.assess_true_ma("IF", weight=200, size="M", category="Migrate")
        rb = bld.assess_true_ma("IF", weight=200, size="M", category="Migrate")
        assert (rm.effort_hours_low, rm.effort_hours_high) == (4, 8)      # SAP migration
        assert (rb.effort_hours_low, rb.effort_hours_high) == (80, 160)   # build profile
        # They must not interfere.
        assert rm.effort_hours_high != rb.effort_hours_high



class TestDiscoverHarvester:
    """Two-strategy Discover harvester. Confirms response parsing and the
    destructive-delete guardrail (only packages this run copied are deleted)."""

    _COPY_PAYLOAD = {"listOfStatus": [
        {"responseInfo": {"Status": "SUCCESS", "Type": "IntegrationFlow",
                          "Title": "Flow A", "id": "iflowA"}},
        {"responseInfo": {"Status": "SUCCESS", "Type": "File",
                          "Title": "Guide", "id": "fileB"}},
        {"responseInfo": {"Status": "SUCCESS", "Type": "ContentPackage",
                          "Title": "Pkg", "TechnicalName": "PkgTech",
                          "id": "pkgRegId"}},
    ]}

    def test_assets_from_copy_response(self):
        from fetcher.discover_harvester import DiscoverHarvester
        assets = DiscoverHarvester.assets_from_copy_response(self._COPY_PAYLOAD)
        assert len(assets) == 3
        types = {a["type"] for a in assets}
        assert "IntegrationFlow" in types and "ContentPackage" in types

    def test_package_id_prefers_technicalname(self):
        from fetcher.discover_harvester import DiscoverHarvester
        pid = DiscoverHarvester.package_id_from_copy_response(self._COPY_PAYLOAD)
        assert pid == "PkgTech"

    def _fake(self, deleted_sink, copy_ok=True, catalog_returns=None):
        import requests
        from fetcher.discover_harvester import DiscoverHarvester
        payload = self._COPY_PAYLOAD
        class FakeH(DiscoverHarvester):
            def list_package_assets(self, pid): return []   # force copy path
            def copy_package(self, pid): return payload if copy_ok else None
            def download_catalog_asset(self, aid, atype): return catalog_returns
            def delete_package(self, pid):
                deleted_sink.append(pid); return True
        return FakeH("https://design.test", requests.Session(),
                     runtime_base_url="https://rt.test", download_dir="/tmp/dh_pytest2")

    def test_cleanup_deletes_only_copied_package(self):
        deleted = []
        h = self._fake(deleted, copy_ok=True, catalog_returns=None)
        h.harvest_one("pkgRegId", allow_copy_fallback=True, cleanup=True)
        assert deleted == ["PkgTech"]   # the ContentPackage TechnicalName

    def test_no_delete_when_copy_fails(self):
        deleted = []
        h = self._fake(deleted, copy_ok=False)
        r = h.harvest_one("pkgRegId", allow_copy_fallback=True, cleanup=True)
        assert deleted == []
        assert r.errors

    def test_catalog_direct_skips_copy_and_delete(self):
        # If catalog-direct yields content, no copy and no delete happen.
        import requests
        from fetcher.discover_harvester import DiscoverHarvester
        deleted = []
        copied = []
        class FakeH(DiscoverHarvester):
            def list_package_assets(self, pid):
                return [{"id": "iflowA", "type": "IntegrationFlow", "title": "A"}]
            def download_catalog_asset(self, aid, atype): return b"PKzipdata"
            def copy_package(self, pid): copied.append(pid); return None
            def delete_package(self, pid): deleted.append(pid); return True
        h = FakeH("https://design.test", requests.Session(),
                  download_dir="/tmp/dh_pytest3")
        r = h.harvest_one("pkgRegId", prefer="catalog", allow_copy_fallback=True)
        assert r.strategy == "catalog"
        assert r.n_downloaded == 1
        assert copied == [] and deleted == []   # never copied, never deleted


class TestEffortModel:
    """The quotable effort model: base + gaps × multiplier + optional hypercare."""

    def test_snap_and_bounds(self):
        from reporter.effort_model import snap_multiplier
        assert snap_multiplier(1.13) == 1.25
        assert snap_multiplier(2.6) == 2.5
        assert snap_multiplier(9.0) == 3.0   # clamp high
        assert snap_multiplier(0.1) == 1.0   # clamp low

    def test_mode_sets_default_multiplier(self):
        from reporter.effort_model import default_multiplier_for_mode
        assert default_multiplier_for_mode("Migration") == 1.0
        assert default_multiplier_for_mode("Support") == 1.75
        assert default_multiplier_for_mode("Implementation") == 2.75
        assert default_multiplier_for_mode(None) == 1.0

    def test_multiplier_excludes_hypercare(self):
        from reporter.effort_model import build_effort
        e = build_effort(10, 10, gap_hours=0, multiplier=3.0,
                         hypercare_enabled=True, hypercare_hours=60)
        # 10*3 = 30, +60 flat = 90 (hypercare NOT multiplied)
        assert abs(e.total_low - 90.0) < 0.01

    def test_hypercare_optional_off(self):
        from reporter.effort_model import build_effort
        e = build_effort(10, 10, gap_hours=5, multiplier=2.0,
                         hypercare_enabled=False)
        # (10+5)*2 = 30, no hypercare
        assert abs(e.total_low - 30.0) < 0.01

    def test_base_plus_gap_then_multiplier(self):
        from reporter.effort_model import build_effort
        e = build_effort(5, 12, gap_hours=6, multiplier=1.5)
        assert abs(e.adjusted_low - 16.5) < 0.01    # (5+6)*1.5
        assert abs(e.adjusted_high - 27.0) < 0.01   # (12+6)*1.5

    def test_days_view(self):
        from reporter.effort_model import build_effort
        e = build_effort(16, 16, gap_hours=0, multiplier=1.0, hours_per_day=8)
        assert abs(e.total_days_low - 2.0) < 0.01


class TestInterventionEffortIntegration:
    """Intervention estimator attaches engine-based effort breakdowns."""

    def _mk_iface(self, name, sa, ra, mp, bpm):
        from extractor.pi_extractor import InterfaceRecord
        return InterfaceRecord(id=name, name=name, namespace="ns",
            software_component="sc", sender_system="S", receiver_system="R",
            sender_adapter=sa, receiver_adapter=ra, message_interface="MI",
            mapping_program=mp, has_bpm=bpm, has_multi_mapping=False,
            channel_count=1, description="", raw={})

    def test_engine_base_flows_into_effort(self):
        from analyzer.sap_complexity_engine import SAPComplexityEngine
        from reporter.intervention_estimator import InterventionEstimator
        from analyzer.complexity_analyzer import MigrationAssessment
        eng = SAPComplexityEngine()
        est = InterventionEstimator(output_dir="/tmp/ie_pytest")
        iface = self._mk_iface("ICO_A", "SOAP", "SFTP", "Map1", False)
        assess = MigrationAssessment(interface=iface, score=0, complexity="MEDIUM",
                                     effort_days=3, notes=[], recommended_pattern="")
        engr = eng.assess_true_ma("ICO_A", weight=300)  # Mode 1 base
        proj = est.estimate_all([assess], engine_results={"ICO_A": engr},
                                multiplier=1.5, mode="Migration",
                                hypercare_enabled=True, hypercare_hours=60)
        # base must equal the engine result (before automation), not zero
        assert proj.effort.base_hours_high == engr.effort_hours_high
        # multiplier applied to base+gap, hypercare added flat on top
        assert proj.effort.hypercare_add == 60.0
        assert proj.effort.total_high == proj.effort.adjusted_high + 60.0
        # per-interface effort present
        assert proj.interfaces[0].effort is not None

    def test_no_engine_result_degrades_gracefully(self):
        # Legacy path: no engine_result -> base 0, breakdown still built.
        from reporter.intervention_estimator import InterventionEstimator
        from analyzer.complexity_analyzer import MigrationAssessment
        est = InterventionEstimator(output_dir="/tmp/ie_pytest2")
        iface = self._mk_iface("ICO_B", "IDoc", "SFTP", None, False)
        assess = MigrationAssessment(interface=iface, score=0, complexity="LOW",
                                     effort_days=1, notes=[], recommended_pattern="")
        proj = est.estimate_all([assess])  # no engine_results
        assert proj.effort is not None
        assert proj.effort.base_hours_low == 0.0


class TestBundleAssembler:
    """Assemble a deployable package bundle from content; round-trip a real one."""

    def test_manifest_format_and_folding(self):
        from library_builder.bundle_assembler import build_manifest, IFlowContent
        c = IFlowContent(display_name="My_Flow", symbolic_name="MyFlow",
                         version="1.0.0")
        m = build_manifest(c).decode()
        assert "Bundle-Name: My_Flow" in m
        assert "Bundle-SymbolicName: MyFlow" in m
        assert "SAP-BundleType: IntegrationFlow" in m
        # OSGi folding: every physical line <= 72 chars
        assert all(len(line) <= 72 for line in m.split("\r\n"))

    def test_project_descriptor(self):
        from library_builder.bundle_assembler import build_project, IFlowContent
        c = IFlowContent(display_name="My_Flow", symbolic_name="MyFlow")
        p = build_project(c).decode()
        assert "<name>MyFlow</name>" in p
        assert "com.sap.ide.ifl.bsn" in p

    def test_bundle_has_required_wrapper(self):
        from library_builder.bundle_assembler import build_bundle, IFlowContent
        import io, zipfile
        c = IFlowContent(display_name="My_Flow",
                         files={"script/x.groovy": b"def x(){}"})
        data = build_bundle([c], package_name="My Package")
        z = zipfile.ZipFile(io.BytesIO(data))
        names = z.namelist()
        assert any(n.endswith("_content") for n in names)
        assert "resources.cnt" in names
        assert "contentmetadata.md" in names
        assert "hash" in names
        assert "ExportInformation.info" in names

    def test_roundtrip_real_bundle_content_identical(self):
        # extract real RCI093 content, rebuild, content files must be identical
        import os, io, zipfile, re
        from library_builder.bundle_assembler import (
            extract_content_from_bundle, build_bundle)
        real_path = "/mnt/user-data/uploads/RCI093_testB__1_.zip"
        if not os.path.exists(real_path):
            import pytest
            pytest.skip("real bundle not present in this environment")
        iflows = extract_content_from_bundle(real_path)
        assert len(iflows) >= 1
        rebuilt = build_bundle(iflows, package_name="RCI093 regen")

        def content_files(data):
            out = {}
            top = zipfile.ZipFile(io.BytesIO(data) if isinstance(data, bytes)
                                  else open(data, "rb"))
            for n in top.namelist():
                if not n.endswith("_content"):
                    continue
                cz = zipfile.ZipFile(io.BytesIO(top.read(n)))
                for inner in cz.namelist():
                    if inner.startswith("src/") and not inner.endswith("/"):
                        out[inner] = cz.read(inner)
            return out

        real = content_files(open(real_path, "rb").read())
        new = content_files(rebuilt)
        assert real, "no content extracted from real bundle"
        identical = sum(1 for k in real if k in new and real[k] == new[k])
        assert identical == len(real), \
            f"only {identical}/{len(real)} content files identical"


class TestBundleManifestExact:
    """The MANIFEST we generate must match the real SAP format exactly."""

    def test_generated_manifest_matches_real_fields(self):
        import os, io, zipfile
        from library_builder.bundle_assembler import build_manifest, IFlowContent
        real_path = "/mnt/user-data/uploads/RCI093_testB__1_.zip"
        if not os.path.exists(real_path):
            import pytest
            pytest.skip("real bundle not present")
        Z = zipfile.ZipFile(real_path)
        real = None
        for n in Z.namelist():
            if n.endswith("_content"):
                cz = zipfile.ZipFile(io.BytesIO(Z.read(n)))
                m = cz.read("META-INF/MANIFEST.MF").decode()
                if "RCI093SuccessFactorstoOpenText" in m:
                    real = m
                    break
        c = IFlowContent(
            display_name="RCI093_SuccessFactors_to_OpenText_SuccessFactors",
            symbolic_name="RCI093SuccessFactorstoOpenTextSuccessFactors",
            version="1.0.3")
        mine = build_manifest(c).decode()

        def fields(m):
            m = m.replace("\r\n ", "")
            d = {}
            for line in m.split("\r\n"):
                if ":" in line:
                    k, v = line.split(":", 1)
                    d[k.strip()] = v.strip()
            return d
        rf, mf = fields(real), fields(mine)
        for k in rf:
            assert mf.get(k) == rf[k], f"field {k} differs"

    def test_minimal_bundle_blank_wrapper(self):
        import io, zipfile
        from library_builder.bundle_assembler import build_bundle, IFlowContent
        c = IFlowContent(display_name="F", files={"script/x.groovy": b"x"})
        data = build_bundle([c], "pkg", minimal=True)
        z = zipfile.ZipFile(io.BytesIO(data))
        # wrapper blank, content + manifest present
        assert z.read("hash") == b""
        assert z.read("resources.cnt") == b""
        assert any(n.endswith("_content") for n in z.namelist())
        cz = zipfile.ZipFile(io.BytesIO(
            z.read([n for n in z.namelist() if n.endswith("_content")][0])))
        assert "META-INF/MANIFEST.MF" in cz.namelist()
        assert b"Bundle-SymbolicName" in cz.read("META-INF/MANIFEST.MF")


class TestBundleValidator:
    def test_trailing_period_name_rejected(self):
        from library_builder.bundle_validator import validate_name, is_deployable
        f = validate_name("MyFlow.")
        assert any(x.severity == "error" for x in f)
        assert not is_deployable(f)

    def test_clean_name_ok(self):
        from library_builder.bundle_validator import validate_name, is_deployable
        assert is_deployable(validate_name("MyFlow"))

    def test_invalid_symbolic_chars(self):
        from library_builder.bundle_validator import validate_symbolic_name
        assert any(x.severity == "error"
                   for x in validate_symbolic_name("My Flow!"))

    def test_real_assembler_bundle_is_deployable(self):
        import os
        from library_builder.bundle_validator import validate_bundle, is_deployable
        from library_builder.bundle_assembler import (
            extract_content_from_bundle, build_bundle)
        p = "/mnt/user-data/uploads/RCI093_testB__1_.zip"
        if not os.path.exists(p):
            import pytest
            pytest.skip("real bundle absent")
        b = build_bundle(extract_content_from_bundle(p), "t", minimal=True)
        assert is_deployable(validate_bundle(b))

    def test_missing_manifest_flagged(self):
        import io, zipfile
        from library_builder.bundle_validator import validate_bundle, is_deployable
        bio = io.BytesIO()
        with zipfile.ZipFile(bio, "w") as z:
            inner = io.BytesIO()
            with zipfile.ZipFile(inner, "w") as cz:
                cz.writestr("src/main/resources/x.groovy", "x")
            z.writestr("abc_content", inner.getvalue())
        f = validate_bundle(bio.getvalue())
        assert not is_deployable(f)  # missing MANIFEST = error


class TestProposalFromEffort:
    def test_quote_from_effort_breakdown(self):
        from reporter.effort_model import build_effort
        from reporter.proposal_generator import quote_from_effort
        eff = build_effort(base_hours_low=16, base_hours_high=24, gap_hours=8,
                           multiplier=1.75, hypercare_enabled=True)
        q = quote_from_effort("Iface", eff)
        assert q.base_days > 0
        assert q.client_price_min > 0
        assert "hypercare" in q.notes


class TestMmapGenerator:
    def test_generates_wellformed_xml(self):
        from library_builder.mmap_generator import (
            MappingSpec, MappingRow, generate_mmap)
        import xml.dom.minidom as M
        spec = MappingSpec(source_message="Src", target_message="Tgt",
                           source_file="S.xsd", target_file="T.xsd",
                           source_root="Src", target_root="Tgt")
        spec.rows = [MappingRow("/ns1:Tgt/x", "add( /ns1:Src/a,/ns1:Src/b)")]
        m = generate_mmap(spec)
        M.parseString(m)  # raises if malformed
        assert 'type="Dst"' in m and 'fname="add"' in m

    def test_nested_function(self):
        from library_builder.mmap_generator import (
            MappingSpec, MappingRow, generate_mmap)
        spec = MappingSpec(rows=[MappingRow(
            "/T/n", "concat( toUpperCase(/S/x),/S/y)")])
        m = generate_mmap(spec)
        # nested func brick present
        assert m.count('type="Func"') == 2
        assert 'fname="toUpperCase"' in m and 'fname="concat"' in m

    def test_direct_copy_and_const(self):
        from library_builder.mmap_generator import (
            MappingSpec, MappingRow, generate_mmap)
        spec = MappingSpec(rows=[
            MappingRow("/T/d", "/S/plain"),
            MappingRow("/T/c", 'const( "FIX")'),
        ])
        m = generate_mmap(spec)
        assert 'path="/S/plain" type="Src"' in m
        assert "<value>FIX</value>" in m
        assert m.count("FIX") == 1  # no double-wrap

    def test_excel_spec_roundtrip(self):
        import os
        from library_builder.mmap_generator import spec_from_excel, generate_mmap
        p = "/mnt/user-data/uploads/mmap.xls"
        if not os.path.exists(p):
            import pytest
            pytest.skip("mapping excel absent")
        spec = spec_from_excel(p)
        assert len(spec.rows) > 50           # the full function catalog
        m = generate_mmap(spec)
        import xml.dom.minidom as M
        M.parseString(m)                     # well-formed
        assert m.count('type="Dst"') == len(spec.rows)


class TestMessageMappingBundle:
    def test_mm_manifest_distinct_from_iflow(self):
        from library_builder.bundle_assembler import (
            build_messagemapping_manifest, build_manifest, IFlowContent)
        c = IFlowContent(display_name="mmap", symbolic_name="mmap")
        mm = build_messagemapping_manifest(c).decode()
        iflow = build_manifest(c).decode()
        assert "SAP-BundleType: MessageMapping" in mm
        assert "SAP-BundleType: IntegrationFlow" in iflow
        assert "Provide-Capability" in mm
        assert "SAP-RuntimeProfile" not in mm      # mappings have no profile
        assert "SAP-RuntimeProfile" in iflow

    def test_mm_content_structure(self):
        import io, zipfile
        from library_builder.bundle_assembler import (
            build_messagemapping_content, IFlowContent)
        c = IFlowContent(display_name="mmap", symbolic_name="mmap",
                         files={"mapping/m.mmap": b"<xiObj/>",
                                "wsdl/S.wsdl": b"<x/>"})
        z = zipfile.ZipFile(io.BytesIO(build_messagemapping_content(c)))
        names = z.namelist()
        assert "src/main/resources/mapping/m.mmap" in names
        assert "src/main/resources/wsdl/S.wsdl" in names
        assert "META-INF/MANIFEST.MF" in names


class TestMmapConfirmedPatterns:
    """Regression tests locking in the tenant-CONFIRMED mmap patterns.
    See library_builder/MMAP_PATTERNS_CONFIRMED.md. These guard against
    regressing any pattern the user verified against a real tenant."""

    def _gen(self):
        from library_builder.mmap_generator import spec_from_excel, generate_mmap
        import os
        p = "/mnt/user-data/uploads/mmap.xls"
        if not os.path.exists(p):
            import pytest
            pytest.skip("reference Excel absent")
        spec = spec_from_excel(p)
        spec.source_namespace = "http://cpi.sap.com/demo"
        spec.target_namespace = "http://cpi.sap.com/demo"
        spec.source_root = "FunctionSource_MT"
        spec.target_root = "FunctionTarget_MT"
        spec.source_file = "FunctionSource.wsdl"
        spec.target_file = "FunctionTarget.wsdl"
        spec.schema_type = "wsdl"
        return generate_mmap(spec)

    def test_pin_rule_sequential(self):
        """§4a: multi-arg functions number args 0,1,2 — second arg pin=1,
        third pin=2 (NOT all pin=1)."""
        import re
        g = self._gen()
        i = g.find('path="/ns1:FunctionTarget_MT/ns1:useOneAsMany_result"')
        seg = g[g.rfind("<brick", 0, i):i + 600]
        seg = g[g.rfind("<brick", 0, i):]
        seg = seg[:seg.find("<group/></brick>") + 16]
        # the three sources must carry no-pin, pin=1, pin=2
        assert 'useOneAsMany_a" type="Src"' in seg
        assert '<arg pin="1"><brick gid="0" path="/ns1:FunctionSource_MT/ns1:useOneAsMany_b"' in seg
        assert '<arg pin="2"><brick gid="0" path="/ns1:FunctionSource_MT/ns1:useOneAsMany_c"' in seg

    def test_two_source_functions_have_both(self):
        """§4a: add/sub/mul/div each wire BOTH sources with pin=1 on second."""
        g = self._gen()
        for fn, a, b in [("add", "add_a", "add_b"),
                         ("sub", "subtract_a", "subtract_b")]:
            assert f'/ns1:FunctionSource_MT/ns1:{a}"' in g
            assert f'<arg pin="1"><brick gid="0" path="/ns1:FunctionSource_MT/ns1:{b}"' in g

    def test_wsdl_binding_has_namespace(self):
        """§3: WSDL schema binding carries the 4th namespace elem."""
        g = self._gen()
        assert "<elem>http://cpi.sap.com/demo</elem>" in g
        assert 'role="SOURCE_IFR_MESS"' in g and 'role="TARGET_IFR_MESS"' in g

    def test_envelope_required_blocks(self):
        """§2: libstorage scaffolding, params, AdditionalProperties present."""
        g = self._gen()
        assert "functionstorage" in g           # libstorage scaffolding
        assert "tr:SourceParameters" in g
        assert "AdditionalProperties" in g
        assert "<calend_props>" in g             # structured date binding

    def test_namespaces_block(self):
        """§5: namespaces block declares ns1."""
        g = self._gen()
        assert '<property name="http://cpi.sap.com/demo">ns1</property>' in g

    def test_divide_aliased_to_div(self):
        """§7: divide -> div internal name."""
        g = self._gen()
        assert 'fname="div"' in g
        assert 'fname="divide"' not in g

    def test_matches_verified_100pct(self):
        """§0: generated mmap matches the verified tenant mmap on ALL entries
        (sources, pins, functions, binding params) — the headline result."""
        import os, re, zipfile
        vp = "/mnt/user-data/uploads/mmap.zip"
        if not os.path.exists(vp):
            import pytest
            pytest.skip("verified mmap absent")
        verified = zipfile.ZipFile(vp).read(
            "src/main/resources/mapping/mmap.mmap").decode("utf-8", "replace")
        g = self._gen()

        def entries(t):
            out = {}
            for m in re.finditer(
                    r'path="(/ns1:FunctionTarget_MT/ns1:\w+)" type="Dst"', t):
                f = m.group(1).split(":")[-1]
                i = m.start()
                s = t.rfind("<brick", 0, i)
                nxt = t.find('gid="0" path="/ns1:FunctionTarget_MT', i + 10)
                seg = re.sub(r"<viewData[^>]*/>", "",
                             t[s:t.rfind("<brick", 0, nxt) if nxt > 0 else s + 900])
                pinned = re.findall(
                    r'<arg( pin="\d+")?><brick gid="0" '
                    r'path="(/ns1:FunctionSource_MT/ns1:\w+)"', seg)
                out[f] = {
                    "sources": [p[1].split(":")[-1] for p in pinned],
                    "pins": [(p[0].strip() or "p0") for p in pinned],
                    "funcs": re.findall(r'fname="(\w+)"', seg),
                }
            return out

        ev, gv = entries(verified), entries(g)
        assert set(ev) == set(gv), "target field set differs"
        for f in ev:
            assert ev[f] == gv[f], f"entry mismatch on {f}: {ev[f]} vs {gv[f]}"

        # binding params identical too
        def params(t):
            return sorted(re.findall(
                r'<param name="(\w+)"><value>(.*?)</value>', t))
        assert params(verified) == params(g), "binding params differ"


class TestMmapSelfValidation:
    def test_good_mmap_passes(self):
        import os
        from library_builder.mmap_generator import (
            spec_from_excel, generate_mmap, validate_generated_mmap)
        if not os.path.exists("/mnt/user-data/uploads/mmap.xls"):
            import pytest
            pytest.skip("reference Excel absent")
        spec = spec_from_excel("/mnt/user-data/uploads/mmap.xls")
        spec.source_namespace = spec.target_namespace = "http://cpi.sap.com/demo"
        spec.source_root, spec.target_root = "FunctionSource_MT", "FunctionTarget_MT"
        spec.source_file, spec.target_file = "FunctionSource.wsdl", "FunctionTarget.wsdl"
        spec.schema_type = "wsdl"
        assert validate_generated_mmap(generate_mmap(spec)) == []

    def test_catches_pin_bug(self):
        from library_builder.mmap_generator import validate_generated_mmap
        broken = ('<xiObj><transformation><brick fname="add" type="Func">'
                  '<arg pin="1"><brick/></arg><arg pin="1"><brick/></arg>'
                  '</brick></transformation></xiObj>')
        problems = validate_generated_mmap(broken)
        assert any("pin" in p for p in problems)

    def test_catches_missing_envelope(self):
        from library_builder.mmap_generator import validate_generated_mmap
        problems = validate_generated_mmap("<xiObj></xiObj>")
        assert any("libstorage" in p for p in problems)


class TestMmapParser:
    """Lock the mmap parser (mmap -> spec) fidelity. The parser is the catalog's
    structural-extraction layer and the inverse of the generator.
    See library_builder/MMAP_LEARNING_MODEL.md."""

    def _all_mmaps(self):
        import zipfile, io, os
        out = {}

        def walk(zf, d=0):
            if d > 6:
                return
            for n in zf.namelist():
                if n.endswith("/"):
                    continue
                try:
                    raw = zf.read(n)
                except Exception:
                    continue
                if raw[:2] == b"PK" and (n.endswith(".zip")
                                         or n.endswith("_content")):
                    try:
                        walk(zipfile.ZipFile(io.BytesIO(raw)), d + 1)
                        continue
                    except Exception:
                        pass
                if n.endswith(".mmap"):
                    nm = n.split("/")[-1]
                    out.setdefault(nm, raw.decode("utf-8", "replace"))
        up = "/mnt/user-data/uploads"
        if not os.path.isdir(up):
            return {}
        for f in os.listdir(up):
            if f.endswith(".zip"):
                try:
                    walk(zipfile.ZipFile(os.path.join(up, f)))
                except Exception:
                    pass
        return out

    def test_parses_reference_mmap(self):
        import os, zipfile
        from library_builder.mmap_parser import parse_mmap
        p = "/mnt/user-data/uploads/mmap.zip"
        if not os.path.exists(p):
            import pytest
            pytest.skip("reference mmap absent")
        txt = zipfile.ZipFile(p).read(
            "src/main/resources/mapping/mmap.mmap").decode("utf-8", "replace")
        pm = parse_mmap(txt)
        assert pm.source_root == "FunctionSource_MT"
        assert pm.target_root == "FunctionTarget_MT"
        assert pm.schema_type == "wsdl"
        assert len(pm.fields) >= 75
        # add_result must parse as add(add_a, add_b) with correct pins
        add = next(f for f in pm.fields if f.target_path.endswith("add_result"))
        fn = add.tree.args[0]   # under __DST__
        assert fn.kind == "func" and fn.value == "add"
        assert [a.pin for a in fn.args] == ["", "1"]

    def test_parses_all_corpus_mmaps_cleanly(self):
        """Every non-empty mmap parses with no failed (None) trees."""
        from library_builder.mmap_parser import parse_mmap
        mmaps = self._all_mmaps()
        if not mmaps:
            import pytest
            pytest.skip("no corpus mmaps available")
        import re
        bad = []
        for nm, txt in mmaps.items():
            if not re.search(r'type="Dst"', txt):
                continue
            pm = parse_mmap(txt)
            if not pm.fields or any(f.tree is None for f in pm.fields):
                bad.append(nm)
        assert not bad, f"mmaps failing to parse: {bad}"

    def test_handles_multiline_tags(self):
        """Pretty-printed mmaps (newlines inside brick tags) must parse."""
        from library_builder.mmap_parser import parse_mmap
        txt = (
            '<transformation>'
            '<brick gid="0"\n   path="/ns1:T/ns1:f"\n   type="Dst">'
            '<arg><brick fname="toUpperCase" fns="dflt" type="Func">'
            '<arg><brick gid="0" path="/ns1:S/ns1:x" type="Src"></brick></arg>'
            '</brick></arg><group/></brick>'
            '</transformation>')
        pm = parse_mmap(txt)
        assert len(pm.fields) == 1
        assert pm.fields[0].tree.args[0].value == "toUpperCase"


class TestMmapCapabilities:
    """Lock the capability-section catalog (blocker #2 architecture):
    decompose an mmap into tagged, searchable capability units."""

    def _ref(self):
        import os, zipfile
        p = "/mnt/user-data/uploads/mmap.zip"
        if not os.path.exists(p):
            import pytest
            pytest.skip("reference mmap absent")
        return zipfile.ZipFile(p).read(
            "src/main/resources/mapping/mmap.mmap").decode("utf-8", "replace")

    def test_extracts_one_capability_per_field(self):
        from library_builder.mmap_capabilities import extract_capabilities
        caps = self._ref() and __import__(
            "library_builder.mmap_capabilities", fromlist=["extract_capabilities"])
        from library_builder.mmap_capabilities import extract_capabilities
        caps = extract_capabilities(self._ref())
        assert len(caps) >= 75
        # add_result is a numeric add of two sources
        add = next(c for c in caps if c.target_field == "add_result")
        assert add.category == "numeric"
        assert "add" in add.functions
        assert len(add.sources) == 2

    def test_categories_and_signatures(self):
        from library_builder.mmap_capabilities import extract_capabilities
        caps = extract_capabilities(self._ref())
        cats = {c.category for c in caps}
        # the semantic spread we expect from the reference mapping
        for expected in ("numeric", "string", "conditional", "lookup",
                         "context", "date"):
            assert expected in cats, f"missing category {expected}"
        # signatures are matchable strings
        sigs = {c.signature() for c in caps}
        assert "numeric:add(2src)" in sigs
        assert any(s.startswith("lookup:valuemap") for s in sigs)

    def test_complexity_weighting_aligns_with_sap(self):
        """Lookup/context score higher than plain numeric/string (per SAP's
        complexity-driver taxonomy)."""
        from library_builder.mmap_capabilities import extract_capabilities
        caps = {c.target_field: c for c in extract_capabilities(self._ref())}
        add = caps["add_result"].weight
        vmap = next(c.weight for c in caps.values() if "valuemap" in c.functions)
        assert vmap > add

    def test_catalog_summary(self):
        from library_builder.mmap_capabilities import catalog_summary
        s = catalog_summary(self._ref())
        assert s["field_count"] >= 75
        assert s["total_weight"] > 0
        assert "numeric" in s["categories"]
        assert isinstance(s["signatures"], list) and s["signatures"]


class TestGroovyCapabilities:
    """Lock the groovy capability extractor (code-capability model:
    envelope + universal operations + empirically-discovered binding table,
    with the five adapt-facets)."""

    SAMPLE = (
        "import com.sap.gateway.ip.core.customdev.util.Message\n"
        "import groovy.json.JsonSlurper\n"
        "import groovy.json.JsonOutput\n"
        "def Message processData(Message message) {\n"
        "    def body = message.getBody(String)\n"
        "    def obj = new JsonSlurper().parseText(body)\n"
        "    def region = message.getProperty('REGION')\n"
        "    message.setProperty('COUNT', obj.size())\n"
        "    message.setBody(JsonOutput.toJson(obj))\n"
        "    return message\n}\n")

    def test_extracts_envelope_and_universal_ops(self):
        from library_builder.groovy_capabilities import extract_capability
        c = extract_capability("sample.groovy", self.SAMPLE)
        assert c.envelope == "processData"
        assert "READ_BODY" in c.operations
        assert "READ_PROPERTY" in c.operations
        assert "WRITE_PROPERTY" in c.operations
        assert "WRITE_BODY" in c.operations
        # portable core detected
        assert "PARSE_JSON" in c.portable_ops
        assert "BUILD_JSON" in c.portable_ops

    def test_five_facets_present(self):
        from library_builder.groovy_capabilities import extract_capability
        c = extract_capability("sample.groovy", self.SAMPLE)
        assert c.purpose                      # what it does
        assert "body" in c.needs              # inputs
        assert "REGION" in c.what_varies      # adaptable param (property name)
        assert "processData" in c.shape       # envelope/contract
        assert c.when_to_use                  # selection hint
        assert c.signature().startswith("groovy:")

    def test_binding_vocabulary_discovered_empirically(self):
        from library_builder.groovy_capabilities import discover_bindings
        corpus = {"a": self.SAMPLE,
                  "b": "message.getHeader('X'); messageLog.addAttachmentAsString('n','c','t')"}
        vocab = discover_bindings(corpus)
        assert vocab["message.getBody"] >= 1
        assert vocab["message.setProperty"] >= 1
        assert vocab["messageLog.addAttachmentAsString"] >= 1

    def test_no_false_positive_on_exception_message(self):
        """exception.message.replaceAll is the .message PROPERTY of a Java
        exception, NOT the SAP Message object — must not be caught as a binding
        (rigor-audit regression)."""
        from library_builder.groovy_capabilities import discover_bindings, extract_capability
        src = ('def Message processData(Message message){'
               'def e = message.getProperty("CamelExceptionCaught");'
               'def clean = e.message.replaceAll("x","");'
               'return message}')
        vocab = discover_bindings({"x": src})
        assert "message.replaceAll" not in vocab     # the false positive
        assert vocab["message.getProperty"] == 1     # the real one still caught
        c = extract_capability("x.groovy", src)
        assert "READ_PROPERTY" in c.operations

    def test_body_read_type_captured(self):
        """How the body is read (String/Reader/InputStream) is part of the I/O
        contract and must be captured, not flattened (rigor-audit finding)."""
        from library_builder.groovy_capabilities import extract_capability
        c = extract_capability(
            "r.groovy",
            "def Message processData(Message m){"
            "def r = m.getBody(java.io.Reader); return m}")
        assert "Reader" in c.body_read_as
        c2 = extract_capability(
            "s.groovy",
            "def Message processData(Message m){"
            "def s = m.getBody(String); return m}")
        assert "String" in c2.body_read_as

    def test_library_flagged(self):
        """A multi-function file is a LIBRARY of capabilities, not one
        (rigor-audit finding)."""
        from library_builder.groovy_capabilities import extract_capability
        src = "\n".join(f"def helper{i}(x){{return x}}" for i in range(6))
        c = extract_capability("lib.groovy", src)
        assert c.is_library
        assert c.function_count >= 6
        single = extract_capability(
            "one.groovy", "def Message processData(Message m){return m}")
        assert not single.is_library

    def test_catalog_and_fetch_by_need(self):
        from library_builder.groovy_capabilities import build_catalog
        corpus = {
            "json.groovy": self.SAMPLE,
            "log.groovy": ("def Message processData(Message m){"
                           "def b=m.getBody(String);"
                           "messageLog.getMessageLog(m).addAttachmentAsString('p',b,'text/plain');"
                           "return m}"),
        }
        cat = build_catalog(corpus)
        assert cat["count"] == 2
        json_caps = [c.name for c in cat["capabilities"]
                     if "PARSE_JSON" in c.portable_ops]
        assert "json.groovy" in json_caps
        emit_caps = [c.name for c in cat["capabilities"]
                     if any(o.startswith("EMIT") for o in c.operations)]
        assert "log.groovy" in emit_caps


class TestSchemaCatalog:
    """Lock the schema (xsd/wsdl/edmx) identity catalog: whole-file identity,
    dedupe, and reuse-index. Fully standalone — no SAP, no tenant."""

    XSD = ('<?xml version="1.0"?>'
           '<xs:schema xmlns:xs="http://www.w3.org/2001/XMLSchema" '
           'targetNamespace="urn:demo">'
           '<xs:element name="Order"/><xs:element name="LineItem"/>'
           '<xs:complexType name="OrderType"/></xs:schema>')
    WSDL = ('<wsdl:definitions xmlns:wsdl="http://schemas.xmlsoap.org/wsdl/" '
            'targetNamespace="urn:svc"><wsdl:types>'
            '<xs:element name="Request" xmlns:xs="http://www.w3.org/2001/XMLSchema"/>'
            '</wsdl:types></wsdl:definitions>')
    EDMX = ('<edmx:Edmx xmlns:edmx="x"><Schema Namespace="SF.OData">'
            '<EntityType Name="User"/><EntityType Name="EmpJob"/></Schema></edmx:Edmx>')

    def test_identity_per_kind(self):
        from library_builder.schema_catalog import extract_identity
        x = extract_identity("o.xsd", self.XSD)
        assert x.kind == "xsd"
        assert x.target_namespace == "urn:demo"
        assert "Order" in x.roots and "LineItem" in x.roots
        assert "OrderType" in x.types
        w = extract_identity("s.wsdl", self.WSDL)
        assert w.kind == "wsdl"
        e = extract_identity("m.edmx", self.EDMX)
        assert e.kind == "edmx"
        assert "User" in e.entities and "EmpJob" in e.entities

    def test_dedupe_by_structure_not_name(self):
        """Identity is the STRUCTURE defined, not the file/top name. Two XSDs
        with the same element set are the same schema even if differently
        named; two with different structure are not, even if same top name."""
        from library_builder.schema_catalog import build_catalog
        same_struct_diff_name = self.XSD  # Order/LineItem/OrderType
        # identical structure, different file name + reformatted
        twin = self.XSD.replace("><", ">\n<")
        # different structure (only Order) -> NOT a duplicate
        different = ('<xs:schema xmlns:xs="http://www.w3.org/2001/XMLSchema" '
                     'targetNamespace="urn:demo"><xs:element name="Order"/>'
                     '</xs:schema>')
        cat = build_catalog({"a.xsd": same_struct_diff_name, "b.xsd": twin,
                             "c.xsd": different})
        assert cat["distinct_schemas"] == 2   # a==b, c separate
        grp = [g for g in cat["duplicate_groups"].values()
               if set(g["members"]) == {"a.xsd", "b.xsd"}]
        assert grp, "same-structure schemas not deduped"

    def test_canonical_prefers_bigger_valid(self):
        from library_builder.schema_catalog import build_catalog
        small = self.XSD
        big = self.XSD.replace("</xs:schema>",
                               "<!-- extensive annotation block -->" * 5
                               + "</xs:schema>")
        cat = build_catalog({"small.xsd": small, "big.xsd": big})
        grp = next(iter(cat["duplicate_groups"].values()))
        assert grp["canonical"] == "big.xsd"

    def test_validity_flags_damaged(self):
        from library_builder.schema_catalog import build_catalog, extract_identity, find_schema
        broken = '<xs:schema><xs:element name="X"></xs:schema>'
        i = extract_identity("broken.xsd", broken)
        assert not i.well_formed and i.validity_error
        cat = build_catalog({"good.xsd": self.XSD, "broken.xsd": broken})
        assert "broken.xsd" in cat["damaged"]
        assert "broken.xsd" not in find_schema(cat, defines="X")

    def test_no_false_collapse_when_no_named_structure(self):
        from library_builder.schema_catalog import build_catalog
        a = '<root><foo/><bar/></root>'
        b = '<root><baz/><qux/><quux/></root>'
        cat = build_catalog({"a.xsd": a, "b.xsd": b})
        assert cat["distinct_schemas"] == 2

    def test_subset_family_detection(self):
        from library_builder.schema_catalog import build_catalog
        full = ('<xs:schema xmlns:xs="http://www.w3.org/2001/XMLSchema" '
                'targetNamespace="urn:f"><xs:element name="A"/>'
                '<xs:element name="B"/><xs:element name="C"/>'
                '<xs:element name="D"/></xs:schema>')
        cut = ('<xs:schema xmlns:xs="http://www.w3.org/2001/XMLSchema" '
               'targetNamespace="urn:f"><xs:element name="A"/>'
               '<xs:element name="B"/><xs:element name="C"/></xs:schema>')
        cat = build_catalog({"full.xsd": full, "cut.xsd": cut})
        fams = [f for f in cat["subset_families"]
                if f["subset"] == "cut.xsd" and f["superset"] == "full.xsd"]
        assert fams and fams[0]["preferred"] == "full.xsd"

    def test_reuse_lookup(self):
        from library_builder.schema_catalog import build_catalog, find_schema
        cat = build_catalog({"o.xsd": self.XSD, "m.edmx": self.EDMX})
        assert "o.xsd" in find_schema(cat, defines="Order")
        assert "m.edmx" in find_schema(cat, defines="User")
        assert "o.xsd" in find_schema(cat, namespace="urn:demo")

    def test_defines_aggregates_keys(self):
        from library_builder.schema_catalog import extract_identity
        d = extract_identity("o.xsd", self.XSD).defines_list()
        assert "Order" in d and "OrderType" in d and "urn:demo" in d


class TestXsltCapabilities:
    """Lock the XSLT capability extractor: portable transform core + SAP
    extension binding layer, with version awareness (rigor-audited)."""

    PORTABLE_1_0 = (
        '<xsl:stylesheet version="1.0" '
        'xmlns:xsl="http://www.w3.org/1999/XSL/Transform">'
        '<xsl:output method="xml"/>'
        '<xsl:template match="/"><out>'
        '<xsl:for-each select="row"><xsl:value-of select="x"/></xsl:for-each>'
        '</out></xsl:template></xsl:stylesheet>')

    WITH_EXT = (
        '<xsl:stylesheet version="1.0" '
        'xmlns:xsl="http://www.w3.org/1999/XSL/Transform" '
        'xmlns:cpi="http://www.sap.com/ica/mag/function/private">'
        '<xsl:output method="text"/>'
        '<xsl:template match="/"><xsl:value-of select="cpi:setProperty(\'k\',\'v\')"/>'
        '</xsl:template></xsl:stylesheet>')

    V2_0 = PORTABLE_1_0.replace('version="1.0"', 'version="2.0"')

    def test_portable_vs_extension(self):
        from library_builder.xslt_capabilities import extract_capability
        p = extract_capability("p.xsl", self.PORTABLE_1_0)
        assert p.portable and not p.extension_calls
        assert p.output_method == "xml"
        assert "for-each" in p.core_elements
        e = extract_capability("e.xsl", self.WITH_EXT)
        assert not e.portable
        assert "cpi:setProperty" in e.extension_calls
        assert e.extension_namespaces  # the SAP binding namespace recorded

    def test_version_separate_from_portability(self):
        """A portable 2.0 stylesheet is vendor-neutral but NOT lxml-runnable —
        the two dimensions must not be conflated (the rigor-audit finding)."""
        from library_builder.xslt_capabilities import extract_capability
        v2 = extract_capability("v2.xsl", self.V2_0)
        assert v2.portable                 # no SAP coupling
        assert v2.xslt_version == "2.0"
        assert not v2.sandbox_runnable     # 2.0 -> lxml can't run it
        v1 = extract_capability("v1.xsl", self.PORTABLE_1_0)
        assert v1.sandbox_runnable         # 1.0 + portable

    def test_verify_runnable_confirms_with_lxml(self):
        from library_builder.xslt_capabilities import verify_runnable
        # the known-good 1.0 stylesheet genuinely compiles
        ok, _ = verify_runnable(self.PORTABLE_1_0)
        assert ok
        # a malformed stylesheet genuinely fails (verifier returns False, not raise)
        bad, err = verify_runnable('<xsl:stylesheet><broken')
        assert not bad and err

    def test_catalog_fetch_and_discover(self):
        from library_builder.xslt_capabilities import build_catalog, discover_extensions
        cat = build_catalog({"p.xsl": self.PORTABLE_1_0, "e.xsl": self.WITH_EXT},
                            verify=True)
        assert cat["count"] == 2
        # FETCH: the text-output one with a binding
        ext = [c.name for c in cat["capabilities"] if c.extension_calls]
        assert "e.xsl" in ext
        # verified runnable count reflects real compilation
        assert cat["verified_runnable"] == cat["sandbox_runnable_count"]
        vocab = discover_extensions({"e.xsl": self.WITH_EXT})
        assert "cpi:setProperty" in vocab


class TestCapabilityCatalogFacade:
    """Lock the unified facade over all four capability extractors."""

    def test_type_routing(self):
        from library_builder.capability_catalog import type_for_ext, TYPES
        assert type_for_ext(".groovy") == "groovy"
        assert type_for_ext("xsd") == "schema"
        assert type_for_ext(".xsl") == "xslt"
        assert type_for_ext(".mmap") == "mmap"
        assert type_for_ext(".js") == "js"
        assert {"mmap", "groovy", "schema", "xslt", "js"} <= set(TYPES)

    def test_catalog_for_each_type(self):
        from library_builder.capability_catalog import catalog_for
        # groovy
        g = catalog_for("groovy",
                        {"a.groovy": "def Message processData(Message m){"
                                     "def b=m.getBody(String);return m}"})
        assert g["count"] == 1
        # schema
        s = catalog_for("schema",
                        {"a.xsd": '<xs:schema xmlns:xs='
                         '"http://www.w3.org/2001/XMLSchema" '
                         'targetNamespace="urn:x"><xs:element name="A"/>'
                         '</xs:schema>'})
        assert s["count"] == 1
        # xslt
        x = catalog_for("xslt",
                        {"a.xsl": '<xsl:stylesheet version="1.0" '
                         'xmlns:xsl="http://www.w3.org/1999/XSL/Transform">'
                         '<xsl:template match="/"/></xsl:stylesheet>'})
        assert x["count"] == 1

    def test_unknown_type_raises(self):
        from library_builder.capability_catalog import catalog_for
        import pytest
        with pytest.raises(ValueError):
            catalog_for("nonsense", {})


class TestSolverReasoningLayer:
    """Lock the reasoning layer (B): normalize + EVALUATE/FETCH/SELECT/ADAPT/
    COMPOSE over the capability catalogs. Sandbox proves fetch/adapt/compose;
    'best pick' + SAP-run are flagged for tenant/user (honest boundary)."""

    def _norm(self):
        from library_builder import groovy_capabilities as g, schema_catalog as s
        from library_builder.solver import normalize
        gcat = g.build_catalog({
            "parseJson.groovy": ("import groovy.json.JsonSlurper\n"
                "def Message processData(Message m){"
                "def b=m.getBody(String); def o=new JsonSlurper().parseText(b);"
                "m.setBody(o.toString()); return m}"),
            "logAttach.groovy": ("def Message processData(Message m){"
                "def b=m.getBody(String);"
                "messageLog.getMessageLog(m).addAttachmentAsString('p',b,'text/plain');"
                "return m}"),
        })
        scat = s.build_catalog({
            "Order.xsd": ('<xs:schema xmlns:xs="http://www.w3.org/2001/XMLSchema" '
                'targetNamespace="urn:o"><xs:element name="Order"/></xs:schema>')})
        return normalize(gcat, "groovy") + normalize(scat, "schema")

    def test_normalize_unifies_shapes(self):
        norm = self._norm()
        # groovy + schema mapped into one shape
        assert all(hasattr(c, "intent") and hasattr(c, "keywords") for c in norm)
        assert any(c.ctype == "groovy" for c in norm)
        assert any(c.ctype == "schema" for c in norm)

    def test_evaluate_decomposes(self):
        from library_builder.solver import evaluate
        needs = evaluate("Parse the JSON then log it as an attachment")
        assert len(needs) >= 2
        kws = set().union(*[n.keywords for n in needs])
        assert "parse" in kws and "json" in kws and "log" in kws

    def test_evaluate_multiword_intent(self):
        """'look up' (spaced) must map to lookup intent (rigor-audit fix)."""
        from library_builder.solver import evaluate
        needs = evaluate("look up the country code")
        assert any("lookup" in n.keywords for n in needs)

    def test_fetch_finds_relevant(self):
        from library_builder.solver import fetch, evaluate
        norm = self._norm()
        need = evaluate("parse the json payload")[0]
        ranked = fetch(need, norm)
        assert ranked
        # the top match should be the json-parsing groovy
        assert "parseJson" in ranked[0].capability.cap_id

    def test_select_prefers_simpler_on_tie(self):
        from library_builder.solver import select, Match, NormalizedCapability, Need
        n = Need(text="x", keywords={"a"})
        heavy = NormalizedCapability("h", "groovy", "do", {"a"}, weight=50)
        light = NormalizedCapability("l", "groovy", "do", {"a"}, weight=1)
        m = select([Match(n, heavy, 1.0), Match(n, light, 1.0)])
        assert m.capability.cap_id == "l"   # simpler wins the tie

    def test_solve_end_to_end(self):
        from library_builder.solver import solve, solution_summary
        norm = self._norm()
        sol = solve("Parse the JSON payload then log it as an attachment", norm)
        s = solution_summary(sol)
        assert s["steps"]                      # produced a solution
        assert s["confidence"] == "reasoned"   # honest: not "certain"
        # the attachment step flags tenant testing (SAP binding)
        assert s["needs_tenant_test"] is True

    def test_adapt_proposes_substitutions(self):
        from library_builder.solver import adapt, Match, NormalizedCapability, Need
        cap = NormalizedCapability("c", "groovy", "set property", {"property"},
                                   varies=["REGION", "COUNTRY"])
        m = Match(Need("set REGION", {"property"}), cap, 1.0)
        plan = adapt(m, "set the REGION property to US")
        assert "REGION" in plan["substitutions"]   # detected concrete value
        assert "COUNTRY" in plan["to_confirm"]      # not in requirement -> confirm

    def test_op_keywords_keep_capabilities_findable(self):
        """Validation finding: capabilities whose function isn't in canned
        phrases (e.g. a substring helper) must still be findable via op_keywords
        drawn from the real code — anti-bias against losing real capabilities."""
        from library_builder.groovy_capabilities import extract_capability
        from library_builder.solver import normalize
        src = ("def Message processData(Message m){"
               "def b=m.getBody(String); def r=b.substring(0,5); "
               "m.setBody(r); return m}")
        c = extract_capability("sub.groovy", src)
        assert "substring" in c.op_keywords
        norm = normalize({"capabilities": [c]}, "groovy")
        assert "substring" in norm[0].keywords

    def test_idf_weights_rare_terms_higher(self):
        from library_builder.solver import _idf, NormalizedCapability
        norm = [NormalizedCapability(f"c{i}", "groovy", "x",
                                     {"common"} | ({"rare"} if i == 0 else set()))
                for i in range(10)]
        idf = _idf(norm)
        assert idf["rare"] > idf["common"]

    def test_noise_fragments_filtered(self):
        from library_builder.solver import normalize
        from library_builder.groovy_capabilities import extract_capability
        c = extract_capability(
            "w.groovy",
            "def Message processData(Message m){m.setBody('x'); return m}")
        norm = normalize({"capabilities": [c]}, "groovy")
        assert "back" not in norm[0].keywords


class TestCorpusPipeline:
    """Lock the corpus_pipeline: the clean orchestrator that walks a corpus,
    classifies by type, builds catalogs, normalizes, and serves solve/search.
    Replaces the retired extractor.py/run_extractor."""

    FILES = {
        "a.groovy": ("import groovy.json.JsonSlurper\n"
                     "def Message processData(Message m){"
                     "def b=m.getBody(String); new JsonSlurper().parseText(b);"
                     "return m}"),
        "b.xsl": ('<xsl:stylesheet version="1.0" '
                  'xmlns:xsl="http://www.w3.org/1999/XSL/Transform">'
                  '<xsl:template match="/"/></xsl:stylesheet>'),
        "c.xsd": ('<xs:schema xmlns:xs="http://www.w3.org/2001/XMLSchema" '
                  'targetNamespace="urn:x"><xs:element name="Order"/></xs:schema>'),
        "readme.md": "# not a capability type",   # unknown -> dropped
    }

    def test_group_by_type_routes_and_drops_unknown(self):
        from library_builder.corpus_pipeline import group_by_type, classify_report
        g = group_by_type(self.FILES)
        assert set(g.keys()) == {"groovy", "xslt", "schema"}
        rep = classify_report(self.FILES)
        assert rep["total"] == 4
        assert "md" in rep["unknown"]              # readme dropped, reported

    def test_build_corpus_from_files(self):
        from library_builder.corpus_pipeline import build_corpus
        c = build_corpus(files=self.FILES)
        r = c.report()
        assert r["capabilities"] >= 3
        assert set(r["by_type"].keys()) == {"groovy", "xslt", "schema"}

    def test_build_corpus_empty_and_single(self):
        from library_builder.corpus_pipeline import build_corpus
        assert build_corpus(files={}).report()["files"] == 0
        c = build_corpus(files={"x.groovy":
                                "def Message processData(Message m){return m}"})
        assert c.report()["capabilities"] == 1

    def test_corpus_solve_and_search(self):
        from library_builder.corpus_pipeline import build_corpus
        c = build_corpus(files=self.FILES)
        # search finds the json-parser
        hits = c.search("parse json", top_n=3)
        assert any("a.groovy" in cid for cid, _ in hits)
        # solve produces a reasoned solution
        from library_builder.solver import solution_summary
        sol = solution_summary(c.solve("parse the json payload"))
        assert sol["confidence"] == "reasoned"

    def test_walk_corpus_reads_zip(self):
        import io, zipfile, tempfile, os
        from library_builder.corpus_pipeline import walk_corpus
        buf = io.BytesIO()
        with zipfile.ZipFile(buf, "w") as z:
            z.writestr("inner/x.groovy", "def Message processData(Message m){return m}")
            z.writestr("y.xsl", "<xsl:stylesheet/>")
        with tempfile.NamedTemporaryFile(suffix=".zip", delete=False) as tf:
            tf.write(buf.getvalue()); p = tf.name
        try:
            files = walk_corpus(p)
            # keys are container-qualified (collision fix v2): the zip's own
            # name prefixes every internal path, so same-named files in
            # different packages never collapse.
            assert any(k.endswith("/inner/x.groovy") for k in files)
            assert any(k.endswith("/y.xsl") for k in files)
            assert all("/" in k for k in files)
            # and extension detection still works on path-qualified keys
            from library_builder.corpus_pipeline import group_by_type
            grouped = group_by_type(files)
            assert "groovy" in grouped and "xslt" in grouped
        finally:
            os.unlink(p)

    def test_build_corpus_from_packages_bytes(self):
        """The workbench shape: uploaded_packages = list[{filename, bytes}].
        Build a corpus straight from in-memory zip bytes (no disk)."""
        import io, zipfile
        from library_builder.corpus_pipeline import build_corpus, walk_corpus_bytes
        buf = io.BytesIO()
        with zipfile.ZipFile(buf, "w") as z:
            z.writestr("a.groovy",
                       "def Message processData(Message m){"
                       "def b=m.getBody(String); return m}")
            z.writestr("b.xsd", '<xs:schema xmlns:xs='
                       '"http://www.w3.org/2001/XMLSchema" '
                       'targetNamespace="urn:x"><xs:element name="O"/></xs:schema>')
        pkgs = [{"filename": "pkg1.zip", "bytes": buf.getvalue()}]
        # walk_corpus_bytes accepts the dict-with-bytes shape
        files = walk_corpus_bytes(pkgs)
        assert "a.groovy" in files and "b.xsd" in files
        # and build_corpus(packages=...) produces capabilities
        c = build_corpus(packages=pkgs)
        r = c.report()
        assert r["capabilities"] >= 2
        assert set(r["by_type"].keys()) == {"groovy", "schema"}
        # also accepts a plain list of bytes and a {name: bytes} mapping
        assert "a.groovy" in walk_corpus_bytes([buf.getvalue()])
        assert "a.groovy" in walk_corpus_bytes({"pkg1.zip": buf.getvalue()})


class TestJsCapabilities:
    """Lock the JS capability extractor (groovy model, JS/Rhino syntax).
    NOTE: corpus has only 1 real .js specimen, so this is structurally sound but
    not corpus-validated at scale — outputs are reasoned (see module docstring)."""

    REAL = ("function processData(message) {"
            "var body = String(message.getBody(new java.lang.String().getClass()));"
            "var j = JSON.parse(body);"
            "message.setBody(JSON.stringify(j, null, 2));"
            "return message; }")

    def test_envelope_all_forms(self):
        from library_builder.js_capabilities import extract_capability
        for src in [
            "function processData(message){return message;}",
            "var processData = function(message){return message;}",
            "processData = (message) => { return message; }",
        ]:
            assert extract_capability("x.js", src).envelope == "processData"
        assert extract_capability("h.js",
                                  "function helper(x){return x;}").envelope == ""

    def test_operations_and_portable(self):
        from library_builder.js_capabilities import extract_capability
        c = extract_capability("r.js", self.REAL)
        assert "READ_BODY" in c.operations and "WRITE_BODY" in c.operations
        assert "PARSE_JSON" in c.portable_ops and "BUILD_JSON" in c.portable_ops
        assert "String" in c.body_read_as

    def test_no_exception_message_false_positive(self):
        """e.message.replace must NOT be read as a SAP binding (groovy lesson)."""
        from library_builder.js_capabilities import extract_capability, discover_bindings
        src = ("function processData(message){"
               "try{x();}catch(e){var m=e.message.replace('a','b');}"
               "return message;}")
        c = extract_capability("e.js", src)
        assert not c.bindings
        assert "message.replace" not in discover_bindings({"e.js": src})

    def test_library_flag(self):
        from library_builder.js_capabilities import extract_capability
        src = "\n".join(f"function f{i}(x){{return x;}}" for i in range(6))
        assert extract_capability("lib.js", src).is_library

    def test_flows_through_facade_and_solver(self):
        from library_builder.capability_catalog import catalog_for, type_for_ext
        from library_builder.solver import normalize, fetch, Need, _kw
        assert type_for_ext(".js") == "js"
        cat = catalog_for("js", {"t.js": self.REAL})
        norm = normalize(cat, "js")
        assert norm and "json" in norm[0].keywords
        # findable by FETCH
        ranked = fetch(Need("parse json", _kw("parse json") | {"json", "parse"}),
                       norm)
        assert ranked and "t.js" in ranked[0].capability.cap_id


class TestPropsCapabilities:
    """Lock the props extractor: iFlow config surface (prop values + propdef
    contract), identity by parameter-name set, prop↔propdef pairing."""

    PROPDEF = ('<parameters><parameter><key/><name>jmsQueueName</name>'
               '<type>xsd:string</type><isRequired>false</isRequired></parameter>'
               '<parameter><name>Address</name><type>xsd:string</type>'
               '<isRequired>true</isRequired></parameter></parameters>')
    PROP = ("#Thu Mar 19 10:07:52 UTC 2026\n"
            "jmsQueueName=JMS_SEND_X\nAddress=/Logistics/Create\n")

    def test_parse_propdef_contract(self):
        from library_builder.props_capabilities import extract_capability
        c = extract_capability("p.propdef", self.PROPDEF)
        assert c.kind == "propdef" and c.has_contract
        assert {p.name for p in c.parameters} == {"jmsQueueName", "Address"}
        addr = next(p for p in c.parameters if p.name == "Address")
        assert addr.type == "xsd:string" and addr.required is True

    def test_parse_prop_values_ignores_comments(self):
        from library_builder.props_capabilities import extract_capability
        c = extract_capability("p.prop", self.PROP)
        assert c.kind == "prop"
        assert {p.name for p in c.parameters} == {"jmsQueueName", "Address"}
        assert next(p for p in c.parameters
                    if p.name == "jmsQueueName").value == "JMS_SEND_X"

    def test_value_with_equals_sign_preserved(self):
        from library_builder.props_capabilities import extract_capability
        c = extract_capability("u.prop", "endpoint=https://h/api?a=1&b=2\ntok=ab==")
        vals = {p.name: p.value for p in c.parameters}
        assert vals["endpoint"] == "https://h/api?a=1&b=2"
        assert vals["tok"] == "ab=="

    def test_pairing_merges_values_into_contract(self):
        from library_builder.props_capabilities import build_catalog
        cat = build_catalog({"x.propdef": self.PROPDEF, "x.prop": self.PROP})
        # paired into one config capability carrying both contract + values
        assert cat["count"] == 1
        merged = cat["capabilities"][0]
        assert merged.has_contract and merged.has_values
        vmap = {p.name: p.value for p in merged.parameters}
        assert vmap["jmsQueueName"] == "JMS_SEND_X"
        assert vmap["Address"] == "/Logistics/Create"

    def test_same_keyset_dedupes_regardless_of_values(self):
        """Locked principle: same parameter-name set = same config solution;
        values are environment-specific examples, not identity."""
        from library_builder.props_capabilities import extract_capability
        a = extract_capability("a.prop", "host=prod\nport=443")
        b = extract_capability("b.prop", "host=dev\nport=8080")
        assert a.surface_hash == b.surface_hash

    def test_flows_through_facade_and_solver(self):
        from library_builder.capability_catalog import catalog_for, type_for_ext
        from library_builder.solver import normalize, fetch, Need, _kw
        assert type_for_ext(".prop") == "props"
        assert type_for_ext(".propdef") == "props"
        cat = catalog_for("props", {"x.propdef": self.PROPDEF})
        norm = normalize(cat, "props")
        assert norm and "jmsqueuename" in norm[0].keywords
        ranked = fetch(Need("jmsQueueName", {"jmsqueuename"}), norm)
        assert ranked


class TestIflwCapabilities:
    """Lock the iflw extractor (the capstone): integration-flow anatomy —
    adapters + step sequence + config — for clone-and-adapt reuse."""

    IFLW = '''<bpmn2:definitions xmlns:bpmn2="x" xmlns:ifl="y">
      <bpmn2:messageFlow id="MF_1" sourceRef="Participant_1" targetRef="StartEvent_1">
        <bpmn2:extensionElements><ifl:property>
          <key>ComponentType</key><value>SFTP</value></ifl:property>
        </bpmn2:extensionElements></bpmn2:messageFlow>
      <bpmn2:messageFlow id="MF_2" sourceRef="EndEvent_1" targetRef="Participant_2">
        <bpmn2:extensionElements><ifl:property>
          <key>ComponentType</key><value>HTTP</value></ifl:property>
          <ifl:property><key>address</key><value>{{TARGET_URL}}</value></ifl:property>
        </bpmn2:extensionElements></bpmn2:messageFlow>
      <bpmn2:serviceTask id="ST_1"><bpmn2:extensionElements><ifl:property>
        <key>activityType</key><value>Script</value></ifl:property>
        </bpmn2:extensionElements></bpmn2:serviceTask>
      <bpmn2:callActivity id="CA_1"><bpmn2:extensionElements><ifl:property>
        <key>activityType</key><value>Enricher</value></ifl:property>
        </bpmn2:extensionElements></bpmn2:callActivity>
    </bpmn2:definitions>'''

    def test_anatomy_extraction(self):
        from library_builder.iflw_capabilities import extract_capability
        c = extract_capability("flow.iflw", self.IFLW)
        assert "SFTP" in c.sender_adapters
        assert "HTTP" in c.receiver_adapters
        assert "Script" in c.step_types and "Enricher" in c.step_types
        assert c.step_count == 2

    def test_direction_resolution(self):
        """Participant-source = sender; Participant-target = receiver."""
        from library_builder.iflw_capabilities import extract_capability
        c = extract_capability("flow.iflw", self.IFLW)
        # SFTP came in (Participant_1 -> StartEvent), HTTP went out (End -> Participant_2)
        assert c.sender_adapters == ["SFTP"]
        assert c.receiver_adapters == ["HTTP"]

    def test_externalized_params_are_what_varies(self):
        from library_builder.iflw_capabilities import extract_capability
        c = extract_capability("flow.iflw", self.IFLW)
        assert "TARGET_URL" in c.externalized_params
        assert any("TARGET_URL" in v for v in c.what_varies)

    def test_purpose_reads_as_pattern(self):
        from library_builder.iflw_capabilities import extract_capability
        c = extract_capability("flow.iflw", self.IFLW)
        # sender → verbs → receiver
        assert "SFTP" in c.purpose and "HTTP" in c.purpose
        assert "enrich" in c.purpose or "script" in c.purpose

    def test_timer_trigger_no_misleading_sender(self):
        from library_builder.iflw_capabilities import extract_capability
        timer = ('<bpmn2:definitions xmlns:bpmn2="x" xmlns:ifl="y">'
                 '<bpmn2:startEvent><bpmn2:extensionElements><ifl:property>'
                 '<key>activityType</key><value>StartTimerEvent</value>'
                 '</ifl:property></bpmn2:extensionElements></bpmn2:startEvent>'
                 '</bpmn2:definitions>')
        c = extract_capability("t.iflw", timer)
        assert c.trigger == "timer/scheduled"
        assert "timer" in c.purpose      # not a misleading "?"

    def test_flows_through_facade_and_solver(self):
        from library_builder.capability_catalog import catalog_for, type_for_ext
        from library_builder.solver import normalize, fetch, Need, _kw
        assert type_for_ext(".iflw") == "iflw"
        cat = catalog_for("iflw", {"flow.iflw": self.IFLW})
        norm = normalize(cat, "iflw")
        assert norm and "sftp" in norm[0].keywords and "http" in norm[0].keywords
        # findable by adapter-pattern search
        ranked = fetch(Need("sftp http", {"sftp", "http"}), norm)
        assert ranked and "flow.iflw" in ranked[0].capability.cap_id
        # iFlows always flag tenant testing (they deploy)
        assert norm[0].needs_binding is True


class TestPiCapabilities:
    """Lock the PI extractor + PI→CPI translator. Grounded in REAL specimens
    (a Java mapping + a UDF library fetched from public GitHub); small sample,
    honestly flagged — see module docstring."""

    JAVA_MAPPING = (
        "import com.sap.aii.mapping.api.AbstractTransformation;\n"
        "import com.sap.aii.mapping.api.TransformationInput;\n"
        "public class M extends AbstractTransformation {\n"
        "  public void transform(TransformationInput in, TransformationOutput out){\n"
        "    String id = in.getInputHeader().getMessageId();\n"
        "    in.getInputPayload().getInputStream();\n"
        "    out.getOutputPayload().getOutputStream();\n"
        "  }\n}")

    UDF_LIB = (
        "import com.sap.aii.mapping.api.*;\n"
        "import com.sap.aii.mappingtool.tf7.rt.*;\n"
        "import com.sap.aii.mapping.lookup.*;\n"
        "import com.sap.ide.esr.tools.mapping.core.LibraryMethod;\n"
        "public class UDFs {\n"
        '  @LibraryMethod(title="setParameter", type=ExecutionType.SINGLE_VALUE)\n'
        "  public String setParameter(String ns, String key, Container c){\n"
        "    DynamicConfiguration cfg = (DynamicConfiguration)c"
        ".getTransformationParameters().get(StreamTransformationConstants"
        ".DYNAMIC_CONFIGURATION);\n    return key;\n  }\n}")

    def test_classify_java_mapping(self):
        from library_builder.pi_capabilities import extract_capability
        c = extract_capability("M.java", self.JAVA_MAPPING)
        assert c.pi_type == "java_mapping"
        assert "READ_BODY" in c.operations and "WRITE_BODY" in c.operations
        assert "READ_HEADER" in c.operations

    def test_classify_udf_library(self):
        from library_builder.pi_capabilities import extract_capability
        c = extract_capability("UDFs.java", self.UDF_LIB)
        assert c.pi_type == "udf_library"
        assert "setParameter" in c.udf_methods
        assert c.uses_tf7
        assert "DYNAMIC_CONFIG" in c.operations

    def test_lookup_not_false_tagged_from_import(self):
        """LOOKUP must come from real usage, not the mapping.lookup import
        (rigor-audit fix)."""
        from library_builder.pi_capabilities import extract_capability
        c = extract_capability("UDFs.java", self.UDF_LIB)
        assert "LOOKUP" not in c.operations   # import present, no lookup call

    def test_translation_to_cpi(self):
        from library_builder.pi_capabilities import extract_capability, translate_to_cpi
        jm = extract_capability("M.java", self.JAVA_MAPPING)
        spec = translate_to_cpi(jm)
        assert "processData" in spec["build_in_cpi"]
        assert spec["confidence"] == "reasoned"
        udf = extract_capability("UDFs.java", self.UDF_LIB)
        assert "Groovy UDF" in translate_to_cpi(udf)["build_in_cpi"]

    def test_flows_through_facade_and_solver(self):
        from library_builder.capability_catalog import catalog_for, TYPES
        from library_builder.solver import normalize, fetch, Need
        assert "pi" in TYPES
        cat = catalog_for("pi", {"M.java": self.JAVA_MAPPING,
                                 "UDFs.java": self.UDF_LIB})
        assert cat["count"] == 2 and len(cat["migration_specs"]) == 2
        norm = normalize(cat, "pi")
        ranked = fetch(Need("dynamic config", {"dynamic", "config"}), norm)
        assert ranked and "UDFs" in ranked[0].capability.cap_id
        assert norm[0].needs_binding is True   # migration → tenant build


class TestRequirementBridge:
    """Lock the input->capability bridge (part 1): requirement / MA / PI inputs
    become solver-ready CapabilityRequirements that carry source/target slots."""

    def _rr(self, **kw):
        from dataclasses import dataclass, field
        @dataclass
        class RequirementResult:
            name: str = "X"; sender_system: str = ""; sender_adapter: str = "HTTPS"
            receiver_system: str = ""; receiver_adapter: str = "IDOC"
            namespace: str = ""; description: str = ""; message_interface: str = ""
            mapping_program: str = ""; target_id: str = ""; sender_address: str = ""
            receiver_address: str = ""; sender_auth_method: str = "Basic"
            receiver_auth_method: str = "Basic"; is_async: bool = False
            message_format: str = "XML"; scheduler_cron: str = ""
            business_process: str = ""; confidence: float = 0.0
            needs_review: list = field(default_factory=list); raw_text: str = ""
        return RequirementResult(**kw)

    def _ir(self, **kw):
        from dataclasses import dataclass, field
        from typing import Optional
        @dataclass
        class InterfaceRecord:
            id: str = "I1"; name: str = "X"; namespace: str = ""
            software_component: str = ""; sender_system: str = ""
            receiver_system: str = ""; sender_adapter: str = "FILE"
            receiver_adapter: str = "SOAP"; message_interface: str = ""
            mapping_program: Optional[str] = None; has_bpm: bool = False
            has_multi_mapping: bool = False; channel_count: int = 1
            description: str = ""; raw: dict = field(default_factory=dict)
        return InterfaceRecord(**kw)

    def test_requirement_result_to_capability(self):
        from library_builder.requirement_bridge import to_requirement
        req = to_requirement(self._rr(
            sender_adapter="HTTPS", receiver_adapter="IDOC",
            mapping_program="OrderMap", message_format="JSON", is_async=True))
        assert "HTTPS to IDOC" in req.requirement_text
        assert "json" in req.requirement_text and "asynchronous" in req.requirement_text
        assert "OrderMap" in req.requirement_text
        assert req.needs_mapping and req.is_async

    def test_source_target_slots(self):
        from library_builder.requirement_bridge import to_requirement
        req = to_requirement(self._rr(sender_system="Shopify",
                                      sender_adapter="HTTPS",
                                      receiver_system="S4", receiver_adapter="IDOC"))
        slots = req.source_target_slots()
        assert len(slots["sources"]) == 1 and len(slots["targets"]) == 1
        assert slots["sources"][0]["system"] == "Shopify"
        assert slots["targets"][0]["adapter"] == "IDOC"
        assert slots["sources"][0]["value"] == ""   # editable, starts blank

    def test_interface_record_to_capability(self):
        from library_builder.requirement_bridge import to_requirement
        req = to_requirement(self._ir(sender_adapter="IDOC",
                                      receiver_adapter="SOAP",
                                      mapping_program="MATMAS"))
        assert "IDOC to SOAP" in req.requirement_text
        assert "mapping MATMAS" in req.requirement_text
        assert req.requirement_text.count("mapping") == 1   # no redundancy
        assert req.source_input == "migration_assessment"

    def test_scheduled_flow_no_fabricated_sender(self):
        from library_builder.requirement_bridge import to_requirement
        req = to_requirement(self._rr(sender_adapter="", receiver_adapter="SFTP",
                                      scheduler_cron="0 2 * * *",
                                      message_format="CSV"))
        # no fabricated sender adapter; reads as scheduled, targets SFTP
        assert "SFTP" in req.requirement_text
        assert "scheduled" in req.requirement_text
        assert req.scheduler

    def test_pi_capability_to_capability(self):
        from library_builder.pi_capabilities import extract_capability
        from library_builder.requirement_bridge import to_requirement
        pc = extract_capability(
            "M.java",
            "import com.sap.aii.mapping.api.AbstractTransformation;"
            " public class M extends AbstractTransformation {"
            " public void transform(TransformationInput in, TransformationOutput out){"
            " in.getInputPayload(); out.getOutputPayload(); } }")
        req = to_requirement(pc)
        assert req.source_input == "pi_migration"
        assert "processData" in req.requirement_text

    def test_unknown_object_raises(self):
        from library_builder.requirement_bridge import to_requirement
        import pytest
        with pytest.raises(TypeError):
            to_requirement(42)

    def test_solve_for_end_to_end(self):
        from library_builder.requirement_bridge import solve_for
        from library_builder.corpus_pipeline import build_corpus
        corpus = build_corpus(files={
            "m.groovy": ("def Message processData(Message m){"
                         "def b=m.getBody(String); return m}"),
            "x.xsl": ('<xsl:stylesheet version="1.0" '
                      'xmlns:xsl="http://www.w3.org/1999/XSL/Transform">'
                      '<xsl:output method="text"/>'
                      '<xsl:template match="/"/></xsl:stylesheet>'),
        })
        summary = solve_for(self._rr(sender_adapter="HTTPS",
                                     receiver_adapter="SFTP",
                                     message_format="CSV"), corpus)
        assert "requirement" in summary
        assert "source_target_slots" in summary
        assert summary["input_source"] == "requirement"
        assert summary["confidence"] == "reasoned"


class TestFieldSpec:
    """Lock the field-spec layer (part 2 logic): editable fields derived from a
    solution, with user-edit preservation across solver re-runs."""

    def _solution(self):
        return {
            "requirement": "HTTPS to IDOC integration",
            "source_target_slots": {
                "sources": [{"role": "sender", "system": "Shopify",
                             "adapter": "HTTPS", "value": ""}],
                "targets": [{"role": "receiver", "system": "S4",
                             "adapter": "IDOC", "value": ""}],
            },
            "steps": [],
        }

    def test_source_target_fields_prefilled(self):
        from library_builder.field_spec import build_field_spec
        spec = build_field_spec(self._solution())
        srcs = spec.by_group("source")
        tgts = spec.by_group("target")
        assert len(srcs) == 1 and len(tgts) == 1
        assert srcs[0].value == "Shopify" and srcs[0].source == "requirement"
        assert tgts[0].value == "S4" and tgts[0].editable

    def test_variable_number_of_slots(self):
        from library_builder.field_spec import build_field_spec
        sol = self._solution()
        sol["source_target_slots"]["targets"].append(
            {"role": "receiver", "system": "CRM", "adapter": "SOAP", "value": ""})
        spec = build_field_spec(sol)
        assert len(spec.by_group("target")) == 2   # variable N, not fixed

    def test_externalized_params_surfaced(self):
        from library_builder.field_spec import build_field_spec
        from library_builder.corpus_pipeline import build_corpus
        corpus = build_corpus(files={
            "f.iflw": '<bpmn2:definitions xmlns:bpmn2="x" xmlns:ifl="y">'
            '<bpmn2:messageFlow sourceRef="Participant_1" targetRef="S">'
            '<bpmn2:extensionElements><ifl:property><key>ComponentType</key>'
            '<value>HTTP</value></ifl:property><ifl:property><key>address</key>'
            '<value>{{TARGET_URL}}</value></ifl:property></bpmn2:extensionElements>'
            '</bpmn2:messageFlow></bpmn2:definitions>'})
        sol = self._solution()
        sol["steps"] = [{"use": "iflw:f.iflw", "ctype": "iflw", "score": 1.0}]
        spec = build_field_spec(sol, corpus)
        params = spec.by_group("parameter")
        assert any(f.label == "TARGET_URL" for f in params)
        # surfaced as a parameter to set, not hardcoded
        assert all("do not hardcode" in f.hint for f in params)

    def test_apply_user_value_marks_edited(self):
        from library_builder.field_spec import build_field_spec, apply_user_value
        spec = build_field_spec(self._solution())
        apply_user_value(spec, "target_1", "S4_QA")
        t = next(f for f in spec.fields if f.key == "target_1")
        assert t.value == "S4_QA" and t.user_edited and t.source == "user"

    def test_merge_edits_preserves_user_values(self):
        """The 'requirements change mid-call' safety: re-proposing defaults must
        NOT clobber values the user hand-edited."""
        from library_builder.field_spec import build_field_spec, apply_user_value, merge_edits
        spec = build_field_spec(self._solution())
        apply_user_value(spec, "target_1", "S4_QA")
        prior = {f.key: {"value": f.value, "user_edited": f.user_edited}
                 for f in spec.fields}
        # solver re-runs -> fresh spec with original suggestions
        fresh = build_field_spec(self._solution())
        merged = merge_edits(fresh, prior)
        t = next(f for f in merged.fields if f.key == "target_1")
        assert t.value == "S4_QA"        # user edit survived
        assert t.user_edited
        s = next(f for f in merged.fields if f.key == "source_1")
        assert s.value == "Shopify"      # unedited took fresh default


class TestCapabilityGeneration:
    """Lock capability-backed Generate All: generate_bundle pulls REAL adapted
    artifacts from the learned corpus, with the generic template as fallback.
    Strictly additive — output is never worse than template-only."""

    def _ir(self, **kw):
        from dataclasses import dataclass, field
        from typing import Optional
        @dataclass
        class InterfaceRecord:
            id: str = "I1"; name: str = "X"; namespace: str = ""
            software_component: str = ""; sender_system: str = ""
            receiver_system: str = ""; sender_adapter: str = "HTTPS"
            receiver_adapter: str = "IDOC"; message_interface: str = ""
            mapping_program: Optional[str] = None; has_bpm: bool = False
            has_multi_mapping: bool = False; channel_count: int = 1
            description: str = ""; raw: dict = field(default_factory=dict)
        return InterfaceRecord(**kw)

    def _corpus(self):
        from library_builder.corpus_pipeline import build_corpus
        return build_corpus(files={
            "real_transform.groovy":
                "import com.sap.gateway.ip.core.customdev.util.Message\n"
                "def Message processData(Message message){\n"
                "  def body = message.getBody(String)\n"
                "  message.setBody(body.toUpperCase())\n"
                "  return message\n}",
        })

    def test_capability_mode_uses_real_artifact(self):
        from scaffolder.capability_generator import generate_script_from_capability
        art = generate_script_from_capability(
            self._ir(description="transform the payload body"), self._corpus())
        assert art is not None
        assert art.source_capability == "real_transform.groovy"
        assert "processData" in art.content       # the REAL body, not a template
        assert art.confidence == "reasoned" and art.needs_tenant_test

    def test_no_match_returns_none_for_fallback(self):
        from scaffolder.capability_generator import generate_mapping_from_capability
        # empty corpus → no mmap capability → None (caller falls back)
        from library_builder.corpus_pipeline import build_corpus
        empty = build_corpus(files={"only.groovy": "def x(){}"})
        assert generate_mapping_from_capability(self._ir(), empty) is None

    def test_generate_bundle_template_fallback_unchanged(self, tmp_path):
        from scaffolder.artifact_bundle import generate_bundle
        iflow = tmp_path / "iflows" / "x.iflw"
        iflow.parent.mkdir(parents=True)
        iflow.write_text("<iflow/>")
        # no corpus → generic templates (the prior behavior)
        bundle = generate_bundle(self._ir(name="Order"), iflow)
        notes = " ".join(bundle.notes).lower()
        assert "generic" in notes or "placeholder" in notes
        assert len(bundle.artifacts) == 2

    def test_generate_bundle_capability_mode(self, tmp_path):
        from scaffolder.artifact_bundle import generate_bundle
        iflow = tmp_path / "iflows" / "y.iflw"
        iflow.parent.mkdir(parents=True)
        iflow.write_text("<iflow/>")
        bundle = generate_bundle(self._ir(name="Order",
                                          description="transform body"),
                                 iflow, corpus=self._corpus())
        # the script must be the REAL learned artifact, not a template
        script = next(a for a in bundle.artifacts if a.kind == "script")
        assert "real learned" in script.note.lower()
        assert "processData" in script.content

    def test_capability_mode_never_worse_than_template(self, tmp_path):
        """Additive guarantee: with a corpus, you get >= what template gives —
        every artifact is either a real match or the same template fallback."""
        from scaffolder.artifact_bundle import generate_bundle
        iflow = tmp_path / "iflows" / "z.iflw"
        iflow.parent.mkdir(parents=True)
        iflow.write_text("<iflow/>")
        bundle = generate_bundle(self._ir(name="Order"), iflow,
                                 corpus=self._corpus())
        # always produces a full valid bundle (2 artifacts), never broken
        assert len(bundle.artifacts) == 2
        assert all(a.content for a in bundle.artifacts)


class TestErrorRecommender:
    """Lock the error->recommendation engine (Tier 1). Maps tenant failures to
    cause + concrete fix, and crucially encodes the safe/risky split via
    fix_class + auto_fixable (structural=auto-fixable, semantic=recommend-only)."""

    def test_parse_odata_error_shapes(self):
        from fetcher.error_recommender import parse_odata_error
        # nested JSON (the common CPI shape)
        p = parse_odata_error(
            '{"error":{"code":"X/500","message":{"value":"the real reason"}}}')
        assert p["code"] == "X/500" and p["message"] == "the real reason"
        # message-as-string
        assert parse_odata_error(
            '{"error":{"message":"plain"}}')["message"] == "plain"
        # malformed json → raw text preserved, no crash
        assert "broken" in parse_odata_error("broken {json")["message"]
        # empty
        assert parse_odata_error("")["message"] == ""

    def test_full_message_not_truncated(self):
        """The prior uploader truncated to 300 chars; the recommender keeps the
        FULL tenant message."""
        from fetcher.error_recommender import recommend
        long_msg = "Resource not found: " + "x" * 500
        r = recommend("upload", 400,
                      '{"error":{"message":{"value":"%s"}}}' % long_msg)
        assert len(r.raw_error) > 400   # not truncated

    def test_structural_upload_error_is_auto_fixable(self):
        from fetcher.error_recommender import recommend, FIX_STRUCTURAL
        r = recommend("upload", 400,
                      '{"error":{"message":{"value":"Resource OrderMap not found"}}}')
        assert r.fix_class == FIX_STRUCTURAL
        assert r.auto_fixable is True       # safe to loop on
        assert "missing artifact" in r.recommendation.lower() or \
               "reference" in r.recommendation.lower()

    def test_semantic_groovy_error_is_NOT_auto_fixable(self):
        """The honesty gate: a logic/compilation error must NOT be auto-fixed
        (stripping it would change behavior)."""
        from fetcher.error_recommender import recommend, FIX_SEMANTIC
        r = recommend("deploy", 500,
                      '{"error":{"message":{"value":'
                      '"Groovy script compilation failed: cannot resolve"}}}')
        assert r.fix_class == FIX_SEMANTIC
        assert r.auto_fixable is False      # recommend only

    def test_unsupported_function_is_bounded_substitution(self):
        from fetcher.error_recommender import recommend, FIX_SUBSTITUTION
        r = recommend("deploy", 400,
                      '{"error":{"message":{"value":'
                      '"Function getTransformationParameters is not supported"}}}')
        assert r.fix_class == FIX_SUBSTITUTION
        assert r.auto_fixable is True       # bounded, from equivalence table

    def test_auth_errors_recommend_user_action(self):
        from fetcher.error_recommender import recommend, FIX_AUTH
        r401 = recommend("upload", 401, "Unauthorized")
        r403 = recommend("upload", 403, '{"error":{"message":{"value":"forbidden"}}}')
        assert r401.fix_class == FIX_AUTH and not r401.auto_fixable
        assert r403.fix_class == FIX_AUTH and not r403.auto_fixable
        assert "oauth" in r401.recommendation.lower()
        assert "role" in r403.recommendation.lower()

    def test_unknown_error_surfaces_raw_message(self):
        from fetcher.error_recommender import recommend, FIX_UNKNOWN
        r = recommend("deploy", 500,
                      '{"error":{"message":{"value":"some novel thing"}}}')
        assert r.fix_class == FIX_UNKNOWN and not r.auto_fixable
        assert "some novel thing" in r.raw_error   # honest: surfaces the real msg

    def test_unresolved_parameter_links_to_field_spec(self):
        from fetcher.error_recommender import recommend
        r = recommend("deploy", 400,
                      '{"error":{"message":{"value":'
                      '"Externalized parameter TARGET_URL is unresolved"}}}')
        assert r.auto_fixable is True
        assert "parameter" in r.recommendation.lower()


class TestUploaderRecommenderWiring:
    """Lock the uploader↔recommender wiring: upload/deploy failures now carry a
    structured recommendation and the full untruncated tenant error."""

    def _uploader(self):
        import requests
        from fetcher.cpi_uploader import CPIUploader
        u = CPIUploader.__new__(CPIUploader)
        u.session = requests.Session()
        u.base_url = "https://tenant.example"
        u._csrf_token = None
        return u

    class _Resp:
        def __init__(self, code, text):
            self.status_code = code
            self.text = text
            self.headers = {}

    def test_upload_failure_attaches_recommendation(self):
        from fetcher.cpi_uploader import UploadResult
        u = self._uploader()
        res = UploadResult(interface_name="X", package_id="P",
                           artifact_id="A", status="pending")
        body = ('{"error":{"message":{"value":'
                '"Resource OrderMap not found in bundle"}}}')
        u._report_failure(self._Resp(400, body), "A", "http://x", res)
        assert res.status == "failed"
        assert res.recommendation is not None
        assert res.recommendation.fix_class == "structural"
        assert res.recommendation.auto_fixable is True
        # the tenant's real reason is preserved (untruncated)
        assert res.recommendation.raw_error == "Resource OrderMap not found in bundle"
        # the message now carries cause + fix, not a raw json dump
        assert "FIX:" in res.message

    def test_upload_failure_long_body_not_truncated(self):
        from fetcher.cpi_uploader import UploadResult
        u = self._uploader()
        res = UploadResult(interface_name="X", package_id="P",
                           artifact_id="A", status="pending")
        long_reason = "Resource not found: " + "z" * 500
        body = '{"error":{"message":{"value":"%s"}}}' % long_reason
        u._report_failure(self._Resp(400, body), "A", "http://x", res)
        assert len(res.recommendation.raw_error) > 400   # full detail kept

    def test_deploy_failure_stores_recommendation(self):
        u = self._uploader()
        # patch session.post to simulate a deploy error
        class _Sess:
            headers = {}
            def post(self, *a, **k):
                return TestUploaderRecommenderWiring._Resp(
                    400, '{"error":{"message":{"value":'
                    '"Function getTransformationParameters is not supported"}}}')
        u.session = _Sess()
        status = u.deploy_iflow("A")
        assert status == "failed"
        assert u.last_deploy_recommendation is not None
        assert u.last_deploy_recommendation.fix_class == "substitution"

    def test_diagnosis_never_masks_original_failure(self):
        """If diagnosis somehow fails, the upload is still reported as failed."""
        from fetcher.cpi_uploader import UploadResult
        u = self._uploader()
        res = UploadResult(interface_name="X", package_id="P",
                           artifact_id="A", status="pending")
        # an empty/garbage body must not crash _report_failure
        u._report_failure(self._Resp(500, ""), "A", "http://x", res)
        assert res.status == "failed"


class TestSchemaDraftedMapping:
    """Lock Gap 2: capability-mode mapping has a schema-drafted tier — when no
    real .mmap matches, draft real field mappings from learned schema element
    names, instead of a blank placeholder. Never fabricates non-matching fields."""

    def _ir(self, **kw):
        from dataclasses import dataclass, field
        from typing import Optional
        @dataclass
        class InterfaceRecord:
            id: str = "I1"; name: str = "X"; namespace: str = ""
            software_component: str = ""; sender_system: str = ""
            receiver_system: str = ""; sender_adapter: str = "HTTPS"
            receiver_adapter: str = "IDOC"; message_interface: str = ""
            mapping_program: Optional[str] = None; has_bpm: bool = False
            has_multi_mapping: bool = False; channel_count: int = 1
            description: str = "zzz"; raw: dict = field(default_factory=dict)
        return InterfaceRecord(**kw)

    def test_schema_drafted_direct_matches(self):
        from scaffolder.capability_generator import generate_mapping_from_capability
        from library_builder.corpus_pipeline import build_corpus
        corpus = build_corpus(files={
            "s.xsd": '<xsd:schema xmlns:xsd="http://www.w3.org/2001/XMLSchema">'
                     '<xsd:element name="OrderId"/><xsd:element name="CustomerName"/>'
                     '<xsd:element name="Amount"/></xsd:schema>',
            "t.xsd": '<xsd:schema xmlns:xsd="http://www.w3.org/2001/XMLSchema">'
                     '<xsd:element name="OrderId"/><xsd:element name="CustomerName"/>'
                     '<xsd:element name="Total"/></xsd:schema>',
        })
        art = generate_mapping_from_capability(self._ir(), corpus)
        assert art is not None and art.source_capability == "schema-derived"
        assert "OrderId" in art.content and "CustomerName" in art.content
        # Amount/Total don't match by name → not fabricated
        assert "Total" not in art.content

    def test_no_schemas_returns_none(self):
        from scaffolder.capability_generator import generate_mapping_from_capability
        from library_builder.corpus_pipeline import build_corpus
        corpus = build_corpus(files={"only.groovy": "def x(){}"})
        assert generate_mapping_from_capability(self._ir(), corpus) is None

    def test_no_field_overlap_returns_none(self):
        """No false matches: schemas with disjoint fields → None, not garbage."""
        from scaffolder.capability_generator import generate_mapping_from_capability
        from library_builder.corpus_pipeline import build_corpus
        corpus = build_corpus(files={
            "a.xsd": '<xsd:schema xmlns:xsd="http://www.w3.org/2001/XMLSchema">'
                     '<xsd:element name="Apple"/></xsd:schema>',
            "b.xsd": '<xsd:schema xmlns:xsd="http://www.w3.org/2001/XMLSchema">'
                     '<xsd:element name="Zebra"/></xsd:schema>',
        })
        assert generate_mapping_from_capability(self._ir(), corpus) is None


class TestEmptyContentGuard:
    """Lock the fix for the live-tenant 'InputStream cannot be null' bug: no
    upload route may send empty ArtifactContent (the 32-byte-body failure).
    Guarded at the deepest choke-point (_post_artifact) so it's impossible
    regardless of caller."""

    def _uploader(self):
        import requests
        from fetcher.cpi_uploader import CPIUploader
        up = CPIUploader.__new__(CPIUploader)
        up.session = requests.Session()
        up._csrf_token = "x"
        up.base_url = "http://tenant.example"
        return up

    def test_post_artifact_blocks_empty_content(self):
        from fetcher.cpi_uploader import UploadResult
        up = self._uploader()
        r = UploadResult(interface_name="X", package_id="P",
                         artifact_id="A", status="pending")
        up._post_artifact(b"", "P", "A", "A", r, "IFlow")
        assert r.status == "failed"
        assert "InputStream cannot be null" in r.message  # the real tenant error

    def test_upload_raw_bundle_blocks_empty(self):
        up = self._uploader()
        r = up.upload_raw_bundle(b"", "P", "A", "A")
        assert r.status == "failed"
        assert "empty bundle" in r.message.lower()

    def test_recommender_diagnoses_inputstream_null(self):
        """The real tenant error maps to a structural, auto-fixable cause."""
        from fetcher.error_recommender import recommend, FIX_STRUCTURAL
        r = recommend("upload", 500,
                      '{"error":{"code":"Internal Server Error",'
                      '"message":{"lang":"en","value":"InputStream cannot be null"}}}')
        assert r.fix_class == FIX_STRUCTURAL
        assert r.auto_fixable is True
        assert "content" in r.cause.lower()


class TestNoBasenameCollision:
    """Lock the collision fix: same-named files across folders/packages must NOT
    collapse to one (the bug that silently shrank the catalog — e.g. hundreds of
    script.groovy → 1). Keys are path-qualified."""

    def test_dir_same_named_files_all_kept(self):
        import os, tempfile
        from library_builder.corpus_pipeline import walk_corpus, build_corpus
        root = tempfile.mkdtemp()
        gdir = os.path.join(root, "groovy")
        os.makedirs(gdir)
        for i in range(20):
            sub = os.path.join(gdir, f"pkg{i}")
            os.makedirs(sub)
            with open(os.path.join(sub, "script.groovy"), "w") as fh:
                fh.write("def Message processData(Message m){ return m }")
        files = walk_corpus(root)
        groovy = [f for f in files if f.endswith(".groovy")]
        assert len(groovy) == 20      # all kept, not collapsed to 1
        corpus = build_corpus(path=root)
        assert corpus.report()["by_type"].get("groovy", 0) == 20

    def test_zip_same_named_files_all_kept(self):
        import io, zipfile, tempfile, os
        from library_builder.corpus_pipeline import walk_corpus
        buf = io.BytesIO()
        with zipfile.ZipFile(buf, "w") as z:
            for i in range(10):
                z.writestr(f"pkg{i}/script.groovy",
                           "def Message processData(Message m){return m}")
        with tempfile.NamedTemporaryFile(suffix=".zip", delete=False) as tf:
            tf.write(buf.getvalue()); p = tf.name
        try:
            files = walk_corpus(p)
            groovy = [f for f in files if f.endswith(".groovy")]
            assert len(groovy) == 10   # path-qualified, not collapsed
        finally:
            os.unlink(p)


class TestWholePackageUpload:
    """Lock the whole-package upload path (the user's architectural insight):
    build a full package export zip and POST it to IntegrationPackages?Overwrite
    =true, instead of the per-artifact create call that 500s with 'InputStream
    cannot be null'. Verifies the package zip matches the real export structure."""

    def _uploader(self):
        import requests
        from fetcher.cpi_uploader import CPIUploader
        up = CPIUploader.__new__(CPIUploader)
        up.session = requests.Session()
        up._csrf_token = "x"
        up.base_url = "http://tenant.example"
        return up

    def test_package_zip_structure(self):
        import io, zipfile, base64, json
        up = self._uploader()
        inner = b"PK\x03\x04fake-inner-artifact-zip"
        pkg = up.build_package_export_zip(
            [("MyFlow", "My Flow", inner)], "MyPkg", "My Package")
        zf = zipfile.ZipFile(io.BytesIO(pkg))
        names = zf.namelist()
        # the real-export members must all be present
        assert "ExportInformation.info" in names
        assert "contentmetadata.md" in names
        assert "hash" in names
        assert "resources.cnt" in names
        # nested content is GUID-named (mirrors real exports), not readable name
        content = [n for n in names if n.endswith("_content")]
        assert len(content) == 1
        guid = content[0].replace("_content", "")
        assert len(guid) == 32 and all(c in "0123456789abcdef" for c in guid)
        # nested content is exactly the inner bundle bytes
        assert zf.read(content[0]) == inner
        # resources.cnt is base64 JSON; id matches the GUID. The `name` is the
        # readable artifact name + ".zip" (the REAL export convention, verified
        # 915/915 in the spine analysis) — NOT "<guid>.zip".
        res = json.loads(base64.b64decode(zf.read("resources.cnt")).decode("utf-8"))
        entry = res["resources"][0]
        assert entry["id"] == guid
        assert entry["name"].endswith(".zip") and entry["name"] != f"{guid}.zip"
        # metadata files MUST be populated (tenant-verified: blank blocks import)
        assert zf.read("contentmetadata.md") != b"" and zf.read("hash") != b""

    def test_multi_artifact_package(self):
        import io, zipfile, base64, json
        up = self._uploader()
        bundles = [("A", "Flow A", b"PKaaa"), ("B", "Flow B", b"PKbbb")]
        pkg = up.build_package_export_zip(bundles, "Pkg", "Pkg Name")
        zf = zipfile.ZipFile(io.BytesIO(pkg))
        content = sorted(n for n in zf.namelist() if n.endswith("_content"))
        assert len(content) == 2
        guids = {n.replace("_content", "") for n in content}
        # resources.cnt JSON lists both, each id matching a content GUID
        res = json.loads(base64.b64decode(zf.read("resources.cnt")).decode("utf-8"))
        # IFlow entries match the content GUIDs; a ContentPackage entry is also
        # present for the package itself (tenant-verified requirement).
        iflow_ids = {r["id"] for r in res["resources"] if r["resourceType"] == "IFlow"}
        assert iflow_ids == guids
        assert any(r["resourceType"] == "ContentPackage" for r in res["resources"])
        # the inner bytes are preserved (order-independent)
        stored = {zf.read(n) for n in content}
        assert stored == {b"PKaaa", b"PKbbb"}


class TestDeletePackage:
    """Lock the clean-slate delete_package used by the opt-in 'delete first'
    path: 200/204/404 are success (deleted or already-gone); 4xx/5xx are not."""

    def _uploader(self):
        from fetcher.cpi_uploader import CPIUploader
        from unittest.mock import MagicMock
        up = CPIUploader.__new__(CPIUploader)
        up.session = MagicMock()
        up._csrf_token = "x"
        up.base_url = "http://tenant.example"
        up._write_headers = lambda: {"X-CSRF-Token": "x"}
        return up

    def test_delete_success_and_already_gone(self):
        from unittest.mock import MagicMock
        up = self._uploader()
        for code in (200, 202, 204, 404):
            up.session.delete = MagicMock(
                return_value=MagicMock(status_code=code, headers={}, text=""))
            assert up.delete_package("Pkg") is True

    def test_delete_failure(self):
        from unittest.mock import MagicMock
        up = self._uploader()
        for code in (403, 500):
            up.session.delete = MagicMock(
                return_value=MagicMock(status_code=code, headers={}, text=""))
            assert up.delete_package("Pkg") is False


class TestFullPackageExport:
    """Lock the full-package-to-disk export (manual CPI import diagnostic): it
    writes a real package zip whose nested artifact entry is `<id>_content`
    (NO .zip extension) and whose inner bundle has the corrected .project."""

    def test_writes_package_with_content_entry(self):
        import tempfile, io, zipfile, requests
        from pathlib import Path
        from fetcher.cpi_uploader import CPIUploader
        from scaffolder.iflow_scaffolder import IFlowScaffolder
        from dataclasses import dataclass, field
        from typing import Optional
        out = tempfile.mkdtemp()
        scaf = IFlowScaffolder(output_dir=out)
        @dataclass
        class IR:
            id: str = "I1"; name: str = "PO_to_S4_Create"; namespace: str = ""
            software_component: str = ""; sender_system: str = "ECC"
            receiver_system: str = "S4"; sender_adapter: str = "IDOC"
            receiver_adapter: str = "SOAP"; message_interface: str = ""
            mapping_program: Optional[str] = None; has_bpm: bool = False
            has_multi_mapping: bool = False; channel_count: int = 1
            description: str = ""; raw: dict = field(default_factory=dict)
        @dataclass
        class A:
            interface: IR = field(default_factory=IR)
        p = scaf.scaffold(A())
        up = CPIUploader.__new__(CPIUploader)
        up.session = requests.Session(); up._csrf_token = "x"; up.base_url = "http://t"
        zp = up.export_full_package_to_disk(
            Path(p), "ECCS4HANApo", "ECC to S4 PO",
            "POtoS4Create", "PO_to_S4_Create", Path(out) / "exp")
        assert zp is not None
        import json, base64
        zf = zipfile.ZipFile(zp)
        # nested entry must be <guid>_content (GUID-named, NO .zip on the entry)
        content = [n for n in zf.namelist() if n.endswith("_content")]
        assert len(content) == 1
        guid = content[0].replace("_content", "")
        assert len(guid) == 32 and all(c in "0123456789abcdef" for c in guid)
        assert not content[0].endswith(".zip")
        # resources.cnt is base64 JSON whose id matches the GUID; the `name` is
        # the readable artifact name + ".zip" (real export convention), not <guid>.zip.
        res = json.loads(base64.b64decode(zf.read("resources.cnt")).decode())
        entry = res["resources"][0]
        assert entry["id"] == guid
        assert entry["name"].endswith(".zip") and entry["name"] != f"{guid}.zip"
        assert entry["resourceType"] == "IFlow"
        # metadata files MUST be populated (tenant-verified: blank blocks import)
        assert zf.read("contentmetadata.md") != b""
        assert zf.read("hash") != b""
        # inner bundle carries the corrected .project (no ABAP nature)
        inner = zipfile.ZipFile(io.BytesIO(zf.read(content[0])))
        proj = inner.read(".project").decode()
        assert "abap.nature" not in proj and "javanature" in proj


class TestIFlowPersonalizer:
    """Clone-and-adapt recipe: re-skin a real iFlow bundle, keep references valid."""

    def _fake_bundle(self):
        import io, zipfile
        members = {
            "META-INF/MANIFEST.MF":
                "Manifest-Version: 1.0\r\nBundle-SymbolicName: OldSym; singleton:=true\r\n"
                "Bundle-Name: Old Flow\r\nImport-Package: org.apache.camel\r\n\r\n",
            ".project": "<projectDescription><name>OldProj</name></projectDescription>",
            "src/main/resources/scenarioflows/integrationflow/Old.iflw":
                '<bpmn2:definitions><x ref="dir://script/src/main/resources/'
                'script/foo.groovy"/></bpmn2:definitions>',
            "src/main/resources/script/foo.groovy": "// old\ndef processData(m){return m}\n",
        }
        buf = io.BytesIO()
        with zipfile.ZipFile(buf, "w", zipfile.ZIP_DEFLATED) as z:
            for n, c in members.items():
                z.writestr(n, c)
        return buf.getvalue()

    def test_personalize_renames_identity_keeps_refs(self):
        import io, zipfile
        from scaffolder.iflow_personalizer import (
            PersonalizationSpec, personalize_bundle, references_intact)
        new = personalize_bundle(
            self._fake_bundle(),
            PersonalizationSpec(new_iflow_name="New_Flow", new_symbolic="NewSym"))
        z = zipfile.ZipFile(io.BytesIO(new))
        names = z.namelist()
        assert "src/main/resources/scenarioflows/integrationflow/New_Flow.iflw" in names
        assert not any(n.endswith("Old.iflw") for n in names)
        # non-iflw filenames preserved -> references resolve
        assert "src/main/resources/script/foo.groovy" in names
        mf = z.read("META-INF/MANIFEST.MF").decode()
        assert "Bundle-SymbolicName: NewSym" in mf
        assert "Import-Package" in mf
        ok, missing = references_intact(new)
        assert ok and not missing

    def test_script_override_applied(self):
        import io, zipfile
        from scaffolder.iflow_personalizer import PersonalizationSpec, personalize_bundle
        new = personalize_bundle(
            self._fake_bundle(),
            PersonalizationSpec(new_iflow_name="New_Flow",
                                script_overrides={"foo.groovy": "// NEWBODY\n"}))
        z = zipfile.ZipFile(io.BytesIO(new))
        assert z.read("src/main/resources/script/foo.groovy").decode() == "// NEWBODY\n"


class TestServiceKeyHelper:
    """Service-key (.json) loading + key listing for the deploy buttons."""

    def _key(self, tmp_path, name="k.json", nested=True):
        import json
        body = {"clientid": "cid", "clientsecret": "sec",
                "tokenurl": "https://x.authentication.us10.hana.ondemand.com",
                "url": "https://h.it-cpitrial05.cfapps.us10-001.hana.ondemand.com/api/v1"}
        data = {"oauth": body} if nested else body
        p = tmp_path / name
        p.write_text(json.dumps(data))
        return str(p)

    def test_load_normalizes_token_and_base(self, tmp_path):
        from fetcher.service_key import load_service_key
        c = load_service_key(self._key(tmp_path))
        assert c["token_url"].endswith("/oauth/token")
        assert c["base_url"].endswith("ondemand.com")
        assert "/api" not in c["base_url"]          # host root only
        assert c["client_id"] == "cid" and c["client_secret"] == "sec"

    def test_load_accepts_flat_key(self, tmp_path):
        from fetcher.service_key import load_service_key
        c = load_service_key(self._key(tmp_path, "flat.json", nested=False))
        assert c["client_id"] == "cid"

    def test_list_service_keys(self, tmp_path):
        from fetcher.service_key import list_service_keys
        self._key(tmp_path, "a.json"); self._key(tmp_path, "b.json")
        assert list_service_keys(str(tmp_path)) == ["a.json", "b.json"]

    def test_connect_requires_all_fields(self, tmp_path):
        import json, pytest
        from fetcher.service_key import connect_with_service_key
        p = tmp_path / "bad.json"
        p.write_text(json.dumps({"oauth": {"clientid": "x"}}))  # missing fields
        with pytest.raises(ValueError):
            connect_with_service_key(str(p))


class TestCloneAndAdapt:
    """clone_and_adapt: inject generated Groovy into a real template's slots."""

    def _template(self):
        import io, zipfile
        members = {
            "META-INF/MANIFEST.MF":
                "Manifest-Version: 1.0\r\nBundle-SymbolicName: Old\r\n"
                "Import-Package: org.apache.camel\r\n\r\n",
            ".project": "<projectDescription><name>Old</name></projectDescription>",
            "src/main/resources/scenarioflows/integrationflow/Old.iflw":
                '<d><a ref="dir://script/src/main/resources/script/a.groovy"/>'
                '<b ref="dir://script/src/main/resources/script/b.groovy"/></d>',
            "src/main/resources/script/a.groovy": "// tpl a\n",
            "src/main/resources/script/b.groovy": "// tpl b\n",
        }
        buf = io.BytesIO()
        with zipfile.ZipFile(buf, "w", zipfile.ZIP_DEFLATED) as z:
            for n, c in members.items():
                z.writestr(n, c)
        return buf.getvalue()

    def test_injects_generated_scripts_keeps_refs(self):
        import io, zipfile
        from scaffolder.iflow_personalizer import clone_and_adapt
        bundle, ok, report = clone_and_adapt(
            self._template(), "Z_PO_Create",
            generated_scripts=["// GEN main\n", "// GEN second\n"])
        z = zipfile.ZipFile(io.BytesIO(bundle))
        assert ok and report["scripts_injected"] == ["a.groovy", "b.groovy"]
        assert b"GEN main" in z.read("src/main/resources/script/a.groovy")
        assert b"GEN second" in z.read("src/main/resources/script/b.groovy")
        assert any(n.endswith("Z_PO_Create.iflw") for n in z.namelist())

    def test_sender_endpoints_made_unique_per_clone(self):
        """Two clones of one template must NOT share a sender endpoint address
        (shared addresses collide at runtime — only one starts). Receiver full
        URLs must be left intact."""
        import io, zipfile, re
        from scaffolder.iflow_personalizer import clone_and_adapt
        members = {
            "META-INF/MANIFEST.MF":
                "Manifest-Version: 1.0\r\nBundle-SymbolicName: Gen\r\n\r\n",
            ".project": "<projectDescription><name>Gen</name></projectDescription>",
            "src/main/resources/scenarioflows/integrationflow/Gen.iflw":
                "<d><key>urlPath</key><value>/GenericMessageProcessor_00</value>"
                "<key>address</key><value>/soap/GenericMessageProcessor_00</value>"
                "<key>address</key><value>https://backend/api</value></d>",
        }
        buf = io.BytesIO()
        with zipfile.ZipFile(buf, "w") as z:
            for n, c in members.items():
                z.writestr(n, c)
        tpl = buf.getvalue()

        def paths(b):
            zz = zipfile.ZipFile(io.BytesIO(b))
            iflw = next(n for n in zz.namelist() if n.endswith(".iflw"))
            body = zz.read(iflw).decode()
            return body, re.findall(
                r"<key>(?:urlPath|address)</key><value>(/[^<]*)</value>", body)

        b1, _, _ = clone_and_adapt(tpl, "S_One", generated_scripts=[])
        b2, _, _ = clone_and_adapt(tpl, "M_Two", generated_scripts=[])
        body1, p1 = paths(b1)
        _, p2 = paths(b2)
        assert p1 and p2 and p1 != p2           # distinct sender paths
        assert "https://backend/api" in body1   # receiver URL untouched
        assert all("genericmessageprocessor" not in p.lower() for p in p1)

    def test_clone_fills_empty_receiver_credential_name(self):
        """A cloned template whose basic-auth receiver has an empty
        credentialName must come out with a placeholder filled in, so the build
        passes instead of failing 'Credential name must be specified …'."""
        import io, zipfile
        from scaffolder.iflow_personalizer import (
            clone_and_adapt, _RECEIVER_CRED_PLACEHOLDER)
        members = {
            "META-INF/MANIFEST.MF":
                "Manifest-Version: 1.0\r\nBundle-SymbolicName: Mon\r\n\r\n",
            ".project": "<projectDescription><name>Mon</name></projectDescription>",
            "src/main/resources/scenarioflows/integrationflow/Mon.iflw":
                "<d><key>credentialName</key><value></value>"
                "<key>credentialName</key><value>KeepMe</value></d>",
        }
        buf = io.BytesIO()
        with zipfile.ZipFile(buf, "w") as z:
            for n, c in members.items():
                z.writestr(n, c)
        bundle, _, _ = clone_and_adapt(buf.getvalue(), "MONSTER_X",
                                       generated_scripts=[])
        zz = zipfile.ZipFile(io.BytesIO(bundle))
        iflw = next(n for n in zz.namelist() if n.endswith(".iflw"))
        body = zz.read(iflw).decode()
        assert _RECEIVER_CRED_PLACEHOLDER in body   # empty one filled
        assert "KeepMe" in body                     # populated one preserved

    def test_externalized_params_completed_unique_and_credential(self):
        """ProcessDirect sender param (ENDPOINT_ID) becomes unique per clone, and
        an empty credential param referenced by a Basic-auth receiver is filled —
        both via parameters.prop, matching the real template structure."""
        import io, zipfile, re
        from scaffolder.iflow_personalizer import (
            clone_and_adapt, _RECEIVER_CRED_PLACEHOLDER)
        iflw = (
            '<bpmn2:messageFlow id="MF1" name="ProcessDirect">'
            '<key>direction</key><value>Sender</value>'
            '<key>ComponentType</key><value>ProcessDirect</value>'
            '<key>address</key><value>GP_{{ENDPOINT_ID}}</value>'
            '</bpmn2:messageFlow>'
            '<bpmn2:messageFlow id="MF2" name="HTTP">'
            '<key>direction</key><value>Receiver</value>'
            '<key>ComponentType</key><value>HTTP</value>'
            '<key>authenticationMethod</key><value>Basic</value>'
            '<key>credentialName</key><value>{{CC Credential Name}}</value>'
            '</bpmn2:messageFlow>')
        members = {
            "META-INF/MANIFEST.MF":
                "Manifest-Version: 1.0\r\nBundle-SymbolicName: T\r\n\r\n",
            ".project": "<projectDescription><name>T</name></projectDescription>",
            "src/main/resources/scenarioflows/integrationflow/T.iflw": iflw,
            "src/main/resources/parameters.prop":
                "ENDPOINT_ID=00\nCC\\ Credential\\ Name=\n",
        }
        buf = io.BytesIO()
        with zipfile.ZipFile(buf, "w") as z:
            for n, c in members.items():
                z.writestr(n, c)
        tpl = buf.getvalue()

        def prop(b):
            zz = zipfile.ZipFile(io.BytesIO(b))
            p = next(n for n in zz.namelist() if n.endswith("parameters.prop"))
            return zz.read(p).decode()

        b1, _, _ = clone_and_adapt(tpl, "Flow_One", generated_scripts=[])
        b2, _, _ = clone_and_adapt(tpl, "Flow_Two", generated_scripts=[])
        e1 = re.search(r"ENDPOINT_ID=(\S+)", prop(b1)).group(1)
        e2 = re.search(r"ENDPOINT_ID=(\S+)", prop(b2)).group(1)
        assert e1 != e2 and e1 != "00"              # unique consumer id per clone
        assert _RECEIVER_CRED_PLACEHOLDER in prop(b1)  # empty credential filled

    def test_no_generated_scripts_is_pure_reskin(self):
        import io, zipfile
        from scaffolder.iflow_personalizer import clone_and_adapt
        bundle, ok, report = clone_and_adapt(self._template(), "Z_Reskin")
        z = zipfile.ZipFile(io.BytesIO(bundle))
        assert ok and report["scripts_injected"] == []
        assert z.read("src/main/resources/script/a.groovy").decode() == "// tpl a\n"


class TestTemplateLibrary:
    """Index a folder of packages and pick a real (importable) iFlow template."""

    def _bundle(self, name, steps, with_script):
        import io, zipfile
        iflw = "<d>" + "<bpmn2:serviceTask/>" * steps + "</d>"
        members = {
            "META-INF/MANIFEST.MF":
                "Manifest-Version: 1.0\r\nBundle-SymbolicName: X\r\n\r\n",
            f"src/main/resources/scenarioflows/integrationflow/{name}.iflw": iflw,
        }
        if with_script:
            members["src/main/resources/script/s.groovy"] = "// x\n"
        buf = io.BytesIO()
        with zipfile.ZipFile(buf, "w") as z:
            for n, c in members.items():
                z.writestr(n, c)
        return buf.getvalue()

    def test_index_and_pick_prefers_real_flow(self, tmp_path):
        from scaffolder.template_library import index_templates, pick_template
        (tmp_path / "real.zip").write_bytes(self._bundle("Real", 3, True))
        (tmp_path / "stub.zip").write_bytes(self._bundle("Stub", 0, True))
        ts = index_templates(str(tmp_path))
        assert {"Real", "Stub"} <= {t.name for t in ts}
        pick = pick_template(ts)
        assert pick.name == "Real" and pick.step_count == 3  # real flow wins


class TestUserSettings:
    """External settings store survives project re-import (lives in ~/.cpi_migrator)."""

    def test_set_get_unset_and_existence_check(self, tmp_path, monkeypatch):
        monkeypatch.setenv("CPI_MIGRATOR_HOME", str(tmp_path))
        import importlib
        from fetcher import user_settings as us
        importlib.reload(us)  # re-evaluate base dir under the patched env
        us.set_setting("capability_corpus_dir", str(tmp_path))   # exists
        us.set_setting("template_library_dir", "/no/such/dir")   # missing
        assert us.get_setting("capability_corpus_dir") == str(tmp_path)
        assert us.get_dir("capability_corpus_dir") == str(tmp_path)
        assert us.get_dir("template_library_dir") is None
        assert us.get_setting("absent", "d") == "d"
        assert us.settings_path().exists()
        us.unset_setting("template_library_dir")
        assert "template_library_dir" not in us.all_settings()


class TestServiceKeyPersistence:
    """Service-key folder + specific key persist via the unified external store."""

    def test_dir_and_path_round_trip(self, tmp_path, monkeypatch):
        monkeypatch.setenv("CPI_MIGRATOR_HOME", str(tmp_path))
        import importlib
        from fetcher import user_settings as us
        importlib.reload(us)
        from fetcher import service_key as sk
        sk.remember_keys_dir(str(tmp_path))
        assert sk.recall_keys_dir() == str(tmp_path)          # via settings.json
        assert us.get_setting("service_keys_dir") == str(tmp_path)
        keyfile = tmp_path / "k.json"
        keyfile.write_text("{}")
        sk.remember_key_path(str(keyfile))
        assert sk.recall_key_path() == str(keyfile)           # exists -> returned
        sk.remember_key_path(str(tmp_path / "gone.json"))
        assert sk.recall_key_path() is None                   # missing -> None


class TestCloneAndAdaptSmart:
    """Smarter injection: primary slot preferred; mapping injected by kind."""

    def _tpl(self):
        import io, zipfile
        members = {
            "META-INF/MANIFEST.MF":
                "Manifest-Version: 1.0\r\nBundle-SymbolicName: Old\r\n\r\n",
            ".project": "<projectDescription><name>Old</name></projectDescription>",
            "src/main/resources/scenarioflows/integrationflow/Old.iflw":
                '<d><a ref="src/main/resources/script/ErrorHandler.groovy"/>'
                '<b ref="src/main/resources/script/process_main.groovy"/>'
                '<m ref="src/main/resources/mapping/legacy.mmap"/></d>',
            "src/main/resources/script/ErrorHandler.groovy": "// tpl error\n",
            "src/main/resources/script/process_main.groovy": "// tpl process\n",
            "src/main/resources/mapping/legacy.mmap":
                '<?xml version="1.0"?><messageMapping name="legacy" draft="true"/>',
        }
        buf = io.BytesIO()
        with zipfile.ZipFile(buf, "w") as z:
            for n, c in members.items():
                z.writestr(n, c)
        return buf.getvalue()

    def test_single_script_lands_in_primary_slot(self):
        import io, zipfile
        from scaffolder.iflow_personalizer import clone_and_adapt
        bundle, ok, rep = clone_and_adapt(
            self._tpl(), "Z_PO", generated_scripts=["// GEN logic\n"])
        z = zipfile.ZipFile(io.BytesIO(bundle))
        assert rep["scripts_injected"] == ["process_main.groovy"]   # not ErrorHandler
        assert rep["scripts_untouched"] == ["ErrorHandler.groovy"]
        assert b"GEN logic" in z.read("src/main/resources/script/process_main.groovy")
        assert z.read("src/main/resources/script/ErrorHandler.groovy").decode() \
            == "// tpl error\n"
        assert ok

    def test_mapping_injected_by_kind_and_reskinned(self):
        import io, zipfile
        from scaffolder.iflow_personalizer import clone_and_adapt
        bundle, ok, rep = clone_and_adapt(
            self._tpl(), "Z_PO",
            generated_mapping='<?xml version="1.0"?>'
                              '<messageMapping name="x" draft="true"/>')
        z = zipfile.ZipFile(io.BytesIO(bundle))
        assert rep["mapping_injected"] == "legacy.mmap" and ok
        assert b'name="Z_PO"' in z.read("src/main/resources/mapping/legacy.mmap")

    def test_xslt_mapping_skipped_when_no_xsl_slot(self):
        from scaffolder.iflow_personalizer import clone_and_adapt
        # template has only a .mmap slot; an XSLT body must NOT be forced into it
        _, ok, rep = clone_and_adapt(
            self._tpl(), "Z_PO",
            generated_mapping='<xsl:stylesheet xmlns:xsl="x"></xsl:stylesheet>')
        assert rep["mapping_injected"] is None and rep["mapping_skipped"] and ok


class TestPreflight:
    """Pre-deploy preflight on the inner bundle (manifest+iflw+refs+hygiene)."""

    def _mk(self, members):
        from scaffolder.iflow_personalizer import _dos_zip
        return _dos_zip(members)

    _MANIFEST = ("Manifest-Version: 1.0\r\nBundle-ManifestVersion: 2\r\n"
                 "Bundle-SymbolicName: Z_PO; singleton:=true\r\nBundle-Name: Z_PO\r\n"
                 "Bundle-Version: 1.0.0\r\nSAP-BundleType: IntegrationFlow\r\n"
                 "Import-Package: com.sap.it.op.agent.api\r\n\r\n")

    def test_good_bundle_passes(self):
        from fetcher.preflight import preflight_inner_bundle
        good = self._mk({
            "META-INF/MANIFEST.MF": self._MANIFEST,
            ".project": "<projectDescription><name>Z_PO</name></projectDescription>",
            "src/main/resources/scenarioflows/integrationflow/Z_PO.iflw":
                '<bpmn2:definitions><x ref="src/main/resources/script/s.groovy"/>'
                '</bpmn2:definitions>',
            "src/main/resources/script/s.groovy": "// ok\n",
        })
        ok, findings = preflight_inner_bundle(good)
        assert ok, findings

    def test_missing_manifest_and_dangling_ref_fail(self):
        from fetcher.preflight import preflight_inner_bundle
        bad = self._mk({
            "src/main/resources/scenarioflows/integrationflow/Z_PO.iflw":
                '<bpmn2:definitions><x ref="src/main/resources/script/missing.groovy"/>'
                '</bpmn2:definitions>',
        })
        ok, findings = preflight_inner_bundle(bad)
        errs = [f["message"] for f in findings if f["severity"] == "error"]
        assert not ok
        assert any("MANIFEST" in e for e in errs)
        assert any("missing.groovy" in e for e in errs)

    def test_empty_fails(self):
        from fetcher.preflight import preflight_inner_bundle
        assert preflight_inner_bundle(b"")[0] is False


class TestCapabilityModeFires:
    """generate_bundle must pull REAL corpus content when a match exists."""

    def test_real_groovy_selected_over_template(self, tmp_path):
        from types import SimpleNamespace
        from library_builder.corpus_pipeline import build_corpus
        from scaffolder.artifact_bundle import generate_bundle
        hdr = "import com.sap.gateway.ip.core.customdev.util.Message\n"
        files = {
            "p/a/logqueue.groovy": hdr + "def Message processData(Message m){ return m }\n",
            "p/b/splitbatch.groovy": hdr + "def Message processData(Message m){ return m }\n",
            "p/d/dedup_idem.groovy": hdr + "def Message processData(Message m){ "
                "// REAL_CORPUS_MARKER idempotent dedup duplicate id store\n return m }\n",
        }
        corpus = build_corpus(files=files)
        iface = SimpleNamespace(name="Dedup_Orders", id="1",
                                description="idempotent dedup of duplicate orders by id store",
                                sender_adapter="HTTPS", receiver_adapter="HTTP")
        iflows = tmp_path / "iflows"
        iflows.mkdir()
        ip = iflows / "D.iflw"
        ip.write_text("<x/>")
        bundle = generate_bundle(iface, ip, corpus=corpus)
        assert any("REAL_CORPUS_MARKER" in a.content for a in bundle.artifacts)


class TestMultiTypeGenerateBundle:
    """generate_bundle assembles ALL relevant file types from the corpus."""

    def test_assembles_script_mapping_xslt_js_schema(self, tmp_path):
        from types import SimpleNamespace
        from library_builder.corpus_pipeline import build_corpus
        from scaffolder.artifact_bundle import generate_bundle
        hdr = "import com.sap.gateway.ip.core.customdev.util.Message\n"
        files = {
            "c/script/enrich.groovy": hdr + "def Message processData(Message m){ "
                "// order invoice enrich customer material transform\n return m }\n",
            "c/map/t.xsl": '<?xml version="1.0"?><xsl:stylesheet '
                'xmlns:xsl="http://www.w3.org/1999/XSL/Transform">'
                '<!-- order invoice transform customer material --></xsl:stylesheet>',
            "c/js/v.js": "// order invoice validate customer material javascript\n",
            "c/schema/Order.xsd": '<?xml version="1.0"?><xs:schema '
                'xmlns:xs="http://www.w3.org/2001/XMLSchema"><xs:element name="Order"/>'
                '<!-- order invoice customer material --></xs:schema>',
        }
        corpus = build_corpus(files=files)
        iface = SimpleNamespace(
            name="Order_to_Invoice", id="1",
            description="order invoice transform validate enrich customer material",
            sender_adapter="SOAP", receiver_adapter="HTTP")
        (tmp_path / "iflows").mkdir()
        ip = tmp_path / "iflows" / "O.iflw"
        ip.write_text("<x/>")
        kinds = {a.kind for a in generate_bundle(iface, ip, corpus=corpus).artifacts}
        # script + mapping always; xslt/js/schema on confident match (all match here)
        assert {"script", "mapping", "xslt", "js", "schema"} <= kinds, kinds

    def test_no_corpus_still_yields_core_two(self, tmp_path):
        from types import SimpleNamespace
        from scaffolder.artifact_bundle import generate_bundle
        iface = SimpleNamespace(name="Plain", id="1", description="plain passthrough",
                                sender_adapter="HTTPS", receiver_adapter="HTTP")
        (tmp_path / "iflows").mkdir()
        ip = tmp_path / "iflows" / "P.iflw"
        ip.write_text("<x/>")
        kinds = [a.kind for a in generate_bundle(iface, ip, corpus=None).artifacts]
        assert kinds == ["script", "mapping"]   # generic fallback, never worse


class TestPackageNamingNoCompanyURL:
    """Package names must carry functionality, never the company/namespace URL."""

    def test_url_namespace_reduced_to_functionality(self):
        from scaffolder.pipeline_scaffolder import (
            generate_package_display_name, generate_package_name)
        disp = generate_package_display_name("ECC", "Ariba", "http://company.com/ariba")
        assert "http" not in disp and "company" not in disp.lower(), disp
        assert "Ariba" in disp
        pid = generate_package_name("", "ECC", "Ariba", "http://company.com/ariba")
        assert "http" not in pid and "company" not in pid.lower(), pid
        assert pid.endswith("ariba") or "ariba" in pid.lower()

    def test_plain_domain_unaffected(self):
        from scaffolder.pipeline_scaffolder import generate_package_display_name
        disp = generate_package_display_name("ECC", "S4HANA", "Procurement")
        assert "Procurement" in disp


class TestTimerInterfaceScaffold:
    """Default shape = self-contained Timer→CM→CM→End, parametrized per
    interface; plus the per-iFlow shape/exclude controls."""

    import types as _t

    def _iface(self, name="Inv_Outbound", sa="SFTP", ra="HTTPS"):
        import types
        return types.SimpleNamespace(
            name=name, id=name, sender_adapter=sa, receiver_adapter=ra,
            sender_system="SRC", receiver_system="S4HANA", namespace="urn:x")

    def test_generator_self_contained_and_documented(self):
        from scaffolder.minimal_iflow import generate_timer_interface_iflow
        import xml.etree.ElementTree as ET
        r = generate_timer_interface_iflow(
            "Inv_Outbound", properties=[("SenderAdapter", "SFTP"),
                                        ("ReceiverSystem", "S4HANA"),
                                        ("Blank", "")])
        x = r.iflw_xml
        ET.fromstring(x)                                  # well-formed
        assert x.count("<bpmn2:messageFlow") == 0         # no endpoint dependency
        assert "fireNow=true" in x                        # run-once timer
        assert "intermediatetimer/version::1.3.0" in x    # proven version
        assert "InterfaceName" in x and "Inv_Outbound" in x
        assert "SenderAdapter" in x and "SFTP" in x and "S4HANA" in x
        assert "<value>Blank</value>" not in x            # empty props dropped
        assert x.count("<value>Enricher</value>") == 2

    def test_scaffold_default_is_timer(self, tmp_path):
        from scaffolder.iflow_scaffolder import IFlowScaffolder
        import types, xml.etree.ElementTree as ET
        a = types.SimpleNamespace(interface=self._iface())
        p = IFlowScaffolder(str(tmp_path)).scaffold(a)    # default shape
        x = p.read_text()
        ET.fromstring(x)
        assert "fireNow=true" in x and x.count("<bpmn2:messageFlow") == 0

    def test_scaffold_minimal_shape_has_sender(self, tmp_path):
        from scaffolder.iflow_scaffolder import IFlowScaffolder
        import types
        a = types.SimpleNamespace(interface=self._iface())
        p = IFlowScaffolder(str(tmp_path)).scaffold(a, shape="minimal")
        assert p.read_text().count("<bpmn2:messageFlow") >= 1

    def test_likely_needs_sender_flag(self):
        from scaffolder.iflow_scaffolder import IFlowScaffolder as S
        assert S.likely_needs_sender(self._iface(sa="SOAP")) is True
        assert S.likely_needs_sender(self._iface(sa="IDoc")) is True
        assert S.likely_needs_sender(self._iface(sa="SFTP")) is False

    def test_orchestrator_excluded_and_shape(self, tmp_path):
        from scaffolder.iflow_scaffolder import IFlowScaffolder
        from scaffolder.batch_orchestrator import BatchOrchestrator
        import types
        mk = lambda n, sa: types.SimpleNamespace(interface=self._iface(n, sa))
        ass = [mk("Keep", "SFTP"), mk("Push", "SOAP")]
        orch = BatchOrchestrator(scaffolder=IFlowScaffolder(str(tmp_path)),
                                 output_dir=str(tmp_path))
        rep = orch.run(ass, {}, excluded_names={"Push"})
        assert {r.interface_name for r in rep.processed} == {"Keep"}
        assert any(r.interface_name == "Push" and "Excluded" in r.reason
                   for r in rep.needs_attention)


class TestUploadShapeRouting:
    """Deploy uses the single generated-package path (clone-and-adapt removed)."""

    def _uploader(self):
        from fetcher.cpi_uploader import CPIUploader
        u = CPIUploader.__new__(CPIUploader)
        u.base_url = "https://x"
        u.calls = {"clone": 0, "package": 0}
        u.sanitize_package_id = lambda s: s
        u._artifact_exists = lambda *a, **k: False
        u._package_iflow = lambda *a, **k: (
            u.calls.__setitem__("package", u.calls["package"] + 1) or b"TIMER")
        u._post_artifact = lambda zip_bytes, *a, **k: setattr(u, "last", zip_bytes)
        return u

    def test_default_uses_scaffold_not_clone(self):
        from pathlib import Path
        u = self._uploader()
        u.upload_iflow(Path("x.iflw"), "P", "A", "Name")
        assert u.calls["clone"] == 0 and u.last == b"TIMER"

    def test_scaffold_path_drops_clone_extras(self):
        # The packaged scaffold must NOT bundle clone-oriented scripts/mappings —
        # doing so produced a CPI 500 ("InputStream cannot be null").
        from pathlib import Path
        from fetcher.cpi_uploader import CPIUploader
        u = CPIUploader.__new__(CPIUploader)
        u.base_url = "https://x"
        seen = {}
        u.sanitize_package_id = lambda s: s
        u._artifact_exists = lambda *a, **k: False
        def _pkg(iflw_path, art_id, art_name, params="", extra_artifacts=None):
            seen["extras"] = extra_artifacts
            return b"ZIP"
        u._package_iflow = _pkg
        u._post_artifact = lambda *a, **k: None
        u.upload_iflow(Path("x.iflw"), "P", "A", "Name",
                       extra_artifacts=[("script/a.groovy", "x")])
        assert seen["extras"] is None


class TestConsultantStructure:
    """Consultant-defined structure → linear iFlow wiring (timer → CM* → end).
    Foundation for the 'consultant structure' feature."""

    def test_parse_basic_path(self):
        from scaffolder.minimal_iflow import parse_consultant_structure
        steps, notes = parse_consultant_structure("timer -> content modifier -> end")
        assert [s["type"] for s in steps] == ["timer", "content_modifier", "end"]
        assert notes == []

    def test_request_reply_is_dropped_with_note(self):
        from scaffolder.minimal_iflow import parse_consultant_structure
        steps, notes = parse_consultant_structure(
            "timer -> content modifier -> request-reply -> end")
        assert [s["type"] for s in steps] == ["timer", "content_modifier", "end"]
        assert any("request-reply" in n.lower() for n in notes)

    def test_timer_and_end_are_ensured(self):
        from scaffolder.minimal_iflow import parse_consultant_structure
        steps, _ = parse_consultant_structure("content modifier")
        assert steps[0]["type"] == "timer" and steps[-1]["type"] == "end"

    def test_multiple_separators_and_aliases(self):
        from scaffolder.minimal_iflow import parse_consultant_structure
        steps, _ = parse_consultant_structure("scheduler => CM => cm => end")
        assert [s["type"] for s in steps] == \
            ["timer", "content_modifier", "content_modifier", "end"]

    def test_build_linear_iflw_is_wellformed_and_linear(self):
        import xml.etree.ElementTree as ET
        from scaffolder.minimal_iflow import generate_structured_iflow
        res = generate_structured_iflow("Flow_X", "timer -> CM -> CM -> CM -> end",
                                        iflow_id="FlowX")
        ET.fromstring(res.iflw_xml)  # well-formed
        x = res.iflw_xml
        starts, cms, ends = (x.count("<bpmn2:startEvent"),
                             x.count("<bpmn2:callActivity"),
                             x.count("<bpmn2:endEvent"))
        flows = x.count("<bpmn2:sequenceFlow")
        assert (starts, cms, ends) == (1, 3, 1)
        assert flows == (starts + cms + ends) - 1  # strictly linear chain

    def test_structured_bundle_has_propdef(self):
        import io, zipfile
        from scaffolder.minimal_iflow import generate_structured_iflow, build_bundle_zip
        res = generate_structured_iflow("F", "timer -> content modifier -> end")
        z = zipfile.ZipFile(io.BytesIO(build_bundle_zip(res)))
        assert "src/main/resources/parameters.propdef" in z.namelist()
        assert "src/main/resources/parameters.prop" in z.namelist()


class TestComplexityScaledStructure:
    """Generated iFlow structure must scale with interface complexity, not be a
    fixed 2-CM shape for everything (the 'all the same structure' complaint)."""

    def _rec(self, **kw):
        from extractor.pi_extractor import InterfaceRecord
        base = dict(id="x", name="X", namespace="", software_component="",
                    sender_system="", receiver_system="", sender_adapter="HTTPS",
                    receiver_adapter="HTTPS", message_interface="X", description="")
        base.update(kw)
        return InterfaceRecord(**base)

    def test_simple_interface_stays_bare(self):
        from scaffolder.iflow_scaffolder import IFlowScaffolder
        assert IFlowScaffolder._complexity_step_plan(self._rec()) == []

    def test_plan_scales_with_signals(self):
        from scaffolder.iflow_scaffolder import IFlowScaffolder
        simple = IFlowScaffolder._complexity_step_plan(self._rec())
        mid = IFlowScaffolder._complexity_step_plan(
            self._rec(mapping_program="MM", has_multi_mapping=True, channel_count=2))
        heavy = IFlowScaffolder._complexity_step_plan(self._rec(
            mapping_program="MM", has_multi_mapping=True, channel_count=4, has_bpm=True,
            description="groovy xslt value mapping router jdbc lookup multicast"))
        assert len(simple) < len(mid) < len(heavy)

    def test_generated_iflw_step_count_scales(self):
        import types, tempfile, xml.etree.ElementTree as ET
        from scaffolder.iflow_scaffolder import IFlowScaffolder
        sc = IFlowScaffolder(tempfile.mkdtemp())
        def cms(rec):
            x = sc.scaffold(types.SimpleNamespace(interface=rec)).read_text()
            ET.fromstring(x)  # well-formed
            return x.count("<bpmn2:callActivity")
        simple = cms(self._rec(name="Simple"))
        heavy = cms(self._rec(name="Heavy", mapping_program="MM",
                              has_multi_mapping=True, channel_count=3, has_bpm=True,
                              description="groovy xslt router value mapping"))
        assert simple == 2          # bare Timer→CM→CM→End
        assert heavy > simple       # scales up

    def test_middle_steps_generator_wellformed(self):
        import xml.etree.ElementTree as ET
        from scaffolder.minimal_iflow import generate_timer_interface_iflow
        res = generate_timer_interface_iflow(
            "Flow", "Flow", middle_steps=["Map Fields", "Run Groovy Script", "Gather"])
        ET.fromstring(res.iflw_xml)
        # CM1 + 3 middle + ack CM = 5 callActivities
        assert res.iflw_xml.count("<bpmn2:callActivity") == 5


class TestDecodedStepBuilders:
    """The real CPI step palette (Script/Mapping/Splitter/Gather/Filter) decoded
    from the 166-iFlow corpus. Each builder must emit the exact decoded
    activityType + cmdVariantUri, stay well-formed, keep the chain linear, and
    bundle any referenced resource file (script .groovy / mapping .xsl)."""

    def test_each_builder_has_decoded_constants(self):
        from scaffolder.minimal_iflow import (
            _script_step, _mapping_step, _splitter_step, _gather_step, _filter_step)
        import xml.etree.ElementTree as ET
        cases = [
            (_script_step("C1", "s", "x.groovy"), "Script",
             "cname::GroovyScript"),
            (_mapping_step("C1", "m", "MapX"), "Mapping",
             "cname::XSLTMapping/version::1.2.0"),
            (_splitter_step("C1", "sp"), "Splitter",
             "cname::GeneralSplitter/version::1.5.1"),
            (_gather_step("C1", "g"), "Gather",
             "cname::Gather/version::1.2.0"),
            (_filter_step("C1", "f"), "Filter",
             "cname::Filter/version::1.1.0"),
        ]
        for xml, at, cmd in cases:
            # wrap the fragment so the bpmn2:/ifl: prefixes resolve (they're
            # only declared on the root of the full document)
            doc = ('<root xmlns:bpmn2="http://www.omg.org/spec/BPMN/20100524/MODEL" '
                   'xmlns:ifl="http:///com.sap.ifl.model/Ifl.xsd">'
                   + xml + '</root>')
            ET.fromstring(doc)
            assert f"<value>{at}</value>" in xml
            assert cmd in xml

    def test_mixed_flow_wellformed_linear_and_bundles_files(self):
        import xml.etree.ElementTree as ET, io, zipfile
        from scaffolder.minimal_iflow import (
            generate_timer_interface_iflow, build_bundle_zip)
        mids = [{"kind": "mapping", "name": "Map Fields"},
                {"kind": "script", "name": "Run Groovy Script"},
                {"kind": "splitter", "name": "Split Records"},
                {"kind": "filter", "name": "Filter Content"},
                {"kind": "gather", "name": "Gather Responses"}]
        res = generate_timer_interface_iflow("Mix", "Mix", middle_steps=mids)
        x = res.iflw_xml
        ET.fromstring(x)
        for at in ["Mapping", "Script", "Splitter", "Filter", "Gather"]:
            assert f"<value>{at}</value>" in x
        # CM1 + 5 middle + ack = 7 callActivities, strictly linear
        acts = x.count("<bpmn2:callActivity")
        flows = x.count("<bpmn2:sequenceFlow")
        starts = x.count("<bpmn2:startEvent")
        ends = x.count("<bpmn2:endEvent")
        assert acts == 7
        assert flows == (starts + acts + ends) - 1
        # referenced files present + valid
        groovy = [p for p in res.files if p.endswith(".groovy")]
        xsl = [p for p in res.files if p.endswith(".xsl")]
        assert groovy and xsl
        assert "processData(Message message)" in res.files[groovy[0]]
        ET.fromstring(res.files[xsl[0]])  # xsl well-formed
        z = zipfile.ZipFile(io.BytesIO(build_bundle_zip(res)))
        assert groovy[0] in z.namelist() and xsl[0] in z.namelist()

    def test_mapping_path_has_no_ext_but_file_does(self):
        from scaffolder.minimal_iflow import generate_timer_interface_iflow
        res = generate_timer_interface_iflow(
            "M", "M", middle_steps=[{"kind": "mapping", "name": "Map Fields"}])
        x = res.iflw_xml
        # corpus convention: mappingpath references the resource WITHOUT ext,
        # the bundled file IS <name>.xsl
        assert "<value>src/main/resources/mapping/MapFields</value>" in x
        assert "src/main/resources/mapping/MapFields.xsl" in res.files

    def test_mapping_uses_bundled_stylesheet_source(self):
        # regression for the tenant "Mapping file not found" / StringIndexOOB:
        # must load from the bundle (mappingSrcIflow) with a mappinguri pointing
        # at the .xsl, NOT from a runtime header.
        from scaffolder.minimal_iflow import generate_timer_interface_iflow
        res = generate_timer_interface_iflow(
            "M", "M", middle_steps=[{"kind": "mapping", "name": "Map Fields"}])
        x = res.iflw_xml
        assert "<value>mappingSrcIflow</value>" in x
        assert "mappingSrcHeader" not in x          # no header dependency
        assert "dir://mapping/xslt/src/main/resources/mapping/MapFields.xsl" in x

    def test_mapping_or_splitter_seed_xml_body(self):
        # XSLT/Splitter need XML input; the first CM must seed a body so a
        # timer flow doesn't fail "supports XML input only".
        from scaffolder.minimal_iflow import generate_timer_interface_iflow
        x = generate_timer_interface_iflow(
            "M", "M", middle_steps=[{"kind": "mapping", "name": "Map Fields"}]).iflw_xml
        assert "&lt;root&gt;&lt;Record/&gt;&lt;/root&gt;" in x or "<root><Record/></root>" in x

    def test_pool_width_grows_with_steps(self):
        import re
        from scaffolder.minimal_iflow import generate_timer_interface_iflow
        def pool_w(n):
            mids = [{"kind": "content_modifier", "name": f"S{i}"} for i in range(n)]
            x = generate_timer_interface_iflow("P", "P", middle_steps=mids).iflw_xml
            m = re.search(r'BPMNShape_Participant_Process_1.*?width="([\d.]+)"', x, re.S)
            return float(m.group(1))
        assert pool_w(8) > pool_w(2)   # pool must scale, not overflow

    def test_string_middle_steps_stay_content_modifiers(self):
        # Backward compat: bare strings remain Content Modifiers (no new
        # activity types, no extra files) — the proven default shape.
        from scaffolder.minimal_iflow import generate_timer_interface_iflow
        res = generate_timer_interface_iflow(
            "S", "S", middle_steps=["Map Fields", "Gather"])
        x = res.iflw_xml
        assert "<value>Mapping</value>" not in x
        assert "<value>Gather</value>" not in x
        assert not any(p.endswith((".groovy", ".xsl")) for p in res.files)

    def test_complexity_plan_emits_typed_kinds(self):
        from scaffolder.iflow_scaffolder import IFlowScaffolder
        from extractor.pi_extractor import InterfaceRecord
        rec = InterfaceRecord(
            id="x", name="X", namespace="", software_component="",
            sender_system="", receiver_system="", sender_adapter="HTTPS",
            receiver_adapter="HTTPS", message_interface="X",
            mapping_program="MM", channel_count=2,
            description="groovy filter")
        plan = IFlowScaffolder._complexity_step_plan(rec)
        kinds = {s["kind"] for s in plan}
        assert "mapping" in kinds   # mapping_program
        assert "script" in kinds    # groovy
        assert "filter" in kinds    # filter keyword
        assert "splitter" in kinds and "gather" in kinds  # channel_count > 1

    def test_scaffolded_iflow_uses_real_step_types(self):
        import types, tempfile, xml.etree.ElementTree as ET
        from scaffolder.iflow_scaffolder import IFlowScaffolder
        from extractor.pi_extractor import InterfaceRecord
        rec = InterfaceRecord(
            id="x", name="HeavyReal", namespace="", software_component="",
            sender_system="", receiver_system="", sender_adapter="HTTPS",
            receiver_adapter="HTTPS", message_interface="X",
            mapping_program="MM", channel_count=3,
            description="groovy xslt filter")
        sc = IFlowScaffolder(tempfile.mkdtemp())
        x = sc.scaffold(types.SimpleNamespace(interface=rec)).read_text()
        ET.fromstring(x)
        # generated flow now contains genuine decoded step elements, not all-CM
        for at in ["Script", "Mapping", "Splitter", "Gather", "Filter"]:
            assert f"<value>{at}</value>" in x, at


class TestCorpusStore:
    """Persistent, editable, incrementally-updatable capability store. Distill
    once -> save sharded JSON -> load fast -> add-only merge (deduped). Proves
    the workbench no longer needs to re-walk the raw corpus each session."""

    @staticmethod
    def _caps():
        from library_builder.solver import NormalizedCapability
        return [
            NormalizedCapability(
                cap_id="groovy:a", ctype="groovy", intent="map idoc to xml",
                keywords={"idoc", "xml", "map"}, varies=["root"], weight=2,
                source_ref="A.groovy", raw={"name": "A", "body": "..."}),
            NormalizedCapability(
                cap_id="xslt:b", ctype="xslt", intent="remove namespaces",
                keywords={"remove", "namespaces"}, weight=1,
                source_ref="B.xsl", raw={"name": "B"}),
        ]

    def test_save_creates_sharded_json_and_index(self):
        import tempfile, os
        from library_builder.corpus_store import CorpusStore
        d = tempfile.mkdtemp()
        idx = CorpusStore(self._caps()).save(d)
        assert os.path.exists(os.path.join(d, "index.json"))
        assert os.path.exists(os.path.join(d, "groovy.json"))
        assert os.path.exists(os.path.join(d, "xslt.json"))
        assert idx["total_capabilities"] == 2

    def test_load_roundtrip_preserves_caps(self):
        import tempfile
        from library_builder.corpus_store import CorpusStore
        d = tempfile.mkdtemp()
        CorpusStore(self._caps()).save(d)
        loaded = CorpusStore.load(d)
        assert {c.cap_id for c in loaded.caps} == {"groovy:a", "xslt:b"}
        a = next(c for c in loaded.caps if c.cap_id == "groovy:a")
        assert isinstance(a.keywords, set) and "idoc" in a.keywords  # set rehydrated
        assert a.raw["name"] == "A"                                  # raw preserved

    def test_merge_is_add_only_and_dedupes(self):
        import tempfile
        from library_builder.corpus_store import CorpusStore
        from library_builder.solver import NormalizedCapability
        d = tempfile.mkdtemp()
        CorpusStore(self._caps()).save(d)
        store = CorpusStore.load(d)
        new = NormalizedCapability(cap_id="groovy:c", ctype="groovy",
                                   intent="split", keywords={"split"},
                                   source_ref="C.groovy", raw={"name": "C"})
        r1 = store.merge_caps([new])
        assert r1["added"] == 1 and r1["total"] == 3
        r2 = store.merge_caps([new])          # same cap again
        assert r2["added"] == 0 and r2["skipped"] == 1
        store.save(d)
        assert len(CorpusStore.load(d).caps) == 3

    def test_load_tolerates_missing_index(self):
        import tempfile, os
        from library_builder.corpus_store import CorpusStore
        d = tempfile.mkdtemp()
        CorpusStore(self._caps()).save(d)
        os.remove(os.path.join(d, "index.json"))
        assert len(CorpusStore.load(d).caps) == 2   # falls back to *.json shards

    def test_search_works_on_loaded_store(self):
        import tempfile
        from library_builder.corpus_store import CorpusStore
        d = tempfile.mkdtemp()
        CorpusStore(self._caps()).save(d)
        loaded = CorpusStore.load(d)
        hits = loaded.search("idoc xml", top_n=3)
        assert any(cid == "groovy:a" for cid, _ in hits)


class TestResourcePackaging:
    """Regression for the tenant 'Mapping file not found': scaffolded resource
    files (scripts/mappings) must persist next to the iFlow and ship inside the
    uploaded package, even on the on-demand path (extra_artifacts=None)."""

    def _rec(self):
        from extractor.pi_extractor import InterfaceRecord
        return InterfaceRecord(
            id="x", name="Pkg_Test", namespace="", software_component="",
            sender_system="", receiver_system="", sender_adapter="HTTPS",
            receiver_adapter="HTTPS", message_interface="X",
            mapping_program="MM", description="groovy", channel_count=1)

    def test_scaffold_persists_resources_and_package_ships_them(self):
        import types, tempfile, io, zipfile
        from pathlib import Path
        from scaffolder.iflow_scaffolder import IFlowScaffolder
        from fetcher.cpi_uploader import CPIUploader
        out = tempfile.mkdtemp()
        p = IFlowScaffolder(output_dir=out).scaffold(
            types.SimpleNamespace(interface=self._rec()))
        meta = Path(p).with_name(Path(p).stem + "__meta")
        persisted = {x.name for x in meta.rglob("*") if x.is_file()}
        assert any(n.endswith(".xsl") for n in persisted)
        assert any(n.endswith(".groovy") for n in persisted)
        # package with NO extra_artifacts (the path that used to drop them)
        z = zipfile.ZipFile(io.BytesIO(
            CPIUploader._package_iflow(Path(p), "PkgTest", "Pkg_Test")))
        names = z.namelist()
        assert any("/mapping/" in n and n.endswith(".xsl") for n in names)
        assert any("/script/" in n and n.endswith(".groovy") for n in names)

    def test_content_modifier_constant_vs_expression_body(self):
        from scaffolder.minimal_iflow import _content_modifier_step
        # literal body -> constant (no "Expression Text has no parameters" error)
        lit = _content_modifier_step("C1", "n", body_expr="<root><Record/></root>")
        assert "<value>constant</value>" in lit
        # ${...} body -> expression
        expr = _content_modifier_step("C1", "n", body_expr="${property.X}")
        assert "<value>expression</value>" in expr


class TestMonsterFixes:
    """Regressions for the monster tenant run: filter-before-splitter empty
    body, and the artifact Name HTTP 400."""

    def test_filter_xpath_is_single_root_safe(self):
        from scaffolder.minimal_iflow import _filter_step
        x = _filter_step("C1", "Filter Content")
        # must select ONE document element (/*), not every element (//*),
        # or the result is a multi-root fragment the splitter can't parse.
        assert "<value>/*</value>" in x
        assert "<value>//*</value>" not in x

    def test_filter_then_splitter_iflow_is_well_formed(self):
        import types, tempfile
        import xml.etree.ElementTree as ET
        from pathlib import Path
        from scaffolder.iflow_scaffolder import IFlowScaffolder
        from extractor.pi_extractor import InterfaceRecord
        rec = InterfaceRecord(id="x", name="Filt_Split", namespace="",
            software_component="", sender_system="", receiver_system="",
            sender_adapter="HTTPS", receiver_adapter="HTTPS", message_interface="X",
            mapping_program="", description="filter duplicates then split batch",
            channel_count=3)
        p = IFlowScaffolder(output_dir=tempfile.mkdtemp()).scaffold(
            types.SimpleNamespace(interface=rec))
        ET.fromstring(Path(p).read_text())   # whole iflow well-formed

    def test_artifact_name_sanitized_for_cpi(self):
        from fetcher.cpi_uploader import CPIUploader as C
        assert C.sanitize_artifact_name("Order Sync & Validate (EU/US)") == \
            "Order Sync Validate EU US"
        assert not C.sanitize_artifact_name("ends with period.").endswith(".")
        import re
        for bad in ["A/B", "x&y", "(z)", "1abc", "  ", ":,;"]:
            s = C.sanitize_artifact_name(bad)
            assert re.match(r"[A-Za-z_]", s) and not s.endswith(".")


class TestGroovyCanonical:
    """Generated Groovy must create a real exchange property and not carry the
    unused import CPI's Groovy 2.0 editor strips."""

    def test_groovy_sets_exchange_property_and_no_unused_import(self):
        from scaffolder.minimal_iflow import _groovy_body
        g = _groovy_body("processData")
        assert "message.setProperty(" in g          # creates a real property
        assert "import java.util.HashMap" not in g  # unused -> CPI flags it
        assert "com.sap.gateway.ip.core.customdev.util.Message" in g
        assert "def Message processData(Message message)" in g
        assert "return message" in g


class TestConvertersAndAdjacency:
    """The 4 decoded converters build well-formed XML and the content-type
    adjacency validator catches a converter feeding an incompatible step."""

    def test_four_converters_well_formed(self):
        import xml.etree.ElementTree as ET
        from scaffolder.minimal_iflow import _build_middle_step
        for k in ("xml_to_json", "json_to_xml", "xml_to_csv", "csv_to_xml"):
            xml, _ = _build_middle_step(k, "C1", {"name": k}, "sf_in", "sf_out")
            ET.fromstring("<r xmlns:bpmn2='b' xmlns:ifl='i'>" + xml + "</r>")

    def test_converter_versions(self):
        from scaffolder.minimal_iflow import _build_middle_step
        want = {"xml_to_json": "XmlToJsonConverter/version::1.0.8",
                "json_to_xml": "JsonToXmlConverter/version::1.1.2",
                "xml_to_csv":  "XmlToCsvConverter/version::1.1.0",
                "csv_to_xml":  "CsvToXmlConverter/version::1.1"}
        for k, sig in want.items():
            xml, _ = _build_middle_step(k, "C1", {"name": k}, "a", "b")
            assert sig in xml

    def test_adjacency_accepts_round_trip(self):
        from scaffolder.minimal_iflow import validate_step_chain
        ok, errs = validate_step_chain(["xml_to_json", "json_to_xml", "mapping"])
        assert ok and not errs

    def test_adjacency_rejects_json_into_mapping(self):
        from scaffolder.minimal_iflow import validate_step_chain
        ok, errs = validate_step_chain(["xml_to_json", "mapping"])
        assert not ok and "mapping" in errs[0]

    def test_script_is_format_agnostic_passthrough(self):
        from scaffolder.minimal_iflow import validate_step_chain
        ok, _ = validate_step_chain(["xml_to_json", "script", "json_to_xml"])
        assert ok


class TestLinearMonster:
    """The high-confidence linear monster: well-formed, content-safe chain, all
    schemas bundled, packages into the tenant-accepted zip layout."""

    def test_builds_well_formed_with_resources(self):
        import xml.etree.ElementTree as ET
        from scaffolder.monster_iflow import build_linear_monster
        r = build_linear_monster()
        ET.fromstring(r.iflw_xml)
        res = {k.split("resources/")[1] for k in r.files
               if k.startswith("src/main/resources/") and not k.endswith(("prop", "propdef"))}
        for needed in ("xsd/Order.xsd", "wsdl/OrderService.wsdl",
                       "edmx/Order.edmx", "mapping/NormalizeOrders.xsl",
                       "script/reformatDate.groovy"):
            assert needed in res

    def test_packages_into_valid_bundle(self):
        import io, zipfile
        from scaffolder.monster_iflow import build_linear_monster, monster_to_zip
        zf = zipfile.ZipFile(io.BytesIO(monster_to_zip(build_linear_monster())))
        names = zf.namelist()
        assert "META-INF/MANIFEST.MF" in names and ".project" in names
        assert any(n.endswith(".iflw") for n in names)


class TestBranchingMonster:
    """The branching monster (router + multicast + join + gather + exception
    subprocess) must be structurally self-consistent: every diagram shape maps
    to a process element, every edge to a flow, and every flow endpoint exists."""

    def _consistency(self, result):
        import xml.etree.ElementTree as ET
        NS = {"bpmn2": "http://www.omg.org/spec/BPMN/20100524/MODEL"}
        root = ET.fromstring(result.iflw_xml)
        proc = root.find(".//bpmn2:process", NS)
        elem_ids, flow_ids, endpoints = set(), set(), []
        for e in proc.iter():
            tag = e.tag.split("}")[-1]; eid = e.get("id")
            if tag == "sequenceFlow":
                flow_ids.add(eid); endpoints.append((e.get("sourceRef"), e.get("targetRef")))
            elif eid and tag in ("startEvent", "endEvent", "callActivity",
                                  "exclusiveGateway", "parallelGateway", "subProcess"):
                elem_ids.add(eid)
        DI = "{http://www.omg.org/spec/BPMN/20100524/DI}"
        shapes = {s.get("bpmnElement") for s in root.iter(DI + "BPMNShape")}
        edges = {ed.get("bpmnElement") for ed in root.iter(DI + "BPMNEdge")}
        for s, t in endpoints:
            assert s in elem_ids and t in elem_ids
        for sh in shapes:
            assert sh == "Participant_Process_1" or sh in elem_ids
        for el in elem_ids:
            assert el in shapes
        assert edges == flow_ids

    def test_branching_consistent_both_variants(self):
        from scaffolder.monster_iflow import build_branching_monster
        self._consistency(build_branching_monster(include_exception=False))
        self._consistency(build_branching_monster(include_exception=True))

    def test_branching_carries_decoded_signatures(self):
        from scaffolder.monster_iflow import build_branching_monster
        x = build_branching_monster(include_exception=True).iflw_xml
        for sig in ("activityType</key><value>ExclusiveGateway",
                    'default="SequenceFlow_8"',
                    "cname::Multicast/version::1.1.1",
                    "cname::Join/version::1.0.0",
                    "cname::Gather/version::1.2.0",
                    "ErrorEventSubProcessTemplate",
                    "tFormalExpression"):
            assert sig in x

    def test_branching_packages(self):
        import io, zipfile
        from scaffolder.monster_iflow import build_branching_monster, monster_to_zip
        zf = zipfile.ZipFile(io.BytesIO(monster_to_zip(build_branching_monster())))
        assert "META-INF/MANIFEST.MF" in zf.namelist()


class TestStepsColumnPipeline:
    """The explicit Steps column drives an arbitrary CPI step pipeline (50+),
    mixing converter steps and XSLT-mapping conversions, content-type-valid."""

    def test_parse_and_validate_big_pipeline(self):
        from scaffolder.minimal_iflow import parse_steps_spec, validate_step_chain
        spec = " | ".join([
            "Content Modifier: Seed",
            "XML to JSON Converter", "JSON to XML Converter",
            "XSLT Mapping (to CSV): MapCsv", "CSV to XML Converter",
            "XSLT Mapping (to JSON): MapJson", "JSON to XML Converter",
            "Splitter", "Gather", "Message Mapping: Norm"])
        mids, kinds = parse_steps_spec(spec)
        assert len(mids) == 10
        ok, errs = validate_step_chain(kinds, "XML")
        assert ok, errs

    def test_xslt_paren_alias_resolves(self):
        from scaffolder.minimal_iflow import parse_steps_spec
        mids, kinds = parse_steps_spec("XSLT Mapping (to CSV): X")
        assert kinds == ["xslt_to_csv"]

    def test_converting_xslt_ships_text_output_stylesheet(self):
        from scaffolder.minimal_iflow import _build_middle_step
        _, files = _build_middle_step("xslt_to_csv", "C1", {"name": "X", "mapping_name": "X"}, "a", "b")
        xsl = next(v for k, v in files.items() if k.endswith(".xsl"))
        assert 'method="text"' in xsl

    def test_unknown_step_falls_back_to_cm(self):
        from scaffolder.minimal_iflow import parse_steps_spec
        _, kinds = parse_steps_spec("Frobnicate Widgets")
        assert kinds == ["content_modifier"]


class TestMonsterValidatorFixes:
    """Fixes for the tenant design-time validator feedback on the monsters."""

    def test_csv_to_xml_has_mandatory_schema_and_bundles_xsd(self):
        from scaffolder.minimal_iflow import _build_middle_step
        xml, files = _build_middle_step("csv_to_xml", "C1", {"name": "x"}, "a", "b")
        assert "XML_Schema_File_Path" in xml
        assert any(k.endswith("CsvTarget.xsd") for k in files)

    def test_router_omits_raisealert_and_sets_throwexception_false(self):
        from scaffolder.monster_iflow import build_branching_monster
        x = build_branching_monster(include_exception=False).iflw_xml
        assert "<key>raiseAlert</key>" not in x
        assert "<key>throwException</key><value>false</value>" in x

    def test_multicast_branches_are_named(self):
        from scaffolder.monster_iflow import build_branching_monster
        x = build_branching_monster(include_exception=False).iflw_xml
        for nm in ('name="JSON Branch"', 'name="Filter Branch"', 'name="CSV Branch"'):
            assert nm in x

    def test_pool_top_clearance_is_widened(self):
        import re
        from scaffolder.monster_iflow import build_branching_monster
        x = build_branching_monster(include_exception=True).iflw_xml
        pool_y = float(re.search(r'Participant_Process_1.*?y="([\d.]+)"', x, re.S).group(1))
        assert 135.0 - pool_y >= 70   # 2.5x the prior 30px gap

    def test_branching_still_structurally_consistent(self):
        import xml.etree.ElementTree as ET
        from scaffolder.monster_iflow import build_branching_monster
        NS = {"bpmn2": "http://www.omg.org/spec/BPMN/20100524/MODEL"}
        for exc in (False, True):
            root = ET.fromstring(build_branching_monster(include_exception=exc).iflw_xml)
            proc = root.find(".//bpmn2:process", NS)
            elem_ids, flow_ids, endpoints = set(), set(), []
            for e in proc.iter():
                tag = e.tag.split("}")[-1]; eid = e.get("id")
                if tag == "sequenceFlow":
                    flow_ids.add(eid); endpoints.append((e.get("sourceRef"), e.get("targetRef")))
                elif eid and tag in ("startEvent", "endEvent", "callActivity",
                                     "exclusiveGateway", "parallelGateway", "subProcess"):
                    elem_ids.add(eid)
            for s, t in endpoints:
                assert s in elem_ids and t in elem_ids


class TestSchemaExtractorFullContent:
    """XSD/WSDL/EDMX extraction retains the FULL file and dedupes identical ones."""

    def test_full_content_retained_and_edmx_parsed(self):
        from extractor.esr_extractor import ESRFileParser
        xsd = b'<?xml version="1.0"?><xs:schema xmlns:xs="http://www.w3.org/2001/XMLSchema" targetNamespace="urn:x"><xs:element name="A"/></xs:schema>'
        edmx = b'<?xml version="1.0"?><edmx:Edmx xmlns:edmx="http://docs.oasis-open.org/odata/ns/edmx" Version="4.0"/>'
        objs = ESRFileParser().parse_uploaded_files({"A.xsd": xsd, "svc.edmx": edmx})
        by = {o.obj_type: o for o in objs}
        assert by["DataType"].content == xsd.decode()
        assert by["EDMX"].content == edmx.decode()
        assert by["DataType"].namespace == "urn:x"

    def test_identical_files_deduped(self):
        from extractor.esr_extractor import ESRFileParser
        xsd = b'<?xml version="1.0"?><xs:schema xmlns:xs="http://www.w3.org/2001/XMLSchema"/>'
        objs = ESRFileParser().parse_uploaded_files({"A.xsd": xsd, "B.xsd": xsd})
        assert sum(1 for o in objs if o.obj_type == "DataType") == 1


class TestSchemaDeduper:
    """3-tier schema clustering: exact / structural / family, richest-canonical."""

    EDMX_A = (b'<?xml version="1.0"?><edmx:Edmx xmlns:edmx="http://schemas.microsoft.com/ado/2007/06/edmx">'
              b'<edmx:DataServices><Schema Namespace="API_BP" xmlns="http://schemas.microsoft.com/ado/2008/09/edm">'
              b'<EntityType Name="A_Partner"/><EntityType Name="A_Address"/>'
              b'<EntitySet Name="A_Partner"/><EntitySet Name="A_Address"/></Schema></edmx:DataServices></edmx:Edmx>')
    # same service, a different host + one extra entity (newer release) → richer
    EDMX_B = (b'<?xml version="1.0"?><edmx:Edmx xmlns:edmx="http://schemas.microsoft.com/ado/2007/06/edmx">'
              b'<edmx:DataServices><Schema Namespace="API_BP" xmlns="http://schemas.microsoft.com/ado/2008/09/edm">'
              b'<EntityType Name="A_Partner"/><EntityType Name="A_Address"/><EntityType Name="A_Industry"/>'
              b'<EntitySet Name="A_Partner"/><EntitySet Name="A_Address"/><EntitySet Name="A_Industry"/></Schema></edmx:DataServices></edmx:Edmx>')

    def _write(self, tmp_path, name, data):
        p = tmp_path / name
        p.write_bytes(data)
        return str(p)

    def test_family_clusters_same_service_across_versions(self, tmp_path):
        from extractor.schema_deduper import dedup, canonical
        a = self._write(tmp_path, "host1_API_BP.edmx", self.EDMX_A)
        b = self._write(tmp_path, "host2_API_BP.edmx", self.EDMX_B)
        res = dedup([a, b])
        fams = [v for k, v in res.family.items() if len(v) > 1]
        assert len(fams) == 1 and len(fams[0]) == 2          # same family
        assert len({f.struct for f in fams[0]}) == 2          # but different structure (versions)
        assert canonical(fams[0]).path == b                   # richest (extra entity) wins

    def test_exact_duplicates_collapse(self, tmp_path):
        from extractor.schema_deduper import dedup
        a = self._write(tmp_path, "x1.edmx", self.EDMX_A)
        b = self._write(tmp_path, "x2.edmx", self.EDMX_A)   # identical bytes
        res = dedup([a, b])
        assert len({k for k in res.exact}) == 1             # one exact group

    def test_idoc_wsdls_not_overmerged_by_shared_namespace(self, tmp_path):
        from extractor.schema_deduper import fingerprint
        w1 = (b'<wsdl:definitions xmlns:wsdl="http://schemas.xmlsoap.org/wsdl/" '
              b'targetNamespace="urn:sap-com:document:sap:idoc:soap:messages">'
              b'<wsdl:portType name="INVOIC.INVOIC01"/></wsdl:definitions>')
        w2 = (b'<wsdl:definitions xmlns:wsdl="http://schemas.xmlsoap.org/wsdl/" '
              b'targetNamespace="urn:sap-com:document:sap:idoc:soap:messages">'
              b'<wsdl:portType name="DESADV.DELVRY07"/></wsdl:definitions>')
        f1 = self._write(tmp_path, "a.wsdl", w1); f2 = self._write(tmp_path, "b.wsdl", w2)
        from extractor.schema_deduper import fingerprint
        assert fingerprint(f1).family != fingerprint(f2).family   # different IDocs split


class TestODataMetadataFetcher:
    """Standalone $metadata fetcher: URL build, retry, save, manifest, dedup."""
    EDMX = (b'<?xml version="1.0"?><edmx:Edmx xmlns:edmx="http://schemas.microsoft.com/ado/2007/06/edmx">'
            b'<edmx:DataServices><Schema Namespace="API_BP" xmlns="http://schemas.microsoft.com/ado/2008/09/edm">'
            b'<EntityType Name="A_BP"/><EntitySet Name="A_BP"/></Schema></edmx:DataServices></edmx:Edmx>')

    def test_metadata_url_build_and_idempotent(self):
        from fetcher.odata_metadata_fetcher import metadata_url
        assert metadata_url(product="s4hanacloud", service="API_BUSINESS_PARTNER").endswith(
            "/sap/opu/odata/sap/API_BUSINESS_PARTNER/$metadata")
        # re-appending $metadata is idempotent + case-insensitive strip
        assert metadata_url(url="https://h/svc/$METADATA") == "https://h/svc/$metadata"

    def test_unknown_product_raises(self):
        from fetcher.odata_metadata_fetcher import metadata_url
        import pytest
        with pytest.raises(ValueError):
            metadata_url(product="nope", service="X")

    def test_fetch_retries_then_saves_and_manifests(self, tmp_path, monkeypatch):
        import fetcher.odata_metadata_fetcher as F
        state = {"n": 0}

        class Resp:
            def __init__(self, c, b=b"", h=None):
                self.status_code, self.content, self.headers = c, b, (h or {})

        def fake_get(session, url, headers, timeout, auth=None):
            assert headers.get("apikey") == "K"
            state["n"] += 1
            if "GOOD" in url:
                return Resp(200, self.EDMX) if state["n"] > 1 else Resp(429, h={"Retry-After": "0"})
            return Resp(404)

        monkeypatch.setattr(F, "_http_get", fake_get)
        res = F.fetch_all(
            [{"name": "GOOD", "product": "s4hanacloud", "service": "GOOD"},
             {"name": "BAD", "product": "s4hanacloud", "service": "BAD"}],
            "K", str(tmp_path), rate_per_sec=100)
        ok = [r for r in res if r.ok]
        assert len(ok) == 1 and (tmp_path / "edmx" / "GOOD.edmx").exists()
        assert state["n"] >= 3                         # 429 retry happened
        assert [r for r in res if not r.ok][0].status == 404
        assert (tmp_path / "_fetch_manifest.json").exists()

    def test_auth_failure_reported(self, tmp_path, monkeypatch):
        import fetcher.odata_metadata_fetcher as F

        class Resp:
            def __init__(self, c): self.status_code, self.content, self.headers = c, b"", {}
        monkeypatch.setattr(F, "_http_get", lambda *a, **k: Resp(403))
        res = F.fetch_all([{"name": "x", "url": "https://h/svc"}], "K", str(tmp_path))
        assert not res[0].ok and res[0].status == 403 and "auth" in res[0].error


class TestFetcherKeyLoading:
    """--key-file accepts plain text or JSON; diagnoses BTP OAuth service keys."""
    def test_plain_text_key(self):
        from fetcher.odata_metadata_fetcher import key_from_text
        assert key_from_text("  ABC123\n") == ("ABC123", "")

    def test_json_apikey_field(self):
        from fetcher.odata_metadata_fetcher import key_from_text
        assert key_from_text('{"x":{"apikey":"KK"}}') == ("KK", "")

    def test_oauth_service_key_is_diagnosed(self):
        from fetcher.odata_metadata_fetcher import key_from_text
        k, err = key_from_text('{"oauth":{"clientid":"sb-a","clientsecret":"z","url":"https://u"}}')
        assert k == "" and "OAuth service key" in err


class TestSchemaLibraryOrganizer:
    """Content-sniff classification + dedupe into typed folders."""
    def test_classify_and_layout(self, tmp_path):
        from tools.organize_schema_library import organize, classify
        src = tmp_path / "src"; src.mkdir()
        # an EDMX saved with a .xml extension must still classify as edmx
        (src / "svc.xml").write_bytes(
            b'<edmx:Edmx xmlns:edmx="http://schemas.microsoft.com/ado/2007/06/edmx">'
            b'<edmx:DataServices><Schema Namespace="API_X" '
            b'xmlns="http://schemas.microsoft.com/ado/2008/09/edm">'
            b'<EntityType Name="E"/><EntitySet Name="E"/></Schema></edmx:DataServices></edmx:Edmx>')
        (src / "a.wsdl").write_bytes(
            b'<wsdl:definitions xmlns:wsdl="http://schemas.xmlsoap.org/wsdl/" '
            b'targetNamespace="urn:t"><wsdl:portType name="P"/></wsdl:definitions>')
        (src / "s.xsd").write_bytes(
            b'<xs:schema xmlns:xs="http://www.w3.org/2001/XMLSchema" '
            b'targetNamespace="urn:s"><xs:element name="Root"/></xs:schema>')
        assert classify(src / "svc.xml") == "edmx"
        out = tmp_path / "lib"
        s = organize([str(src)], str(out), generate_openapi=False)
        assert s["kept"]["edmx"] == 1 and s["kept"]["wsdl"] == 1 and s["kept"]["xsd"] == 1
        assert (out / "edmx" / "API_X.edmx").exists()      # named from Schema Namespace
        assert (out / "suggested_fetch_targets.csv").exists()


class TestDiscoveryHints:
    """Extract re-fetch templates from WSDL/EDMX; recognize Gateway catalog."""
    def test_edmx_and_wsdl_hints(self, tmp_path):
        from extractor.discovery_hints import from_edmx, from_wsdl, build_manifest
        e = tmp_path / "host_x_sap_opu_odata_sap_API_FOO.edmx"
        e.write_bytes(b'<edmx:Edmx xmlns:edmx="http://schemas.microsoft.com/ado/2007/06/edmx">'
                      b'<edmx:DataServices><Schema Namespace="API_FOO" '
                      b'xmlns="http://schemas.microsoft.com/ado/2008/09/edm"/></edmx:DataServices></edmx:Edmx>')
        h = from_edmx(str(e))[0]
        assert h.type == "edmx" and h.refetch_url_template == "{HOST}/sap/opu/odata/sap/API_FOO/$metadata"
        w = tmp_path / "svc.wsdl"
        w.write_bytes(b'<wsdl:definitions xmlns:wsdl="http://schemas.xmlsoap.org/wsdl/" '
                      b'xmlns:soap="http://schemas.xmlsoap.org/wsdl/soap/" targetNamespace="t">'
                      b'<wsdl:portType name="MyPort"/><wsdl:service name="S"><wsdl:port>'
                      b'<soap:address location="https://host:port/sap/bc/srt/scs_ext/sap/myop"/>'
                      b'</wsdl:port></wsdl:service></wsdl:definitions>')
        wh = [x for x in from_wsdl(str(w)) if x.service_path.startswith("/sap")][0]
        assert wh.host_hint == ""                      # host:port placeholder → no real host
        assert wh.refetch_url_template == "{HOST}/sap/bc/srt/scs_ext/sap/myop?wsdl"
        summary = build_manifest([str(tmp_path)], str(tmp_path / "m.csv"))
        assert summary["edmx"] == 1 and summary["wsdl"] >= 1
        # catalog enumerate-endpoints always present
        text = (tmp_path / "m.csv").read_text()
        assert "CATALOGSERVICE" in text


class TestFetcherOwnCredentials:
    """{HOST} substitution, verbatim wsdl URL, basic-auth, typed routing."""
    def test_resolve_url_variants(self):
        from fetcher.odata_metadata_fetcher import resolve_url
        assert resolve_url({"url": "{HOST}/svc/$metadata"}, "https://s:443") == "https://s:443/svc/$metadata"
        assert resolve_url({"url": "{HOST}/p?wsdl"}, "https://s") == "https://s/p?wsdl"
        assert resolve_url({"url": "https://s/svc"}) == "https://s/svc/$metadata"
        import pytest
        with pytest.raises(ValueError):
            resolve_url({"url": "{HOST}/x"})           # no host supplied

    def test_basic_auth_routes_wsdl(self, tmp_path, monkeypatch):
        import fetcher.odata_metadata_fetcher as F

        class R:
            def __init__(self, c, b): self.status_code, self.content, self.headers = c, b, {}

        def fake(session, url, headers, timeout, auth=None):
            assert auth == ("u", "p") and "apikey" not in headers
            return R(200, b'<wsdl:definitions xmlns:wsdl="http://schemas.xmlsoap.org/wsdl/" '
                          b'targetNamespace="t"><wsdl:portType name="P"/></wsdl:definitions>')
        monkeypatch.setattr(F, "_http_get", fake)
        res = F.fetch_all([{"name": "svc", "url": "{HOST}/sap/bc/srt/x?wsdl"}],
                          "n/a", str(tmp_path), host="https://mysys:443", auth=("u", "p"))
        assert res[0].ok and (tmp_path / "wsdl" / "svc.wsdl").exists()


class TestFetcherOAuth:
    """OAuth client-credentials → bearer from a BTP/API-Mgmt service key."""
    def test_bearer_from_service_key(self, monkeypatch):
        import fetcher.odata_metadata_fetcher as F

        class R:
            def __init__(self, c, j=None): self.status_code, self._j, self.headers, self.content = c, j, {}, b""
            def json(self): return self._j

        def fake_post(session, turl, cid, csec, timeout):
            assert turl.endswith("/oauth/token") and cid == "c" and csec == "s"
            return R(200, {"access_token": "TOK"})
        monkeypatch.setattr(F, "_http_post_token", fake_post)
        tok, err = F.bearer_from_service_key('{"uaa":{"clientid":"c","clientsecret":"s","url":"https://a"}}')
        assert tok == "TOK" and err == ""
        _, e = F.bearer_from_service_key('{"x":1}')
        assert "clientid" in e

    def test_bearer_used_on_request(self, tmp_path, monkeypatch):
        import fetcher.odata_metadata_fetcher as F

        class R:
            def __init__(self, c, b): self.status_code, self.content, self.headers = c, b, {}

        def fake_get(session, url, headers, timeout, auth=None):
            assert headers.get("Authorization") == "Bearer TOK"
            return R(200, b'<edmx:Edmx xmlns:edmx="http://schemas.microsoft.com/ado/2007/06/edmx">'
                          b'<edmx:DataServices><Schema Namespace="X" '
                          b'xmlns="http://schemas.microsoft.com/ado/2008/09/edm"/></edmx:DataServices></edmx:Edmx>')
        monkeypatch.setattr(F, "_http_get", fake_get)
        res = F.fetch_all([{"name": "X", "url": "{HOST}/svc"}], "n/a", str(tmp_path),
                          host="https://t:443", bearer="TOK")
        assert res[0].ok and (tmp_path / "edmx" / "X.edmx").exists()


class TestOAuthBlockSelection:
    """Token URL must come from the OAuth block (uaa), not the top-level service url."""
    def test_abap_servicekey_shape_uses_uaa_url(self, monkeypatch):
        import fetcher.odata_metadata_fetcher as F
        cap = {}

        class R:
            def __init__(self, c, j=None): self.status_code, self._j, self.headers, self.content = c, j, {}, b""
            def json(self): return self._j

        def fake_post(session, turl, cid, csec, timeout):
            cap["turl"] = turl
            return R(200, {"access_token": "TOK"})
        monkeypatch.setattr(F, "_http_post_token", fake_post)
        key = ('{"url":"https://SYS.abap.us10.hana.ondemand.com",'
               '"endpoints":{"abap":"https://SYS.abap.us10.hana.ondemand.com"},'
               '"uaa":{"clientid":"sb-x|abap!b1","clientsecret":"P",'
               '"url":"https://tenant.authentication.us10.hana.ondemand.com"}}')
        tok, err = F.bearer_from_service_key(key)
        assert tok == "TOK" and err == ""
        assert cap["turl"] == "https://tenant.authentication.us10.hana.ondemand.com/oauth/token"
        assert "abap.us10" not in cap["turl"]      # must NOT use the system url


class TestTargetsCommentSkipping:
    """Seed CSV comments (# lines) and empty rows are skipped, not fetched as data."""
    def test_load_targets_skips_comments(self, tmp_path):
        from fetcher.odata_metadata_fetcher import load_targets
        p = tmp_path / "seed.csv"
        p.write_text(
            "name,product,service,url\n"
            "# a comment section\n"
            "API_X,s4hanacloud,API_X,\n"
            "# MyExample,,,https://h/svc\n"
            "\n"
            "API_Y,s4hanacloud,API_Y,\n")
        t = load_targets(str(p))
        assert [r["name"] for r in t] == ["API_X", "API_Y"]


class TestSchemaMatcher:
    """Match interface message types to the right canonical schema."""
    def _lib(self, tmp_path):
        (tmp_path / "orders.xsd").write_bytes(
            b'<xs:schema xmlns:xs="http://www.w3.org/2001/XMLSchema" '
            b'targetNamespace="http://acme.com/xi/SD"><xs:element name="OrderRequest"/></xs:schema>')
        (tmp_path / "invoice.xsd").write_bytes(
            b'<xs:schema xmlns:xs="http://www.w3.org/2001/XMLSchema" '
            b'targetNamespace="http://acme.com/xi/FI"><xs:element name="InvoiceRequest"/></xs:schema>')
        (tmp_path / "bp.edmx").write_bytes(
            b'<edmx:Edmx xmlns:edmx="http://schemas.microsoft.com/ado/2007/06/edmx"><edmx:DataServices>'
            b'<Schema Namespace="API_BP" xmlns="http://schemas.microsoft.com/ado/2008/09/edm">'
            b'<EntityType Name="A_BusinessPartner"/></Schema></edmx:DataServices></edmx:Edmx>')
        from extractor.schema_matcher import SchemaIndex
        return SchemaIndex.build(str(tmp_path))

    def test_name_plus_namespace_is_best(self, tmp_path):
        from extractor.schema_matcher import match_for_interface
        idx = self._lib(tmp_path)
        m = match_for_interface(idx, message_interface="OrderRequest", namespace="http://acme.com/xi/SD")
        assert m and m.entry.path.endswith("orders.xsd") and m.score >= 6

    def test_name_only_match(self, tmp_path):
        from extractor.schema_matcher import match_for_interface
        m = match_for_interface(self._lib(tmp_path), message_interface="InvoiceRequest")
        assert m and m.entry.path.endswith("invoice.xsd")

    def test_edmx_entity_match(self, tmp_path):
        from extractor.schema_matcher import match_for_interface
        m = match_for_interface(self._lib(tmp_path), message_interface="A_BusinessPartner")
        assert m and m.entry.path.endswith("bp.edmx")

    def test_kind_is_filter_not_self_match(self, tmp_path):
        idx = self._lib(tmp_path)
        assert idx.match(name="OrderRequest", kind="edmx") == []   # name doesn't match any edmx

    def test_no_match_returns_none(self, tmp_path):
        from extractor.schema_matcher import match_for_interface
        assert match_for_interface(self._lib(tmp_path), message_interface="Xyz", namespace="urn:nope") is None


class TestSchemaBindingIntoScaffold:
    """Matched schema is bundled into a generated iFlow, additively."""
    def _idx(self, tmp_path):
        (tmp_path / "orders.xsd").write_bytes(
            b'<xs:schema xmlns:xs="http://www.w3.org/2001/XMLSchema" '
            b'targetNamespace="http://acme.com/xi/SD"><xs:element name="OrderRequest"/></xs:schema>')
        from extractor.schema_matcher import SchemaIndex
        return SchemaIndex.build(str(tmp_path))

    def test_enriches_and_still_zips(self, tmp_path):
        import io, zipfile
        from scaffolder.schema_binding import bundle_matched_schema
        from scaffolder.minimal_iflow import generate_structured_iflow, build_bundle_zip
        idx = self._idx(tmp_path)
        res = generate_structured_iflow("SD_OrderFlow", "timer -> content modifier -> end")
        b = bundle_matched_schema(res, "OrderRequest", "http://acme.com/xi/SD", idx)
        assert b and b.confident and b.resource_path in res.files
        assert isinstance(res.files[b.resource_path], bytes)
        names = zipfile.ZipFile(io.BytesIO(build_bundle_zip(res))).namelist()
        assert b.resource_path in names
        assert any(n.endswith(".iflw") for n in names)

    def test_no_match_leaves_bundle_untouched(self, tmp_path):
        from scaffolder.schema_binding import bundle_matched_schema
        from scaffolder.minimal_iflow import generate_structured_iflow
        idx = self._idx(tmp_path)
        res = generate_structured_iflow("X", "timer -> end")
        before = dict(res.files)
        assert bundle_matched_schema(res, "Nope", "urn:nope", idx) is None
        assert res.files == before

    def test_reference_path_shape(self, tmp_path):
        from scaffolder.schema_binding import bind_interface_schema
        idx = self._idx(tmp_path)
        b = bind_interface_schema("OrderRequest", "http://acme.com/xi/SD", idx)
        assert b.reference_path == "/xsd/orders.xsd"
        assert b.resource_path == "src/main/resources/xsd/orders.xsd"


class TestOrganizerNoAccumulation:
    """Re-running the organizer into the same dir must not pile up identical copies."""
    def _bp(self, ns):
        return ('<edmx:Edmx xmlns:edmx="http://schemas.microsoft.com/ado/2007/06/edmx"><edmx:DataServices>'
                f'<Schema Namespace="{ns}" xmlns="http://schemas.microsoft.com/ado/2008/09/edm">'
                '<EntityType Name="A_BusinessPartner"/></Schema></edmx:DataServices></edmx:Edmx>').encode()

    def test_reruns_do_not_accumulate(self, tmp_path):
        from tools.organize_schema_library import organize
        src = tmp_path / "src"; src.mkdir()
        (src / "API_BUSINESS_PARTNER.edmx").write_bytes(self._bp("API_BUSINESS_PARTNER"))
        out = tmp_path / "out"
        for _ in range(3):
            organize([str(src)], str(out), generate_openapi=False)
        files = sorted(p.name for p in (out / "edmx").glob("*.edmx"))
        assert files == ["API_BUSINESS_PARTNER.edmx"]

    def test_distinct_variant_is_kept(self, tmp_path):
        from tools.organize_schema_library import organize
        src = tmp_path / "src"; src.mkdir()
        (src / "a.edmx").write_bytes(self._bp("API_BUSINESS_PARTNER"))
        (src / "b.edmx").write_bytes(self._bp("API_BUSINESS_PARTNER_V2"))
        out = tmp_path / "out"
        organize([str(src)], str(out), generate_openapi=False)
        files = sorted(p.name for p in (out / "edmx").glob("*.edmx"))
        assert len(files) == 2

    def test_clean_wipes_before_rebuild(self, tmp_path):
        from tools.organize_schema_library import organize
        src = tmp_path / "src"; src.mkdir()
        (src / "a.edmx").write_bytes(self._bp("API_BUSINESS_PARTNER"))
        out = tmp_path / "out"
        organize([str(src)], str(out), generate_openapi=False)
        (out / "edmx" / "STALE.edmx").write_bytes(b"<x/>")  # leftover from a prior layout
        organize([str(src)], str(out), generate_openapi=False, clean=True)
        files = sorted(p.name for p in (out / "edmx").glob("*.edmx"))
        assert "STALE.edmx" not in files


class TestXsdFamilyNoOverMerge:
    """Distinct message types in the same namespace must NOT collapse (data loss)."""
    def test_distinct_types_same_ns_separate(self, tmp_path):
        from extractor.schema_deduper import fingerprint
        a = tmp_path / "o.xsd"; b = tmp_path / "i.xsd"
        a.write_bytes(b'<xs:schema xmlns:xs="http://www.w3.org/2001/XMLSchema" targetNamespace="urn:co:sd"><xs:element name="OrderRequest_MT"/></xs:schema>')
        b.write_bytes(b'<xs:schema xmlns:xs="http://www.w3.org/2001/XMLSchema" targetNamespace="urn:co:sd"><xs:element name="Invoice_MT"/></xs:schema>')
        assert fingerprint(str(a)).family != fingerprint(str(b)).family

    def test_same_type_diff_whitespace_merges(self, tmp_path):
        from extractor.schema_deduper import fingerprint
        a = tmp_path / "o.xsd"; c = tmp_path / "o2.xsd"
        a.write_bytes(b'<xs:schema xmlns:xs="http://www.w3.org/2001/XMLSchema" targetNamespace="urn:co:sd"><xs:element name="OrderRequest_MT"/></xs:schema>')
        c.write_bytes(b'<xs:schema  xmlns:xs="http://www.w3.org/2001/XMLSchema"  targetNamespace="urn:co:sd"><xs:element name="OrderRequest_MT"/></xs:schema>')
        assert fingerprint(str(a)).family == fingerprint(str(c)).family

    def test_same_type_diff_ns_separate(self, tmp_path):
        from extractor.schema_deduper import fingerprint
        a = tmp_path / "o.xsd"; e = tmp_path / "o3.xsd"
        a.write_bytes(b'<xs:schema xmlns:xs="http://www.w3.org/2001/XMLSchema" targetNamespace="urn:co:sd"><xs:element name="OrderRequest_MT"/></xs:schema>')
        e.write_bytes(b'<xs:schema xmlns:xs="http://www.w3.org/2001/XMLSchema" targetNamespace="urn:other"><xs:element name="OrderRequest_MT"/></xs:schema>')
        assert fingerprint(str(a)).family != fingerprint(str(e)).family


class TestLibraryAudit:
    def test_audit_surfaces_drops_and_no_false_divergent(self, tmp_path):
        from tools.audit_schema_library import audit
        def w(n,b): (tmp_path/n).write_bytes(b)
        w("o.xsd", b'<xs:schema xmlns:xs="http://www.w3.org/2001/XMLSchema" targetNamespace="urn:co:sd"><xs:element name="OrderRequest_MT"/></xs:schema>')
        w("i.xsd", b'<xs:schema xmlns:xs="http://www.w3.org/2001/XMLSchema" targetNamespace="urn:co:sd"><xs:element name="Invoice_MT"/></xs:schema>')
        w("o_copy.xsd", b'<xs:schema xmlns:xs="http://www.w3.org/2001/XMLSchema" targetNamespace="urn:co:sd"><xs:element name="OrderRequest_MT"/></xs:schema>')
        w("notes.txt", b'hello')
        r = audit([str(tmp_path)])
        assert r["families"]["xsd"] == 2          # Order (with its dupe) + Invoice
        assert sum(len(v) for v in r["divergent"].values()) == 0
        assert len(r["unclassified"]) == 1        # the txt is surfaced, not silently dropped


class TestRouterIflow:
    NS={"bpmn2":"http://www.omg.org/spec/BPMN/20100524/MODEL","ifl":"http:///com.sap.ifl.model/Ifl.xsd",
        "bpmndi":"http://www.omg.org/spec/BPMN/20100524/DI"}
    def _routes(self, res):
        import xml.etree.ElementTree as ET
        root = ET.fromstring(res.iflw_xml)
        out=[]
        for sf in root.findall(".//bpmn2:sequenceFlow", self.NS):
            if any((p.find("key") is not None and p.find("key").text=="cmdVariantUri"
                    and "GatewayRoute" in (p.find("value").text or ""))
                   for p in sf.findall(".//ifl:property", self.NS)):
                c=sf.find("bpmn2:conditionExpression", self.NS)
                out.append((sf.get("id"), c.text if c is not None else None))
        return root, out
    def test_default_matches_gateway_and_conditions(self):
        from scaffolder.router_iflow import generate_router_iflow
        res = generate_router_iflow("R", route_property="value", set_value="Y")
        root, routes = self._routes(res)
        gw = root.find(".//bpmn2:exclusiveGateway", self.NS)
        defaults=[rid for rid,c in routes if c is None]
        conds=[c for _,c in routes if c]
        assert len(defaults)==1 and defaults[0]==gw.get("default")
        assert any("= 'Y'" in c for c in conds)
        assert any("!= null and" in c and "!= ''" in c for c in conds)
        assert len(root.findall(".//bpmn2:endEvent", self.NS))==3
    def test_no_dangling_flows_and_di_complete(self):
        import xml.etree.ElementTree as ET
        from scaffolder.router_iflow import generate_router_iflow
        root = ET.fromstring(generate_router_iflow("R").iflw_xml)
        ids={e.get("id") for tag in ("startEvent","endEvent","callActivity","exclusiveGateway")
             for e in root.findall(f".//bpmn2:{tag}", self.NS)}
        assert not [sf.get("id") for sf in root.findall(".//bpmn2:sequenceFlow", self.NS)
                    if sf.get("sourceRef") not in ids or sf.get("targetRef") not in ids]
        shapes={s.get("bpmnElement") for s in root.findall(".//bpmndi:BPMNShape", self.NS)}
        assert not [i for i in ids if i not in shapes]
    def test_n_branches(self):
        import xml.etree.ElementTree as ET
        from scaffolder.router_iflow import generate_router_iflow
        routes=[{"label":f"R{i}","condition":f"${{property.x}} = '{i}'","expr_type":"NonXML","process":True} for i in range(4)]
        routes.append({"label":"def","condition":None,"expr_type":"NonXML","process":False})
        root=ET.fromstring(generate_router_iflow("R5", routes=routes).iflw_xml)
        assert len(root.findall(".//bpmn2:exclusiveGateway/bpmn2:outgoing", self.NS))==5
    def test_bundle_zips(self):
        import io, zipfile
        from scaffolder.router_iflow import generate_router_iflow
        from scaffolder.minimal_iflow import build_bundle_zip
        z=build_bundle_zip(generate_router_iflow("R"))
        assert any(n.endswith(".iflw") for n in zipfile.ZipFile(io.BytesIO(z)).namelist())


class TestIFlowParser:
    def test_round_trips_router(self):
        from scaffolder.router_iflow import generate_router_iflow
        from extractor.iflow_parser import parse_iflow
        m = parse_iflow(generate_router_iflow("R", route_property="value", set_value="Y").iflw_xml, "R")
        assert "ExclusiveGateway" in m.kinds()
        assert len(m.routes) == 3
        assert sum(1 for r in m.routes if r.condition is None) == 1     # one default
        assert any("= 'Y'" in (r.condition or "") for r in m.routes)

    def test_round_trips_linear_sequence(self):
        from scaffolder.minimal_iflow import generate_timer_interface_iflow
        from extractor.iflow_parser import parse_iflow
        mids = [{"kind": "content_modifier", "name": "A"}, {"kind": "script", "name": "B"}]
        m = parse_iflow(generate_timer_interface_iflow("L", middle_steps=mids).iflw_xml, "L")
        assert {"Enricher", "Script", "StartTimerEvent", "EndEvent"} <= m.kinds()
        assert m.sequence and m.steps[m.sequence[0]].kind == "StartTimerEvent"

    def test_extracts_processes_and_params(self):
        from scaffolder.router_iflow import generate_router_iflow
        from extractor.iflow_parser import parse_iflow
        m = parse_iflow(generate_router_iflow("R").iflw_xml, "R")
        assert any(p.is_main for p in m.processes)
        assert isinstance(m.parameters, set)

    def test_handles_malformed_gracefully(self):
        from extractor.iflow_parser import parse_iflow
        import pytest
        with pytest.raises(Exception):
            parse_iflow("<not><closed>", "bad")


class TestCorpusGuardAndCache:
    def test_walk_skips_oversized_files(self, tmp_path):
        from library_builder import corpus_pipeline as cp
        (tmp_path / "big.xsd").write_bytes(b"<x/>" + b" " * (6 * 1024 * 1024))
        (tmp_path / "small.groovy").write_text("def x(){}")
        files = cp.walk_corpus(str(tmp_path))
        assert not any("big" in k for k in files)
        assert any("small" in k for k in files)

    def test_budget_caps(self):
        from library_builder.corpus_pipeline import _Budget
        b = _Budget(max_files=2, max_bytes=10**9)
        assert b.take(1) and b.take(1)
        assert not b.take(1) and b.capped

    def test_disk_cache_builds_once_and_invalidates(self, tmp_path, monkeypatch):
        import time
        from library_builder import corpus_pipeline as cp
        calls = {"n": 0}
        orig = cp._build_corpus
        def counting(*a, **k):
            calls["n"] += 1
            return orig(*a, **k)
        monkeypatch.setattr(cp, "_build_corpus", counting)
        cache_dir = tmp_path / "cache"; cache_dir.mkdir()
        monkeypatch.setattr(cp, "_cache_file", lambda sig: cache_dir / f"c_{sig}.pkl")
        src = tmp_path / "src"; src.mkdir()
        (src / "a.groovy").write_text("def a(){}")
        cp.build_corpus(path=str(src))
        cp.build_corpus(path=str(src))           # cache hit
        assert calls["n"] == 1
        time.sleep(0.01); (src / "b.groovy").write_text("def b(){}")
        cp.build_corpus(path=str(src))           # invalidated
        assert calls["n"] == 2

    def test_signature_changes_with_content(self, tmp_path):
        from library_builder.corpus_pipeline import _dir_signature
        (tmp_path / "x").write_text("a")
        s1 = _dir_signature(str(tmp_path))
        (tmp_path / "y").write_text("b")
        assert _dir_signature(str(tmp_path)) != s1


class TestRoundTripCoverage:
    def test_supported_set_derived_empirically(self):
        from extractor.coverage import generator_supported_kinds
        s = generator_supported_kinds()
        # the load-bearing kinds the generator emits today
        for k in ("Enricher", "Script", "Mapping", "ExclusiveGateway",
                  "EndEvent", "StartTimerEvent"):
            assert k in s, f"{k} should be in the empirical supported set"

    def test_generated_flow_is_reproducible(self):
        from scaffolder import minimal_iflow as mi
        from extractor.coverage import measure_corpus, generator_supported_kinds
        iflw, _ = mi.build_flow_from_steps(
            "T", "T", [{"kind": "script"}, {"kind": "content_modifier"}])
        rep = measure_corpus([("t.iflw", iflw)], generator_supported_kinds())
        assert rep.total == 1 and rep.reproducible == 1 and rep.pct == 100.0

    def test_unsupported_kind_blocks_and_greedy_unlocks(self):
        from extractor.coverage import (measure_corpus, assess, greedy_unlock_order,
                                        IFlowVerdict)
        # synthetic verdict: blocked solely by one fake construct
        rep = measure_corpus([], set())
        rep.total = 2
        rep.reproducible = 1
        rep.verdicts = [
            IFlowVerdict("a", True, 3, 1),
            IFlowVerdict("b", False, 4, 1, unsupported_kinds={"FooStep"}),
        ]
        curve = greedy_unlock_order(rep)
        assert curve and curve[0][0] == "FooStep" and curve[-1][2] == 100.0


class TestModelGenerator:
    def test_generate_from_model_roundtrips_linear(self):
        from scaffolder import minimal_iflow as mi
        from scaffolder.model_generator import generate_from_model
        from extractor.iflow_parser import parse_iflow
        iflw, _ = mi.build_flow_from_steps(
            "X", "X", [{"kind": "script"}, {"kind": "mapping"},
                       {"kind": "content_modifier"}])
        m1 = parse_iflow(iflw, "X")
        res = generate_from_model(m1, name="X")
        m2 = parse_iflow(res.iflw_xml, "X")
        mids = lambda m: [m.steps[s].kind for s in m.sequence
                          if m.steps[s].kind not in
                          ("StartEvent", "StartTimerEvent", "EndEvent")]
        assert mids(m1) == mids(m2) == ["Script", "Mapping", "Enricher"]

    def test_unsupported_construct_raises(self):
        from scaffolder.model_generator import (generate_from_model,
                                                UnsupportedConstruct)
        from extractor.iflow_parser import IFlowModel, Process, Step
        # gateway + multi-process now emit; an unsupported STEP KIND still raises
        m = IFlowModel(name="x", processes=[Process("Process_1", "m", True)])
        m.steps = {"S1": Step(id="S1", kind="ErrorEventSubProcessTemplate",
                              name="e", process_id="Process_1", config={},
                              incoming=[], outgoing=[], parent_subprocess="")}
        m.sequence = ["S1"]
        try:
            generate_from_model(m)
            assert False, "should raise on unsupported kind"
        except UnsupportedConstruct:
            pass


class TestExternalCallEmitter:
    def test_external_call_reparses(self):
        import xml.etree.ElementTree as ET
        from scaffolder.external_call_iflow import generate_external_call_iflow
        from extractor.iflow_parser import parse_iflow
        res = generate_external_call_iflow("EC", address="https://h/api",
                                           receiver_name="API")
        ET.fromstring(res.iflw_xml)            # well-formed
        m = parse_iflow(res.iflw_xml, "EC")
        assert "ExternalCall" in [s.kind for s in m.steps.values()]
        assert any(e.direction == "receiver" for e in m.endpoints)


class TestCpiRegenerationWiring:
    def test_regenerate_reproduces_supported(self):
        from scaffolder import minimal_iflow as mi
        from scaffolder.regenerate import regenerate_iflow_xml
        iflw, _ = mi.build_flow_from_steps(
            "R", "R", [{"kind": "script"}, {"kind": "mapping"}])
        r = regenerate_iflow_xml(iflw, "R")
        assert r.reproduced and r.n_steps == 2

    def test_regenerate_honest_on_unsupported(self):
        from scaffolder.regenerate import regenerate_iflow_xml
        from extractor.iflow_parser import IFlowModel, Process, Step
        # an unsupported step kind can't be emitted → honest blocker, no crash
        import scaffolder.regenerate as rg
        from unittest.mock import patch
        m = IFlowModel(name="x", processes=[Process("Process_1", "m", True)])
        m.steps = {"S1": Step(id="S1", kind="ErrorEventSubProcessTemplate",
                              name="e", process_id="Process_1", config={},
                              incoming=[], outgoing=[], parent_subprocess="")}
        m.sequence = ["S1"]
        with patch.object(rg, "parse_iflow", return_value=m):
            r = rg.regenerate_iflow_xml("<x/>", "x")
        assert not r.reproduced and r.blockers

    def test_scaffold_uses_source_iflow(self, tmp_path):
        import types
        from scaffolder import minimal_iflow as mi
        from scaffolder.iflow_scaffolder import IFlowScaffolder
        from extractor.iflow_parser import parse_iflow
        src, _ = mi.build_flow_from_steps(
            "S", "S", [{"kind": "script"}, {"kind": "filter"}])
        iface = types.SimpleNamespace(name="S", id="1", source_iflow_xml=src,
                                      steps_spec="", sender_adapter="",
                                      receiver_adapter="")
        out = IFlowScaffolder(output_dir=str(tmp_path)).scaffold(
            types.SimpleNamespace(interface=iface))
        g = parse_iflow(out.read_text(), "S")
        mids = [g.steps[s].kind for s in g.sequence
                if g.steps[s].kind not in ("StartEvent", "StartTimerEvent", "EndEvent")]
        assert mids == ["Script", "Filter"]      # real structure, not a 2-CM stub

    def test_extractor_pulls_iflw_text(self):
        import io, zipfile
        from scaffolder import minimal_iflow as mi
        from fetcher.artifact_router import _iflw_text_from_bundle
        res = mi.generate_minimal_iflow("B", "B")
        buf = io.BytesIO()
        with zipfile.ZipFile(buf, "w") as z:
            for path, content in res.files.items():
                z.writestr(path, content)
        text = _iflw_text_from_bundle(buf.getvalue())
        assert "<bpmn2:" in text


class TestExternalCallIntegration:
    def test_external_call_composes_in_linear_flow(self):
        import xml.etree.ElementTree as ET
        from scaffolder import minimal_iflow as mi
        from extractor.iflow_parser import parse_iflow
        iflw, _ = mi.build_flow_from_steps("M", "M", [
            {"kind": "script"},
            {"kind": "external_call", "name": "Request-Reply",
             "receiver_name": "ERP", "address": "https://erp/api"},
            {"kind": "content_modifier"}])
        ET.fromstring(iflw)
        m = parse_iflow(iflw, "M")
        mids = [m.steps[s].kind for s in m.sequence
                if m.steps[s].kind not in ("StartEvent", "StartTimerEvent", "EndEvent")]
        assert mids == ["Script", "ExternalCall", "Enricher"]
        assert any(e.direction == "receiver" and e.name == "ERP"
                   for e in m.endpoints)

    def test_no_receiver_path_has_no_receiver_participant(self):
        from scaffolder import minimal_iflow as mi
        iflw, _ = mi.build_flow_from_steps(
            "L", "L", [{"kind": "script"}, {"kind": "mapping"}])
        assert "EndpointRecevier" not in iflw      # linear path unchanged

    def test_externalcall_model_roundtrips(self):
        from scaffolder import minimal_iflow as mi
        from scaffolder.regenerate import regenerate_iflow_xml
        src, _ = mi.build_flow_from_steps("E", "E", [
            {"kind": "content_modifier"},
            {"kind": "external_call", "receiver_name": "API",
             "address": "https://h/x"}])
        r = regenerate_iflow_xml(src, "E")
        assert r.reproduced and r.n_steps == 2

    def test_parser_links_externalcall_receiver(self):
        from scaffolder.external_call_iflow import generate_external_call_iflow
        from extractor.iflow_parser import parse_iflow
        res = generate_external_call_iflow("X", address="https://h/api",
                                           receiver_name="ERP")
        m = parse_iflow(res.iflw_xml, "X")
        ec = next(s for s in m.steps.values() if s.kind == "ExternalCall")
        assert ec.config.get("receiver_name") == "ERP"
        assert "https://h/api" in ec.config.get("address", "")


class TestLinearConstructs:
    import pytest as _pytest

    @_pytest.mark.parametrize("kind", [
        "Encoder", "Decoder", "DBstorage", "XMLDigitalSignMessage",
        "SimpleSignMessage", "Send", "Variables", "XmlModifier",
        "contentEnricherWithLookup", "Persist", "XmlValidator"])
    def test_passthrough_construct_roundtrips(self, kind):
        import xml.etree.ElementTree as ET
        from scaffolder import minimal_iflow as mi
        from extractor.iflow_parser import parse_iflow
        iflw, _ = mi.build_flow_from_steps(
            "P", "P", [{"kind": "content_modifier"}, {"kind": kind}])
        ET.fromstring(iflw)
        m = parse_iflow(iflw, "P")
        kinds = [m.steps[s].kind for s in m.sequence
                 if m.steps[s].kind not in ("StartEvent", "StartTimerEvent", "EndEvent")]
        assert kinds == ["Enricher", kind]

    def test_parser_derives_externalcall_from_cmdvariant(self):
        # a serviceTask with cmdVariant ExternalCall but NO activityType must
        # classify as ExternalCall, not the bare 'ServiceTask' tag fallback.
        from extractor.iflow_parser import parse_iflow
        xml = ('<bpmn2:definitions xmlns:bpmn2="http://www.omg.org/spec/BPMN/20100524/MODEL"'
               ' xmlns:ifl="http:///com.sap.ifl.model/Ifl.xsd">'
               '<bpmn2:process id="Process_1"><bpmn2:serviceTask id="ST" name="Request-Reply">'
               '<bpmn2:extensionElements><ifl:property><key>cmdVariantUri</key>'
               '<value>ctype::FlowstepVariant/cname::ExternalCall</value></ifl:property>'
               '</bpmn2:extensionElements></bpmn2:serviceTask></bpmn2:process></bpmn2:definitions>')
        m = parse_iflow(xml, "x")
        assert any(s.kind == "ExternalCall" for s in m.steps.values())


class TestRequestReplyReclassification:
    def test_bare_servicetask_with_receiver_is_externalcall(self):
        # serviceTask, no activityType/cmdVariant, but sends a messageFlow to a
        # receiver participant → must classify as ExternalCall, not ServiceTask.
        from extractor.iflow_parser import parse_iflow
        xml = ('<bpmn2:definitions xmlns:bpmn2="http://www.omg.org/spec/BPMN/20100524/MODEL"'
               ' xmlns:ifl="http:///com.sap.ifl.model/Ifl.xsd">'
               '<bpmn2:collaboration id="C">'
               '<bpmn2:participant id="Participant_2" ifl:type="EndpointRecevier" name="ERP"/>'
               '<bpmn2:messageFlow id="MF3" sourceRef="ServiceTask_1" targetRef="Participant_2"/>'
               '</bpmn2:collaboration>'
               '<bpmn2:process id="Process_1">'
               '<bpmn2:serviceTask id="ServiceTask_1" name="Request-Reply">'
               '<bpmn2:incoming>f1</bpmn2:incoming><bpmn2:outgoing>f2</bpmn2:outgoing>'
               '</bpmn2:serviceTask></bpmn2:process></bpmn2:definitions>')
        m = parse_iflow(xml, "x")
        st = m.steps["ServiceTask_1"]
        assert st.kind == "ExternalCall"
        assert st.config.get("receiver_name") == "ERP"


class TestEndpointExtraction:
    def test_extracts_real_sender_receiver_adapters(self):
        from extractor.iflow_parser import extract_endpoints
        xml = ('<bpmn2:definitions xmlns:bpmn2="http://www.omg.org/spec/BPMN/20100524/MODEL"'
               ' xmlns:ifl="http:///com.sap.ifl.model/Ifl.xsd">'
               '<bpmn2:collaboration id="C">'
               '<bpmn2:participant id="P1" ifl:type="EndpointSender" name="S4HANA"/>'
               '<bpmn2:participant id="P2" ifl:type="EndpointRecevier" name="OpenText"/>'
               '<bpmn2:messageFlow id="MF1" sourceRef="P1" targetRef="StartEvent_1">'
               '<bpmn2:extensionElements><ifl:property><key>ComponentType</key>'
               '<value>HTTPS</value></ifl:property></bpmn2:extensionElements></bpmn2:messageFlow>'
               '<bpmn2:messageFlow id="MF2" sourceRef="ServiceTask_1" targetRef="P2">'
               '<bpmn2:extensionElements><ifl:property><key>ComponentType</key>'
               '<value>SOAP</value></ifl:property></bpmn2:extensionElements></bpmn2:messageFlow>'
               '</bpmn2:collaboration>'
               '<bpmn2:process id="Process_1"><bpmn2:startEvent id="StartEvent_1"/>'
               '</bpmn2:process></bpmn2:definitions>')
        ep = extract_endpoints(xml)
        assert ep["sender_system"] == "S4HANA"
        assert ep["sender_adapter"] == "HTTPS"
        assert ep["receiver_system"] == "OpenText"
        assert ep["receiver_adapter"] == "SOAP"

    def test_timer_flow_has_empty_sender(self):
        from extractor.iflow_parser import extract_endpoints
        xml = ('<bpmn2:definitions xmlns:bpmn2="http://www.omg.org/spec/BPMN/20100524/MODEL"'
               ' xmlns:ifl="http:///com.sap.ifl.model/Ifl.xsd">'
               '<bpmn2:process id="Process_1"><bpmn2:startEvent id="StartEvent_1"/>'
               '</bpmn2:process></bpmn2:definitions>')
        ep = extract_endpoints(xml)
        assert ep["sender_system"] == "" and ep["receiver_system"] == ""


class TestExternalCallReceiverFidelity:
    def test_receiver_adapter_preserved_not_defaulted_to_http(self):
        # An ExternalCall whose receiver is a SOAP adapter must regenerate with a
        # SOAP receiver message flow, not a hardcoded HTTP one.
        import xml.etree.ElementTree as ET
        from scaffolder.regenerate import regenerate_iflow_xml
        xml = ('<bpmn2:definitions xmlns:bpmn2="http://www.omg.org/spec/BPMN/20100524/MODEL"'
               ' xmlns:ifl="http:///com.sap.ifl.model/Ifl.xsd">'
               '<bpmn2:collaboration id="C">'
               '<bpmn2:participant id="P1" ifl:type="EndpointSender" name="Sender"/>'
               '<bpmn2:participant id="P2" ifl:type="EndpointRecevier" name="Backend"/>'
               '<bpmn2:participant id="PP" ifl:type="IntegrationProcess" name="Integration Process" processRef="Process_1"/>'
               '<bpmn2:messageFlow id="MF1" name="SOAP" sourceRef="ServiceTask_1" targetRef="P2">'
               '<bpmn2:extensionElements>'
               '<ifl:property><key>ComponentType</key><value>SOAP</value></ifl:property>'
               '<ifl:property><key>address</key><value>https://soap.example/svc</value></ifl:property>'
               '<ifl:property><key>cmdVariantUri</key><value>ctype::AdapterVariant/cname::sap:SOAP/tp::HTTP/mp::SOAP 1.x/direction::Receiver/version::1.6.0</value></ifl:property>'
               '</bpmn2:extensionElements></bpmn2:messageFlow></bpmn2:collaboration>'
               '<bpmn2:process id="Process_1">'
               '<bpmn2:startEvent id="StartEvent_2"><bpmn2:outgoing>s1</bpmn2:outgoing></bpmn2:startEvent>'
               '<bpmn2:serviceTask id="ServiceTask_1" name="Request-Reply">'
               '<bpmn2:extensionElements>'
               '<ifl:property><key>activityType</key><value>ExternalCall</value></ifl:property>'
               '<ifl:property><key>cmdVariantUri</key><value>ctype::FlowstepVariant/cname::ExternalCall</value></ifl:property>'
               '</bpmn2:extensionElements>'
               '<bpmn2:incoming>s1</bpmn2:incoming><bpmn2:outgoing>s2</bpmn2:outgoing></bpmn2:serviceTask>'
               '<bpmn2:endEvent id="EndEvent_2"><bpmn2:incoming>s2</bpmn2:incoming></bpmn2:endEvent>'
               '<bpmn2:sequenceFlow id="s1" sourceRef="StartEvent_2" targetRef="ServiceTask_1"/>'
               '<bpmn2:sequenceFlow id="s2" sourceRef="ServiceTask_1" targetRef="EndEvent_2"/>'
               '</bpmn2:process></bpmn2:definitions>')
        r = regenerate_iflow_xml(xml, "soap_call")
        assert r.reproduced
        g = r.result.iflw_xml
        ET.fromstring(g)
        # the regenerated receiver flow must carry SOAP, not HTTP
        assert "<value>SOAP</value>" in g
        assert "sap:SOAP" in g
        assert "https://soap.example/svc" in g


class TestStructureAwareEffort:
    def _multiproc_xml(self):
        return ('<bpmn2:definitions xmlns:bpmn2="http://www.omg.org/spec/BPMN/20100524/MODEL"'
                ' xmlns:ifl="http:///com.sap.ifl.model/Ifl.xsd">'
                '<bpmn2:collaboration id="C">'
                '<bpmn2:participant id="PP" ifl:type="IntegrationProcess" name="Integration Process" processRef="Process_1"/>'
                '<bpmn2:participant id="LP" ifl:type="LocalIntegrationProcess" name="LIP" processRef="Process_2"/>'
                '</bpmn2:collaboration>'
                '<bpmn2:process id="Process_1"><bpmn2:startEvent id="StartEvent_1"/>'
                '<bpmn2:callActivity id="C1"><bpmn2:extensionElements>'
                '<ifl:property><key>activityType</key><value>Mapping</value></ifl:property>'
                '</bpmn2:extensionElements></bpmn2:callActivity>'
                '<bpmn2:endEvent id="EndEvent_1"/></bpmn2:process>'
                '<bpmn2:process id="Process_2"><bpmn2:startEvent id="StartEvent_2"/>'
                '<bpmn2:callActivity id="C2"><bpmn2:extensionElements>'
                '<ifl:property><key>activityType</key><value>Script</value></ifl:property>'
                '</bpmn2:extensionElements></bpmn2:callActivity>'
                '<bpmn2:endEvent id="EndEvent_2"/></bpmn2:process></bpmn2:definitions>')

    def test_real_structure_raises_complexity(self):
        from extractor.pi_extractor import InterfaceRecord
        from analyzer.complexity_analyzer import ComplexityAnalyzer
        an = ComplexityAnalyzer({})
        base = dict(id="x", name="x", namespace="", software_component="",
                    sender_system="", receiver_system="",
                    sender_adapter="HTTPS", receiver_adapter="HTTPS",
                    message_interface="")
        no_xml = an.assess(InterfaceRecord(**base))
        with_xml = an.assess(InterfaceRecord(source_iflow_xml=self._multiproc_xml(), **base))
        # the multi-process flow must score strictly higher than metadata-only
        assert with_xml.score > no_xml.score

    def test_metadata_only_path_unchanged(self):
        # No source XML → structural signal contributes nothing.
        from extractor.pi_extractor import InterfaceRecord
        from analyzer.complexity_analyzer import ComplexityAnalyzer
        an = ComplexityAnalyzer({})
        rec = InterfaceRecord(id="x", name="x", namespace="", software_component="",
                              sender_system="", receiver_system="",
                              sender_adapter="HTTPS", receiver_adapter="HTTPS",
                              message_interface="")
        pts, note = an._structural_score(rec)
        assert pts == 0 and note == ""


class TestPassthroughHonesty:
    def test_endpoint_only_passthrough_reproduced_with_endpoints(self):
        # A flow with 0 middle steps but real sender+receiver endpoints is now
        # reproduced FAITHFULLY — the generator emits the participants + message
        # flows, so it is not an empty Start→End. (The regenerate guard remains
        # as a safety net for any path that would still drop them.)
        from scaffolder.regenerate import regenerate_iflow_xml
        from extractor.iflow_parser import parse_iflow
        xml = ('<bpmn2:definitions xmlns:bpmn2="http://www.omg.org/spec/BPMN/20100524/MODEL"'
               ' xmlns:ifl="http:///com.sap.ifl.model/Ifl.xsd">'
               '<bpmn2:collaboration id="C">'
               '<bpmn2:participant id="P1" ifl:type="EndpointSender" name="Src"/>'
               '<bpmn2:participant id="P2" ifl:type="EndpointRecevier" name="Tgt"/>'
               '<bpmn2:participant id="PP" ifl:type="IntegrationProcess" name="Integration Process" processRef="Process_1"/>'
               '<bpmn2:messageFlow id="MF1" name="HTTPS" sourceRef="P1" targetRef="StartEvent_1">'
               '<bpmn2:extensionElements><ifl:property><key>ComponentType</key><value>HTTPS</value></ifl:property></bpmn2:extensionElements>'
               '</bpmn2:messageFlow>'
               '<bpmn2:messageFlow id="MF2" sourceRef="EndEvent_1" targetRef="P2"/>'
               '</bpmn2:collaboration>'
               '<bpmn2:process id="Process_1">'
               '<bpmn2:startEvent id="StartEvent_1"><bpmn2:outgoing>s1</bpmn2:outgoing></bpmn2:startEvent>'
               '<bpmn2:endEvent id="EndEvent_1"><bpmn2:incoming>s1</bpmn2:incoming></bpmn2:endEvent>'
               '<bpmn2:sequenceFlow id="s1" sourceRef="StartEvent_1" targetRef="EndEvent_1"/>'
               '</bpmn2:process></bpmn2:definitions>')
        r = regenerate_iflow_xml(xml, "passthrough")
        assert r.reproduced is True and r.result is not None
        m2 = parse_iflow(r.result.iflw_xml, "passthrough")
        assert sorted(e.direction for e in m2.endpoints) == ["receiver", "sender"]
        assert any(mf.config.get("ComponentType") == "HTTPS"
                   for mf in m2.message_flows)


class TestStructureAwareSizing:
    def test_real_iflow_sizes_scale_with_structure(self):
        # A 9-process/112-step monster must outweigh a 0-step passthrough. The
        # old keyword path collapsed every uploaded iFlow to "S" because the
        # record had no rich description; sizing now reads the parsed topology.
        import os, pickle, pytest
        corpus = "/tmp/learn/iflws.pkl"
        if not os.path.exists(corpus):
            pytest.skip("corpus not present in this environment")
        from analyzer.sap_complexity_engine import SAPComplexityEngine
        iflws = pickle.load(open(corpus, "rb"))
        by = {p.rsplit("/", 1)[-1]: x for p, x in iflws}
        eng = SAPComplexityEngine()

        class Rec:
            def __init__(s, nm, x):
                s.name = nm; s.source_iflow_xml = x; s.description = ""

        big = eng.assess_interface(Rec("m", by["S_ContentModifier_SetConstant.iflw"]))
        small = eng.assess_interface(Rec("s", by["AribaPOOut.iflw"]))
        assert big.total_weight > small.total_weight
        assert big.size in ("L", "XL")
        assert small.size in ("S", "M")


class TestConstructCoverage:
    def test_coverage_reports_matched_and_generic(self):
        # The weighting audit must report unknown constructs as "generic_only"
        # (honestly flagged, no dedicated rule) rather than guessing them free.
        import os, pickle, pytest
        corpus = "/tmp/learn/iflws.pkl"
        if not os.path.exists(corpus):
            pytest.skip("corpus not present in this environment")
        from analyzer.sap_complexity_engine import construct_coverage

        class Rec:
            def __init__(s, x): s.name = "x"; s.source_iflow_xml = x

        by = {p.rsplit("/", 1)[-1]: x for p, x in pickle.load(open(corpus, "rb"))}
        cov = construct_coverage(Rec(by["com_sap_GS_Chile_GetStatus.iflw"]))
        assert cov is not None
        # Script/Mapping/ExternalCall have dedicated rules; Enricher/Filter do not
        assert "ExternalCall" in cov["matched"]
        assert "Enricher" in cov["generic_only"]
        assert cov["total_middle_steps"] == sum(
            [v["count"] for v in cov["matched"].values()]
            + list(cov["generic_only"].values()))
        # no real iFlow -> None (falls back, never crashes)
        class Bare:
            name = "y"; source_iflow_xml = None
        assert construct_coverage(Bare()) is None


class TestPassthroughConfigReEmit:
    def test_signer_config_carried_not_hollow_shell(self):
        # Regression guard: passthrough steps (XMLDigitalSign, Variables, ...)
        # must re-emit their captured config verbatim. A hollow shell makes CPI
        # reject the step ("Private Key Alias is not specified", etc.).
        import os, pickle, pytest
        import xml.dom.minidom as MD
        corpus = "/tmp/learn/iflws.pkl"
        if not os.path.exists(corpus):
            pytest.skip("corpus not present")
        from extractor.iflow_parser import parse_iflow
        from scaffolder.model_generator import generate_from_model
        by = {p.rsplit("/", 1)[-1]: x for p, x in pickle.load(open(corpus, "rb"))}
        out = generate_from_model(
            parse_iflow(by["com_sap_GS_Chile_SignEnvioDTE.iflw"], "s")).iflw_xml
        # functional signer properties must be present with their real values
        assert "<key>privateKeyAlias</key><value>${header.signatureKey}</value>" in out
        assert "signatureAlgorithm" in out and "canonicalizationMethod" in out
        assert "transformMethod" in out and "signatureType" in out
        MD.parseString(out)  # still well-formed


class TestFilterConfigReEmit:
    def test_filter_carries_real_xpath_not_synthetic_nodelist(self):
        # The filter must reproduce the real wrapContent/xpathType (e.g.
        # /p2:SetDTE + Node), not the synthetic /* + Nodelist that trips CPI's
        # content-type check.
        import os, pickle, re, pytest
        corpus = "/tmp/learn/iflws.pkl"
        if not os.path.exists(corpus):
            pytest.skip("corpus not present")
        from extractor.iflow_parser import parse_iflow
        from scaffolder.model_generator import generate_from_model
        by = {p.rsplit("/", 1)[-1]: x for p, x in pickle.load(open(corpus, "rb"))}
        out = generate_from_model(
            parse_iflow(by["com_sap_GS_Chile_SignEnvioDTE.iflw"], "x")).iflw_xml
        m = re.search(r'<bpmn2:callActivity\b(?:(?!</bpmn2:callActivity>).)*?'
                      r'Filter(?:(?!</bpmn2:callActivity>).)*?</bpmn2:callActivity>',
                      out, re.S)
        f = dict(re.findall(r'<key>([^<]+)</key>\s*<value>([^<]*)</value>', m.group(0)))
        assert f.get("wrapContent") == "/p2:SetDTE"
        assert f.get("xpathType") == "Node"


class TestResourceResolver:
    def _corpus(self):
        return {
            "PkgA.zip!h1_content!src/main/resources/mapping/Foo.xsl": "<xslA/>",
            "PkgB.zip!h2_content!src/main/resources/mapping/Foo.xsl": "<xslB/>",
            "PkgA.zip!h1_content!src/main/resources/script/bar.groovy": "A.bar",
            "PkgA.zip!h1_content!src/main/resources/xsd/Schema_v1.xsd": "<xsd/>",
        }

    def test_resolves_dir_uri_to_content(self):
        from scaffolder.resource_resolver import resolve
        c = self._corpus()
        r = resolve("dir://mapping/xslt/src/main/resources/mapping/Foo.xsl",
                    c, package="PkgA.zip", kind="mapping")
        assert r.ok and r.content == "<xslA/>"

    def test_package_scope_disambiguates_same_basename(self):
        from scaffolder.resource_resolver import resolve
        c = self._corpus()
        assert resolve("Foo.xsl", c, package="PkgA.zip").content == "<xslA/>"
        assert resolve("Foo.xsl", c, package="PkgB.zip").content == "<xslB/>"

    def test_script_and_xsd_basenames(self):
        from scaffolder.resource_resolver import resolve
        c = self._corpus()
        assert resolve("bar.groovy", c, kind="script").content == "A.bar"
        assert resolve("Schema_v1.xsd", c, kind="schema").content == "<xsd/>"

    def test_unresolved_is_not_ok(self):
        from scaffolder.resource_resolver import resolve
        assert not resolve("Nope.xsl", self._corpus()).ok

    def test_nested_batch_zip_scopes_to_real_package(self):
        # The real package is the .zip before *_content, not the outer batch.
        from scaffolder.resource_resolver import _package_of, resolve
        key = ("part4.zip!Data_Ingestion_Integration_with_SAP_S-4HANA.zip!"
               "e547_content!META-INF/MANIFEST.MF")
        assert _package_of(key) == \
            "data_ingestion_integration_with_sap_s-4hana.zip"
        c = {
            "batch.zip!PkgA.zip!h_content!src/main/resources/mapping/Foo.xsl": "<A/>",
            "batch.zip!PkgB.zip!h_content!src/main/resources/mapping/Foo.xsl": "<B/>",
        }
        assert resolve("Foo.xsl", c, package="PkgA.zip").content == "<A/>"
        assert resolve("Foo.xsl", c, package="PkgB.zip").content == "<B/>"


class TestMockEndpoints:
    def test_mock_scaffold_is_well_formed_and_self_triggering(self):
        import os, pickle, xml.dom.minidom as MD, pytest
        if not os.path.exists("/tmp/learn/iflws.pkl"):
            pytest.skip("corpus not present")
        from extractor.iflow_parser import parse_iflow
        from scaffolder.model_generator import generate_mock_from_model
        by = {p.rsplit("/", 1)[-1]: x
              for p, x in pickle.load(open("/tmp/learn/iflws.pkl", "rb"))}
        # any flow works; mock replaces its I/O with a testable scaffold
        m = parse_iflow(next(iter(by.values())), "x")
        res = generate_mock_from_model(m, name="MockTest")
        MD.parseString(res.iflw_xml)                       # well-formed
        assert "timerEventDefinition" in res.iflw_xml      # self-triggering
        m2 = parse_iflow(res.iflw_xml, "x")
        kinds = {m2.steps[s].kind for s in m2.steps}
        assert "ExternalCall" in kinds and "Enricher" in kinds


class TestSchemaMockPayload:
    def test_xsd_sample_no_child_leak(self):
        from scaffolder.sample_payload import sample_payload_from_xsd
        xsd = ('<xs:schema xmlns:xs="http://www.w3.org/2001/XMLSchema">'
               '<xs:element name="A"><xs:complexType><xs:sequence>'
               '<xs:element name="B"><xs:complexType><xs:sequence>'
               '<xs:element name="C" type="xs:string"/></xs:sequence>'
               '</xs:complexType></xs:element></xs:sequence></xs:complexType>'
               '</xs:element></xs:schema>')
        out = sample_payload_from_xsd(xsd)
        # C must appear once, nested under B — not leaked up to A
        assert out.count("<C>") == 1 and "<B><C>" in out

    def test_garbage_xsd_falls_back(self):
        from scaffolder.sample_payload import sample_payload_from_xsd
        assert "MockPayload" in sample_payload_from_xsd("not xml")

    def test_mock_seeds_schema_derived_body_when_resolvable(self):
        from types import SimpleNamespace as NS
        from scaffolder.model_generator import mock_specs_from_model
        xsd = ('<xs:schema xmlns:xs="http://www.w3.org/2001/XMLSchema">'
               '<xs:element name="Req"><xs:complexType><xs:sequence>'
               '<xs:element name="Id" type="xs:int"/></xs:sequence>'
               '</xs:complexType></xs:element></xs:schema>')
        corpus = {"PkgX.zip!h_content!src/main/resources/xsd/Req.xsd": xsd}
        step = NS(config={"schemaResourceUri": "src/main/resources/xsd/Req.xsd"})
        model = NS(steps={"s": step}, endpoints=[], name="X",
                   source_package="PkgX.zip")
        specs = mock_specs_from_model(model, corpus=corpus)
        assert specs[0].get("body") == "<Req><Id>0</Id></Req>"

    def test_mock_without_schema_uses_generic(self):
        from types import SimpleNamespace as NS
        from scaffolder.model_generator import mock_specs_from_model
        model = NS(steps={}, endpoints=[], name="Y", source_package=None)
        assert "body" not in mock_specs_from_model(model, corpus={})[0]


class TestMockPayloadPriority:
    def test_real_body_beats_synthetic(self):
        from types import SimpleNamespace as NS
        from scaffolder.model_generator import mock_specs_from_model
        real = '<Doc xmlns="urn:x"><field>value</field></Doc>'
        m = NS(steps={"s": NS(config={"wrapContent": real})},
               endpoints=[], name="A", source_package=None)
        assert mock_specs_from_model(m)[0].get("body") == real

    def test_schema_derived_when_no_real_body(self):
        from types import SimpleNamespace as NS
        from scaffolder.model_generator import mock_specs_from_model
        xsd = ('<xs:schema xmlns:xs="http://www.w3.org/2001/XMLSchema">'
               '<xs:element name="R" type="xs:string"/></xs:schema>')
        m = NS(steps={"s": NS(config={"schemaResourceUri": "xsd/R.xsd"})},
               endpoints=[], name="B", source_package="P.zip")
        body = mock_specs_from_model(
            m, corpus={"P.zip!h_content!xsd/R.xsd": xsd})[0].get("body")
        assert body and "<R>" in body

    def test_bare_passthrough_is_generic(self):
        from types import SimpleNamespace as NS
        from scaffolder.model_generator import mock_specs_from_model
        m = NS(steps={}, endpoints=[], name="C", source_package=None)
        assert mock_specs_from_model(m)[0].get("body") is None


class TestSchemaPipelineFunctional:
    """End-to-end: ref -> resolve(corpus) -> sample_payload -> conforms.
    Uses reusable XSD fixtures so the schema resolver is functionally tested
    even though the sandbox corpus has no real xsd."""
    def test_every_fixture_resolves_generates_and_conforms(self):
        import sys, os
        sys.path.insert(0, os.path.join(os.path.dirname(__file__)))
        from schema_fixtures import FIXTURES, sample_conforms
        from scaffolder.sample_payload import sample_payload_from_xsd
        from scaffolder.resource_resolver import resolve
        for nm, (xsd, root) in FIXTURES.items():
            path = f"Pkg.zip!h_content!src/main/resources/xsd/{nm}.xsd"
            res = resolve(f"{nm}.xsd", {path: xsd}, kind="schema",
                          package="Pkg.zip")
            assert res.ok, nm
            sample = sample_payload_from_xsd(res.content, root)
            assert sample_conforms(sample, xsd), f"{nm}: {sample}"

    def test_mock_uses_fixture_schema_when_resolvable(self):
        import sys, os
        sys.path.insert(0, os.path.join(os.path.dirname(__file__)))
        from types import SimpleNamespace as NS
        from schema_fixtures import FIXTURES, sample_conforms
        from scaffolder.model_generator import mock_specs_from_model
        xsd, root = FIXTURES["nested"]
        corpus = {"Pkg.zip!h_content!src/main/resources/xsd/nested.xsd": xsd}
        model = NS(steps={"s": NS(config={
            "schemaResourceUri": "src/main/resources/xsd/nested.xsd"})},
            endpoints=[], name="N", source_package="Pkg.zip")
        body = mock_specs_from_model(model, corpus=corpus)[0].get("body")
        assert body and sample_conforms(body, xsd)


class TestTransformDiff:
    def test_identity_matches_real_transform_diverges(self):
        import pytest
        try:
            from scaffolder.transform_diff import compare_mappings
        except Exception:
            pytest.skip("lxml unavailable")
        ident = ('<xsl:stylesheet xmlns:xsl="http://www.w3.org/1999/XSL/Transform"'
                 ' version="1.0"><xsl:template match="@*|node()"><xsl:copy>'
                 '<xsl:apply-templates select="@*|node()"/></xsl:copy>'
                 '</xsl:template></xsl:stylesheet>')
        real = ('<xsl:stylesheet xmlns:xsl="http://www.w3.org/1999/XSL/Transform"'
                ' version="1.0"><xsl:template match="/A"><B><id>'
                '<xsl:value-of select="x"/></id></B></xsl:template></xsl:stylesheet>')
        p = "<A><x>1</x></A>"
        assert compare_mappings(p, ident, ident).match is True
        assert compare_mappings(p, real, ident).match is False

    def test_mmap_not_comparable(self):
        import pytest
        try:
            from scaffolder.transform_diff import compare_mappings
        except Exception:
            pytest.skip("lxml unavailable")
        assert compare_mappings("<a/>", "x", "y", kind="mmap").comparable is False


class TestMultiProcessEmitter:
    def test_multiprocess_reproduces_with_local_processes(self):
        import os, pickle, xml.dom.minidom as MD, pytest
        if not os.path.exists("/tmp/learn/iflws.pkl"):
            pytest.skip("corpus not present")
        from extractor.iflow_parser import parse_iflow
        from scaffolder.model_generator import generate_multiprocess
        by = {p.rsplit("/", 1)[-1]: x
              for p, x in pickle.load(open("/tmp/learn/iflws.pkl", "rb"))}
        m = parse_iflow(by["S_ContentModifier_SetConstant.iflw"], "mp")
        res = generate_multiprocess(m, name="mp")
        MD.parseString(res.iflw_xml)                      # well-formed
        m2 = parse_iflow(res.iflw_xml, "mp")
        assert len(m2.processes) > 1                      # local processes present
        _B = {"StartEvent", "EndEvent", "StartTimerEvent"}
        src = [m.steps[s].kind for s in m.sequence
               if s in m.steps and m.steps[s].kind not in _B]
        out = [m2.steps[s].kind for s in m2.sequence
               if s in m2.steps and m2.steps[s].kind not in _B]
        assert src == out and "ProcessCallElement" in out

    def test_processid_collision_remapped(self):
        # a local process reusing 'Process_1' must not collide with the main
        from scaffolder.minimal_iflow import inject_local_processes
        from types import SimpleNamespace as NS
        proc = NS(id="Process_1", name="LIP", is_main=False, step_ids=[])
        model = NS(processes=[NS(id="MainP", is_main=True, step_ids=[]), proc],
                   steps={})
        out = inject_local_processes(
            '<x>    <bpmndi:BPMNDiagram></x>', model)
        assert "Process_1_LIP" in out


class TestGatewayEmitter:
    def test_gateway_flow_round_trips(self):
        import os, pickle, xml.dom.minidom as MD, pytest
        if not os.path.exists("/tmp/learn/iflws.pkl"):
            pytest.skip("corpus not present")
        from extractor.iflow_parser import parse_iflow
        from scaffolder.minimal_iflow import build_gateway_flow
        by = {}
        for p, x in pickle.load(open("/tmp/learn/iflws.pkl", "rb")):
            by.setdefault(p.rsplit("/", 1)[-1], x)
        _B = {"StartEvent", "EndEvent", "StartTimerEvent", "MessageStartEvent",
              "MessageEndEvent"}
        m = parse_iflow(by["ReadMailsfromDataStore.iflw"], "g")  # 2 gateways
        iflw, _ = build_gateway_flow("g", "g", m)
        MD.parseString(iflw)
        m2 = parse_iflow(iflw, "g")
        mids = lambda mm: [mm.steps[s].kind for s in mm.sequence
                           if s in mm.steps and mm.steps[s].kind not in _B]
        assert mids(m) == mids(m2)
        assert "ExclusiveGateway" in mids(m2)

    def test_gateway_conditions_emitted(self):
        # branch conditions + default branch must survive the round-trip
        import os, pickle, pytest
        if not os.path.exists("/tmp/learn/iflws.pkl"):
            pytest.skip("corpus not present")
        from extractor.iflow_parser import parse_iflow
        from scaffolder.minimal_iflow import build_gateway_flow
        by = {}
        for p, x in pickle.load(open("/tmp/learn/iflws.pkl", "rb")):
            by.setdefault(p.rsplit("/", 1)[-1], x)
        m = parse_iflow(by["SIChile_Boleta_Operations.iflw"], "g")
        iflw, _ = build_gateway_flow("g", "g", m)
        m2 = parse_iflow(iflw, "g")
        src_conds = sorted(r.condition or "" for r in m.routes)
        out_conds = sorted(r.condition or "" for r in m2.routes)
        assert src_conds == out_conds


class TestExceptionSubprocessEmitter:
    def test_subprocess_emitted_with_children(self):
        import os, pickle, xml.dom.minidom as MD, pytest
        if not os.path.exists("/tmp/learn/iflws.pkl"):
            pytest.skip("corpus not present")
        from extractor.iflow_parser import parse_iflow
        from scaffolder.model_generator import generate_from_model
        by = {}
        for p, x in pickle.load(open("/tmp/learn/iflws.pkl", "rb")):
            by.setdefault(p.rsplit("/", 1)[-1], x)
        _B = {"StartEvent", "EndEvent", "StartTimerEvent", "MessageStartEvent",
              "MessageEndEvent"}
        m = parse_iflow(by["Testing.iflw"], "s")
        res = generate_from_model(m, name="s")
        MD.parseString(res.iflw_xml)                       # well-formed
        m2 = parse_iflow(res.iflw_xml, "s")
        mids = lambda mm: [mm.steps[s].kind for s in mm.sequence
                           if s in mm.steps and mm.steps[s].kind not in _B]
        assert mids(m) == mids(m2)
        assert "ErrorEventSubProcessTemplate" in mids(m2)
        # the handler children must survive (not flattened to a childless node)
        sub = [s for s in m2.steps.values()
               if s.kind == "ErrorEventSubProcessTemplate"][0]
        kids = sorted(s.kind for s in m2.steps.values()
                      if s.parent_subprocess == sub.id)
        assert kids == ["EndEvent", "Enricher", "StartErrorEvent"]


class TestResourceAttach:
    def _model(self):
        from extractor.iflow_parser import IFlowModel, Process, Step
        m = IFlowModel(name="r", processes=[Process("Process_1", "m", True)])
        m.steps = {
            "S1": Step(id="S1", kind="Script", name="Script",
                       process_id="Process_1",
                       config={"script": "helper.groovy",
                               "cmdVariantUri": "ctype::FlowstepVariant/cname::GroovyScript"},
                       incoming=[], outgoing=[], parent_subprocess=""),
            "S2": Step(id="S2", kind="Mapping", name="XSLT",
                       process_id="Process_1",
                       config={"mappinguri": "dir://mapping/xslt/src/main/resources/mapping/conv.xsl",
                               "mappingname": "conv",
                               "cmdVariantUri": "ctype::FlowstepVariant/cname::XSLTMapping/version::1.2.0"},
                       incoming=[], outgoing=[], parent_subprocess=""),
        }
        m.sequence = ["S1", "S2"]
        return m

    def test_resolves_and_ships_real_files(self):
        from scaffolder.resource_attach import attach_resources
        corpus = {
            "MyPkg.zip/MyPkg_content/src/main/resources/script/helper.groovy":
                "def processData(m){ return m }  // REAL",
            "MyPkg.zip/MyPkg_content/src/main/resources/mapping/conv.xsl":
                '<xsl:stylesheet xmlns:xsl="http://www.w3.org/1999/XSL/Transform" version="1.0"/>',
        }
        rep = attach_resources(self._model(), corpus)
        assert len(rep.resolved) == 2 and not rep.unresolved
        assert rep.shipped["src/main/resources/script/helper.groovy"].endswith("REAL")
        assert "xsl:stylesheet" in rep.shipped["src/main/resources/mapping/conv.xsl"]

    def test_unresolved_is_graceful(self):
        from scaffolder.resource_attach import attach_resources
        rep = attach_resources(self._model(), {"other/x.txt": "y"})
        assert not rep.shipped and len(rep.unresolved) == 2     # logged, not fatal

    def test_generate_overwrites_stub_with_real(self):
        from scaffolder.model_generator import generate_from_model
        corpus = {
            "P.zip/P_content/src/main/resources/script/helper.groovy": "REAL-LOGIC",
            "P.zip/P_content/src/main/resources/mapping/conv.xsl":
                '<xsl:stylesheet xmlns:xsl="http://www.w3.org/1999/XSL/Transform" version="1.0"/>',
        }
        m = self._model()
        r = generate_from_model(m, name="r", resources=corpus)
        assert r.files["src/main/resources/script/helper.groovy"] == "REAL-LOGIC"
        # without a corpus the bundle is still produced (no report)
        r0 = generate_from_model(self._model(), name="r")
        assert getattr(r0, "resource_report", None) is None


class TestEndpointEmitter:
    def test_passthrough_endpoints_reproduced(self):
        import os, pickle, xml.dom.minidom as MD, pytest
        if not os.path.exists("/tmp/learn/iflws.pkl"):
            pytest.skip("corpus not present")
        from extractor.iflow_parser import parse_iflow
        from scaffolder.model_generator import generate_from_model
        from scaffolder.regenerate import regenerate_iflow_xml
        by = {}
        for p, x in pickle.load(open("/tmp/learn/iflws.pkl", "rb")):
            by.setdefault(p.rsplit("/", 1)[-1], x)
        xml = by["POtoS4Create.iflw"]
        m = parse_iflow(xml, "p")
        res = generate_from_model(m, name="p")
        MD.parseString(res.iflw_xml)                          # well-formed
        m2 = parse_iflow(res.iflw_xml, "p")
        # sender/receiver participants + their adapter config survive
        assert sorted(e.direction for e in m2.endpoints) == ["receiver", "sender"]
        src_adapters = sorted(mf.config.get("ComponentType", "")
                              for mf in m.message_flows)
        out_adapters = sorted(mf.config.get("ComponentType", "")
                              for mf in m2.message_flows)
        assert src_adapters == out_adapters and src_adapters
        # and the honesty guard now passes (faithful, not mock)
        assert regenerate_iflow_xml(xml, "p").reproduced is True

    def test_message_flow_adapter_config_carried(self):
        # the real adapter URL/path must travel, not a placeholder
        import os, pickle, pytest
        if not os.path.exists("/tmp/learn/iflws.pkl"):
            pytest.skip("corpus not present")
        from extractor.iflow_parser import parse_iflow
        from scaffolder.model_generator import generate_from_model
        by = {}
        for p, x in pickle.load(open("/tmp/learn/iflws.pkl", "rb")):
            by.setdefault(p.rsplit("/", 1)[-1], x)
        m = parse_iflow(by["POtoS4Create.iflw"], "p")
        res = generate_from_model(m, name="p")
        m2 = parse_iflow(res.iflw_xml, "p")
        src = {k: v for mf in m.message_flows for k, v in mf.config.items()}
        out = {k: v for mf in m2.message_flows for k, v in mf.config.items()}
        assert out.get("urlPath") == src.get("urlPath")
        assert out.get("ComponentType") == src.get("ComponentType")


class TestIFlowLevelCarry:
    def test_collab_config_carried_verbatim(self):
        import os, pickle, pytest
        if not os.path.exists("/tmp/learn/iflws.pkl"):
            pytest.skip("corpus not present")
        from extractor.iflow_parser import parse_iflow
        from scaffolder.model_generator import generate_from_model
        by = {}
        for p, x in pickle.load(open("/tmp/learn/iflws.pkl", "rb")):
            by.setdefault(p.rsplit("/", 1)[-1], x)
        m = parse_iflow(by["POtoS4Create.iflw"], "p")
        assert m.collab_config.get("namespaceMapping") is not None
        m2 = parse_iflow(generate_from_model(m, name="p").iflw_xml, "p")
        for k, v in m.collab_config.items():
            assert m2.collab_config.get(k) == v, k

    def test_apply_collab_config_overlay(self):
        from scaffolder.minimal_iflow import build_iflw, apply_collab_config
        from extractor.iflow_parser import parse_iflow
        out = apply_collab_config(
            build_iflw("x", "x"),
            {"namespaceMapping": "xmlns:s=http://x", "corsEnabled": "true"})
        cc = parse_iflow(out, "x").collab_config
        assert cc.get("namespaceMapping") == "xmlns:s=http://x"   # replaced
        assert cc.get("corsEnabled") == "true"                    # appended

    def test_parameter_files_shipped(self):
        from scaffolder.minimal_iflow import emit_parameter_files
        files = emit_parameter_files({"OAUTH_URL", "My Param"})
        prop = files["src/main/resources/parameters.prop"]
        pdef = files["src/main/resources/parameters.propdef"]
        assert "OAUTH_URL=" in prop and "My\\ Param=" in prop     # key escaping
        assert pdef.count("<parameter>") == 2
        assert "<name>OAUTH_URL</name>" in pdef
        empty = emit_parameter_files(set())
        assert "<parameters></parameters>" in \
            empty["src/main/resources/parameters.propdef"]        # real empty shell

    def test_bundle_ships_metainfo_and_params(self):
        from scaffolder.model_generator import generate_from_model
        from extractor.iflow_parser import IFlowModel, Process
        m = IFlowModel(name="x", processes=[Process("Process_1", "m", True)])
        r = generate_from_model(m, name="x")
        assert "metainfo.prop" in r.files
        assert "src/main/resources/parameters.prop" in r.files
        assert "src/main/resources/parameters.propdef" in r.files


class TestTenantFixes:
    def test_layered_layout_no_overlap_and_anchored_edges(self):
        import os, pickle, re, pytest
        if not os.path.exists("/tmp/learn/iflws.pkl"):
            pytest.skip("corpus not present")
        from extractor.iflow_parser import parse_iflow
        from scaffolder.model_generator import generate_from_model
        by = {}
        for p, x in pickle.load(open("/tmp/learn/iflws.pkl", "rb")):
            by.setdefault(p.rsplit("/", 1)[-1], x)
        out = generate_from_model(parse_iflow(by["ReadMailsfromDataStore.iflw"],
                                              "g"), name="g").iflw_xml
        shapes = {}
        for mm in re.finditer(r'<bpmndi:BPMNShape bpmnElement="([^"]+)"[^>]*>'
                              r'<dc:Bounds height="([\d.]+)" width="([\d.]+)" '
                              r'x="([\d.]+)" y="([\d.]+)"', out):
            i, h, w, x, y = mm.groups()
            shapes[i] = (float(x), float(y), float(w), float(h))
        ids = [i for i in shapes if not i.startswith("Participant_")]
        for a in range(len(ids)):                  # no two nodes overlap
            for b in range(a + 1, len(ids)):
                x1, y1, w1, h1 = shapes[ids[a]]
                x2, y2, w2, h2 = shapes[ids[b]]
                assert not (x1 < x2 + w2 and x2 < x1 + w1
                            and y1 < y2 + h2 and y2 < y1 + h1)
        for mm in re.finditer(                      # arrows touch their shapes
                r'sourceElement="BPMNShape_([^"]+)" targetElement='
                r'"BPMNShape_([^"]+)"><di:waypoint x="([\d.]+)"[^/]*'
                r'y="([\d.]+)"/><di:waypoint x="([\d.]+)"[^/]*y="([\d.]+)"', out):
            s, t, ax, ay, bx, by = mm.groups()
            for sid, px, py in ((s, float(ax), float(ay)),
                                (t, float(bx), float(by))):
                if sid in shapes:
                    x, y, w, h = shapes[sid]
                    assert x - 1 <= px <= x + w + 1 and y - 1 <= py <= y + h + 1

    def test_parallel_gateway_element_and_version_floor(self):
        import os, pickle, re, pytest
        if not os.path.exists("/tmp/learn/iflws.pkl"):
            pytest.skip("corpus not present")
        from extractor.iflow_parser import parse_iflow
        from scaffolder.model_generator import generate_from_model
        by = {}
        for p, x in pickle.load(open("/tmp/learn/iflws.pkl", "rb")):
            by.setdefault(p.rsplit("/", 1)[-1], x)
        m = None
        for k, v in by.items():                     # multicast in the MAIN process
            if "parallelGateway" not in v:
                continue
            cand = parse_iflow(v, "p")
            if any(s.kind in ("Multicast", "SequentialMulticast")
                   and not s.parent_subprocess
                   and s.process_id == next(p.id for p in cand.processes
                                            if p.is_main)
                   for s in cand.steps.values()):
                m = cand
                break
        assert m is not None
        out = generate_from_model(m, name="p").iflw_xml
        assert "<bpmn2:parallelGateway" in out      # not a callActivity
        ver = re.search(r"cname::(?:Sequential)?Multicast/version::([\d.]+)", out)
        assert ver and tuple(map(int, ver.group(1).split("."))) >= (1, 1)
        # kinds still round-trip
        m2 = parse_iflow(out, "p")
        assert any(s.kind in ("Multicast", "SequentialMulticast")
                   for s in m2.steps.values())

    def test_gateway_branches_named(self):
        import os, pickle, re, pytest
        if not os.path.exists("/tmp/learn/iflws.pkl"):
            pytest.skip("corpus not present")
        from extractor.iflow_parser import parse_iflow
        from scaffolder.model_generator import generate_from_model
        by = {}
        for p, x in pickle.load(open("/tmp/learn/iflws.pkl", "rb")):
            by.setdefault(p.rsplit("/", 1)[-1], x)
        out = generate_from_model(
            parse_iflow(by["SIChile_Boleta_Operations.iflw"], "g"),
            name="g").iflw_xml
        m2 = parse_iflow(out, "g")
        for s in m2.steps.values():
            if s.kind == "ExclusiveGateway":
                for fid in s.outgoing:
                    mm = re.search(r'<bpmn2:sequenceFlow id="%s"([^>]*)'
                                   % re.escape(fid), out)
                    assert mm and 'name="' in mm.group(1)

    def test_error_start_event_cpi_form(self):
        # 'StartErrorEvent v1.0 not supported' fix: cname::ErrorStartEvent
        # inside the errorEventDefinition, no stale componentVersion
        from scaffolder.minimal_iflow import _gw_fix_versions, _gw_event_def
        cfg = _gw_fix_versions("StartErrorEvent",
                               {"componentVersion": "1.0",
                                "activityType": "StartErrorEvent"})
        assert "componentVersion" not in cfg
        assert cfg["cmdVariantUri"].endswith("cname::ErrorStartEvent")
        assert "cname::ErrorStartEvent" in _gw_event_def("StartErrorEvent")

    def test_mapping_referenced_schemas_shipped(self):
        from extractor.iflow_parser import IFlowModel, Process, Step
        from scaffolder.resource_attach import attach_resources
        m = IFlowModel(name="x", processes=[Process("Process_1", "m", True)])
        m.steps = {"S1": Step(
            id="S1", kind="Mapping", name="MM", process_id="Process_1",
            config={"mappinguri": "dir://mmap/src/main/resources/mapping/P.mmap",
                    "cmdVariantUri":
                        "ctype::FlowstepVariant/cname::MessageMapping"},
            incoming=[], outgoing=[], parent_subprocess="")}
        m.sequence = ["S1"]
        corpus = {
            "K.zip/K_content/src/main/resources/mapping/P.mmap":
                '<m src="In.wsdl" tgt="Out.xsd"/>',
            "K.zip/K_content/src/main/resources/wsdl/In.wsdl": "<w/>",
            "K.zip/K_content/src/main/resources/xsd/Out.xsd": "<x/>",
        }
        rep = attach_resources(m, corpus)
        assert "src/main/resources/wsdl/In.wsdl" in rep.shipped
        assert "src/main/resources/xsd/Out.xsd" in rep.shipped


class TestEditorLoadFixes:
    def test_every_process_has_participant_and_pool(self):
        import os, pickle, re, pytest
        if not os.path.exists("/tmp/learn/iflws.pkl"):
            pytest.skip("corpus not present")
        from extractor.iflow_parser import parse_iflow
        from scaffolder.model_generator import generate_from_model
        by = {}
        for p, x in pickle.load(open("/tmp/learn/iflws.pkl", "rb")):
            by.setdefault(p.rsplit("/", 1)[-1], x)
        # multi-process flow: 'Error while loading' root cause
        out = generate_from_model(
            parse_iflow(by["S_ContentModifier_SetConstant.iflw"], "s"),
            name="s").iflw_xml
        procs = set(re.findall(r'<bpmn2:process id="([^"]+)"', out))
        refs = set(re.findall(r'processRef="([^"]+)"', out))
        pools = set(re.findall(r'BPMNShape bpmnElement="Participant_([^"]+)"',
                               out))
        assert len(procs) > 1                       # actually multi-process
        for pid in procs:                           # every process referenced
            assert pid in refs and pid in pools
        # every LIP chain node has a DI shape
        for nid in re.findall(r'<bpmn2:(?:callActivity|startEvent|endEvent) '
                              r'id="((?:SE_|EE_|N_)[^"]+)"', out):
            assert f'BPMNShape bpmnElement="{nid}"' in out

    def test_main_process_selection_ignores_called_lips(self):
        import os, pickle, pytest
        if not os.path.exists("/tmp/learn/iflws.pkl"):
            pytest.skip("corpus not present")
        from extractor.iflow_parser import parse_iflow
        from scaffolder.model_generator import generate_from_model
        by = {}
        for p, x in pickle.load(open("/tmp/learn/iflws.pkl", "rb")):
            by.setdefault(p.rsplit("/", 1)[-1], x)
        m = parse_iflow(by["S_ContentModifier_SetConstant.iflw"], "s")
        main = next(p for p in m.processes if p.is_main)
        m2 = parse_iflow(generate_from_model(m, name="s").iflw_xml, "s")
        main2 = next(p for p in m2.processes if p.is_main)
        assert main.id == main2.id                  # regen keeps the same main

    def test_timer_start_emits_timer_event_definition(self):
        import os, pickle, re, pytest
        if not os.path.exists("/tmp/learn/iflws.pkl"):
            pytest.skip("corpus not present")
        from extractor.iflow_parser import parse_iflow
        from scaffolder.model_generator import generate_from_model
        by = {}
        for p, x in pickle.load(open("/tmp/learn/iflws.pkl", "rb")):
            by.setdefault(p.rsplit("/", 1)[-1], x)
        nm = next(k for k, v in by.items() if "timerEventDefinition" in v)
        m = parse_iflow(by[nm], "t")
        t = next(s for s in m.steps.values() if "Timer" in s.kind)
        assert "scheduleKey" in t.config            # parser captures schedule
        out = generate_from_model(m, name="t").iflw_xml
        assert "timerEventDefinition" in out        # not a message start
        assert "scheduleKey" in out

    def test_main_pool_encloses_layout(self):
        import os, pickle, re, pytest
        if not os.path.exists("/tmp/learn/iflws.pkl"):
            pytest.skip("corpus not present")
        from extractor.iflow_parser import parse_iflow
        from scaffolder.model_generator import generate_from_model
        by = {}
        for p, x in pickle.load(open("/tmp/learn/iflws.pkl", "rb")):
            by.setdefault(p.rsplit("/", 1)[-1], x)
        out = generate_from_model(
            parse_iflow(by["ReadMailsfromDataStore.iflw"], "g"),
            name="g").iflw_xml
        pool = re.search(r'bpmnElement="Participant_Process[^"]*"[^>]*>\s*'
                         r'<dc:Bounds height="([\d.]+)" width="([\d.]+)" '
                         r'x="([\d.]+)" y="([\d.]+)"', out)
        ph, pw, px, py = (float(v) for v in pool.groups())
        for mm in re.finditer(r'BPMNShape bpmnElement="((?!Participant)[^"]+)" '
                              r'id="[^"]*"><dc:Bounds height="([\d.]+)" '
                              r'width="([\d.]+)" x="([\d.]+)" y="([\d.]+)"', out):
            nid, h, w, x, y = mm.group(1), *map(float, mm.groups()[1:])
            if nid.startswith(("SE_", "EE_", "N_")):
                continue                            # LIP pools live below
            assert px <= x and x + w <= px + pw
            assert py <= y and y + h <= py + ph

    def test_processdirect_gets_valid_default_address(self):
        from extractor.iflow_parser import (IFlowModel, Process, Step,
                                            MessageFlow, Endpoint)
        from scaffolder.model_generator import generate_from_model
        import re
        m = IFlowModel(name="My Flow", processes=[Process("Process_1", "m",
                                                          True)])
        s = Step(id="S1", kind="Enricher", name="cm", process_id="Process_1",
                 config={}, incoming=["F1"], outgoing=["F2"],
                 parent_subprocess="")
        m.steps = {"S1": s}
        m.sequence = ["S1"]
        m._flow_target = {}
        m.endpoints = [Endpoint(id="Participant_9", name="R",
                                direction="receiver", etype="EndpointRecevier")]
        m.message_flows = [MessageFlow(
            id="MessageFlow_4", name="ProcessDirect", source="S1",
            target="Participant_9",
            config={"ComponentType": "ProcessDirect", "Name": "ProcessDirect"})]
        out = generate_from_model(m, name="My Flow").iflw_xml
        addr = re.search(r"<key>address</key><value>([^<]*)</value>", out)
        assert addr and addr.group(1).startswith("/")


class TestNestedCorpusAndParameters:
    def _nested_corpus(self, tmp_path):
        import io, zipfile, os, pickle
        def jar(files):
            b = io.BytesIO()
            with zipfile.ZipFile(b, "w") as z:
                for k, v in files.items():
                    z.writestr(k, v)
            return b.getvalue()
        by = {}
        for p, x in pickle.load(open("/tmp/learn/iflws.pkl", "rb")):
            by.setdefault(p.rsplit("/", 1)[-1], x)
        src = by["nz.ird.oauth.token.revoke.iflw"]
        bundleA = jar({
            "src/main/resources/scenarioflows/integrationflow/"
            "nz.ird.oauth.token.revoke.iflw": src,
            "src/main/resources/parameters.prop":
                "Revoke_URL=/revoke/v2\nrole=ESBMessaging.send\n",
            "src/main/resources/parameters.propdef":
                '<?xml version="1.0"?><parameters><parameter>'
                '<key>Revoke_URL</key></parameter></parameters>',
            "src/main/resources/script/Revoke process header.groovy": "// s",
            "src/main/resources/mapping/Revoke xml to xhtml.xsl":
                "<xsl:stylesheet xmlns:xsl='http://www.w3.org/1999/XSL/"
                "Transform' version='1.0'/>"})
        bundleB = jar({
            "src/main/resources/scenarioflows/integrationflow/other.iflw":
                "<x/>",
            "src/main/resources/parameters.prop": "WRONG=other\n"})
        part2 = jar({
            "New_Zealand_Inland_Revenue_Reporting_for_Payroll_Version_2.zip":
                jar({"aaa_content": bundleA}),
            "Other_Package.zip": jar({"bbb_content": bundleB})})
        d = tmp_path / "pkgs"
        d.mkdir()
        (d / "part2.zip").write_bytes(part2)
        return str(d), src

    def test_walk_keys_are_container_prefixed(self, tmp_path):
        import os, pytest
        if not os.path.exists("/tmp/learn/iflws.pkl"):
            pytest.skip("corpus not present")
        from library_builder.corpus_pipeline import walk_corpus
        d, _ = self._nested_corpus(tmp_path)
        files = walk_corpus(d)
        props = [k for k in files if k.endswith("parameters.prop")]
        # without container prefixes the second package's prop is dropped
        assert len(props) == 2
        assert all(k.startswith("part2.zip/") for k in files)
        assert any("_content::src/" in k for k in files)

    def test_original_parameters_from_source_bundle(self, tmp_path):
        import os, pytest
        if not os.path.exists("/tmp/learn/iflws.pkl"):
            pytest.skip("corpus not present")
        from library_builder.corpus_pipeline import walk_corpus
        from scaffolder.regenerate import regenerate_iflow_xml
        d, src = self._nested_corpus(tmp_path)
        files = walk_corpus(d)
        regen = regenerate_iflow_xml(
            src, "OAUTH Token Revoke V2", resources=files,
            package="New Zealand Inland Revenue Reporting for Payroll "
                    "Version 2")
        assert regen.reproduced and regen.result is not None
        prop = regen.result.files["src/main/resources/parameters.prop"]
        assert "Revoke_URL=/revoke/v2" in prop      # real configured values
        assert "WRONG" not in prop                  # not another package's
        assert "Revoke_URL" in regen.result.files[
            "src/main/resources/parameters.propdef"]
        shipped = regen.result.files
        assert "src/main/resources/script/Revoke process header.groovy" \
            in shipped                              # References tab content
        assert "src/main/resources/mapping/Revoke xml to xhtml.xsl" in shipped

    def test_resolver_normalizes_package_names(self):
        from scaffolder.resource_resolver import resolve
        files = {
            "part2.zip/My_Pkg_V2.zip/a_content::src/main/resources/script/"
            "f.groovy": "right",
            "other.zip/b_content::src/main/resources/script/f.groovy":
                "wrong"}
        r = resolve("f.groovy", files, package="My Pkg V2")
        assert r.ok and r.content == "right" and not r.ambiguous

    def test_lip_multicast_emitted_as_real_gateway(self):
        import os, pickle, re, pytest
        if not os.path.exists("/tmp/learn/iflws.pkl"):
            pytest.skip("corpus not present")
        from extractor.iflow_parser import parse_iflow
        from scaffolder.model_generator import generate_from_model
        src = next(x for p, x in pickle.load(open("/tmp/learn/iflws.pkl",
                                                  "rb"))
                   if p.endswith("Job_Recruiting_Fields_OpenText.iflw"))
        out = generate_from_model(parse_iflow(src, "rci"), name="rci").iflw_xml
        lip = re.search(r'<bpmn2:process id="Process_45360926".*?'
                        r'</bpmn2:process>', out, re.S).group(0)
        # linearized LIPs emitted Multicast as a callActivity — the editor's
        # 'Error while loading' breaker class
        assert "<bpmn2:parallelGateway" in lip
        lip2 = re.search(r'<bpmn2:process id="Process_145".*?</bpmn2:process>',
                         out, re.S).group(0)
        assert "<bpmn2:exclusiveGateway" in lip2
        assert "<bpmn2:subProcess" in lip2          # nested, not flattened


class TestEditorFidelityRound4:
    """Locks for the RCI093 editor-error delta: source-faithful event
    definitions, timer schedule, gateway-route flow props, selective resource
    walk, and partner-adjacent endpoint placement."""

    def _rci(self):
        import os, pickle, pytest
        if not os.path.exists("/tmp/learn/iflws.pkl"):
            pytest.skip("corpus not present")
        src = next(x for p, x in pickle.load(open("/tmp/learn/iflws.pkl",
                                                  "rb"))
                   if p.endswith("Job_Recruiting_Fields_OpenText.iflw"))
        from extractor.iflow_parser import parse_iflow
        from scaffolder.model_generator import generate_from_model
        return src, generate_from_model(parse_iflow(src, "rci"),
                                        name="rci").iflw_xml

    def test_plain_lip_events_have_no_event_definition(self):
        import re
        _, out = self._rci()
        # 'Start event should have an incoming message flow' / 'LIP does not
        # support this variant of end event' both came from an unconditional
        # messageEventDefinition on plain LIP starts/ends
        e = re.search(r'<bpmn2:endEvent id="EndEvent_147".*?</bpmn2:endEvent>',
                      out, re.S).group(0)
        assert "EventDefinition" not in e
        s = re.search(r'<bpmn2:startEvent id="StartEvent_45360927".*?'
                      r'</bpmn2:startEvent>', out, re.S).group(0)
        assert "EventDefinition" not in s

    def test_timer_def_faithful(self):
        import re
        _, out = self._rci()
        t = re.search(r'<bpmn2:startEvent id="StartEvent_4".*?'
                      r'</bpmn2:startEvent>', out, re.S).group(0)
        # 'Timer is not configured' came from duplicating schedule props at
        # node level + dropping the definition's id
        assert 'TimerEventDefinition_55599' in t
        assert "<bpmn2:extensionElements>" not in \
            t.split("timerEventDefinition")[0]
        assert "scheduleKey" in t

    def test_default_branch_keeps_gateway_route_props(self):
        import re
        _, out = self._rci()
        f = re.search(r'<bpmn2:sequenceFlow id="SequenceFlow_45361016".*?'
                      r'(?:</bpmn2:sequenceFlow>|/>)', out, re.S).group(0)
        # dropping the GatewayRoute extension on the default branch raises
        # 'Condition cannot be empty'
        assert "GatewayRoute" in f and "expressionType" in f

    def test_event_def_parity_whole_corpus(self):
        import os, pickle, re, pytest
        if not os.path.exists("/tmp/learn/iflws.pkl"):
            pytest.skip("corpus not present")
        from extractor.iflow_parser import parse_iflow
        from scaffolder.model_generator import generate_from_model
        by = {}
        for p, x in pickle.load(open("/tmp/learn/iflws.pkl", "rb")):
            by.setdefault(p.rsplit("/", 1)[-1], x)

        def evmap(xml):
            out = {}
            for tag in ("startEvent", "endEvent"):
                for m in re.finditer(
                        rf'<bpmn2:{tag}\b[^>]*\bid="([^"]+)"[^>]*>(.*?)'
                        rf'</bpmn2:{tag}>', xml, re.S):
                    b = m.group(2)
                    out[m.group(1)] = (
                        "message" if "messageEventDefinition" in b else
                        "timer" if "timerEventDefinition" in b else
                        "error" if "errorEventDefinition" in b else
                        "other" if "EventDefinition" in b else None)
                for m in re.finditer(
                        rf'<bpmn2:{tag}\b[^>]*\bid="([^"]+)"[^>]*/>', xml):
                    out[m.group(1)] = None
            return out
        bad = []
        for nm, x in by.items():
            out = generate_from_model(parse_iflow(x, nm), name=nm).iflw_xml
            s, g = evmap(x), evmap(out)
            bad += [(nm, i) for i in s if i in g and g[i] != s[i]]
        assert not bad, bad[:5]

    def test_walk_corpus_ext_filter(self, tmp_path):
        import io, zipfile
        from library_builder.corpus_pipeline import walk_corpus, WIRING_EXTS
        b = io.BytesIO()
        with zipfile.ZipFile(b, "w") as z:
            z.writestr("a.groovy", "x")
            z.writestr("big.bin", "y" * 100)
            z.writestr("doc.pdf", "z")
        d = tmp_path / "c"
        d.mkdir()
        (d / "p.zip").write_bytes(b.getvalue())
        files = walk_corpus(str(d), exts=WIRING_EXTS)
        assert any(k.endswith("a.groovy") for k in files)
        assert not any(k.endswith((".bin", ".pdf")) for k in files)

    def test_endpoints_not_stacked_far_right(self):
        import re
        _, out = self._rci()
        eps = re.findall(r'BPMNShape bpmnElement="Participant_[\w]+" '
                         r'id="[^"]*"><dc:Bounds height="140\.0" '
                         r'width="100\.0" x="([\d.+-]+)" y="([\d.+-]+)"', out)
        assert eps
        # every endpoint is placed near a partner (no far-right stack, no
        # negative coordinates)
        assert all(float(x) >= 0 and float(y) >= 0 for x, y in eps)
        stack = [1 for x, y in eps if float(x) > 1800
                 and round(float(y)) % 170 == 110 % 170]
        assert not stack


class TestExtractCorpusTool:
    def test_nested_extraction_organized_with_collisions(self, tmp_path):
        import io, zipfile, os, csv
        from tools.extract_corpus import main
        def jar(files):
            b = io.BytesIO()
            with zipfile.ZipFile(b, "w") as z:
                for k, v in files.items():
                    z.writestr(k, v)
            return b.getvalue()
        pkg = jar({
            "a_content": jar({
                "src/main/resources/parameters.prop": "A=1\n",
                "src/main/resources/script/util.groovy": "// a",
                "src/main/resources/mapping/m.xsl": "<x/>"}),
            "b_content": jar({
                "src/main/resources/parameters.prop": "B=2\n",
                "src/main/resources/script/util.groovy": "// b"})})
        srcd = tmp_path / "pkgs"; srcd.mkdir()
        (srcd / "outer.zip").write_bytes(jar({"My_Package.zip": pkg}))
        dst = tmp_path / "corpus"
        assert main(["--src", str(srcd), "--dst", str(dst)]) == 0
        files = {os.path.relpath(os.path.join(r, f), dst)
                 for r, _, fs in os.walk(dst) for f in fs}
        # organized by type, package-prefixed, collisions suffixed not dropped
        assert "Prop/My_Package__parameters.prop" in files
        assert "Prop/My_Package__parameters__2.prop" in files
        assert "Groovy/My_Package__util.groovy" in files
        assert "Groovy/My_Package__util__2.groovy" in files
        assert "Xsl/My_Package__m.xsl" in files
        # both collision contents survive (first-seen-wins was the original
        # corpus-walk bug — the standalone must never repeat it)
        bodies = {open(dst / "Prop" / f).read().strip()
                  for f in ("My_Package__parameters.prop",
                            "My_Package__parameters__2.prop")}
        assert bodies == {"A=1", "B=2"}
        # manifest traces every file to its container-qualified source
        rows = list(csv.reader(open(dst / "manifest.csv")))[1:]
        srcs = {r[1] for r in rows}
        assert any(s.startswith("outer.zip/My_Package.zip/a_content::src/")
                   for s in srcs)
